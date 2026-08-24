from __future__ import annotations

from pathlib import Path
import tempfile
from unittest.mock import patch
import os
import json

from openpyxl import Workbook
from PIL import Image,ImageDraw
from docx import Document

from modules.idp_documental.repository import IDPRepository
from modules.idp_documental.services import attendance_official_payload, canonicalize, classify_document, connect, read_document, read_document_azure, read_document_ocr, sha256_file
from modules.idp_documental.worker import process_next


def require(condition, message):
    if not condition:
        raise AssertionError(message)


def create(repo, tenant, path, user=10):
    digest=sha256_file(path)
    document_id=repo.create_document({'fundacion_id':tenant,'nombre_original':path.name,'nombre_guardado':path.name,'ruta_privada':str(path),'extension':path.suffix,'mime_type':'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet','tamano_bytes':path.stat().st_size,'sha256':digest,'usuario_id':user})
    raw=read_document(path); classification=classify_document(raw.get('texto') or '',path.name); canonical,fields=canonicalize(raw,classification[0])
    canonical['fundacion']['id']=tenant
    repo.complete_extraction(document_id,tenant,raw,canonical,fields,classification,user)
    return document_id,digest


def main():
    with tempfile.TemporaryDirectory() as temporary:
        root=Path(temporary); db=root/'idp.sqlite3'; book=root/'LISTADO_ASISTENCIA.xlsx'
        wb=Workbook(); ws=wb.active; ws.title='ASISTENCIA'; ws.append(['LISTADO DE ASISTENCIA']); ws.append(['Nombre completo','Documento','UDS','Asistió','Firma']); ws.append(['ANA PEREZ','1001','UCA 1','SI','X']); ws.append(['LUIS DIAZ','1002','UCA 1','NO','']); wb.save(book)
        repo=IDPRepository(str(db))
        conn=connect(str(db))
        conn.execute("CREATE TABLE master_ninos(id INTEGER PRIMARY KEY,documento TEXT,nombre_completo TEXT,unidad_servicio TEXT,estado TEXT,activo INTEGER,fundacion_id INTEGER)")
        conn.execute("INSERT INTO master_ninos VALUES(1,'1001','ANA PEREZ','UCA 1','ACTIVO',1,1)")
        conn.execute("INSERT INTO master_ninos VALUES(2,'1002','LUIS DIAZ','UCA 1','ACTIVO',1,1)")
        conn.execute("INSERT INTO master_ninos VALUES(3,'1001','OTRA FUNDACION','UCA X','ACTIVO',1,2)")
        conn.execute("CREATE TABLE plantillas_oficiales_versiones(id INTEGER PRIMARY KEY,tipo_formato TEXT,codigo TEXT,nombre TEXT,version TEXT,fecha_vigencia TEXT,fecha_vigencia_fin TEXT,estado TEXT,hash_sha256 TEXT,mapeo_json TEXT,fundacion_id INTEGER,updated_at TEXT)")
        conn.execute("INSERT INTO plantillas_oficiales_versiones VALUES(1,'LISTADO_ASISTENCIA_USUARIOS','ICBF-ASIS','Asistencia oficial','2026.03','2026-03-01',NULL,'vigente','hash-tenant-1','{\"campos\":{\"documento\":{}}}',1,'2026-03-01')")
        conn.execute("INSERT INTO plantillas_oficiales_versiones VALUES(2,'LISTADO_ASISTENCIA_USUARIOS','ICBF-ASIS','Otra fundacion','2099.01','2099-01-01',NULL,'vigente','hash-tenant-2','{}',2,'2099-01-01')")
        ram_mapping=[{'field':'tipo_documento','sheet':'FORMATO RAM','col_letter':'B','col':2,'data_start_row':15,'fila_fin':16},{'field':'documento_beneficiario','sheet':'FORMATO RAM','col_letter':'C','col':3,'data_start_row':15,'fila_fin':16},{'field':'primer_nombre','sheet':'FORMATO RAM','col_letter':'D','col':4,'data_start_row':15,'fila_fin':16},{'field':'primer_apellido','sheet':'FORMATO RAM','col_letter':'F','col':6,'data_start_row':15,'fila_fin':16},{'field':'control_asistencia','sheet':'FORMATO RAM','col_letter':'J:K','col':10,'data_start_row':15,'fila_fin':16},{'field':'total_asistencias','sheet':'FORMATO RAM','col_letter':'AI','col':35,'data_start_row':15,'fila_fin':16}]
        conn.execute("INSERT INTO plantillas_oficiales_versiones(id,tipo_formato,codigo,nombre,version,fecha_vigencia,estado,mapeo_json,fundacion_id,updated_at) VALUES(3,'RAM','F27.MT1.PP','RAM oficial','3','2026-08-01','vigente',?,1,'2026-08-01')",(json.dumps(ram_mapping),))
        conn.commit(); conn.close()
        document_id,digest=create(repo,1,book)
        item=repo.get_document(document_id,1)
        require(item and item['tipo_documento']=='LISTADO_ASISTENCIA','No clasifico asistencia')
        require(item['estado']=='REQUIERE_REVISION','Estado incorrecto')
        require(len(item['resultado_canonico']['participantes'])==2,'No extrajo participantes')
        require(item['resultado_canonico']['participantes'][0]['documento']=='1001','Documento mal mapeado')
        require(item['resultado_canonico']['version_plantilla']=='2026.03','No vinculo la version oficial del tenant')
        require(item['resultado_canonico']['metadatos']['plantilla_oficial']['id']==1,'Mezclo versiones oficiales entre fundaciones')
        require(item['validaciones']['semaforo']=='VERDE','La planilla valida no quedo en verde')
        require(item['validaciones']['coincidencias']==2,'No valido los participantes contra Base Maestra')
        generation_blocked=False
        try: attendance_official_payload(item)
        except ValueError: generation_blocked=True
        require(generation_blocked,'Permitio generar formato oficial sin aprobacion')
        require(repo.get_document(document_id,2) is None,'Fallo aislamiento por fundacion')
        require(repo.find_duplicate(1,digest)['id']==document_id,'No detecto duplicado del tenant')
        require(repo.find_duplicate(2,digest) is None,'Bloqueo incorrectamente el mismo archivo en otro tenant')
        name_field=next(field for field in item['campos'] if field['ruta_canonica']=='participantes.0.nombre_completo')
        repo.correct_field(document_id,name_field['id'],1,'ANA MARIA PEREZ',10,'Correccion contra original')
        corrected=repo.get_document(document_id,1)
        require(corrected['resultado_canonico']['participantes'][0]['nombre_completo']=='ANA MARIA PEREZ','Correccion no actualizo canonico')
        repo.approve(document_id,1,10)
        approved=repo.get_document(document_id,1)
        require(approved['estado']=='APROBADO' and approved['progreso']==100,'No aprobo documento revisado')
        official_users,official_metadata=attendance_official_payload(approved)
        require(len(official_users)==2 and official_users[0]['documento']=='1001','No preparo usuarios para el listado oficial')
        require(official_metadata['unidad']=='UCA 1','No preparo la UDS para el listado oficial')
        require(any(event['evento']=='CAMPO_CORREGIDO' for event in approved['eventos']),'No audito correccion')
        require(any(event['evento']=='DOCUMENTO_APROBADO' for event in approved['eventos']),'No audito aprobacion')
        imported=repo.import_attendance(document_id,1,10,'2026-08-20','Encuentro educativo')
        require(imported['total_registros']==2 and not imported['ya_importado'],'No importo el lote de asistencia')
        imported_again=repo.import_attendance(document_id,1,10,'2026-08-20','Encuentro educativo')
        require(imported_again['ya_importado'],'No hizo idempotente la segunda importacion')
        imported_doc=repo.get_document(document_id,1)
        require(imported_doc['estado']=='IMPORTADO','No actualizo el estado importado')
        conn=connect(str(db)); imported_rows=conn.execute('SELECT asistio FROM idp_asistencias_importadas WHERE documento_id=? AND fundacion_id=? ORDER BY indice_participante',(document_id,1)).fetchall(); attendance_count=len(imported_rows); other_tenant_count=conn.execute('SELECT COUNT(*) total FROM idp_asistencias_importadas WHERE documento_id=? AND fundacion_id=?',(document_id,2)).fetchone()['total']; conn.close()
        require(attendance_count==2 and other_tenant_count==0,'Fallo persistencia o aislamiento del lote importado')
        require([row['asistio'] for row in imported_rows]==[1,0],'Interpreto incorrectamente SI/NO al importar')
        require(any(event['evento']=='ASISTENCIA_IMPORTADA' for event in imported_doc['eventos']),'No audito la importacion')
        learned_book=root/'LISTADO_ASISTENCIA_APRENDIZAJE.xlsx'; learned_wb=Workbook(); learned_ws=learned_wb.active; learned_ws.append(['LISTADO DE ASISTENCIA - SEGUNDO DOCUMENTO']); learned_ws.append(['Nombre completo','Documento','UDS','Asistio']); learned_ws.append(['ANA PEREZ','1001','UCA 1','SI']); learned_wb.save(learned_book)
        learned_id,_=create(repo,1,learned_book); learned_doc=repo.get_document(learned_id,1); learned_name=next(field for field in learned_doc['campos'] if field['ruta_canonica']=='participantes.0.nombre_completo')
        require(learned_name['valor']=='ANA PEREZ','La memoria aplico una correccion automaticamente')
        require(learned_name['sugerencia_correccion']['valor']=='ANA MARIA PEREZ' and learned_name['sugerencia_correccion']['aplicacion']=='MANUAL','No sugirio la correccion aprobada')
        require(repo.get_document(learned_id,2) is None,'Expuso aprendizaje entre fundaciones')
        attendance_mapping={'hoja':'ASISTENCIA OFICIAL','fila_datos':5,'campos':{'nombre':2,'documento':3,'asistencia':4,'firma':5}}
        conn=connect(str(db)); conn.execute("INSERT INTO plantillas_oficiales_versiones(id,tipo_formato,codigo,nombre,version,fecha_vigencia,estado,mapeo_json,fundacion_id,updated_at) VALUES(4,'LISTADO_ASISTENCIA_USUARIOS','ICBF-ASIS','Asistencia mapeada','2026.08','2026-08-01','vigente',?,1,'2026-08-02')",(json.dumps(attendance_mapping),)); conn.commit(); conn.close()
        mapped_attendance=root/'LISTADO_ASISTENCIA_SIN_ENCABEZADOS.xlsx'; mapped_att_wb=Workbook(); mapped_att_ws=mapped_att_wb.active; mapped_att_ws.title='ASISTENCIA OFICIAL'; mapped_att_ws['A1']='LISTADO DE ASISTENCIA OFICIAL'; mapped_att_ws.cell(5,2,'ANA PEREZ'); mapped_att_ws.cell(5,3,'1001'); mapped_att_ws.cell(5,4,'SI'); mapped_att_ws.cell(5,5,'X'); mapped_att_wb.save(mapped_attendance)
        mapped_att_id,_=create(repo,1,mapped_attendance); mapped_att_doc=repo.get_document(mapped_att_id,1); mapped_att_participant=mapped_att_doc['resultado_canonico']['participantes'][0]
        require(mapped_att_participant['documento']=='1001' and mapped_att_participant['asistio'] is True and mapped_att_participant['firma_presente'] is True,'No aplico mapeo versionado de asistencia')
        queued_book=root/'LISTADO_ASISTENCIA_COLA.xlsx'
        queue_wb=Workbook(); queue_ws=queue_wb.active; queue_ws.append(['LISTADO DE ASISTENCIA EN COLA']); queue_ws.append(['Nombre completo','Documento','UDS','Asistio']); queue_ws.append(['ANA PEREZ','1001','UCA 1','SI']); queue_wb.save(queued_book)
        queue_id=repo.create_document({'fundacion_id':1,'nombre_original':queued_book.name,'nombre_guardado':queued_book.name,'ruta_privada':str(queued_book),'extension':'.xlsx','mime_type':'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet','tamano_bytes':queued_book.stat().st_size,'sha256':sha256_file(queued_book),'usuario_id':10})
        queued_job=repo.enqueue_extraction(queue_id,1)
        require(repo.get_document(queue_id,1)['estado']=='EN_COLA','No marco el documento en cola')
        require(process_next(str(db),'WORKER-PRUEBA'),'El worker no reclamo el trabajo pendiente')
        queue_doc=repo.get_document(queue_id,1)
        require(queue_doc['estado']=='REQUIERE_REVISION' and queue_doc['resultado_canonico']['participantes'][0]['documento']=='1001','El worker no completo la extraccion')
        conn=connect(str(db)); queue_state=conn.execute('SELECT estado,intentos FROM idp_trabajos_cola WHERE id=?',(queued_job['id'],)).fetchone(); conn.close()
        require(queue_state['estado']=='COMPLETADO' and queue_state['intentos']==1,'La cola no persistio el resultado del worker')
        schedule_book=root/'CRONOGRAMA_MENSUAL.xlsx'; schedule_wb=Workbook(); schedule_ws=schedule_wb.active; schedule_ws.append(['CRONOGRAMA MENSUAL']); schedule_ws.append(['Fecha','Actividad','Responsable','Entregable','Modulo']); schedule_ws.append(['20/08/2026','Encuentro familiar','Psicosocial','Acta y listado','Familias']); schedule_wb.save(schedule_book)
        schedule_id,_=create(repo,1,schedule_book); schedule_doc=repo.get_document(schedule_id,1)
        require(schedule_doc['tipo_documento']=='CRONOGRAMA' and len(schedule_doc['resultado_canonico']['actividades'])==1,'No estructuro el cronograma Excel')
        require(schedule_doc['resultado_canonico']['actividades'][0]['fecha']=='2026-08-20','No normalizo la fecha del cronograma')
        require(schedule_doc['validaciones']['semaforo']=='VERDE','Marco incorrectamente el cronograma completo')
        deliverables_classification=classify_document('RAM, listado de asistencia, informe y acta', 'entregables agosto.jpeg')
        require(deliverables_classification[0]=='CRONOGRAMA','Clasifico la imagen de entregables como un listado con participantes requeridos')
        schedule_word=root/'CRONOGRAMA_WORD.docx'; word=Document(); table=word.add_table(rows=2,cols=3)
        for column,value in enumerate(['Fecha','Actividad','Responsable']): table.rows[0].cells[column].text=value
        for column,value in enumerate(['21/08/2026','Taller pedagógico','Docente']): table.rows[1].cells[column].text=value
        word.save(schedule_word); word_id,_=create(repo,1,schedule_word); word_doc=repo.get_document(word_id,1)
        require(word_doc['tipo_documento']=='CRONOGRAMA' and word_doc['resultado_canonico']['actividades'][0]['responsable']=='Docente','No estructuro la tabla Word del cronograma')
        nutrition_book=root/'FORMATO_PESO_TALLA.xlsx'; nutrition_wb=Workbook(); nutrition_ws=nutrition_wb.active; nutrition_ws.append(['VALORACION NUTRICIONAL PESO KG TALLA CM']); nutrition_ws.append(['Fecha','Nombre completo','Documento','Peso kg','Talla cm','Perimetro braquial cm','UDS']); nutrition_ws.append(['20/08/2026','ANA PEREZ','1001','12,5','90.2','15','UCA 1']); nutrition_wb.save(nutrition_book)
        nutrition_id,_=create(repo,1,nutrition_book); nutrition_doc=repo.get_document(nutrition_id,1); valuation=nutrition_doc['resultado_canonico']['valoraciones'][0]
        require(nutrition_doc['tipo_documento']=='PESO_TALLA' and valuation['peso_kg']==12.5 and valuation['talla_cm']==90.2,'No estructuro peso y talla')
        require(valuation['validado_base_maestra'] and nutrition_doc['validaciones']['semaforo']=='VERDE','No valido la valoracion contra Base Maestra')
        invalid_nutrition=root/'FORMATO_PESO_TALLA_INVALIDO.xlsx'; invalid_nutrition_wb=Workbook(); invalid_nutrition_ws=invalid_nutrition_wb.active; invalid_nutrition_ws.append(['PESO KG TALLA CM VALORACION NUTRICIONAL']); invalid_nutrition_ws.append(['Fecha','Nombre','Documento','Peso','Talla']); invalid_nutrition_ws.append(['20/08/2026','ANA PEREZ','1001','999','10']); invalid_nutrition_wb.save(invalid_nutrition)
        invalid_nutrition_id,_=create(repo,1,invalid_nutrition); invalid_nutrition_doc=repo.get_document(invalid_nutrition_id,1)
        require(invalid_nutrition_doc['validaciones']['semaforo']=='ROJO' and invalid_nutrition_doc['validaciones']['errores_criticos']>=2,'No bloqueo rangos nutricionales imposibles')
        planning_word=root/'PLANEACION_PEDAGOGICA.docx'; planning=Document(); planning.add_heading('PLANEACION PEDAGOGICA',0); planning.add_paragraph('UDS: UCA 1'); planning.add_paragraph('Fecha programada: 22/08/2026'); planning.add_paragraph('Tema: Los colores de mi entorno'); planning.add_paragraph('Intencionalidad pedagógica: Fortalecer la exploración y expresión creativa'); planning.add_paragraph('Experiencia pedagógica: Recorrido guiado y creación con materiales del entorno'); planning.add_paragraph('Responsable: Docente titular'); planning.add_paragraph('Recursos: Papel, colores y elementos naturales'); planning.save(planning_word)
        planning_id,_=create(repo,1,planning_word); planning_doc=repo.get_document(planning_id,1); planning_data=planning_doc['resultado_canonico']['planeacion']
        require(planning_doc['tipo_documento']=='PLANEACION_PEDAGOGICA','No clasifico la planeacion pedagogica')
        require(planning_data['fecha_programada']=='2026-08-22' and 'exploración' in planning_data['objetivo'],'No estructuro los campos de planeacion')
        require(planning_doc['validaciones']['semaforo']=='VERDE','Marco incorrectamente la planeacion completa')
        incomplete_planning=root/'PLANEACION_PEDAGOGICA_INCOMPLETA.docx'; incomplete=Document(); incomplete.add_paragraph('PLANEACION PEDAGOGICA'); incomplete.add_paragraph('Tema: Actividad sin desarrollo'); incomplete.save(incomplete_planning)
        incomplete_id,_=create(repo,1,incomplete_planning); incomplete_doc=repo.get_document(incomplete_id,1)
        require(incomplete_doc['validaciones']['semaforo']=='ROJO' and incomplete_doc['validaciones']['errores_criticos']>=2,'No bloqueo planeacion sin intencionalidad ni experiencia')
        minutes_word=root/'ACTA_REUNION.docx'; minutes=Document(); minutes.add_heading('ACTA DE REUNION',0); minutes.add_paragraph('Fecha: 20/08/2026'); minutes.add_paragraph('Lugar: UCA 1'); minutes.add_paragraph('Tema: Seguimiento mensual'); minutes.add_paragraph('Asistentes: Docente, coordinador y familias'); minutes.add_paragraph('Desarrollo: Se revisaron avances y necesidades del grupo'); minutes.add_paragraph('Compromisos: Entregar evidencias el 25 de agosto'); minutes.add_paragraph('Responsable: Coordinación'); minutes.add_paragraph('Firmas: Se observan espacios diligenciados'); minutes.save(minutes_word)
        minutes_id,_=create(repo,1,minutes_word); minutes_doc=repo.get_document(minutes_id,1)
        require(minutes_doc['tipo_documento']=='ACTA' and minutes_doc['resultado_canonico']['acta']['fecha']=='2026-08-20','No estructuro el acta')
        require(minutes_doc['validaciones']['semaforo']=='VERDE','Marco incorrectamente el acta completa')
        report_word=root/'INFORME_MENSUAL.docx'; report=Document(); report.add_heading('INFORME MENSUAL',0); report.add_paragraph('Periodo: Agosto 2026'); report.add_paragraph('Objetivo: Presentar resultados del acompañamiento'); report.add_paragraph('Actividades realizadas: Visitas, talleres y seguimiento'); report.add_paragraph('Resultados: Participación de las familias y cumplimiento del cronograma'); report.add_paragraph('Conclusiones: Se cumplieron los objetivos programados'); report.add_paragraph('Responsable: Equipo interdisciplinario'); report.save(report_word)
        report_id,_=create(repo,1,report_word); report_doc=repo.get_document(report_id,1)
        require(report_doc['tipo_documento']=='INFORME' and 'cumplieron' in report_doc['resultado_canonico']['informe']['conclusiones'],'No estructuro el informe')
        require(report_doc['validaciones']['semaforo']=='VERDE','Marco incorrectamente el informe completo')
        rpp_word=root/'RPP_RACION_PREPARADA.docx'; rpp=Document(); rpp.add_heading('RPP RACION PREPARADA',0); rpp.add_paragraph('Fecha: 20/08/2026'); rpp.add_paragraph('UDS: UCA 1'); rpp.add_paragraph('Tiempo de comida: Almuerzo'); rpp.add_paragraph('Preparación: Arroz, pollo guisado y ensalada'); rpp.add_paragraph('Minuta patrón: Ciclo 3 día 4'); rpp.add_paragraph('Número de porciones: 25'); rpp.add_paragraph('Responsable: Manipuladora de alimentos'); rpp.save(rpp_word)
        rpp_id,_=create(repo,1,rpp_word); rpp_doc=repo.get_document(rpp_id,1); rpp_data=rpp_doc['resultado_canonico']['rpp']
        require(rpp_doc['tipo_documento']=='RPP' and rpp_data['fecha']=='2026-08-20' and rpp_data['porciones']=='25','No estructuro el RPP')
        require(rpp_doc['validaciones']['semaforo']=='VERDE','Marco incorrectamente el RPP completo')
        incomplete_rpp=root/'RPP_INCOMPLETO.docx'; rpp_incomplete=Document(); rpp_incomplete.add_heading('RPP RACION PREPARADA',0); rpp_incomplete.add_paragraph('Fecha: 20/08/2026'); rpp_incomplete.save(incomplete_rpp)
        incomplete_rpp_id,_=create(repo,1,incomplete_rpp); incomplete_rpp_doc=repo.get_document(incomplete_rpp_id,1)
        require(incomplete_rpp_doc['validaciones']['semaforo']=='ROJO' and incomplete_rpp_doc['validaciones']['errores_criticos']>=4,'No bloqueo el RPP incompleto')
        ram_book=root/'FORMATO_RAM_V3.xlsx'; ram_wb=Workbook(); ram_ws=ram_wb.active; ram_ws.title='FORMATO RAM'; ram_ws.append(['FORMATO RAM F27 MT1 PP']); ram_ws.append(['Tipo documento','Documento beneficiario','Primer nombre','Segundo nombre','Primer apellido','Segundo apellido','Día 1','Día 2','Total asistencias','Total inasistencias','UDS']); ram_ws.append(['RC','1001','ANA','','PEREZ','','A','I','1','1','UCA 1']); ram_wb.save(ram_book)
        ram_id,_=create(repo,1,ram_book); ram_doc=repo.get_document(ram_id,1); ram_participant=ram_doc['resultado_canonico']['participantes'][0]
        require(ram_doc['tipo_documento']=='RAM' and ram_participant['nombre_completo']=='ANA PEREZ','No estructuro la identidad RAM')
        require(ram_participant['asistencia_dias']=={'1':'A','2':'I'} and ram_participant['total_asistencias']=='1','No conservo dias y totales RAM')
        require(ram_doc['validaciones']['semaforo']=='VERDE','No valido RAM contra Base Maestra')
        mapped_ram=root/'RAM_SIN_ENCABEZADOS.xlsx'; mapped_wb=Workbook(); mapped_ws=mapped_wb.active; mapped_ws.title='FORMATO RAM'; mapped_ws['A1']='FORMATO RAM F27 MT1 PP'; mapped_ws.cell(15,2,'RC'); mapped_ws.cell(15,3,'1001'); mapped_ws.cell(15,4,'ANA'); mapped_ws.cell(15,6,'PEREZ'); mapped_ws.cell(15,10,'A'); mapped_ws.cell(15,11,'I'); mapped_ws.cell(15,35,1); mapped_wb.save(mapped_ram)
        mapped_ram_id,_=create(repo,1,mapped_ram); mapped_ram_doc=repo.get_document(mapped_ram_id,1); mapped_participant=mapped_ram_doc['resultado_canonico']['participantes'][0]
        require(mapped_participant['documento']=='1001' and mapped_participant['asistencia_dias']=={'1':'A','2':'I'},'No aplico el mapeo oficial versionado')
        require(any(field['regla']=='mapeo_oficial_versionado' for field in mapped_ram_doc['campos']),'No audito evidencia del mapeo versionado')
        bien_book=root/'ENTREGA_BIENESTARINA.xlsx'; bien_wb=Workbook(); bien_ws=bien_wb.active; bien_ws.title='BIENESTARINA'; bien_ws.append(['ENTREGA DE ALIMENTO BIENESTARINA']); bien_ws.append(['Tipo documento','NUI','Primer nombre','Primer apellido','Fecha de entrega','Lote','Cantidad entregada','UDS']); bien_ws.append(['RC','1001','ANA','PEREZ','20/08/2026','LT-2026-08','2','UCA 1']); bien_wb.save(bien_book)
        bien_id,_=create(repo,1,bien_book); bien_doc=repo.get_document(bien_id,1); bien_delivery=bien_doc['resultado_canonico']['entregas'][0]
        require(bien_doc['tipo_documento']=='BIENESTARINA' and bien_delivery['fecha_entrega']=='2026-08-20' and bien_delivery['lote']=='LT-2026-08','No estructuro Bienestarina')
        require(bien_doc['validaciones']['semaforo']=='VERDE','No valido Bienestarina contra Base Maestra')
        bien_mapping=[{'field':'nombre_completo','sheet':'BIENESTARINA','col':2,'data_start_row':7,'fila_fin':7},{'field':'documento_beneficiario','sheet':'BIENESTARINA','col':3,'data_start_row':7,'fila_fin':7},{'field':'fecha_entrega','sheet':'BIENESTARINA','col':4,'data_start_row':7,'fila_fin':7},{'field':'lote','sheet':'BIENESTARINA','col':5,'data_start_row':7,'fila_fin':7},{'field':'cantidad','sheet':'BIENESTARINA','col':6,'data_start_row':7,'fila_fin':7}]
        conn=connect(str(db)); conn.execute("INSERT INTO plantillas_oficiales_versiones(id,tipo_formato,codigo,nombre,version,fecha_vigencia,estado,mapeo_json,fundacion_id,updated_at) VALUES(5,'BIENESTARINA','BIEN-01','Bienestarina mapeada','2026.08','2026-08-01','vigente',?,1,'2026-08-03')",(json.dumps(bien_mapping),)); conn.commit(); conn.close()
        mapped_bien=root/'BIENESTARINA_SIN_ENCABEZADOS.xlsx'; mapped_bien_wb=Workbook(); mapped_bien_ws=mapped_bien_wb.active; mapped_bien_ws.title='BIENESTARINA'; mapped_bien_ws['A1']='ENTREGA DE ALIMENTO BIENESTARINA'; mapped_bien_ws.cell(7,2,'ANA PEREZ'); mapped_bien_ws.cell(7,3,'1001'); mapped_bien_ws.cell(7,4,'20/08/2026'); mapped_bien_ws.cell(7,5,'LT-MAP-01'); mapped_bien_ws.cell(7,6,2); mapped_bien_wb.save(mapped_bien)
        mapped_bien_id,_=create(repo,1,mapped_bien); mapped_bien_doc=repo.get_document(mapped_bien_id,1); mapped_bien_delivery=mapped_bien_doc['resultado_canonico']['entregas'][0]
        require(mapped_bien_delivery['fecha_entrega']=='2026-08-20' and mapped_bien_doc['validaciones']['semaforo']=='VERDE','No aplico mapeo versionado de Bienestarina')
        invalid_bien=root/'BIENESTARINA_INCOMPLETA.xlsx'; invalid_bien_wb=Workbook(); invalid_bien_ws=invalid_bien_wb.active; invalid_bien_ws.append(['BIENESTARINA ENTREGA DE ALIMENTO']); invalid_bien_ws.append(['NUI','Nombre','Cantidad']); invalid_bien_ws.append(['1001','ANA PEREZ','0']); invalid_bien_wb.save(invalid_bien)
        invalid_bien_id,_=create(repo,1,invalid_bien); invalid_bien_doc=repo.get_document(invalid_bien_id,1)
        require(invalid_bien_doc['validaciones']['semaforo']=='ROJO' and invalid_bien_doc['validaciones']['errores_criticos']>=3,'No bloqueo Bienestarina incompleta')
        invalid_book=root/'LISTADO_ASISTENCIA_INVALIDO.xlsx'
        invalid_wb=Workbook(); invalid_ws=invalid_wb.active; invalid_ws.append(['LISTADO DE ASISTENCIA']); invalid_ws.append(['Nombre completo','Documento','UDS','Asistio']); invalid_ws.append(['PERSONA INEXISTENTE','9999','UCA 1','SI']); invalid_wb.save(invalid_book)
        invalid_id,_=create(repo,1,invalid_book)
        invalid_doc=repo.get_document(invalid_id,1)
        require(invalid_doc['validaciones']['semaforo']=='ROJO','No marco en rojo el documento inexistente')
        require(invalid_doc['validaciones']['errores_criticos']>0,'No genero error critico de Base Maestra')
        invalid_blocked=False
        try: repo.approve(invalid_id,1,10)
        except ValueError: invalid_blocked=True
        require(invalid_blocked,'Permitio aprobar inconsistencias criticas')
        image_path=root/'foto_asistencia.jpg'; image=Image.new('RGB',(1200,1600),'white'); drawing=ImageDraw.Draw(image); drawing.rectangle((80,80,1120,1520),outline='black',width=6)
        for line,text in enumerate(['LISTADO DE ASISTENCIA','Nombre Documento UDS Firma','ANA PEREZ 1001 UCA 1 SI']*8): drawing.text((120,140+line*52),text,fill='black')
        image.save(image_path)
        image_id,_=create(repo,1,image_path)
        image_doc=repo.get_document(image_id,1)
        require(image_doc['estado']=='REQUIERE_OCR','La imagen no quedo pendiente de OCR')
        blank_path=root/'foto_blanca.jpg'; Image.new('RGB',(1200,1600),'white').save(blank_path); blank_raw=read_document_ocr(blank_path)
        require(blank_raw['motor']=='CONTROL_CALIDAD' and blank_raw['calidad']['rechazo_automatico'],'No rechazo una imagen sin contraste')
        low_path=root/'foto_pequena.jpg'; Image.new('RGB',(400,500),'gray').save(low_path); low_raw=read_document_ocr(low_path)
        require('RESOLUCION_INSUFICIENTE' in low_raw['calidad']['problemas'],'No rechazo resolucion insuficiente')
        class AzureResponse:
            def __init__(self,status_code,payload=None,headers=None): self.status_code=status_code; self._payload=payload or {}; self.headers=headers or {}
            def json(self): return self._payload
        azure_payload={'status':'succeeded','analyzeResult':{'apiVersion':'2024-11-30','modelId':'prebuilt-layout','content':'LISTADO DE ASISTENCIA ANA PEREZ 1001 UCA 1','pages':[{'pageNumber':1,'width':1000,'height':1400,'unit':'pixel','lines':[{'content':'LISTADO DE ASISTENCIA'}]}],'tables':[{'rowCount':2,'columnCount':4,'cells':[{'rowIndex':0,'columnIndex':0,'content':'Nombre completo','confidence':.99},{'rowIndex':0,'columnIndex':1,'content':'Documento','confidence':.99},{'rowIndex':0,'columnIndex':2,'content':'UDS','confidence':.99},{'rowIndex':0,'columnIndex':3,'content':'Asistio','confidence':.99},{'rowIndex':1,'columnIndex':0,'content':'ANA PEREZ','confidence':.93},{'rowIndex':1,'columnIndex':1,'content':'1001','confidence':.97},{'rowIndex':1,'columnIndex':2,'content':'UCA 1','confidence':.91},{'rowIndex':1,'columnIndex':3,'content':'SI','confidence':.89}]}]}}
        with patch.dict(os.environ,{'AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT':'https://example.cognitiveservices.azure.com','AZURE_DOCUMENT_INTELLIGENCE_KEY':'secret'},clear=False),patch('requests.post',return_value=AzureResponse(202,headers={'Operation-Location':'https://example/result'})),patch('requests.get',return_value=AzureResponse(200,azure_payload)):
            azure_raw=read_document_azure(image_path)
        azure_classification=classify_document(azure_raw['texto'],image_path.name); azure_canonical,azure_fields=canonicalize(azure_raw,azure_classification[0])
        require(azure_raw['motor']=='AZURE_DOCUMENT_INTELLIGENCE' and len(azure_raw['hojas'])==1,'No convirtio la respuesta Azure')
        require(azure_canonical['participantes'][0]['documento']=='1001','No mapeo la tabla Azure al canonico')
        require(any(field['regla']=='tabla_azure' and field['confianza']>.9 for field in azure_fields),'No conservo confianza Azure por campo')
        blocked=False
        try: repo.approve(image_id,1,10)
        except ValueError: blocked=True
        require(blocked,'Permitio aprobar una imagen sin OCR')
        repo.restart_extraction(image_id,1,10)
        with patch('modules.idp_documental.services._ocr_image_text',return_value='LISTADO DE ASISTENCIA\nNombre Documento UDS Firma\nANA PEREZ  1001  UCA 1  SI'):
            ocr_raw=read_document_ocr(image_path)
        ocr_classification=classify_document(ocr_raw['texto'],image_path.name); ocr_canonical,ocr_fields=canonicalize(ocr_raw,ocr_classification[0]); ocr_canonical['fundacion']['id']=1
        repo.complete_extraction(image_id,1,ocr_raw,ocr_canonical,ocr_fields,ocr_classification,10)
        ocr_doc=repo.get_document(image_id,1)
        require(ocr_doc['estado']=='REQUIERE_REVISION' and ocr_doc['motor_lectura']=='TESSERACT_LOCAL','El reintento OCR no avanzo a revision')
        require(len(ocr_doc['resultado_canonico']['participantes'])==1,'El OCR no estructuro la fila del participante')
        require(ocr_doc['resultado_canonico']['participantes'][0]['documento']=='1001','El OCR no conservo el documento detectado')
        require(ocr_doc['validaciones']['coincidencias']==1,'El participante OCR no se valido contra Base Maestra')
        require(any(field['regla']=='fila_ocr_con_documento' for field in ocr_doc['campos']),'El OCR no guardo evidencia editable')
        require(any(event['evento']=='OCR_REINTENTADO' for event in ocr_doc['eventos']),'No audito el reintento OCR')
        print('IDP core PASS: clasificacion, canonico, tenant, Base Maestra, correccion y aprobacion')


if __name__=='__main__':
    main()
