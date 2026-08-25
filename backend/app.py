import os
import sys
import json
import re
import unicodedata
import pandas as pd
from datetime import datetime, timedelta
import calendar
import copy
import socket
import traceback
import uuid
from pathlib import Path
from difflib import SequenceMatcher

from flask import Flask, request, jsonify, send_from_directory, g, has_request_context, redirect
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Alignment
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.exceptions import HTTPException
from werkzeug.middleware.proxy_fix import ProxyFix

# Importar módulos personalizados
sys.path.insert(0, os.path.dirname(__file__))
from models import (
    Rol, EstadoUsuario, TipoMovimiento, EstadoNutricion,
    AlertaNivel, AlertaConfiguracion, ConfiguracionSistema, Schema
)
from generador_formatos import GeneradorFormatos
from motor_alertas import MotorAlertas
from config import get_config, validate_runtime_config
from extensions import init_extensions
from database import configure_database, get_db_connection as database_connection, database
from services.observability import configure_observability
from modules.print_master import aplicar_configuracion_impresion_libro, infer_print_format
from modules.plantillas_oficiales import iter_plantillas_oficiales_para_generacion
from modules.operational_jobs import configure as configure_operational_jobs, start_job, get_job, list_jobs
from modules.seguridad.tenant_context import tenant_path, resolve_tenant_path
from modules.seguridad.runtime_diagnostics import (
    configure_application_logging,
    logging_health,
    project_instance_id,
    write_exception_report,
)
from services.uds_catalog import (
    INVALID_UNIT_VALUES as UDS_INVALID_UNIT_VALUES,
    aliases_upper as uds_aliases_upper,
    canonical_units as uds_canonical_units,
    equivalent_values as uds_equivalent_values,
    normalization_map as uds_normalization_map,
    normalize_unit as uds_normalize_unit,
)

app = Flask(__name__)
_APP_CONFIGURED = False


def parse_allowed_origins():
    """Orígenes CORS por ambiente, conservando modo local y red privada."""
    raw = str(app.config.get('ALLOWED_ORIGINS') or os.environ.get('ALLOWED_ORIGINS', '')).strip()
    if raw == '*':
        # Se conserva únicamente para desarrollo explícito; producción debe usar allowlist.
        if str(app.config.get('APP_ENV', '')).lower() == 'production':
            raise RuntimeError('ALLOWED_ORIGINS=* no está permitido en producción.')
        return '*'
    if raw:
        return [origin.strip() for origin in raw.split(',') if origin.strip()]

    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        # El frontend y la API se sirven desde el mismo dominio; CORS no es necesario.
        frontend_origin = str(app.config.get('FRONTEND_ORIGIN') or '').strip()
        return [frontend_origin] if frontend_origin else None

    return [
        'null',
        r"http://localhost(:\d+)?",
        r"http://127\.0\.0\.1(:\d+)?",
        r"http://192\.168\.\d{1,3}\.\d{1,3}(:\d+)?",
        r"http://10\.\d{1,3}\.\d{1,3}\.\d{1,3}(:\d+)?",
        r"http://172\.(1[6-9]|2\d|3[0-1])\.\d{1,3}\.\d{1,3}(:\d+)?",
    ]


def create_app(config_name=None):
    """Application factory transicional y compatible con las rutas históricas.

    En esta fase las rutas siguen registradas sobre una instancia única para evitar
    regresiones. Las siguientes fases extraerán dominios de app.py a blueprints.
    """
    global _APP_CONFIGURED
    config_class = get_config(config_name)
    app.config.from_object(config_class)
    validate_runtime_config(app.config)

    if not _APP_CONFIGURED:
        proxy_count = max(0, int(app.config.get('TRUSTED_PROXY_COUNT', 0)))
        if proxy_count:
            app.wsgi_app = ProxyFix(
                app.wsgi_app,
                x_for=proxy_count,
                x_proto=proxy_count,
                x_host=proxy_count,
                x_port=proxy_count,
            )
        init_extensions(app, origins=parse_allowed_origins())
        configure_database(app)
        configure_observability(app, database)
        _APP_CONFIGURED = True
    return app


# Configuración temprana: rutas y módulos históricos consumen estas constantes.
create_app(os.environ.get('APP_ENV') or os.environ.get('FLASK_ENV'))
configure_application_logging(app)

BASE_DIR = app.config['BASE_DIR']
# Las carpetas operativas se resuelven por fundación durante cada petición.
# Plantillas oficiales y backups globales permanecen compartidos y solo el
# SUPERADMIN puede administrarlos.
UPLOAD_FOLDER = tenant_path(app.config['UPLOAD_FOLDER'])
TEMPLATES_FOLDER = app.config['TEMPLATES_FOLDER']
OUTPUT_FOLDER = tenant_path(app.config['OUTPUT_FOLDER'])
BACKUPS_FOLDER = app.config['BACKUPS_FOLDER']
DOCUMENTOS_FOLDER = tenant_path(app.config['DOCUMENTOS_FOLDER'])
CUENTAS_COBRO_FOLDER = tenant_path(app.config['CUENTAS_COBRO_FOLDER'])
LOG_FOLDER = tenant_path(app.config['LOG_FOLDER'])
LOCAL_STORAGE_PATH = tenant_path(app.config['LOCAL_STORAGE_PATH'])
DATABASE_PATH = app.config['DATABASE_PATH']
DATABASE_URL = app.config['DATABASE_URL']

# Crear directorios
for folder in [UPLOAD_FOLDER, TEMPLATES_FOLDER, OUTPUT_FOLDER, BACKUPS_FOLDER, DOCUMENTOS_FOLDER, CUENTAS_COBRO_FOLDER, LOG_FOLDER, LOCAL_STORAGE_PATH]:
    os.makedirs(folder, exist_ok=True)

# Jobs operativos en segundo plano para evitar timeouts 524 en túnel online.
configure_operational_jobs(tenant_path(LOG_FOLDER, 'jobs'))

app.config.setdefault('MAX_CONTENT_LENGTH', AlertaConfiguracion.TAMAÑO_MAX_MB * 1024 * 1024)
ALLOWED_BASE_EXTENSIONS = {'.xlsx', '.xls', '.xlsm', '.csv', '.txt', '.tsv', '.tab', '.dat', '.ods', '.html', '.htm', '.json', '.docx', '.pdf'}
ALLOWED_TEMPLATE_EXTENSIONS = {'.xlsx', '.xls', '.xlsm', '.csv', '.txt', '.doc', '.docx', '.pdf', '.png', '.jpg', '.jpeg', '.zip', '.rar'}
ALLOWED_NUTRICION_EXTENSIONS = {'.xlsx', '.xls', '.xlsm', '.csv', '.txt'}
ALLOWED_TALENTO_EXTENSIONS = {'.xlsx', '.xls', '.xlsm', '.csv', '.txt', '.zip', '.docx'}
ALLOWED_DOCUMENT_EXTENSIONS = {'.pdf', '.doc', '.docx', '.xlsx', '.xls', '.xlsm', '.csv', '.ppt', '.pptx', '.txt', '.png', '.jpg', '.jpeg', '.zip', '.rar'}

# Inicializar generadores y motores
generador = GeneradorFormatos(DATABASE_PATH, TEMPLATES_FOLDER, OUTPUT_FOLDER)
motor_alertas = MotorAlertas(DATABASE_PATH)

# Seguridad multi-fundación, login, roles y permisos.
# Se registra antes de los módulos nuevos para proteger todas las rutas /api.
try:
    from modules.seguridad import register_security
    register_security(app, DATABASE_PATH)
except Exception as exc:
    # Nunca se permite que el servicio productivo arranque sin autenticación,
    # autorización por roles y aislamiento multi-fundación.
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        raise RuntimeError('La capa de seguridad no pudo registrarse; se cancela el arranque.') from exc
    print(f'Seguridad multi-fundación no pudo registrarse: {exc}')

# Fase 2C.7: endpoints auxiliares del núcleo Talento Humano migrado a SQLAlchemy Core.
try:
    from modules.talento_humano.routes import bp as talento_humano_core_bp
    app.register_blueprint(talento_humano_core_bp)
except Exception as exc:
    print(f'Talento Humano Core no pudo registrarse: {exc}')

# FASE 4: Facturación, planes, mensualidades, créditos y suscripciones.
# Se registra después de seguridad para validar sesión y antes de módulos operativos para aplicar control.
try:
    from modules.facturacion_suscripcion import register_facturacion
    register_facturacion(app, DATABASE_PATH, UPLOAD_FOLDER)
except Exception as exc:
    print(f'Facturación y suscripciones no pudo registrarse: {exc}')

# Panel Comercial y Soporte: control de clientes, suscripciones, ingresos, créditos y tickets.
try:
    from modules.panel_comercial import register_panel_comercial
    register_panel_comercial(app, DATABASE_PATH)
except Exception as exc:
    print(f'Panel Comercial no pudo registrarse: {exc}')

# Sistema Integral Administrativo y Financiero.
try:
    from modules.administrativo_financiero import register_administrativo_financiero
    register_administrativo_financiero(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    print(f'Administrativo y Financiero no pudo registrarse: {exc}')

# Gerencia General: tablero ejecutivo comercial, operativo y de licencias.
try:
    from modules.gerencia_general import register_gerencia_general
    register_gerencia_general(app, DATABASE_PATH, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Gerencia General no pudo registrarse: {exc}')

# Módulo independiente de Gestión Pedagógica.
# Si el módulo no carga, el resto de la plataforma continúa funcionando.
try:
    from modules.gestion_pedagogica.routes import register_gestion_pedagogica
    register_gestion_pedagogica(app, DATABASE_PATH, UPLOAD_FOLDER)
except Exception as exc:
    print(f'Gestión Pedagógica no pudo registrarse: {exc}')

# Módulo independiente de Salud y Nutrición Inteligente.
try:
    from modules.salud_nutricion.routes import register_salud_nutricion
    register_salud_nutricion(app, DATABASE_PATH, UPLOAD_FOLDER, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Salud y Nutrición Inteligente no pudo registrarse: {exc}')

# FASE 2: Gestión por Coordinador y Calendario Inteligente.
try:
    from modules.gestion_coordinador.routes import register_gestion_coordinador
    register_gestion_coordinador(app, DATABASE_PATH, UPLOAD_FOLDER)
except Exception as exc:
    print(f'Gestión por Coordinador no pudo registrarse: {exc}')

# Alpha15: Calendario Inteligente central de entregables y alertas operativas.
try:
    from modules.calendario_inteligente import register_calendario_inteligente
    register_calendario_inteligente(app, DATABASE_PATH, UPLOAD_FOLDER)
except Exception as exc:
    print(f'Calendario Inteligente no pudo registrarse: {exc}')

# FASE 3: Planeación Pedagógica y Gestión Automática de Informes.
try:
    from modules.planeacion_pedagogica.routes import register_planeacion_pedagogica
    register_planeacion_pedagogica(app, DATABASE_PATH, UPLOAD_FOLDER, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Planeación Pedagógica no pudo registrarse: {exc}')

# Consola de Auditoría: cruce mensual de bases, novedades y reportes por unidad.
try:
    from modules.cruce_bases import register_cruce_bases
    register_cruce_bases(app, DATABASE_PATH, UPLOAD_FOLDER, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Cruce mensual de bases no pudo registrarse: {exc}')

# ALPHA28: Base Maestra única, validación, consolidación, publicación y trazabilidad multi-corporación.
# Se registra como módulo independiente; no reemplaza login, menú, RPP, Bienestarina ni reportes existentes.
try:
    from modules.base_maestra import register_base_maestra
    register_base_maestra(app, DATABASE_PATH, UPLOAD_FOLDER, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Base Maestra no pudo registrarse: {exc}')

# Motor Universal Tabular: permanece detrás de bandera durante la transición.
try:
    from modules.importaciones_universales import register_importaciones_universales
    register_importaciones_universales(app, DATABASE_PATH, UPLOAD_FOLDER)
except Exception as exc:
    print(f'Motor Universal de Mapeo no pudo registrarse: {exc}')

# Fase comercial 1: Backups automáticos y restauración.
try:
    from modules.backups import register_backups
    register_backups(app, DATABASE_PATH, BACKUPS_FOLDER)
except Exception as exc:
    print(f'Backups y restauración no pudo registrarse: {exc}')


# Fase comercial 2: Centro de Calidad de Datos.
try:
    from modules.calidad_datos import register_calidad_datos
    register_calidad_datos(app, DATABASE_PATH, UPLOAD_FOLDER, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Calidad de Datos no pudo registrarse: {exc}')

# Motor de Plantillas Oficiales ICBF.
try:
    from modules.motor_plantillas import register_motor_plantillas
    register_motor_plantillas(app, DATABASE_PATH, TEMPLATES_FOLDER, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Motor de Plantillas Oficiales no pudo registrarse: {exc}')



# Plantillas oficiales Excel activas: RPP y Bienestarina se diligencian desde archivo oficial.
try:
    from modules.plantillas_oficiales import register_plantillas_oficiales
    register_plantillas_oficiales(app, TEMPLATES_FOLDER)
except Exception as exc:
    print(f'Plantillas oficiales Excel no pudo registrarse: {exc}')

# Paquete Mensual Completo: consolida formatos y reportes operativos.
try:
    from modules.paquete_mensual import register_paquete_mensual
    register_paquete_mensual(app, DATABASE_PATH, OUTPUT_FOLDER, app.config['DATA_DIR'])
except Exception as exc:
    print(f'Paquete Mensual no pudo registrarse: {exc}')

# Reportes Gerenciales Profesionales: informes ejecutivos con indicadores, hallazgos, alertas y recomendaciones.
try:
    from modules.reportes_gerenciales import register_reportes_gerenciales
    register_reportes_gerenciales(app, DATABASE_PATH, OUTPUT_FOLDER)
except Exception as exc:
    print(f'Reportes Gerenciales no pudo registrarse: {exc}')

# Fase UX/UI: Ajustes visuales, temas y auditoría de componentes quemados.
try:
    from modules.ajustes_ui import register_ajustes_ui
    register_ajustes_ui(app, DATABASE_PATH, os.path.dirname(BASE_DIR))
except Exception as exc:
    print(f'Ajustes UI no pudo registrarse: {exc}')

# ALPHA29: Theme Manager centralizado para múltiples diseños, preferencias por usuario y configuración por corporación.
# Se agrega como módulo independiente; conserva el dashboard actual como tema base y no toca login, menú funcional ni formatos.
try:
    from modules.theme_manager import register_theme_manager
    register_theme_manager(app, DATABASE_PATH)
except Exception as exc:
    print(f'Theme Manager no pudo registrarse: {exc}')


# ALPHA45: Configuración Institucional / Marca blanca y Motor Normativo base.
# Integración modular sobre Pack35. No toca carga Cuéntame, CoreCursor, Base Maestra ni formatos oficiales.
try:
    from modules.institucional_normativo import register_institucional_normativo
    register_institucional_normativo(app, DATABASE_PATH, BASE_DIR)
except Exception as exc:
    print(f'Configuración Institucional / Motor Normativo no pudo registrarse: {exc}')

# V2.5.0: Expediente Operativo por UCA, Ruta Operativa y Biblioteca Oficial ICBF.
# Se integra de forma no destructiva: reutiliza los módulos existentes y crea
# únicamente tablas con prefijo giu_/biblioteca_icbf_.
try:
    from modules.gestion_integral_uca import register_gestion_integral_uca
    register_gestion_integral_uca(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Gestión Integral por UCA no pudo registrarse')
    else:
        print(f'Gestión Integral por UCA no pudo registrarse: {exc}')

# V2.5.3: Motor Inteligente de Gestión del Proyecto.
# Orquesta referencias a Ruta Operativa, calendario y entregables existentes;
# no duplica datos misionales ni aprueba productos sin intervención humana.
try:
    from modules.motor_gestion_proyecto import register_motor_gestion_proyecto
    register_motor_gestion_proyecto(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Motor Inteligente de Gestión del Proyecto no pudo registrarse')
    else:
        print(f'Motor Inteligente de Gestión del Proyecto no pudo registrarse: {exc}')


# V2.5.4: Centro Inteligente de Supervisión, Auditoría y Calidad.
try:
    from modules.supervision_calidad import register_supervision_calidad
    register_supervision_calidad(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Centro de Supervisión, Auditoría y Calidad no pudo registrarse')
    else:
        print(f'Centro de Supervisión, Auditoría y Calidad no pudo registrarse: {exc}')

# V2.5.4: Gestión Integral de Familias, Comunidad y Redes de Apoyo.
try:
    from modules.familias_redes import register_familias_redes
    register_familias_redes(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Gestión Integral de Familias, Comunidad y Redes no pudo registrarse')
    else:
        print(f'Gestión Integral de Familias, Comunidad y Redes no pudo registrarse: {exc}')


# V2.7.0: Centro Inteligente de Planeación y Calendario Operativo.
# Consume referencias de los módulos existentes y conserva un único calendario.
try:
    from modules.centro_planeacion import register_centro_planeacion
    register_centro_planeacion(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Centro Inteligente de Planeación no pudo registrarse')
    else:
        print(f'Centro Inteligente de Planeación no pudo registrarse: {exc}')

# V2.7.0: Componente Psicosocial especializado sobre Familias y Redes.
try:
    from modules.componente_psicosocial import register_componente_psicosocial
    register_componente_psicosocial(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Componente Psicosocial no pudo registrarse')
    else:
        print(f'Componente Psicosocial no pudo registrarse: {exc}')

# V2.7.0: Sistema Integral de Ambientes Educativos y Protectores.
try:
    from modules.ambientes_protectores import register_ambientes_protectores
    register_ambientes_protectores(app, DATABASE_PATH, app.config['DATA_DIR'], OUTPUT_FOLDER)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Ambientes Educativos y Protectores no pudo registrarse')
    else:
        print(f'Ambientes Educativos y Protectores no pudo registrarse: {exc}')

# V2.7.0: conserva Integridad y agrega Planeación Operativa/Psicosocial.
try:
    from modules.integraciones_configuracion import register_integraciones_configuracion
    register_integraciones_configuracion(app, DATABASE_PATH, app.config['PROJECT_DIR'], app.config['DATA_DIR'])
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Integraciones y Configuración no pudo registrarse')
    else:
        print(f'Integraciones y Configuración no pudo registrarse: {exc}')

# V2.7.0: conserva Integridad y agrega Planeación Operativa/Psicosocial.
try:
    from modules.asistente_capacitacion import register_asistente_capacitacion
    register_asistente_capacitacion(app, DATABASE_PATH)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Asistente de ayuda y capacitación no pudo registrarse')
    else:
        print(f'Asistente de ayuda y capacitación no pudo registrarse: {exc}')

try:
    from modules.integrity_stability import register_integrity_stability
    register_integrity_stability(app, app.config['PROJECT_DIR'], app.config['DATA_DIR'], DATABASE_PATH)
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        raise RuntimeError('El Motor de Integridad no pudo registrarse; se bloquea el arranque productivo.') from exc
    print(f'Motor de Integridad no pudo registrarse: {exc}')

try:
    from modules.idp_documental import register_idp_documental
    register_idp_documental(app, DATABASE_PATH, app.config['DATA_DIR'])

    from modules.centro_documental import register_centro_documental
    register_centro_documental(app, DATABASE_PATH, app.config['DATA_DIR'])
except Exception as exc:
    if str(app.config.get('APP_ENV', '')).lower() == 'production':
        app.logger.exception('Motor Universal Documental IDP no pudo registrarse')
    else:
        print(f'Motor Universal Documental IDP no pudo registrarse: {exc}')


KNOWN_UNITS = uds_canonical_units()

UNIT_NORMALIZATION_MAP = uds_normalization_map()

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def get_db_connection():
    """Adaptador SQLAlchemy Core para operación central.

    Fase 2C.6: beneficiarios, usuarios, unidades, movimientos y auditoría
    operativa pasan por el Engine central. En SQLite mantiene compatibilidad
    con la API histórica tipo sqlite3; en PostgreSQL evita abrir conexiones
    sqlite3 directas.
    """
    try:
        from modules.sqlalchemy_compat import CoreConnection
        return CoreConnection()
    except Exception:
        # Fallback de contingencia para arranque local si la capa Core todavía
        # no está configurada durante importaciones tempranas.
        return database_connection()


PROCESSING_LOG_FILE = os.path.join(LOG_FOLDER, 'procesamiento_base_maestra.log')


def log_procesamiento_base_maestra(etapa, detalle='', **extra):
    """Registra trazas livianas del procesamiento Cuéntame/Base Maestra.

    No reemplaza el log de jobs; deja una bitácora legible cuando un proceso
    en segundo plano falla antes de devolver respuesta al navegador.
    """
    payload = ' '.join(f'{k}={v}' for k, v in extra.items() if v is not None)
    linea = f"{datetime.now().isoformat(timespec='seconds')} | {etapa}"
    if detalle:
        linea += f" | {detalle}"
    if payload:
        linea += f" | {payload}"
    linea = linea[:3000]
    print(f'[PROCESSING] {linea}', flush=True)
    try:
        os.makedirs(os.path.dirname(PROCESSING_LOG_FILE), exist_ok=True)
        with open(PROCESSING_LOG_FILE, 'a', encoding='utf-8') as fh:
            fh.write(linea + '\n')
    except Exception as log_exc:
        print(f'[PROCESSING] No se pudo escribir bitácora persistente: {log_exc}', flush=True)


def log_beneficiarios_sincronizacion_batch(registros):
    """Escribe trazabilidad por registro con una sola apertura de archivo."""
    if not registros:
        return
    try:
        ruta = os.path.join(LOG_FOLDER, 'beneficiarios_sincronizacion.log')
        os.makedirs(os.path.dirname(ruta), exist_ok=True)
        fecha = datetime.now().isoformat(timespec='seconds')
        with open(ruta, 'a', encoding='utf-8') as fh:
            for registro in registros:
                payload = {'fecha': fecha, **dict(registro or {})}
                fh.write(json.dumps(payload, ensure_ascii=False, default=str)[:3000] + '\n')
    except Exception:
        pass


def safe_executemany(conn, sql, rows, batch_size=500, logger=None):
    """Ejecuta operaciones masivas de forma compatible y por lotes.

    ALPHA35: evita llamar ``cursor.executemany`` sobre cursores que no lo
    soportan. Funciona con sqlite3.Connection y con el wrapper CoreConnection.
    El commit queda a cargo del llamador para conservar la transacción existente.
    """
    rows = list(rows or [])
    if not rows:
        return 0

    try:
        batch_size = int(batch_size or 500)
    except Exception:
        batch_size = 500
    batch_size = max(1, batch_size)

    total = 0
    try:
        for start in range(0, len(rows), batch_size):
            batch = rows[start:start + batch_size]
            if hasattr(conn, 'executemany'):
                conn.executemany(sql, batch)
            else:
                cursor = conn.cursor() if hasattr(conn, 'cursor') else None
                if cursor is not None and hasattr(cursor, 'executemany'):
                    cursor.executemany(sql, batch)
                else:
                    for row in batch:
                        conn.execute(sql, row)
            total += len(batch)
            if logger:
                try:
                    logger(total, len(rows))
                except Exception:
                    pass
        return total
    except Exception:
        try:
            if hasattr(conn, 'rollback'):
                conn.rollback()
        except Exception:
            pass
        raise


def usuario_actual():
    user = getattr(g, 'current_user', None) if 'g' in globals() else None
    if user:
        return user
    # Los jobs se ejecutan fuera del request Flask. El gestor conserva allí el
    # tenant autenticado mediante ContextVar; nunca debe caer a fundación 1.
    try:
        from modules.seguridad.tenant_context import current_tenant_context
        context = current_tenant_context()
        if context.tenant_id:
            return {
                'fundacion_id': int(context.tenant_id),
                'username': context.username or 'sistema',
                'rol': context.role or 'SYSTEM',
            }
    except Exception:
        pass
    return None


def fundacion_actual_id():
    user = usuario_actual()
    if user and user.get('fundacion_id'):
        return int(user['fundacion_id'])
    return 1


def usuario_actual_id():
    user = usuario_actual()
    if user and user.get('id'):
        return int(user['id'])
    return None


def crear_backup_operativo(motivo: str, descripcion: str = ''):
    """Crea backup preventivo sin interrumpir el flujo si ocurre un error.

    Se usa antes de importar Cuéntame, actualizar Talento Humano y generar formatos.
    Nunca modifica formatos oficiales ICBF.
    """
    try:
        from modules.backups.services import BackupService
        service = BackupService(DATABASE_PATH, BACKUPS_FOLDER)
        service.init()
        user = usuario_actual() or {'username': 'sistema', 'fundacion_id': fundacion_actual_id(), 'id': usuario_actual_id()}
        return service.create_backup(motivo, descripcion, user=user, ip=getattr(request, 'remote_addr', None))
    except Exception as exc:
        print(f'No se pudo crear backup preventivo ({motivo}): {exc}')
        return None


def rol_actual():
    user = usuario_actual()
    return (user or {}).get('rol', 'SYSTEM')


def filtro_fundacion_sql(alias=None):
    # SUPERADMIN también trabaja dentro de su fundación por defecto. El acceso
    # global se concede únicamente en rutas centrales explícitas y auditadas.
    if rol_actual() == 'SUPERADMIN' and bool(getattr(g, 'allow_global_tenant_access', False)):
        return '1=1', []
    pref = f'{alias}.' if alias else ''
    return f'COALESCE({pref}fundacion_id, 1) = ?', [fundacion_actual_id()]


def aplicar_metadatos_tenant(datos=None):
    datos = dict(datos or {})
    ahora = datetime.now().isoformat()
    datos.setdefault('fundacion_id', fundacion_actual_id())
    datos.setdefault('usuario_creador_id', usuario_actual_id())
    datos.setdefault('fecha_creacion', ahora)
    datos.setdefault('fecha_actualizacion', ahora)
    return datos


def table_columns(cursor, table_name):
    if database.is_sqlite:
        cursor.execute(f'PRAGMA table_info("{str(table_name)}")')
        return {row['name'] for row in cursor.fetchall()}
    cursor.execute(
        """
        SELECT column_name AS name
        FROM information_schema.columns
        WHERE table_schema = current_schema() AND table_name = ?
        """,
        (str(table_name),),
    )
    return {row['name'] for row in cursor.fetchall()}


def ensure_column(cursor, table_name, column_name, definition):
    if column_name not in table_columns(cursor, table_name):
        cursor.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {definition}")


def ensure_unidades_upsert_constraint(cursor):
    """Garantiza el conflicto lógico usado por los upserts de unidades.

    Las bases PostgreSQL creadas antes del aislamiento multi-tenant conservan
    ``UNIQUE(nombre)``. Añadir ``fundacion_id`` no crea automáticamente la
    clave compuesta que exige ``ON CONFLICT(fundacion_id, nombre)``.
    """
    ensure_column(cursor, 'unidades', 'fundacion_id', 'INTEGER DEFAULT 1')
    cursor.execute('UPDATE unidades SET fundacion_id = 1 WHERE fundacion_id IS NULL')
    cursor.execute(
        'CREATE UNIQUE INDEX IF NOT EXISTS idx_unidades_fundacion_nombre '
        'ON unidades(fundacion_id, nombre)'
    )


def ensure_runtime_schema(cursor):
    """Aplica compatibilidad incremental para bases creadas con versiones previas."""
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            documento TEXT,
            nombre TEXT,
            unidad TEXT,
            fecha_nacimiento TEXT,
            estado TEXT,
            peso_talla_al_dia TEXT,
            docente TEXT,
            tipo_beneficiario TEXT,
            fecha_carga TEXT
        )
    """)
    for col, definition in {
        'docente': 'TEXT',
        'tipo_beneficiario': "TEXT DEFAULT 'NINO'",
        'fecha_carga': 'TEXT',
        'fecha_ingreso': 'TEXT',
        'nui': 'TEXT',
        'tipo_documento': 'TEXT',
        'primer_nombre': 'TEXT',
        'segundo_nombre': 'TEXT',
        'primer_apellido': 'TEXT',
        'segundo_apellido': 'TEXT',
        'sexo': 'TEXT',
        'nombre_acudiente': 'TEXT',
        'documento_acudiente': 'TEXT',
        'tipo_documento_acudiente': 'TEXT',
        'parentesco': 'TEXT',
        'primer_nombre_acudiente': 'TEXT',
        'segundo_nombre_acudiente': 'TEXT',
        'primer_apellido_acudiente': 'TEXT',
        'segundo_apellido_acudiente': 'TEXT',
        'fecha_modificacion_cuentame': 'TEXT',
        'edad_meses': 'INTEGER DEFAULT 0',
        'grupo_edad': 'TEXT',
        'telefono': 'TEXT',
        'regional': 'TEXT',
        'centro_zonal': 'TEXT',
        'municipio': 'TEXT',
        'modalidad': 'TEXT',
        'numero_contrato': 'TEXT',
        'vigencia': 'TEXT',
        'nombre_eas': 'TEXT',
        'nit_eas': 'TEXT',
        'servicio_atencion': 'TEXT',
        'direccion_unidad': 'TEXT',
        'codigo_unidad_servicio': 'TEXT'
    }.items():
        ensure_column(cursor, 'usuarios', col, definition)

    for col, definition in {
        'direccion': 'TEXT',
        'telefono': 'TEXT',
        'coordinador_id': 'INTEGER',
        'docente_asignado': 'TEXT',
        'docente_documento': 'TEXT',
        'coordinador_nombre': 'TEXT',
        'contrato': 'TEXT',
        'total_usuarios': 'INTEGER DEFAULT 0',
        'total_gestantes': 'INTEGER DEFAULT 0',
        'fecha_actualizacion': 'TEXT'
    }.items():
        ensure_column(cursor, 'unidades', col, definition)
    ensure_unidades_upsert_constraint(cursor)

    for col, definition in {
        'fecha_ingreso': 'TEXT',
        'nui': 'TEXT',
        'tipo_documento': 'TEXT',
        'primer_nombre': 'TEXT',
        'segundo_nombre': 'TEXT',
        'primer_apellido': 'TEXT',
        'segundo_apellido': 'TEXT',
        'nombre_acudiente': 'TEXT',
        'documento_acudiente': 'TEXT',
        'tipo_documento_acudiente': 'TEXT',
        'parentesco': 'TEXT',
        'primer_nombre_acudiente': 'TEXT',
        'segundo_nombre_acudiente': 'TEXT',
        'primer_apellido_acudiente': 'TEXT',
        'segundo_apellido_acudiente': 'TEXT',
        'fecha_modificacion_cuentame': 'TEXT',
        'edad_meses': 'INTEGER DEFAULT 0',
        'grupo_edad': 'TEXT',
        'telefono': 'TEXT',
        'regional': 'TEXT',
        'centro_zonal': 'TEXT',
        'municipio': 'TEXT',
        'modalidad': 'TEXT',
        'numero_contrato': 'TEXT',
        'vigencia': 'TEXT',
        'nombre_eas': 'TEXT',
        'nit_eas': 'TEXT',
        'servicio_atencion': 'TEXT',
        'direccion_unidad': 'TEXT',
        'codigo_unidad_servicio': 'TEXT'
    }.items():
        ensure_column(cursor, 'beneficiarios', col, definition)

    for col, definition in {
        'beneficiario_id': 'INTEGER DEFAULT 0',
        'nombre': 'TEXT',
        'fecha_toma': 'TEXT',
        'estado': 'TEXT',
        'documento': 'TEXT',
        'unidad': 'TEXT',
        'fecha_medicion': 'TEXT',
        'responsable': 'TEXT',
        'estado_nutricional': "TEXT DEFAULT 'PENDIENTE'",
        'fecha_proximo_control': 'TEXT',
        'fecha_carga': 'TEXT'
    }.items():
        ensure_column(cursor, 'peso_talla', col, definition)

    for col, definition in {
        'beneficiario_id': 'INTEGER DEFAULT 0',
        'documento': 'TEXT',
        'nombre': 'TEXT',
        'fecha': 'TEXT',
        'detalle': 'TEXT',
        'fecha_movimiento': 'TEXT',
        'razon': 'TEXT',
        'usuario_registra': 'TEXT',
        'fecha_registro': 'TEXT'
    }.items():
        ensure_column(cursor, 'movimientos', col, definition)

    for col, definition in {
        'nombre': 'TEXT',
        'tipo': 'TEXT',
        'ruta_archivo': 'TEXT',
        'nombre_original': 'TEXT',
        'nombre_guardado': 'TEXT',
        'version': 'TEXT',
        'estado': "TEXT DEFAULT 'activo'",
        'activa': 'INTEGER DEFAULT 1',
        'fecha_carga': 'TEXT',
        'fecha_ultima_actualizacion': 'TEXT'
    }.items():
        ensure_column(cursor, 'plantillas', col, definition)

    for col, definition in {
        'accion': 'TEXT',
        'tabla': 'TEXT',
        'registro_id': 'INTEGER',
        'datos_anteriores': 'TEXT',
        'datos_nuevos': 'TEXT',
        'archivo_cargado': 'TEXT',
        'formato_generado': 'TEXT',
        'fecha_accion': 'TEXT',
        'direccion_ip': 'TEXT',
        'fecha': 'TEXT',
        'archivo': 'TEXT',
        'total_registros': 'INTEGER',
        'cambios_detectados': 'TEXT'
    }.items():
        ensure_column(cursor, 'auditoria', col, definition)

    for col, definition in {
        'documento': 'TEXT',
        'nombre': 'TEXT',
        'nombres': 'TEXT',
        'apellidos': 'TEXT',
        'cargo': 'TEXT',
        'unidad': 'TEXT',
        'unidades': 'TEXT',
        'direccion': 'TEXT',
        'telefono': 'TEXT',
        'coordinador': 'TEXT',
        'tipo_equipo': 'TEXT',
        'contrato': 'TEXT',
        'perfil': 'TEXT',
        'estado': "TEXT DEFAULT 'activo'",
        'activo': 'INTEGER DEFAULT 1',
        'archivo': 'TEXT',
        'fecha_carga': 'TEXT',
        'fecha_ultima_actualizacion': 'TEXT'
    }.items():
        ensure_column(cursor, 'coordinadores', col, definition)

    # Talento Humano como fuente maestra global.
    # Estas tablas NO reemplazan las tablas existentes; consolidan la información
    # para que otros módulos puedan consumir la misma fuente de verdad.
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS th_personas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            documento TEXT,
            nombre TEXT NOT NULL,
            nombres TEXT,
            apellidos TEXT,
            cargo TEXT,
            tipo_equipo TEXT,
            rol_normalizado TEXT,
            unidad TEXT,
            direccion TEXT,
            telefono TEXT,
            coordinador TEXT,
            contrato TEXT,
            perfil TEXT,
            estado TEXT DEFAULT 'activo',
            activo INTEGER DEFAULT 1,
            origen_tabla TEXT DEFAULT 'coordinadores',
            origen_id INTEGER,
            archivo TEXT,
            fundacion_id INTEGER DEFAULT 1,
            usuario_creador_id INTEGER,
            fecha_creacion TEXT,
            fecha_actualizacion TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS th_asignaciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            persona_id INTEGER,
            coordinador_id INTEGER,
            coordinador_nombre TEXT,
            unidad TEXT,
            rol TEXT,
            cargo TEXT,
            estado TEXT DEFAULT 'ACTIVO',
            fecha_inicio TEXT,
            fecha_fin TEXT,
            observaciones TEXT,
            fundacion_id INTEGER DEFAULT 1,
            usuario_creador_id INTEGER,
            fecha_creacion TEXT,
            fecha_actualizacion TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS th_historial (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            persona_id INTEGER,
            accion TEXT NOT NULL,
            datos_anteriores TEXT,
            datos_nuevos TEXT,
            usuario TEXT,
            fundacion_id INTEGER DEFAULT 1,
            fecha_accion TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS th_sincronizaciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            origen TEXT,
            total_personas INTEGER DEFAULT 0,
            total_asignaciones INTEGER DEFAULT 0,
            resultado_json TEXT,
            usuario TEXT,
            fundacion_id INTEGER DEFAULT 1,
            fecha_sincronizacion TEXT NOT NULL
        )
    """)
    for stmt in [
        "CREATE INDEX IF NOT EXISTS idx_th_personas_doc_fund ON th_personas(documento, fundacion_id)",
        "CREATE INDEX IF NOT EXISTS idx_th_personas_unidad ON th_personas(unidad)",
        "CREATE INDEX IF NOT EXISTS idx_th_personas_rol ON th_personas(rol_normalizado)",
        "CREATE INDEX IF NOT EXISTS idx_th_asignaciones_persona ON th_asignaciones(persona_id)",
        "CREATE INDEX IF NOT EXISTS idx_th_asignaciones_coord ON th_asignaciones(coordinador_id, fundacion_id)",
    ]:
        cursor.execute(stmt)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS documentos_institucionales (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo TEXT NOT NULL,
            titulo TEXT NOT NULL,
            nombre_original TEXT NOT NULL,
            nombre_guardado TEXT NOT NULL,
            ruta_archivo TEXT NOT NULL,
            version TEXT DEFAULT '1.0',
            texto_indexado TEXT,
            estado TEXT DEFAULT 'vigente',
            fecha_carga TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS reglas_cumplimiento (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            documento_id INTEGER,
            codigo TEXT NOT NULL,
            componente TEXT NOT NULL,
            descripcion TEXT NOT NULL,
            frecuencia TEXT DEFAULT 'MENSUAL',
            criterio TEXT,
            nivel_alerta TEXT DEFAULT 'AMARILLO',
            activa INTEGER DEFAULT 1,
            fundacion_id INTEGER DEFAULT 1,
            fecha_creacion TEXT NOT NULL
        )
    """)
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_reglas_cumplimiento_fundacion ON reglas_cumplimiento(fundacion_id)")
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS estandares_icbf (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codigo TEXT UNIQUE NOT NULL,
            componente TEXT NOT NULL,
            descripcion TEXT NOT NULL,
            evidencia_requerida TEXT,
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS evaluaciones_cumplimiento (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            periodo TEXT NOT NULL,
            cumplimiento_general REAL NOT NULL,
            resultado_json TEXT NOT NULL,
            usuario TEXT,
            fecha_evaluacion TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS entregables_operacion (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo TEXT NOT NULL,
            periodo TEXT NOT NULL,
            unidad TEXT,
            ruta_archivo TEXT,
            estado TEXT DEFAULT 'pendiente',
            observaciones TEXT,
            fecha_carga TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS cuentas_cobro_plantillas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre_original TEXT NOT NULL,
            nombre_guardado TEXT NOT NULL,
            ruta_archivo TEXT NOT NULL,
            docente_nombre TEXT,
            documento TEXT,
            unidad TEXT,
            estado TEXT DEFAULT 'activo',
            fecha_carga TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS cuentas_cobro_generadas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plantilla_id INTEGER,
            docente_nombre TEXT,
            documento TEXT,
            unidad TEXT,
            periodo TEXT NOT NULL,
            numero_cuenta INTEGER,
            ciudad TEXT,
            nombre_archivo TEXT NOT NULL,
            ruta_archivo TEXT NOT NULL,
            fecha_generacion TEXT NOT NULL,
            FOREIGN KEY (plantilla_id) REFERENCES cuentas_cobro_plantillas(id)
        )
    """)
    for col, definition in {
        'fecha_limite': 'TEXT',
        'responsable': 'TEXT',
        'categoria': 'TEXT',
        'documento_analizado': 'TEXT'
    }.items():
        ensure_column(cursor, 'entregables_operacion', col, definition)


def seed_compliance_catalog(cursor, fundacion_id=1):
    ahora = datetime.now().isoformat()
    estandares = [
        ('EST-01', 'Proceso Pedagógico', 'Planeaciones e informes pedagógicos cargados para el periodo.', 'Planeación, informe pedagógico y evidencias'),
        ('EST-02', 'Familia y Comunidad', 'Actas o evidencias de encuentros comunitarios disponibles.', 'Actas y evidencias'),
        ('EST-03', 'Salud y Nutrición', 'Control de peso y talla vigente y sin vencimientos críticos.', 'Registros de peso y talla'),
        ('EST-04', 'Talento Humano', 'Talento humano registrado y asociado a unidades.', 'Archivo de talento humano'),
        ('EST-05', 'Ambientes Educativos', 'Soportes de ambientes, dotación o protocolos cargados.', 'Protocolos, evidencias o inventario'),
        ('EST-06', 'Administrativo y Gestión', 'Formatos, asistencia, RPP/RAN y auditoría del periodo disponibles.', 'Formatos generados y auditoría')
    ]
    for codigo, componente, descripcion, evidencia in estandares:
        cursor.execute("""
            INSERT OR IGNORE INTO estandares_icbf
            (codigo, componente, descripcion, evidencia_requerida, activo, fecha_creacion)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (codigo, componente, descripcion, evidencia, 1, ahora))

    reglas = [
        ('REG-EDAD-71', 'Salud y Nutrición', 'Detectar niños y niñas activos con 71 meses o más para alerta de retiro.', 'DIARIA', 'edad_meses >= 71', 'CRITICA'),
        ('REG-PESO-TALLA', 'Salud y Nutrición', 'Verificar control de peso y talla vigente según control periódico.', 'MENSUAL', 'peso_talla <= 90 dias', 'ROJO'),
        ('REG-PLANEACION', 'Proceso Pedagógico', 'Verificar planeación pedagógica o informe del periodo.', 'MENSUAL', 'informe_pedagogico_periodo', 'AMARILLO'),
        ('REG-EVIDENCIAS', 'Familia y Comunidad', 'Verificar actas, listados o evidencias obligatorias cargadas.', 'MENSUAL', 'evidencias_periodo', 'AMARILLO'),
        ('REG-TALENTO', 'Talento Humano', 'Verificar talento humano registrado.', 'MENSUAL', 'coordinadores_o_docentes', 'AMARILLO'),
        ('REG-FORMATOS', 'Administrativo y Gestión', 'Verificar generación o carga de formatos de asistencia, RPP y RAN.', 'MENSUAL', 'formatos_operacion', 'ROJO')
    ]
    for codigo, componente, descripcion, frecuencia, criterio, nivel in reglas:
        cursor.execute("""
            INSERT INTO reglas_cumplimiento
            (codigo, componente, descripcion, frecuencia, criterio, nivel_alerta, activa, fundacion_id, fecha_creacion)
            SELECT ?, ?, ?, ?, ?, ?, 1, ?, ?
            WHERE NOT EXISTS (SELECT 1 FROM reglas_cumplimiento WHERE codigo = ? AND fundacion_id = ?)
        """, (codigo, componente, descripcion, frecuencia, criterio, nivel, int(fundacion_id or 1), ahora, codigo, int(fundacion_id or 1)))


def init_db():
    """Inicializa la base de datos con el esquema completo"""
    conn = get_db_connection()
    cursor = conn.cursor()

    schema = Schema.get_schema_sql()
    if database.is_postgresql:
        # El esquema histórico se conserva compatible con SQLite para pruebas,
        # pero producción no debe depender de que el cursor concreto traduzca
        # DDL implícitamente. Normalizar cada sentencia aquí garantiza que
        # AUTOINCREMENT/BLOB/REAL nunca lleguen sin convertir a PostgreSQL.
        from modules.dbapi_compat import (
            _split_script, _translate_ddl, order_schema_statements_by_foreign_keys,
        )
        statements = order_schema_statements_by_foreign_keys(_split_script(schema))
        statements = [_translate_ddl(statement) for statement in statements]
        schema = ';\n'.join(statements) + ';\n'
        if re.search(r'\bAUTOINCREMENT\b', schema, re.I):
            raise RuntimeError('El esquema PostgreSQL conserva AUTOINCREMENT después de normalizarse.')
    cursor.executescript(schema)
    ensure_runtime_schema(cursor)
    seed_compliance_catalog(cursor)
    
    # Las credenciales iniciales se crean exclusivamente desde init_hosting.py
    # usando variables privadas. Nunca se insertan usuarios o contraseñas fijas.
    cursor.execute("""
        INSERT INTO configuracion (clave, valor, tipo, fecha_actualizacion)
        SELECT ?, ?, ?, ?
        WHERE NOT EXISTS (SELECT 1 FROM configuracion WHERE clave = ?)
    """, ('VERSION', app.config.get('APP_VERSION', '2.3.7-railway-operativa'), 'STRING', datetime.now().isoformat(), 'VERSION'))
    cursor.execute(
        "UPDATE configuracion SET valor=?, fecha_actualizacion=? WHERE clave='VERSION'",
        (app.config.get('APP_VERSION', '2.3.7-railway-operativa'), datetime.now().isoformat())
    )

    conn.commit()
    conn.close()

    # Completa el esquema de seguridad y aplica la migración tenant de forma
    # idempotente tanto en desarrollo local como en Railway. Así una base nueva
    # nunca queda con restricciones UNIQUE globales incompatibles con varias
    # fundaciones.
    from modules.seguridad.services import ensure_security_schema
    ensure_security_schema(DATABASE_PATH)
    if database.is_sqlite:
        from migrations.migrate_multitenant_phase3 import migrate as migrate_multitenant_phase3
        migrate_multitenant_phase3(DATABASE_PATH)


def es_extension_valida(filename, allowed_ext):
    _, ext = os.path.splitext(filename.lower())
    return ext in allowed_ext


def leer_tabla_desde_texto(texto):
    """Convierte texto con estructura tabular en DataFrame usando separadores comunes."""
    if not texto:
        raise ValueError('El archivo no contiene texto legible.')

    lineas = [linea.strip() for linea in str(texto).splitlines() if linea and linea.strip()]
    if not lineas:
        raise ValueError('No se encontraron filas útiles en el archivo.')

    # Quitar líneas decorativas o demasiado cortas.
    lineas_utiles = [linea for linea in lineas if len(linea) > 2]
    muestra = '\n'.join(lineas_utiles[:60])

    separadores = ['\t', ';', ',', '|']
    for sep in separadores:
        if sep in muestra:
            try:
                from io import StringIO
                df = pd.read_csv(StringIO('\n'.join(lineas_utiles)), sep=sep, dtype=str, engine='python')
                if df is not None and not df.empty and len(df.columns) > 1:
                    return df
            except Exception:
                pass

    # Último recurso: columnas separadas por 2 o más espacios.
    try:
        from io import StringIO
        df = pd.read_csv(StringIO('\n'.join(lineas_utiles)), sep=r'\s{2,}', dtype=str, engine='python')
        if df is not None and not df.empty and len(df.columns) > 1:
            return df
    except Exception:
        pass

    raise ValueError('No se pudo detectar una tabla. Usa Excel, CSV, TXT delimitado, DOCX con tabla o PDF con texto tabular.')


def leer_base_datos_flexible(ruta_archivo, filename=None):
    """Lee bases en formatos comunes sin exigir una sola extensión.

    Formatos soportados:
    Excel: xlsx, xls, xlsm, ods
    Texto/tablas: csv, txt, tsv, tab, dat
    Web/datos: html, htm, json
    Documentos: docx con tablas o texto tabular, pdf con texto tabular
    """
    nombre = filename or os.path.basename(ruta_archivo)
    _, ext = os.path.splitext(str(nombre).lower())

    if ext in {'.xlsx', '.xls', '.xlsm', '.ods'}:
        try:
            hojas = pd.read_excel(ruta_archivo, sheet_name=None, dtype=str)
            mejor = None
            mejor_score = -1
            for _, df_hoja in hojas.items():
                if df_hoja is None or df_hoja.empty:
                    continue
                columnas = ' '.join([normalizar_texto_clave(c) for c in df_hoja.columns])
                score = len(df_hoja) + len(df_hoja.columns) * 5
                for clave in ['documento', 'beneficiario', 'unidad', 'nacimiento', 'acudiente', 'responsable']:
                    if clave in columnas:
                        score += 100
                if score > mejor_score:
                    mejor = df_hoja
                    mejor_score = score
            if mejor is not None:
                return mejor
            raise ValueError('El libro no contiene hojas con datos.')
        except Exception as exc:
            raise ValueError(f'No se pudo leer el archivo Excel/ODS: {exc}')

    if ext in {'.csv', '.txt', '.tsv', '.tab', '.dat'}:
        # Probar codificaciones y autodetección de separador.
        ultimo_error = None
        for encoding in ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']:
            try:
                if ext in {'.tsv', '.tab'}:
                    df = pd.read_csv(ruta_archivo, sep='\t', dtype=str, encoding=encoding)
                else:
                    df = pd.read_csv(ruta_archivo, sep=None, dtype=str, engine='python', encoding=encoding)
                if df is not None and not df.empty:
                    return df
            except Exception as exc:
                ultimo_error = exc
        try:
            with open(ruta_archivo, 'r', encoding='latin-1', errors='ignore') as fh:
                return leer_tabla_desde_texto(fh.read())
        except Exception as exc:
            raise ValueError(f'No se pudo leer archivo de texto/csv: {ultimo_error or exc}')

    if ext in {'.html', '.htm'}:
        try:
            tablas = pd.read_html(ruta_archivo)
            if tablas:
                return tablas[0].astype(str)
        except Exception as exc:
            raise ValueError(f'No se pudo leer tabla HTML: {exc}')

    if ext == '.json':
        try:
            df = pd.read_json(ruta_archivo, dtype=str)
            if df is not None and not df.empty:
                return df
        except Exception:
            try:
                with open(ruta_archivo, 'r', encoding='utf-8', errors='ignore') as fh:
                    data = json.load(fh)
                if isinstance(data, dict):
                    for value in data.values():
                        if isinstance(value, list):
                            return pd.DataFrame(value).astype(str)
                    return pd.DataFrame([data]).astype(str)
                if isinstance(data, list):
                    return pd.DataFrame(data).astype(str)
            except Exception as exc:
                raise ValueError(f'No se pudo leer JSON: {exc}')

    if ext == '.docx':
        try:
            from docx import Document
            document = Document(ruta_archivo)
            tablas = []
            for table in document.tables:
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if len(rows) >= 2 and len(rows[0]) > 1:
                    headers = rows[0]
                    data = rows[1:]
                    tablas.append(pd.DataFrame(data, columns=headers))
            if tablas:
                tablas.sort(key=lambda df: len(df) * max(1, len(df.columns)), reverse=True)
                return tablas[0].astype(str)
            texto = '\n'.join([p.text for p in document.paragraphs if p.text])
            return leer_tabla_desde_texto(texto)
        except Exception as exc:
            raise ValueError(f'No se pudo leer DOCX tabular: {exc}')

    if ext == '.pdf':
        try:
            from pypdf import PdfReader
            reader = PdfReader(ruta_archivo)
            texto = '\n'.join(page.extract_text() or '' for page in reader.pages)
            return leer_tabla_desde_texto(texto)
        except Exception as exc:
            raise ValueError(f'No se pudo leer PDF tabular: {exc}')

    raise ValueError(f'Formato no soportado para base de datos tabular: {ext or "sin extensión"}')


def parse_fecha_cuentame(valor):
    """Convierte fechas de Cuéntame priorizando formato colombiano día/mes/año."""
    if valor is None:
        return None
    try:
        if pd.isna(valor):
            return None
    except Exception:
        pass

    if isinstance(valor, datetime):
        return valor

    texto = str(valor).strip()
    if not texto or texto.lower() in {'nan', 'nat', 'none', 'null'}:
        return None

    # Números seriales de Excel.
    if re.fullmatch(r'\d+(\.\d+)?', texto):
        try:
            numero = float(texto)
            if numero > 20000:
                return pd.to_datetime(numero, unit='D', origin='1899-12-30').to_pydatetime()
        except Exception:
            pass

    for dayfirst in (True, False):
        try:
            fecha = pd.to_datetime(texto, errors='coerce', dayfirst=dayfirst)
            if pd.notna(fecha):
                return fecha.to_pydatetime()
        except Exception:
            continue
    return None


def calcular_edad_meses(fecha_nacimiento, fecha_referencia=None):
    """Calcula edad en meses usando fechas de Cuéntame en formato colombiano."""
    try:
        fecha = parse_fecha_cuentame(fecha_nacimiento)
        if not fecha:
            return 0

        hoy = fecha_referencia or datetime.now()
        meses = (hoy.year - fecha.year) * 12 + (hoy.month - fecha.month)
        if hoy.day < fecha.day:
            meses -= 1
        return max(0, meses)
    except Exception:
        return 0


def formatear_edad_completa(edad_meses, tipo_beneficiario=''):
    """Convierte edad en meses a un texto legible: '3 años y 4 meses'.

    Se usa solo para visualización del dashboard; los cálculos internos
    siguen trabajando en meses para no romper filtros ni formatos.
    """
    tipo = normalizar_texto_clave(tipo_beneficiario)
    if 'gestante' in tipo:
        return 'Gestante'

    try:
        total_meses = int(float(edad_meses or 0))
    except Exception:
        total_meses = 0

    total_meses = max(0, total_meses)
    anios = total_meses // 12
    meses = total_meses % 12

    partes = []
    if anios:
        partes.append(f"{anios} año" + ('' if anios == 1 else 's'))
    if meses or not partes:
        partes.append(f"{meses} mes" + ('' if meses == 1 else 'es'))

    return ' y '.join(partes)


def inferir_edad_meses_desde_valor(valor):
    """Usa la columna Edad del beneficiario como respaldo cuando la fecha no se pueda leer."""
    texto = limpiar_valor(valor)
    if not texto:
        return 0
    clave = normalizar_texto_clave(texto)
    numeros = re.findall(r'\d+', clave)
    if not numeros:
        return 0
    numero = int(numeros[0])

    if 'mes' in clave:
        return numero
    if 'ano' in clave or 'anio' in clave or 'años' in str(texto).lower():
        return numero * 12

    # En Cuéntame, la columna "Edad del beneficiario" suele venir en años.
    if 0 <= numero <= 6:
        return numero * 12

    # Si viene 12, 24, 60, etc., asumimos meses.
    return numero


def inferir_reglas_desde_texto(texto, documento_id=None):
    """Convierte menciones frecuentes del manual en reglas verificables."""
    texto_normalizado = (texto or '').lower()
    catalogo = [
        {
            'codigo': 'DOC-EDAD-71',
            'componente': 'Salud y Nutrición',
            'descripcion': 'Detectar niños y niñas mayores de 71 meses y generar alerta de retiro.',
            'frecuencia': 'DIARIA',
            'criterio': 'edad_meses >= 71',
            'nivel_alerta': 'CRITICA',
            'keywords': ['71 meses', 'edad de retiro', 'retiro']
        },
        {
            'codigo': 'DOC-PLANEACION',
            'componente': 'Proceso Pedagógico',
            'descripcion': 'Verificar que la planeación pedagógica del periodo esté cargada.',
            'frecuencia': 'MENSUAL',
            'criterio': 'entregable:Planeación',
            'nivel_alerta': 'AMARILLO',
            'keywords': ['planeación', 'planeacion', 'proceso pedagógico', 'proceso pedagogico']
        },
        {
            'codigo': 'DOC-ACTAS',
            'componente': 'Familia y Comunidad',
            'descripcion': 'Verificar actas y listados de asistencia exigidos por el manual.',
            'frecuencia': 'MENSUAL',
            'criterio': 'entregable:Acta',
            'nivel_alerta': 'AMARILLO',
            'keywords': ['acta', 'actas', 'listados de asistencia', 'listado de asistencia']
        },
        {
            'codigo': 'DOC-EVIDENCIAS',
            'componente': 'Familia y Comunidad',
            'descripcion': 'Verificar evidencias obligatorias cargadas para el periodo.',
            'frecuencia': 'MENSUAL',
            'criterio': 'entregable:Evidencia',
            'nivel_alerta': 'AMARILLO',
            'keywords': ['evidencia', 'evidencias']
        },
        {
            'codigo': 'DOC-ENCUENTRO-COMUNITARIO',
            'componente': 'Familia y Comunidad',
            'descripcion': 'Verificar realización y soporte de encuentros comunitarios.',
            'frecuencia': 'MENSUAL',
            'criterio': 'entregable:Encuentro comunitario',
            'nivel_alerta': 'AMARILLO',
            'keywords': ['encuentro comunitario', 'encuentros comunitarios']
        },
        {
            'codigo': 'DOC-PESO-TALLA',
            'componente': 'Salud y Nutrición',
            'descripcion': 'Verificar registro vigente de peso y talla.',
            'frecuencia': 'MENSUAL',
            'criterio': 'peso_talla <= 90 dias',
            'nivel_alerta': 'ROJO',
            'keywords': ['peso y talla', 'peso/talla', 'talla']
        },
        {
            'codigo': 'DOC-RPP',
            'componente': 'Administrativo y Gestión',
            'descripcion': 'Verificar entrega o generación de RPP del periodo.',
            'frecuencia': 'MENSUAL',
            'criterio': 'entregable:RPP',
            'nivel_alerta': 'ROJO',
            'keywords': ['rpp']
        }
    ]
    reglas = []
    for item in catalogo:
        if any(keyword in texto_normalizado for keyword in item['keywords']):
            regla = item.copy()
            regla.pop('keywords', None)
            regla['documento_id'] = documento_id
            reglas.append(regla)
    return reglas


def guardar_reglas_documentales(cursor, reglas):
    ahora = datetime.now().isoformat()
    for regla in reglas:
        cursor.execute("""
            UPDATE reglas_cumplimiento
            SET documento_id = ?, componente = ?, descripcion = ?, frecuencia = ?,
                criterio = ?, nivel_alerta = ?, activa = 1
            WHERE codigo = ?
        """, (
            regla.get('documento_id'),
            regla['componente'],
            regla['descripcion'],
            regla['frecuencia'],
            regla['criterio'],
            regla['nivel_alerta'],
            regla['codigo']
        ))
        if cursor.rowcount == 0:
            cursor.execute("""
                INSERT INTO reglas_cumplimiento
                (documento_id, codigo, componente, descripcion, frecuencia, criterio, nivel_alerta, activa, fundacion_id, fecha_creacion)
                VALUES (?, ?, ?, ?, ?, ?, ?, 1, ?, ?)
            """, (
                regla.get('documento_id'),
                regla['codigo'],
                regla['componente'],
                regla['descripcion'],
                regla['frecuencia'],
                regla['criterio'],
                regla['nivel_alerta'],
                fundacion_actual_id(),
                ahora
            ))


def clasificar_nutricional(peso, talla, edad_meses):
    """Clasifica estado nutricional"""
    if not peso or not talla or not edad_meses:
        return EstadoNutricion.PENDIENTE
    
    talla_m = talla / 100
    imc = peso / (talla_m ** 2)
    
    if edad_meses < 24:
        if imc < 14:
            return EstadoNutricion.DESNUTRICION
        elif imc < 16:
            return EstadoNutricion.RIESGO
        elif imc <= 18:
            return EstadoNutricion.ADECUADO
        else:
            return EstadoNutricion.SOBREPESO
    else:
        if imc < 15:
            return EstadoNutricion.DESNUTRICION
        elif imc < 16:
            return EstadoNutricion.RIESGO
        elif imc <= 20:
            return EstadoNutricion.ADECUADO
        else:
            return EstadoNutricion.SOBREPESO


MESES_ES = {
    1: 'ENERO',
    2: 'FEBRERO',
    3: 'MARZO',
    4: 'ABRIL',
    5: 'MAYO',
    6: 'JUNIO',
    7: 'JULIO',
    8: 'AGOSTO',
    9: 'SEPTIEMBRE',
    10: 'OCTUBRE',
    11: 'NOVIEMBRE',
    12: 'DICIEMBRE'
}

ALIAS_UNIDADES_CUENTAME = uds_aliases_upper()

UNIDADES_INVALIDAS = set(UDS_INVALID_UNIT_VALUES)



def normalizar_texto_clave(valor):
    """Normaliza textos para comparar encabezados sin depender de tildes o mayúsculas."""
    texto = str(valor or '').strip().lower()
    texto = unicodedata.normalize('NFKD', texto)
    texto = ''.join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r'[^a-z0-9]+', ' ', texto)
    return ' '.join(texto.split())


def limpiar_valor(valor, default=''):
    if valor is None or pd.isna(valor):
        return default
    texto = str(valor).strip()
    if texto.lower() in {'nan', 'nat', 'none', 'null'}:
        return default
    return texto


def unir_partes(*partes):
    return ' '.join([limpiar_valor(p) for p in partes if limpiar_valor(p)]).strip()


def limpiar_documento_talento(valor):
    """Normaliza documentos leídos desde Excel evitando valores tipo 11811380.0."""
    texto = limpiar_valor(valor)
    if re.fullmatch(r'\d+\.0+', texto):
        return texto.split('.')[0]
    return texto




def abreviar_tipo_documento(valor):
    """Convierte tipos de documento largos a siglas usadas por los formatos oficiales."""
    texto = normalizar_texto_clave(valor)
    if not texto:
        return ''
    reglas = [
        (['registro civil', 'reg civil', 'rc'], 'RC'),
        (['tarjeta de identidad', 't i', 'ti'], 'TI'),
        (['cedula de ciudadania', 'cedula ciudadania', 'cc', 'c c'], 'CC'),
        (['cedula de extranjeria', 'ce'], 'CE'),
        (['pasaporte', 'pa'], 'PA'),
        (['permiso especial', 'pep'], 'PEP'),
        (['permiso por proteccion temporal', 'ppt'], 'PPT')
    ]
    for aliases, sigla in reglas:
        if any(alias in texto for alias in aliases):
            return sigla
    return limpiar_valor(valor).upper()


def fecha_entrega_bienestarina_desde_request(mes=None, año=None, options=None):
    """Fecha de entrega para Bienestarina.

    Alpha24 permite usar esta función también en trabajos de segundo plano,
    donde no existe request activo. Por eso primero lee options y solo después
    intenta request.form/request.args si hay contexto Flask.
    """
    options = dict(options or {})

    def opt(*keys):
        for key in keys:
            valor = limpiar_valor(options.get(key))
            if valor:
                return valor
        if has_request_context():
            for key in keys:
                valor = limpiar_valor(request.form.get(key) or request.args.get(key))
                if valor:
                    return valor
        return ''

    fecha = opt('fecha_entrega_bienestarina')
    if fecha:
        try:
            return datetime.fromisoformat(fecha[:10]).strftime('%d/%m/%Y')
        except Exception:
            return fecha
    dia = opt('dia_entrega_bienestarina')
    mes_form = opt('mes_entrega_bienestarina')
    año_form = opt('anio_entrega_bienestarina', 'año_entrega_bienestarina', 'anio', 'año')
    if dia and mes_form and año_form:
        return f'{dia.zfill(2)}/{mes_form.zfill(2)}/{año_form}'
    if mes and año:
        ultimo_dia = calendar.monthrange(int(año), int(mes))[1]
        return f'{ultimo_dia:02d}/{int(mes):02d}/{int(año)}'
    return datetime.now().strftime('%d/%m/%Y')


def partes_fecha_ddmmaaaa(fecha):
    """Devuelve día, mes, año desde fechas tipo dd/mm/aaaa, yyyy-mm-dd o similares."""
    parsed = parse_fecha_cuentame(fecha)
    if parsed:
        return f'{parsed.day:02d}', f'{parsed.month:02d}', str(parsed.year)
    texto = limpiar_valor(fecha)
    numeros = re.findall(r'\d+', texto)
    if len(numeros) >= 3:
        if len(numeros[0]) == 4:
            return numeros[2].zfill(2), numeros[1].zfill(2), numeros[0]
        return numeros[0].zfill(2), numeros[1].zfill(2), numeros[2]
    return '', '', ''


def nombre_documento_acudiente(user):
    """Formato: NOMBRE COMPLETO - CC 900000001."""
    nombre = limpiar_valor(user.get('Acudiente') or user.get('nombre_acudiente'))
    tipo = abreviar_tipo_documento(user.get('TipoDocumentoAcudiente') or user.get('tipo_documento_acudiente') or 'CC')
    documento = limpiar_valor(user.get('DocumentoAcudiente') or user.get('documento_acudiente'))
    partes = [nombre]
    doc = ' '.join([p for p in [tipo, documento] if p])
    if doc:
        partes.append(doc)
    return ' - '.join([p for p in partes if p])


def _talento_posibles_unidades(datos):
    """Devuelve todas las unidades asociadas a un registro de talento humano."""
    posibles = set()
    for key in ['unidad', 'nombre_uds', 'unidad_servicio', 'comunidad', 'direccion']:
        valor = datos.get(key) if isinstance(datos, dict) else None
        unidad = normalize_unidad(valor or '')
        if unidad:
            posibles.add(unidad)
    unidades_json = datos.get('unidades') if isinstance(datos, dict) else None
    if unidades_json:
        try:
            if isinstance(unidades_json, str):
                cargadas = json.loads(unidades_json)
            else:
                cargadas = unidades_json
            if isinstance(cargadas, (list, tuple, set)):
                for item in cargadas:
                    unidad = normalize_unidad(item)
                    if unidad:
                        posibles.add(unidad)
        except Exception:
            # Algunas bases guardan las unidades como texto separado por coma.
            for item in re.split(r'[,;/|]+', str(unidades_json)):
                unidad = normalize_unidad(item)
                if unidad:
                    posibles.add(unidad)
    return posibles


def _talento_coincide_unidad(datos, unidad):
    """Compara una unidad normalizada sin mezclar registros de otra UDS."""
    unidad_norm = normalize_unidad(unidad)
    if not unidad_norm:
        return False
    buscadas = {unidad_norm}
    buscadas.update(normalize_unidad(u) for u in equivalentes_unidad(unidad_norm))
    buscadas_texto = {normalizar_texto_clave(u) for u in buscadas if u}
    for posible in _talento_posibles_unidades(datos):
        if posible == unidad_norm:
            return True
        posible_txt = normalizar_texto_clave(posible)
        if posible_txt and posible_txt in buscadas_texto:
            return True
    return False


def _orden_talento_para_encabezado(datos):
    cargo = normalizar_texto_clave((datos or {}).get('cargo') or (datos or {}).get('tipo_equipo') or '')
    if 'agente' in cargo or 'docente' in cargo or 'educativo' in cargo:
        return 0
    if 'coordin' in cargo:
        return 1
    if 'pedagog' in cargo:
        return 2
    if 'psicosocial' in cargo or 'psicolog' in cargo:
        return 3
    if 'enfermer' in cargo or 'nutric' in cargo:
        return 4
    if 'suplente' in cargo or 'apoyo' in cargo or 'auxiliar' in cargo:
        return 5
    return 9


def obtener_talentos_por_unidad(unidad):
    """Devuelve todos los registros de talento humano asociados a una unidad.

    Fuente principal: tabla coordinadores, con respaldo en th_personas cuando existe.
    La comparación usa alias normalizados para impedir que una UDS herede datos de otra.
    """
    registros = []
    vistos = set()
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        ensure_runtime_schema(cursor)
        consultas = [
            """
            SELECT *
            FROM coordinadores
            WHERE COALESCE(activo, 1) = 1 AND lower(COALESCE(estado, 'activo')) <> 'inactivo'
            """,
            """
            SELECT *
            FROM th_personas
            WHERE COALESCE(activo, 1) = 1 AND lower(COALESCE(estado, 'activo')) <> 'inactivo'
            """,
        ]
        for sql in consultas:
            try:
                filas = cursor.execute(sql).fetchall()
            except Exception:
                continue
            for fila in filas:
                datos = dict(fila)
                nombre = limpiar_valor(datos.get('nombre') or unir_partes(datos.get('nombres'), datos.get('apellidos'))).upper()
                if nombre and not datos.get('nombre'):
                    datos['nombre'] = nombre
                if not _talento_coincide_unidad(datos, unidad):
                    continue
                clave = (limpiar_documento_talento(datos.get('documento')), normalizar_texto_clave(nombre), normalizar_texto_clave(datos.get('cargo')))
                if clave in vistos:
                    continue
                vistos.add(clave)
                registros.append(datos)
        conn.close()
    except Exception:
        pass
    registros.sort(key=lambda item: (_orden_talento_para_encabezado(item), normalizar_texto_clave(item.get('nombre'))))
    return registros


def obtener_talento_por_unidad(unidad):
    """Devuelve el agente educativo/docente correcto asociado a una unidad."""
    registros = obtener_talentos_por_unidad(unidad)
    if not registros:
        return {}
    for datos in registros:
        cargo = normalizar_texto_clave(datos.get('cargo') or datos.get('tipo_equipo') or '')
        if 'agente' in cargo or 'docente' in cargo or 'educativo' in cargo:
            return datos
    return registros[0]


def crear_registro_talento(data, archivo='manual'):
    """Normaliza un registro de talento para guardarlo o editarlo."""
    nombre = limpiar_valor(
        data.get('nombre') or data.get('nombres_y_apellidos') or data.get('NOMBRES Y APELLIDOS')
        or unir_partes(data.get('primer_nombre') or data.get('Primer Nombre'),
                       data.get('segundo_nombre') or data.get('Segundo Nombre'),
                       data.get('primer_apellido') or data.get('Primer Apellido'),
                       data.get('segundo_apellido') or data.get('Segundo Apellido'))
    ).upper()
    documento = limpiar_documento_talento(
        data.get('documento') or data.get('cedula') or data.get('cédula') or data.get('CEDULA')
        or data.get('numero_documento') or data.get('Número de Documento') or data.get('Numero de Documento')
        or data.get('identificacion') or data.get('identificación')
    )
    cargo = limpiar_valor(data.get('cargo') or data.get('CARGO') or data.get('Cargo') or 'AGENTE EDUCATIVO').upper()
    unidad = normalize_unidad(
        data.get('unidad') or data.get('Nombre UDS') or data.get('nombre_uds')
        or data.get('comunidad') or data.get('COMUNIDAD') or data.get('uca')
        or data.get('unidad_servicio') or data.get('direccion') or data.get('DIRECCION')
    )
    direccion = limpiar_valor(data.get('direccion') or data.get('Dirección de Residencia') or data.get('DIRECCION') or unidad)
    telefono = limpiar_documento_talento(data.get('telefono') or data.get('Número de Teléfono') or data.get('Numero de Telefono') or data.get('TELEFONO') or data.get('celular'))
    coordinador = limpiar_valor(data.get('coordinador') or data.get('COORDINADOR') or data.get('coordinador_responsable')).upper()
    tipo_equipo = limpiar_valor(data.get('tipo_equipo') or data.get('tipo') or data.get('equipo') or data.get('TIPO EQUIPO'))
    if not tipo_equipo:
        cargo_norm = normalizar_texto_clave(cargo)
        if 'coordin' in cargo_norm:
            tipo_equipo = 'COORDINADOR'
        elif 'psicosocial' in cargo_norm or 'psicolog' in cargo_norm:
            tipo_equipo = 'PSICOSOCIAL'
        elif 'enfermer' in cargo_norm or 'nutricion' in cargo_norm:
            tipo_equipo = 'ENFERMERIA'
        elif 'pedagog' in cargo_norm:
            tipo_equipo = 'PEDAGOGIA'
        elif 'agente' in cargo_norm or 'docente' in cargo_norm:
            tipo_equipo = 'DOCENTE'
        else:
            tipo_equipo = 'OTRO'
    tipo_equipo = tipo_equipo.upper()
    contrato = limpiar_valor(data.get('contrato') or data.get('CONTRATO') or data.get('numero_contrato'))
    perfil = limpiar_valor(data.get('perfil') or data.get('PERFIL') or data.get('profesion') or data.get('PROFESION'))
    estado = limpiar_valor(data.get('estado') or 'activo').lower() or 'activo'
    nombres, apellidos = dividir_nombre(nombre)
    return {
        'documento': documento,
        'nombre': nombre,
        'nombres': nombres,
        'apellidos': apellidos,
        'cargo': cargo,
        'unidad': unidad,
        'unidades': json.dumps([unidad]) if unidad else json.dumps([]),
        'direccion': direccion,
        'telefono': telefono,
        'coordinador': coordinador,
        'tipo_equipo': tipo_equipo,
        'contrato': contrato,
        'perfil': perfil,
        'estado': estado,
        'activo': 0 if estado in {'inactivo', 'eliminado', 'retirado'} else 1,
        'archivo': archivo
    }


def parsear_talento_dataframe(df, archivo='talento', coordinador_respaldo=''):
    """Lee archivos de talento humano con encabezados simples o institucionales.

    Soporta bases tipo ICBF/Cuéntame de Talento Humano UDS, donde el nombre
    puede venir separado en Primer Nombre, Segundo Nombre, Primer Apellido y
    Segundo Apellido, y donde el documento real aparece como Número de
    Documento. Se evita tomar columnas auxiliares como Tipo de Documento,
    Tipo de Unidad, Regional UDS o Código UDS.
    """
    registros = []
    if df is None or df.empty:
        return registros

    def fila_parece_encabezado(valores):
        textos = [normalizar_texto_clave(v) for v in valores]
        tiene_documento = any(
            t in {'numero de documento', 'nro documento', 'no documento', 'documento', 'cedula', 'cedula de ciudadania', 'identificacion'}
            or (('documento' in t or 'cedula' in t or 'identificacion' in t) and 'tipo de documento' not in t)
            for t in textos
        )
        tiene_nombre = any(
            t in {'nombre', 'nombre completo', 'nombres y apellidos', 'primer nombre'}
            or 'nombres y apellidos' in t
            or 'primer nombre' in t
            for t in textos
        )
        tiene_cargo = any(t == 'cargo' or t.endswith(' cargo') or ' cargo ' in f' {t} ' for t in textos)
        tiene_uds = any(t in {'nombre uds', 'unidad', 'comunidad'} or 'nombre uds' in t for t in textos)
        return tiene_documento and (tiene_nombre or tiene_cargo or tiene_uds)

    # Detectar fila de encabezados institucionales cuando el archivo fue leído
    # con header=None. Si ya llegó con encabezados reales, se conserva.
    columnas_actuales = [normalizar_texto_clave(c) for c in df.columns]
    columnas_son_posicionales = all(str(c).isdigit() for c in df.columns) or all(c.startswith('unnamed') or c == '' for c in columnas_actuales)

    header_row = None
    if columnas_son_posicionales:
        for idx in range(min(len(df), 25)):
            if fila_parece_encabezado(df.iloc[idx].tolist()):
                header_row = idx
                break

    if header_row is not None:
        headers = [limpiar_valor(v) or f'col_{i}' for i, v in enumerate(df.iloc[header_row].tolist())]
        data = df.iloc[header_row + 1:].copy()
        data.columns = headers
    else:
        data = df.copy()

    columnas_items = []
    for i, c in enumerate(data.columns):
        norm = normalizar_texto_clave(c)
        if norm:
            columnas_items.append((norm, c, i))

    def col(*aliases, excluir=()):
        alias_norms = [normalizar_texto_clave(a) for a in aliases if normalizar_texto_clave(a)]
        excluir_norms = [normalizar_texto_clave(e) for e in excluir if normalizar_texto_clave(e)]
        candidatos = []
        for norm, original, orden in columnas_items:
            if any(ex and ex in norm for ex in excluir_norms):
                continue
            for alias_idx, alias_norm in enumerate(alias_norms):
                if norm == alias_norm:
                    score = 0
                elif norm.endswith(' ' + alias_norm) or norm.startswith(alias_norm + ' '):
                    score = 10
                elif alias_norm in norm:
                    score = 20
                else:
                    continue
                candidatos.append((score, alias_idx, len(norm), orden, original))
        if not candidatos:
            return None
        candidatos.sort()
        return candidatos[0][-1]

    col_nombre = col('nombres y apellidos', 'nombre completo', 'nombre talento humano', 'funcionario', 'talento humano', 'nombre', excluir=('primer nombre', 'segundo nombre', 'nombre uds', 'nombre zona'))
    col_primer_nombre = col('primer nombre')
    col_segundo_nombre = col('segundo nombre')
    col_primer_apellido = col('primer apellido')
    col_segundo_apellido = col('segundo apellido')
    col_documento = col('numero de documento', 'número de documento', 'no documento', 'nro documento', 'documento de identidad', 'identificacion', 'identificación', 'cedula', 'cédula', 'documento', excluir=('tipo de documento',))
    col_cargo = col('cargo')
    col_unidad = col('nombre uds', 'nombre unidad', 'unidad de servicio', 'unidad servicio', 'comunidad', 'uca', 'unidad', excluir=('tipo de unidad', 'regional uds', 'codigo uds', 'código uds'))
    col_direccion = col('direccion de residencia', 'dirección de residencia', 'direccion', 'dirección')
    col_telefono = col('numero de telefono', 'número de teléfono', 'telefono', 'teléfono', 'celular', 'contacto')
    col_coordinador = col('coordinador responsable', 'coordinador', 'supervisor', 'responsable')
    col_tipo_equipo = col('tipo equipo', 'tipo de equipo', 'equipo', 'rol')
    col_contrato = col('numero contrato', 'número contrato', 'numero de contrato', 'número de contrato', 'contrato')
    col_perfil = col('perfil', 'profesion', 'profesión', 'titulo obtenido de educacion superior', 'título obtenido de educación superior')
    col_estado = col('estado de vinculacion', 'estado de vinculación', 'estado')

    for _, fila in data.iterrows():
        nombre_compuesto = unir_partes(
            fila.get(col_primer_nombre) if col_primer_nombre else '',
            fila.get(col_segundo_nombre) if col_segundo_nombre else '',
            fila.get(col_primer_apellido) if col_primer_apellido else '',
            fila.get(col_segundo_apellido) if col_segundo_apellido else '',
        )
        nombre = nombre_compuesto or (limpiar_valor(fila.get(col_nombre)) if col_nombre else '')
        documento = limpiar_documento_talento(fila.get(col_documento)) if col_documento else ''
        cargo = limpiar_valor(fila.get(col_cargo)) if col_cargo else 'AGENTE EDUCATIVO'
        unidad = limpiar_valor(fila.get(col_unidad)) if col_unidad else ''
        if not nombre or not documento:
            continue
        registro = crear_registro_talento({
            'nombre': nombre,
            'documento': documento,
            'cargo': cargo or 'AGENTE EDUCATIVO',
            'unidad': unidad,
            'direccion': fila.get(col_direccion) if col_direccion else unidad,
            'telefono': fila.get(col_telefono) if col_telefono else '',
            'coordinador': fila.get(col_coordinador) if col_coordinador else coordinador_respaldo,
            'tipo_equipo': fila.get(col_tipo_equipo) if col_tipo_equipo else '',
            'contrato': fila.get(col_contrato) if col_contrato else '',
            'perfil': fila.get(col_perfil) if col_perfil else '',
            'estado': fila.get(col_estado) if col_estado else 'activo',
            'primer_nombre': fila.get(col_primer_nombre) if col_primer_nombre else '',
            'segundo_nombre': fila.get(col_segundo_nombre) if col_segundo_nombre else '',
            'primer_apellido': fila.get(col_primer_apellido) if col_primer_apellido else '',
            'segundo_apellido': fila.get(col_segundo_apellido) if col_segundo_apellido else '',
        }, archivo=archivo)
        registros.append(registro)
    return registros


def guardar_registros_talento(registros):
    """Guarda Talento Humano mediante el módulo migrado a SQLAlchemy Core."""
    from modules.talento_humano.services import TalentoHumanoService
    resultado = TalentoHumanoService().guardar_registros(registros, origen='guardar_registros_talento')
    return int(resultado.get('total', 0))

def _talento_current_context():
    user = getattr(g, 'current_user', None) or {}
    return {
        'fundacion_id': user.get('fundacion_id') or 1,
        'usuario_id': user.get('id'),
        'username': user.get('username') or 'sistema'
    }


def _talento_table_columns(cursor, table):
    try:
        return {row['name'] for row in cursor.execute(f"PRAGMA table_info({table})").fetchall()}
    except Exception:
        return set()


def _talento_ensure_column(cursor, table, column, definition):
    try:
        cols = _talento_table_columns(cursor, table)
        if column not in cols:
            cursor.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")
    except Exception:
        pass


def _talento_ensure_global_schema(cursor):
    """Asegura tablas mínimas de integración sin depender de que el usuario abra cada módulo."""
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS gp_coordinadores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contrato_id INTEGER,
            contrato TEXT,
            nombre TEXT NOT NULL,
            documento TEXT,
            telefono TEXT,
            email TEXT,
            cargo TEXT DEFAULT 'COORDINADOR',
            zona TEXT,
            unidades_json TEXT,
            observaciones TEXT,
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT NOT NULL,
            fecha_actualizacion TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS gp_docentes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            coordinador_id INTEGER,
            nombre TEXT NOT NULL,
            documento TEXT,
            unidad TEXT,
            telefono TEXT,
            email TEXT,
            cargo TEXT DEFAULT 'DOCENTE',
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT NOT NULL,
            fecha_actualizacion TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS gp_equipos_interdisciplinarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            coordinador_id INTEGER,
            nombre TEXT NOT NULL,
            documento TEXT,
            rol TEXT NOT NULL,
            profesion TEXT,
            telefono TEXT,
            email TEXT,
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT NOT NULL,
            fecha_actualizacion TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS gp_unidades_asignadas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            coordinador_id INTEGER NOT NULL,
            unidad TEXT NOT NULL,
            estado TEXT DEFAULT 'activo',
            fecha_creacion TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS gp_asignaciones_coordinador (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            coordinador_id INTEGER,
            tipo_talento TEXT,
            origen_tabla TEXT,
            origen_id INTEGER,
            nombre TEXT,
            documento TEXT,
            cargo TEXT,
            rol TEXT,
            unidad TEXT,
            telefono TEXT,
            email TEXT,
            estado TEXT DEFAULT 'ACTIVO',
            fecha_inicio TEXT,
            fecha_fin TEXT,
            observaciones TEXT,
            fecha_creacion TEXT,
            fecha_actualizacion TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS gp_historial_acciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT,
            accion TEXT NOT NULL,
            entidad_tipo TEXT,
            entidad_id INTEGER,
            datos_anteriores TEXT,
            datos_nuevos TEXT,
            fecha_accion TEXT NOT NULL
        )
    """)

    common_columns = {
        'fundacion_id': 'INTEGER',
        'usuario_creador_id': 'INTEGER',
        'fecha_actualizacion': 'TEXT'
    }
    for table in [
        'gp_coordinadores', 'gp_docentes', 'gp_equipos_interdisciplinarios',
        'gp_unidades_asignadas', 'gp_asignaciones_coordinador', 'gp_historial_acciones'
    ]:
        for col, definition in common_columns.items():
            _talento_ensure_column(cursor, table, col, definition)

    # Compatibilidad de tablas principales ya existentes.
    for col, definition in {
        'direccion': 'TEXT',
        'telefono': 'TEXT',
        'docente_asignado': 'TEXT',
        'docente_documento': 'TEXT',
        'coordinador_nombre': 'TEXT',
        'contrato': 'TEXT',
        'fecha_actualizacion': 'TEXT',
    }.items():
        _talento_ensure_column(cursor, 'unidades', col, definition)

    for col, definition in {
        'fundacion_id': 'INTEGER',
        'usuario_creador_id': 'INTEGER',
        'fecha_creacion': 'TEXT',
        'fecha_actualizacion': 'TEXT',
    }.items():
        _talento_ensure_column(cursor, 'coordinadores', col, definition)


def _talento_norm(value):
    return normalizar_texto_clave(value)


def _talento_row(row, key, default=''):
    try:
        value = row[key]
    except Exception:
        value = default
    return limpiar_valor(value, default)


def _talento_tipo(row):
    texto = _talento_norm(' '.join([
        _talento_row(row, 'tipo_equipo'),
        _talento_row(row, 'cargo'),
        _talento_row(row, 'perfil')
    ]))
    if 'coordin' in texto:
        return 'COORDINADOR'
    if 'psicosocial' in texto or 'psicolog' in texto:
        return 'PSICOSOCIAL'
    if 'nutricion' in texto or 'nutricionista' in texto:
        return 'NUTRICIONISTA'
    if 'enfermer' in texto or 'salud' in texto:
        return 'ENFERMERIA'
    if 'pedagog' in texto:
        return 'PEDAGOGIA'
    if 'administr' in texto or 'auxiliar' in texto:
        return 'ADMINISTRATIVO'
    if 'docente' in texto or 'agente educativo' in texto or 'agente' in texto:
        return 'DOCENTE'
    return _talento_row(row, 'tipo_equipo') or 'OTRO'


def _talento_unidades(row):
    unidades = set()
    unidad = normalize_unidad(_talento_row(row, 'unidad'))
    if unidad:
        unidades.add(unidad)
    raw = _talento_row(row, 'unidades')
    if raw:
        try:
            parsed = json.loads(raw)
            if isinstance(parsed, list):
                unidades.update(normalize_unidad(x) for x in parsed if normalize_unidad(x))
            elif isinstance(parsed, str):
                unidades.add(normalize_unidad(parsed))
        except Exception:
            unidades.update(normalize_unidad(x) for x in re.split(r'[,;/|]+', raw) if normalize_unidad(x))
    return sorted(unidades)


def _talento_find_coordinador_id(cursor, nombre='', documento='', contrato='', fundacion_id=1):
    nombre_norm = _talento_norm(nombre)
    documento = limpiar_valor(documento)
    contrato = limpiar_valor(contrato)
    params = [fundacion_id, fundacion_id]
    if documento:
        row = cursor.execute("""
            SELECT id FROM gp_coordinadores
            WHERE COALESCE(documento,'') = ? AND (fundacion_id = ? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (documento, fundacion_id)).fetchone()
        if row:
            return row['id']
    if nombre_norm:
        rows = cursor.execute("""
            SELECT id, nombre FROM gp_coordinadores
            WHERE (fundacion_id = ? OR fundacion_id IS NULL) AND activo = 1
        """, (fundacion_id,)).fetchall()
        for row in rows:
            if _talento_norm(row['nombre']) == nombre_norm:
                return row['id']
    if contrato:
        row = cursor.execute("""
            SELECT id FROM gp_coordinadores
            WHERE COALESCE(contrato,'') = ? AND (fundacion_id = ? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (contrato, fundacion_id)).fetchone()
        if row:
            return row['id']
    return None


def _talento_upsert_coordinador(cursor, row, fundacion_id, usuario_id, ahora, placeholder_name=None):
    nombre = limpiar_valor(placeholder_name or _talento_row(row, 'nombre') or 'SIN COORDINADOR ASIGNADO').upper()
    documento = _talento_row(row, 'documento') if not placeholder_name else ''
    telefono = _talento_row(row, 'telefono') if not placeholder_name else ''
    contrato = _talento_row(row, 'contrato') if not placeholder_name else _talento_row(row, 'contrato')
    cargo = _talento_row(row, 'cargo') if not placeholder_name else 'COORDINADOR'
    if not cargo:
        cargo = 'COORDINADOR'
    unidades = _talento_unidades(row)
    unidades_json = json.dumps(unidades, ensure_ascii=False)

    existing_id = _talento_find_coordinador_id(cursor, nombre, documento, contrato, fundacion_id)
    if existing_id:
        anterior = cursor.execute("SELECT unidades_json FROM gp_coordinadores WHERE id=?", (existing_id,)).fetchone()
        prev_units = []
        try:
            prev_units = json.loads(anterior['unidades_json'] or '[]') if anterior else []
        except Exception:
            prev_units = []
        merged_units = sorted({normalize_unidad(u) for u in list(prev_units) + unidades if normalize_unidad(u)})
        cursor.execute("""
            UPDATE gp_coordinadores
            SET nombre = ?, documento = COALESCE(NULLIF(?,''), documento), telefono = COALESCE(NULLIF(?,''), telefono),
                cargo = COALESCE(NULLIF(?,''), cargo), contrato = COALESCE(NULLIF(?,''), contrato),
                unidades_json = ?, activo = 1, fundacion_id = COALESCE(fundacion_id, ?),
                usuario_creador_id = COALESCE(usuario_creador_id, ?), fecha_actualizacion = ?
            WHERE id = ?
        """, (
            nombre, documento, telefono, cargo, contrato,
            json.dumps(merged_units, ensure_ascii=False), fundacion_id, usuario_id, ahora, existing_id
        ))
        return existing_id, False

    cursor.execute("""
        INSERT INTO gp_coordinadores
        (contrato, nombre, documento, telefono, email, cargo, unidades_json, activo,
         fecha_creacion, fecha_actualizacion, fundacion_id, usuario_creador_id)
        VALUES (?, ?, ?, ?, '', ?, ?, 1, ?, ?, ?, ?)
    """, (contrato, nombre, documento, telefono, cargo, unidades_json, ahora, ahora, fundacion_id, usuario_id))
    return cursor.lastrowid, True


def _talento_upsert_docente(cursor, row, coordinador_id, fundacion_id, usuario_id, ahora):
    nombre = _talento_row(row, 'nombre')
    documento = _talento_row(row, 'documento')
    unidad = normalize_unidad(_talento_row(row, 'unidad'))
    telefono = _talento_row(row, 'telefono')
    cargo = _talento_row(row, 'cargo') or 'DOCENTE'
    if documento:
        existente = cursor.execute("""
            SELECT id FROM gp_docentes
            WHERE documento=? AND (fundacion_id=? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (documento, fundacion_id)).fetchone()
    else:
        existente = cursor.execute("""
            SELECT id FROM gp_docentes
            WHERE upper(nombre)=upper(?) AND COALESCE(unidad,'')=? AND (fundacion_id=? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (nombre, unidad, fundacion_id)).fetchone()
    if existente:
        cursor.execute("""
            UPDATE gp_docentes
            SET coordinador_id=?, nombre=?, unidad=?, telefono=?, cargo=?, activo=1,
                fundacion_id=COALESCE(fundacion_id, ?), usuario_creador_id=COALESCE(usuario_creador_id, ?),
                fecha_actualizacion=?
            WHERE id=?
        """, (coordinador_id, nombre, unidad, telefono, cargo, fundacion_id, usuario_id, ahora, existente['id']))
        return existente['id'], False
    cursor.execute("""
        INSERT INTO gp_docentes
        (coordinador_id, nombre, documento, unidad, telefono, email, cargo, activo,
         fecha_creacion, fecha_actualizacion, fundacion_id, usuario_creador_id)
        VALUES (?, ?, ?, ?, ?, '', ?, 1, ?, ?, ?, ?)
    """, (coordinador_id, nombre, documento, unidad, telefono, cargo, ahora, ahora, fundacion_id, usuario_id))
    return cursor.lastrowid, True


def _talento_upsert_equipo(cursor, row, coordinador_id, fundacion_id, usuario_id, ahora):
    nombre = _talento_row(row, 'nombre')
    documento = _talento_row(row, 'documento')
    rol = _talento_tipo(row)
    profesion = _talento_row(row, 'perfil') or _talento_row(row, 'cargo')
    telefono = _talento_row(row, 'telefono')
    if documento:
        existente = cursor.execute("""
            SELECT id FROM gp_equipos_interdisciplinarios
            WHERE documento=? AND (fundacion_id=? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (documento, fundacion_id)).fetchone()
    else:
        existente = cursor.execute("""
            SELECT id FROM gp_equipos_interdisciplinarios
            WHERE upper(nombre)=upper(?) AND (fundacion_id=? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (nombre, fundacion_id)).fetchone()
    if existente:
        cursor.execute("""
            UPDATE gp_equipos_interdisciplinarios
            SET coordinador_id=?, nombre=?, rol=?, profesion=?, telefono=?, activo=1,
                fundacion_id=COALESCE(fundacion_id, ?), usuario_creador_id=COALESCE(usuario_creador_id, ?),
                fecha_actualizacion=?
            WHERE id=?
        """, (coordinador_id, nombre, rol, profesion, telefono, fundacion_id, usuario_id, ahora, existente['id']))
        return existente['id'], False
    cursor.execute("""
        INSERT INTO gp_equipos_interdisciplinarios
        (coordinador_id, nombre, documento, rol, profesion, telefono, email, activo,
         fecha_creacion, fecha_actualizacion, fundacion_id, usuario_creador_id)
        VALUES (?, ?, ?, ?, ?, ?, '', 1, ?, ?, ?, ?)
    """, (coordinador_id, nombre, documento, rol, profesion, telefono, ahora, ahora, fundacion_id, usuario_id))
    return cursor.lastrowid, True


def _talento_upsert_asignacion(cursor, row, coordinador_id, tipo, fundacion_id, usuario_id, ahora):
    nombre = _talento_row(row, 'nombre')
    documento = _talento_row(row, 'documento')
    unidad = normalize_unidad(_talento_row(row, 'unidad'))
    origen_id = row['id']
    existente = cursor.execute("""
        SELECT id FROM gp_asignaciones_coordinador
        WHERE origen_tabla='coordinadores' AND origen_id=?
        ORDER BY id LIMIT 1
    """, (origen_id,)).fetchone()
    payload = (
        coordinador_id, tipo, nombre, documento, _talento_row(row, 'cargo'), tipo, unidad,
        _talento_row(row, 'telefono'), 'ACTIVO', fundacion_id, usuario_id, ahora
    )
    if existente:
        cursor.execute("""
            UPDATE gp_asignaciones_coordinador
            SET coordinador_id=?, tipo_talento=?, nombre=?, documento=?, cargo=?, rol=?, unidad=?,
                telefono=?, estado=?, fundacion_id=COALESCE(fundacion_id, ?),
                usuario_creador_id=COALESCE(usuario_creador_id, ?), fecha_actualizacion=?
            WHERE id=?
        """, payload + (existente['id'],))
        return False
    cursor.execute("""
        INSERT INTO gp_asignaciones_coordinador
        (coordinador_id, tipo_talento, origen_tabla, origen_id, nombre, documento, cargo, rol, unidad,
         telefono, email, estado, fecha_inicio, observaciones, fundacion_id, usuario_creador_id, fecha_creacion, fecha_actualizacion)
        VALUES (?, ?, 'coordinadores', ?, ?, ?, ?, ?, ?, ?, '', ?, ?, 'Sincronizado desde Talento Humano', ?, ?, ?, ?)
    """, (
        coordinador_id, tipo, origen_id, nombre, documento, _talento_row(row, 'cargo'), tipo, unidad,
        _talento_row(row, 'telefono'), 'ACTIVO', ahora[:10], fundacion_id, usuario_id, ahora, ahora
    ))
    return True


def _talento_upsert_unidad_asignada(cursor, coordinador_id, unidad, fundacion_id, usuario_id, ahora):
    unidad = normalize_unidad(unidad)
    if not coordinador_id or not unidad:
        return False
    existente = cursor.execute("""
        SELECT id FROM gp_unidades_asignadas
        WHERE coordinador_id=? AND upper(unidad)=upper(?) AND estado='activo'
        LIMIT 1
    """, (coordinador_id, unidad)).fetchone()
    if existente:
        cursor.execute("""
            UPDATE gp_unidades_asignadas
            SET fundacion_id=COALESCE(fundacion_id, ?), usuario_creador_id=COALESCE(usuario_creador_id, ?),
                fecha_actualizacion=?
            WHERE id=?
        """, (fundacion_id, usuario_id, ahora, existente['id']))
        return False
    cursor.execute("""
        INSERT INTO gp_unidades_asignadas
        (coordinador_id, unidad, estado, fecha_creacion, fecha_actualizacion, fundacion_id, usuario_creador_id)
        VALUES (?, ?, 'activo', ?, ?, ?, ?)
    """, (coordinador_id, unidad, ahora, ahora, fundacion_id, usuario_id))
    return True


def sincronizar_talento_global(origen='manual'):
    """Sincroniza Talento Humano como fuente maestra usando SQLAlchemy Core."""
    from modules.talento_humano.services import TalentoHumanoService
    return TalentoHumanoService().sincronizar_global(origen=origen)

def resumen_integracion_talento():
    """Resumen de integración de Talento Humano desde el servicio migrado."""
    from modules.talento_humano.services import TalentoHumanoService
    return TalentoHumanoService().resumen_integracion()

def _th_safe_json(data):
    try:
        return json.dumps(data, ensure_ascii=False, default=str)
    except Exception:
        return '{}'


def _th_normalizar_rol(row):
    cargo = normalizar_texto_clave(_talento_row(row, 'cargo'))
    tipo = normalizar_texto_clave(_talento_row(row, 'tipo_equipo'))
    perfil = normalizar_texto_clave(_talento_row(row, 'perfil'))
    texto = f'{cargo} {tipo} {perfil}'
    if 'coordinador' in texto or 'coordinadora' in texto:
        return 'COORDINADOR'
    if 'psicosocial' in texto or 'psicolog' in texto or 'trabajador social' in texto:
        return 'PSICOSOCIAL'
    if 'enfermer' in texto or 'salud' in texto or 'nutricion' in texto or 'nutrición' in texto:
        if 'nutric' in texto:
            return 'NUTRICIONISTA'
        return 'ENFERMERA'
    if 'pedagog' in texto or 'pedagoga' in texto or 'pedagogo' in texto:
        return 'PEDAGOGA'
    if 'admin' in texto or 'contad' in texto or 'auxiliar' in texto:
        return 'AUXILIAR_ADMINISTRATIVO'
    if 'agente' in texto or 'docente' in texto or 'educativo' in texto or 'educativa' in texto:
        return 'DOCENTE'
    return 'APOYO'


def _th_find_gp_coordinador_id(cursor, nombre, documento, fundacion_id):
    nombre = limpiar_valor(nombre)
    documento = limpiar_valor(documento)
    if documento:
        row = cursor.execute("""
            SELECT id FROM gp_coordinadores
            WHERE documento=? AND (fundacion_id=? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (documento, fundacion_id)).fetchone()
        if row:
            return row['id']
    if nombre:
        row = cursor.execute("""
            SELECT id FROM gp_coordinadores
            WHERE upper(nombre)=upper(?) AND (fundacion_id=? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (nombre, fundacion_id)).fetchone()
        if row:
            return row['id']
    return None


def _th_upsert_persona(cursor, row, fundacion_id, usuario_id, ahora):
    rol = _th_normalizar_rol(row)
    documento = _talento_row(row, 'documento')
    nombre = _talento_row(row, 'nombre')
    if not nombre:
        return None, False
    unidad = normalize_unidad(_talento_row(row, 'unidad'))
    try:
        origen_id = row['id']
    except Exception:
        origen_id = None

    existente = None
    if documento:
        existente = cursor.execute("""
            SELECT * FROM th_personas
            WHERE documento=? AND (fundacion_id=? OR fundacion_id IS NULL)
            ORDER BY id LIMIT 1
        """, (documento, fundacion_id)).fetchone()
    if not existente and origen_id:
        existente = cursor.execute("""
            SELECT * FROM th_personas
            WHERE origen_tabla='coordinadores' AND origen_id=?
            ORDER BY id LIMIT 1
        """, (origen_id,)).fetchone()

    payload = {
        'documento': documento,
        'nombre': nombre,
        'nombres': _talento_row(row, 'nombres') or dividir_nombre(nombre)[0],
        'apellidos': _talento_row(row, 'apellidos') or dividir_nombre(nombre)[1],
        'cargo': _talento_row(row, 'cargo'),
        'tipo_equipo': _talento_row(row, 'tipo_equipo'),
        'rol_normalizado': rol,
        'unidad': unidad,
        'direccion': _talento_row(row, 'direccion'),
        'telefono': _talento_row(row, 'telefono'),
        'coordinador': _talento_row(row, 'coordinador'),
        'contrato': _talento_row(row, 'contrato'),
        'perfil': _talento_row(row, 'perfil'),
        'estado': _talento_row(row, 'estado') or 'activo',
        'activo': int(row['activo']) if 'activo' in row.keys() and row['activo'] is not None else 1,
        'origen_tabla': 'coordinadores',
        'origen_id': origen_id,
        'archivo': _talento_row(row, 'archivo'),
        'fundacion_id': fundacion_id,
        'usuario_creador_id': usuario_id,
        'fecha_actualizacion': ahora,
    }

    if existente:
        datos_anteriores = dict(existente)
        cursor.execute("""
            UPDATE th_personas
            SET documento=:documento, nombre=:nombre, nombres=:nombres, apellidos=:apellidos,
                cargo=:cargo, tipo_equipo=:tipo_equipo, rol_normalizado=:rol_normalizado,
                unidad=:unidad, direccion=:direccion, telefono=:telefono, coordinador=:coordinador,
                contrato=:contrato, perfil=:perfil, estado=:estado, activo=:activo,
                origen_tabla=:origen_tabla, origen_id=:origen_id, archivo=:archivo,
                fundacion_id=COALESCE(fundacion_id, :fundacion_id),
                usuario_creador_id=COALESCE(usuario_creador_id, :usuario_creador_id),
                fecha_actualizacion=:fecha_actualizacion
            WHERE id=:id
        """, {**payload, 'id': existente['id']})
        cursor.execute("""
            INSERT INTO th_historial (persona_id, accion, datos_anteriores, datos_nuevos, usuario, fundacion_id, fecha_accion)
            VALUES (?, 'ACTUALIZAR_DESDE_TALENTO_BASE', ?, ?, ?, ?, ?)
        """, (existente['id'], _th_safe_json(datos_anteriores), _th_safe_json(payload), _talento_current_context().get('username'), fundacion_id, ahora))
        return existente['id'], False

    cursor.execute("""
        INSERT INTO th_personas
        (documento, nombre, nombres, apellidos, cargo, tipo_equipo, rol_normalizado, unidad,
         direccion, telefono, coordinador, contrato, perfil, estado, activo, origen_tabla,
         origen_id, archivo, fundacion_id, usuario_creador_id, fecha_creacion, fecha_actualizacion)
        VALUES (:documento, :nombre, :nombres, :apellidos, :cargo, :tipo_equipo, :rol_normalizado,
                :unidad, :direccion, :telefono, :coordinador, :contrato, :perfil, :estado,
                :activo, :origen_tabla, :origen_id, :archivo, :fundacion_id,
                :usuario_creador_id, :fecha_creacion, :fecha_actualizacion)
    """, {**payload, 'fecha_creacion': ahora})
    persona_id = cursor.lastrowid
    cursor.execute("""
        INSERT INTO th_historial (persona_id, accion, datos_anteriores, datos_nuevos, usuario, fundacion_id, fecha_accion)
        VALUES (?, 'CREAR_DESDE_TALENTO_BASE', NULL, ?, ?, ?, ?)
    """, (persona_id, _th_safe_json(payload), _talento_current_context().get('username'), fundacion_id, ahora))
    return persona_id, True


def _th_upsert_asignacion(cursor, persona_id, row, fundacion_id, usuario_id, ahora):
    if not persona_id:
        return False
    rol = _th_normalizar_rol(row)
    unidad = normalize_unidad(_talento_row(row, 'unidad'))
    coordinador_nombre = _talento_row(row, 'coordinador')
    coordinador_id = _th_find_gp_coordinador_id(cursor, coordinador_nombre, '', fundacion_id)
    existente = cursor.execute("""
        SELECT id FROM th_asignaciones
        WHERE persona_id=? AND COALESCE(unidad,'')=COALESCE(?, '') AND COALESCE(rol,'')=COALESCE(?, '')
          AND COALESCE(estado,'ACTIVO')='ACTIVO'
        ORDER BY id LIMIT 1
    """, (persona_id, unidad, rol)).fetchone()
    payload = {
        'persona_id': persona_id,
        'coordinador_id': coordinador_id,
        'coordinador_nombre': coordinador_nombre,
        'unidad': unidad,
        'rol': rol,
        'cargo': _talento_row(row, 'cargo'),
        'estado': 'ACTIVO'
    }
    # activo se normaliza aparte para filas SQLite y registros manuales.
    try:
        activo = int(row['activo']) if 'activo' in row.keys() and row['activo'] is not None else 1
    except Exception:
        activo = 1
    payload['estado'] = 'ACTIVO' if activo else 'INACTIVO'

    if existente:
        cursor.execute("""
            UPDATE th_asignaciones
            SET coordinador_id=?, coordinador_nombre=?, cargo=?, estado=?,
                fundacion_id=COALESCE(fundacion_id, ?), usuario_creador_id=COALESCE(usuario_creador_id, ?),
                fecha_actualizacion=?
            WHERE id=?
        """, (coordinador_id, coordinador_nombre, payload['cargo'], payload['estado'], fundacion_id, usuario_id, ahora, existente['id']))
        return False
    cursor.execute("""
        INSERT INTO th_asignaciones
        (persona_id, coordinador_id, coordinador_nombre, unidad, rol, cargo, estado, fecha_inicio,
         observaciones, fundacion_id, usuario_creador_id, fecha_creacion, fecha_actualizacion)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'Sincronizado desde Talento Humano fuente maestra', ?, ?, ?, ?)
    """, (persona_id, coordinador_id, coordinador_nombre, unidad, rol, payload['cargo'], payload['estado'], ahora[:10], fundacion_id, usuario_id, ahora, ahora))
    return True


def sincronizar_talento_maestro(origen='manual', cursor=None):
    """Compatibilidad histórica: delega en la sincronización global SQLAlchemy Core."""
    from modules.talento_humano.services import TalentoHumanoService
    return TalentoHumanoService().sincronizar_global(origen=origen)

def resumen_talento_maestro():
    """Compatibilidad histórica: devuelve el resumen de la fuente maestra th_*."""
    from modules.talento_humano.services import TalentoHumanoService
    return TalentoHumanoService().resumen_integracion()

def normalize_unidad(unidad):
    """Normaliza una UDS mediante el catálogo central y conserva desconocidas."""
    return uds_normalize_unit(unidad, preserve_unknown=True)


def unidad_es_valida(unidad):
    unidad_norm = normalize_unidad(unidad)
    return bool(unidad_norm and unidad_norm not in UNIDADES_INVALIDAS)


def clasificar_grupo_edad(tipo_beneficiario='', fecha_nacimiento=None, edad_meses=None):
    """Clasifica beneficiarios en los grupos operativos del tablero y del RPP.

    Reglas solicitadas:
    - 0 a 6 meses y gestantes se muestran en una sola categoría.
    - 6 a 11 meses y 29 días.
    - 1 a 2 años y 11 meses.
    - 3 a 5 años y 11 meses.
    """
    tipo = normalizar_texto_clave(tipo_beneficiario)
    if 'gestante' in tipo:
        return '0 A 6 MESES Y GESTANTES'
    if edad_meses is None:
        edad_meses = calcular_edad_meses(fecha_nacimiento)
    try:
        edad_meses = int(edad_meses)
    except Exception:
        edad_meses = 0

    if 'menor de seis meses' in tipo or edad_meses <= 5:
        return '0 A 6 MESES Y GESTANTES'
    if 6 <= edad_meses <= 11:
        return '6 A 11 MESES 29 DÍAS'
    if 12 <= edad_meses <= 35:
        return '1 A 2 AÑOS 11 MESES'
    if 36 <= edad_meses <= 71:
        return '3 A 5 AÑOS 11 MESES'
    if edad_meses >= 72:
        return '5 AÑOS EN ADELANTE'
    return 'FUERA DE RANGO'


def equivalentes_unidad(unidad):
    """Devuelve claves equivalentes de la UDS según el catálogo central."""
    return {normalizar_texto_clave(valor) for valor in uds_equivalent_values(unidad) if valor}


def detectar_columna_unidad(df):
    """
    Detecta la columna real de unidad por encabezado y por el contenido de las filas.
    Evita tomar columnas de estado con valores como ACTIVO.
    """
    preferida = buscar_columna(df, [
        'Nombre de la unidad de servicio',
        'Unidad de servicio',
        'Nombre unidad de servicio',
        'Unidad de atención',
        'Unidad de Atencion',
        'UDS',
        'UCA'
    ])

    mejores = []
    for col in df.columns:
        encabezado = normalizar_texto_clave(col)
        score = 0

        if encabezado == 'nombre de la unidad de servicio':
            score += 1000
        elif 'nombre de la unidad' in encabezado:
            score += 700
        elif 'unidad de servicio' in encabezado and 'tipo de unidad' not in encabezado:
            score += 500
        elif encabezado in {'unidad', 'uds', 'uca', 'unidad de atencion', 'unidad de atención'}:
            score += 250

        if 'estado' in encabezado or encabezado == 'tipo de unidad':
            score -= 700

        serie = df[col].dropna().astype(str).head(150)
        for valor in serie:
            raw = valor.strip().upper()
            raw_norm = normalizar_texto_clave(raw)
            unidad_norm = normalize_unidad(raw)

            if not raw_norm or raw in UNIDADES_INVALIDAS:
                score -= 3
                continue

            if raw.startswith('UCA '):
                score += 8
            if unidad_norm in KNOWN_UNITS:
                score += 5
            if unidad_norm != raw and unidad_norm:
                score += 4
            if 'unidad demo' in raw_norm:
                score += 4

        mejores.append((score, col))

    mejores.sort(reverse=True, key=lambda item: item[0])
    if mejores and mejores[0][0] > 0:
        return mejores[0][1]
    return preferida


def extraer_unidad_desde_fila(fila):
    """Busca dentro de una fila el primer valor que parezca unidad real."""
    for valor in fila.values:
        unidad = normalize_unidad(valor)
        raw = str(valor or '').strip().upper()
        if unidad and unidad not in UNIDADES_INVALIDAS and (raw.startswith('UCA ') or unidad in KNOWN_UNITS):
            return unidad
    for valor in fila.values:
        unidad = normalize_unidad(valor)
        if unidad and unidad not in UNIDADES_INVALIDAS:
            raw = normalizar_texto_clave(valor)
            if 'unidad demo' in raw:
                return unidad
    return 'SIN UNIDAD'


def obtener_unidades_registradas():
    """Devuelve unidades guardadas con actividad real.

    No se agregan unidades conocidas con cero usuarios al diagnóstico, porque eso
    generaba filas rojas innecesarias y confundía la cobertura real del archivo cargado.
    """
    unidades = set()
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        ensure_runtime_schema(cursor)
        filas = cursor.execute("""
            SELECT nombre, COALESCE(total_usuarios, 0) AS total_usuarios
            FROM unidades
            WHERE nombre IS NOT NULL AND TRIM(nombre) != ''
        """).fetchall()
        conn.close()
        for fila in filas:
            unidad = normalize_unidad(fila['nombre'])
            if unidad and int(fila['total_usuarios'] or 0) > 0:
                unidades.add(unidad)
    except Exception:
        pass
    return sorted(unidades)


def sincronizar_unidades_desde_dataframe(df):
    """Actualiza el catálogo de unidades con los nombres detectados en la base."""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        ensure_runtime_schema(cursor)
        ahora = datetime.now().isoformat()
        conteos = df.groupby('unidad').size().to_dict() if 'unidad' in df.columns else {}
        gestantes = {}
        if 'tipo_beneficiario' in df.columns:
            gestantes = df[df['tipo_beneficiario'].astype(str).str.upper().str.contains('GESTANTE', na=False)].groupby('unidad').size().to_dict()

        for unidad, total in conteos.items():
            unidad_norm = normalize_unidad(unidad)
            if not unidad_norm:
                continue
            cursor.execute("""
                INSERT INTO unidades (nombre, total_usuarios, total_gestantes, fecha_actualizacion, fundacion_id)
                VALUES (?, ?, ?, ?, ?)
                ON CONFLICT(fundacion_id, nombre) DO UPDATE SET
                    total_usuarios = excluded.total_usuarios,
                    total_gestantes = excluded.total_gestantes,
                    fecha_actualizacion = excluded.fecha_actualizacion
            """, (unidad_norm, int(total), int(gestantes.get(unidad, 0)), ahora, fundacion_actual_id()))
        conn.commit()
        conn.close()
    except Exception as exc:
        print(f'No se pudo sincronizar catálogo de unidades: {exc}')



def obtener_ultimas_valoraciones_salud():
    """Devuelve últimas valoraciones del módulo Salud y Nutrición por documento.

    No obliga a que el módulo exista. Si no hay tablas sn_*, retorna un diccionario vacío.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        tables = {row['name'] for row in cursor.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
        if 'sn_valoraciones' not in tables:
            conn.close()
            return {}
        fid = fundacion_actual_id()
        rows = cursor.execute("""
            SELECT v.*
            FROM sn_valoraciones v
            INNER JOIN (
                SELECT sv.documento, MAX(sv.fecha_valoracion || printf('%010d', sv.id)) AS max_key
                FROM sn_valoraciones sv
                WHERE sv.activo = 1 AND COALESCE(sv.fundacion_id, 1) = ?
                GROUP BY sv.documento
            ) ult ON ult.documento = v.documento
                 AND ult.max_key = (v.fecha_valoracion || printf('%010d', v.id))
            WHERE v.activo = 1 AND COALESCE(v.fundacion_id, 1) = ?
        """, (fid, fid)).fetchall()
        conn.close()
        return {str(row['documento']).strip(): dict(row) for row in rows if row['documento']}
    except Exception as exc:
        print(f'No se pudieron consultar valoraciones de Salud y Nutrición: {exc}')
        return {}

def log_auditoria(usuario, accion, tabla=None, registro_id=None, datos_anteriores=None, datos_nuevos=None):
    """Registra auditoría en la BD"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO auditoria
        (usuario, accion, tabla, registro_id, datos_anteriores, datos_nuevos, fecha_accion, direccion_ip)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (usuario, accion, tabla, registro_id,
          json.dumps(datos_anteriores) if datos_anteriores else None,
          json.dumps(datos_nuevos) if datos_nuevos else None,
          datetime.now().isoformat(),
          request.remote_addr if request else '127.0.0.1'))
    conn.commit()
    conn.close()


def buscar_columna(df, alias):
    """
    Busca una columna del archivo cargado usando alias normalizados.
    Prioriza coincidencia exacta para evitar tomar columnas de entidad contratista
    o acudiente cuando se necesita el dato del beneficiario.
    """
    columnas = [(normalizar_texto_clave(col), col) for col in df.columns]
    alias_norm = [normalizar_texto_clave(a) for a in alias]

    for objetivo in alias_norm:
        for normalizada, original in columnas:
            if normalizada == objetivo:
                return original

    for objetivo in alias_norm:
        for normalizada, original in columnas:
            if objetivo and objetivo in normalizada:
                return original

    return None


def serie_texto(df, columna, default=''):
    if columna and columna in df.columns:
        return df[columna].apply(lambda v: limpiar_valor(v, default))
    return pd.Series([default] * len(df), index=df.index)


def limpiar_y_normalizar_dataframe(df):
    """
    Convierte la base de Cuéntame a nombres de columnas internos.

    Corrección crítica:
    - Ya no crea "Activo" como nombre cuando el archivo no trae una columna llamada "nombre".
    - Usa los campos oficiales de Cuéntame: primer nombre, segundo nombre, apellidos,
      documento, acudiente, parentesco y unidad.
    """
    df = df.copy()
    df.columns = [str(col).strip() for col in df.columns]

    col_documento = buscar_columna(df, [
        'Documento del beneficiario',
        'Número de documento del beneficiario',
        'Numero de documento del beneficiario',
        'No. de documento de identidad',
        'Nº DOC. IDENT.',
        'Documento',
        'NUI',
        'NUIP'
    ])
    col_tipo_doc = buscar_columna(df, [
        'Tipo de documento del beneficiario',
        'Tipo documento beneficiario',
        'Tipo Doc'
    ])
    # La unidad se detecta con una función especializada para evitar que columnas
    # como Estado=ACTIVO terminen tomadas como unidad.
    col_unidad = detectar_columna_unidad(df)
    col_nombre = buscar_columna(df, [
        'Nombre completo del beneficiario',
        'Nombres y apellidos del participante',
        'Nombre',
        'Nombres'
    ])
    col_primer_nombre = buscar_columna(df, [
        'Primer Nombre del beneficiario',
        'Primer nombre beneficiario',
        'Primer Nombre'
    ])
    col_segundo_nombre = buscar_columna(df, [
        'Segundo Nombre del beneficiario',
        'Segundo nombre beneficiario',
        'Segundo Nombre'
    ])
    col_primer_apellido = buscar_columna(df, [
        'Primer apellido del beneficiario',
        'Primer apellido beneficiario',
        'Primer Apellido'
    ])
    col_segundo_apellido = buscar_columna(df, [
        'Segundo apellido del beneficiario',
        'Segundo apellido beneficiario',
        'Segundo Apellido'
    ])
    col_fecha_nacimiento = buscar_columna(df, [
        'Fecha de nacimiento del beneficiario',
        'Fecha nacimiento beneficiario',
        'Fecha de nacimiento',
        'fecha_nacimiento'
    ])
    col_edad_beneficiario = buscar_columna(df, [
        'Edad del beneficiario',
        'Edad beneficiario',
        'Edad'
    ])
    col_estado = buscar_columna(df, [
        'Estado del beneficiario',
        'Estado'
    ])
    col_peso_talla = buscar_columna(df, [
        'Peso y talla al día',
        'Peso talla al dia',
        'peso_talla_al_dia'
    ])
    col_tipo_beneficiario = buscar_columna(df, [
        'Nombre Tipo de beneficiario',
        'Tipo de beneficiario',
        'tipo_beneficiario'
    ])
    col_sexo = buscar_columna(df, [
        'Sexo del beneficiario',
        'Sexo'
    ])
    col_telefono = buscar_columna(df, [
        'Teléfono del beneficiario',
        'Telefono del beneficiario',
        'Celular del beneficiario',
        'Número de celular',
        'Numero de celular',
        'Teléfono',
        'Telefono',
        'Celular'
    ])
    col_fecha_modificacion = buscar_columna(df, [
        'Fecha de modificación del Beneficiario',
        'Fecha de modificacion del Beneficiario',
        'Fecha modificación beneficiario'
    ])
    col_fecha_ingreso = buscar_columna(df, [
        'Fecha de atención del beneficiario a la UDS',
        'Fecha de atencion del beneficiario a la UDS',
        'Fecha de vinculación del beneficiario',
        'Fecha de vinculacion del beneficiario',
        'Fecha de ingreso',
        'Fecha ingreso'
    ])

    col_regional = buscar_columna(df, [
        'Regional del Contrato',
        'Nombre de la Regional de la Unidad de servicio',
        'Regional'
    ])
    col_centro_zonal = buscar_columna(df, [
        'Nombre del Centro Zonal',
        'Centro Zonal',
        'Centro zonal'
    ])
    col_municipio = buscar_columna(df, [
        'Nombre Municipio de la Unidad de servicio',
        'Municipio de residencia del beneficiario',
        'Municipio'
    ])
    col_modalidad = buscar_columna(df, [
        'Modalidad',
        'Modalidad de atención',
        'Servicio'
    ])
    col_numero_contrato = buscar_columna(df, [
        'Número del Contrato',
        'Numero del Contrato',
        'Contrato'
    ])
    col_vigencia = buscar_columna(df, [
        'Año',
        'Vigencia',
        'Año del contrato',
        'Ano'
    ])
    col_nombre_eas = buscar_columna(df, [
        'Nombre de la Entidad Contratista',
        'Nombre de la EAS',
        'EAS'
    ])
    col_nit_eas = buscar_columna(df, [
        'Número de documento de la Entidad Contratista',
        'Numero de documento de la Entidad Contratista',
        'NIT de la Entidad Contratista',
        'NIT de la EAS',
        'NIT'
    ])
    col_servicio_atencion = buscar_columna(df, [
        'Tipo de Unidad',
        'Servicio de atención',
        'Servicio de atencion',
        'Tipo de servicio'
    ])
    col_direccion_unidad = buscar_columna(df, [
        'Dirección de la unidad de servicio',
        'Direccion de la unidad de servicio',
        'Dirección de la UDS',
        'Direccion de la UDS',
        'Dirección de la unidad',
        'Direccion de la unidad'
    ])
    col_codigo_unidad_servicio = buscar_columna(df, [
        'Código de la unidad de servicio',
        'Codigo de la unidad de servicio',
        'Código UDS',
        'Codigo UDS',
        'Código CUENTAME UDS',
        'Codigo CUENTAME UDS'
    ])

    col_parentesco = buscar_columna(df, [
        'Tipo de responsable',
        'Parentesco',
        'Parentesco acudiente'
    ])
    col_tipo_doc_acudiente = buscar_columna(df, [
        'Tipo de documento del acudiente o responsable',
        'Tipo documento acudiente',
        'Tipo documento responsable'
    ])
    col_doc_acudiente = buscar_columna(df, [
        'Número de documento del acudiente o responsable',
        'Numero de documento del acudiente o responsable',
        'Documento del acudiente o responsable',
        'Documento acudiente',
        'Documento responsable'
    ])
    col_primer_nombre_acudiente = buscar_columna(df, [
        'Primer nombre del acudiente o responsable',
        'Primer nombre acudiente',
        'Primer nombre responsable'
    ])
    col_segundo_nombre_acudiente = buscar_columna(df, [
        'Segundo nombre del acudiente o responsable',
        'Segundo nombre acudiente',
        'Segundo nombre responsable'
    ])
    col_primer_apellido_acudiente = buscar_columna(df, [
        'Primer apellido del acudiente o responsable',
        'Primer apellido acudiente',
        'Primer apellido responsable'
    ])
    col_segundo_apellido_acudiente = buscar_columna(df, [
        'Segundo apellido del acudiente o responsable',
        'Segundo apellido acudiente',
        'Segundo apellido responsable'
    ])

    normalizado = pd.DataFrame(index=df.index)
    normalizado['documento'] = serie_texto(df, col_documento)
    normalizado['nui'] = normalizado['documento']
    normalizado['tipo_documento'] = serie_texto(df, col_tipo_doc)
    normalizado['primer_nombre'] = serie_texto(df, col_primer_nombre).str.upper()
    normalizado['segundo_nombre'] = serie_texto(df, col_segundo_nombre).str.upper()
    normalizado['primer_apellido'] = serie_texto(df, col_primer_apellido).str.upper()
    normalizado['segundo_apellido'] = serie_texto(df, col_segundo_apellido).str.upper()

    nombre_completo_componentes = (
        normalizado['primer_nombre'] + ' ' +
        normalizado['segundo_nombre'] + ' ' +
        normalizado['primer_apellido'] + ' ' +
        normalizado['segundo_apellido']
    ).str.replace(r'\s+', ' ', regex=True).str.strip()

    nombre_simple = serie_texto(df, col_nombre).str.upper()
    normalizado['nombre'] = nombre_completo_componentes
    normalizado.loc[normalizado['nombre'] == '', 'nombre'] = nombre_simple
    normalizado.loc[normalizado['nombre'] == '', 'nombre'] = 'SIN NOMBRE'

    normalizado['unidad'] = serie_texto(df, col_unidad, 'SIN UNIDAD').apply(normalize_unidad)
    unidades_invalidas = normalizado['unidad'].isin(UNIDADES_INVALIDAS) | (normalizado['unidad'] == '')
    if unidades_invalidas.any():
        normalizado.loc[unidades_invalidas, 'unidad'] = df.loc[unidades_invalidas].apply(extraer_unidad_desde_fila, axis=1)
    normalizado['unidad'] = normalizado['unidad'].apply(normalize_unidad)
    normalizado.loc[normalizado['unidad'] == '', 'unidad'] = 'SIN UNIDAD'
    normalizado['fecha_nacimiento'] = serie_texto(df, col_fecha_nacimiento)
    normalizado['estado'] = serie_texto(df, col_estado, EstadoUsuario.ACTIVO).str.upper()
    normalizado['estado'] = normalizado['estado'].replace({
        'ACTIVO': EstadoUsuario.ACTIVO,
        'ACTIVA': EstadoUsuario.ACTIVO,
        'FALLECIDO': EstadoUsuario.FALLECIDO,
        'FALLECIDA': EstadoUsuario.FALLECIDO,
        'RETIRADO': EstadoUsuario.RETIRADO,
        'RETIRADA': EstadoUsuario.RETIRADO,
        'TRASLADADO': EstadoUsuario.TRASLADADO,
        'TRASLADADA': EstadoUsuario.TRASLADADO
    })
    normalizado.loc[~normalizado['estado'].isin(EstadoUsuario.ESTADOS_VALIDOS), 'estado'] = EstadoUsuario.ACTIVO
    normalizado['peso_talla_al_dia'] = serie_texto(df, col_peso_talla, 'Pendiente')
    normalizado['tipo_beneficiario'] = serie_texto(df, col_tipo_beneficiario, 'NINO').str.upper()

    edad_desde_fecha = normalizado['fecha_nacimiento'].apply(calcular_edad_meses)
    edad_desde_columna = serie_texto(df, col_edad_beneficiario).apply(inferir_edad_meses_desde_valor)
    normalizado['edad_meses'] = edad_desde_fecha
    usar_edad_respaldo = (normalizado['edad_meses'].fillna(0).astype(int) == 0) & (edad_desde_columna.fillna(0).astype(int) > 0)
    normalizado.loc[usar_edad_respaldo, 'edad_meses'] = edad_desde_columna.loc[usar_edad_respaldo]
    normalizado['grupo_edad'] = normalizado.apply(
        lambda fila: clasificar_grupo_edad(fila.get('tipo_beneficiario', ''), fila.get('fecha_nacimiento', ''), fila.get('edad_meses', 0)),
        axis=1
    )
    normalizado['sexo'] = serie_texto(df, col_sexo)
    normalizado['telefono'] = serie_texto(df, col_telefono)
    normalizado['docente'] = serie_texto(df, buscar_columna(df, ['Docente', 'Agente educativo']), '')

    normalizado['tipo_documento_acudiente'] = serie_texto(df, col_tipo_doc_acudiente)
    normalizado['documento_acudiente'] = serie_texto(df, col_doc_acudiente)
    normalizado['primer_nombre_acudiente'] = serie_texto(df, col_primer_nombre_acudiente).str.upper()
    normalizado['segundo_nombre_acudiente'] = serie_texto(df, col_segundo_nombre_acudiente).str.upper()
    normalizado['primer_apellido_acudiente'] = serie_texto(df, col_primer_apellido_acudiente).str.upper()
    normalizado['segundo_apellido_acudiente'] = serie_texto(df, col_segundo_apellido_acudiente).str.upper()
    normalizado['nombre_acudiente'] = (
        normalizado['primer_nombre_acudiente'] + ' ' +
        normalizado['segundo_nombre_acudiente'] + ' ' +
        normalizado['primer_apellido_acudiente'] + ' ' +
        normalizado['segundo_apellido_acudiente']
    ).str.replace(r'\s+', ' ', regex=True).str.strip()
    normalizado['parentesco'] = serie_texto(df, col_parentesco).str.upper()
    normalizado['fecha_modificacion_cuentame'] = serie_texto(df, col_fecha_modificacion)
    normalizado['fecha_ingreso'] = serie_texto(df, col_fecha_ingreso)

    normalizado['regional'] = serie_texto(df, col_regional).str.upper()
    normalizado['centro_zonal'] = serie_texto(df, col_centro_zonal).str.upper()
    normalizado['municipio'] = serie_texto(df, col_municipio).str.upper()
    normalizado['modalidad'] = serie_texto(df, col_modalidad).str.upper()
    normalizado['numero_contrato'] = serie_texto(df, col_numero_contrato)
    normalizado['vigencia'] = serie_texto(df, col_vigencia)
    normalizado['nombre_eas'] = serie_texto(df, col_nombre_eas).str.upper()
    normalizado['nit_eas'] = serie_texto(df, col_nit_eas)
    normalizado['servicio_atencion'] = serie_texto(df, col_servicio_atencion).str.upper()
    normalizado['direccion_unidad'] = serie_texto(df, col_direccion_unidad).str.upper()
    normalizado['codigo_unidad_servicio'] = serie_texto(df, col_codigo_unidad_servicio)

    # No permitir que estados operativos terminen escritos como nombres.
    estados_invalidos_como_nombre = {'ACTIVO', 'ACTIVA', 'INACTIVO', 'INACTIVA', 'PENDIENTE', 'RETIRADO', 'RETIRADA'}
    for campo_nombre in ['primer_nombre', 'segundo_nombre', 'primer_apellido', 'segundo_apellido', 'nombre']:
        normalizado.loc[normalizado[campo_nombre].isin(estados_invalidos_como_nombre), campo_nombre] = ''

    normalizado.loc[normalizado['nombre'] == '', 'nombre'] = 'SIN NOMBRE'
    normalizado['documento'] = normalizado['documento'].astype(str).str.strip()

    return normalizado[normalizado['documento'].notna() & (normalizado['documento'] != '')]


def dividir_nombre(nombre):
    partes = str(nombre or '').strip().split()
    if not partes:
        return 'SIN NOMBRE', ''
    if len(partes) == 1:
        return partes[0], ''
    mitad = max(1, len(partes) // 2)
    return ' '.join(partes[:mitad]), ' '.join(partes[mitad:])


def guardar_beneficiarios_actuales(df, archivo_origen=''):
    conn = database_connection()
    cursor = conn.cursor()
    fundacion_id = fundacion_actual_id()
    cursor.execute("UPDATE beneficiarios SET estado = ? WHERE estado = ? AND COALESCE(fundacion_id, 1) = ?", (EstadoUsuario.RETIRADO, EstadoUsuario.ACTIVO, fundacion_id))
    # Precargar la clave que realmente protege PostgreSQL. Usar solo documento
    # colapsaba unidades diferentes y permitía insertar dos veces la misma
    # combinación dentro de un único Excel.
    cursor.execute(
        "SELECT id, documento, unidad FROM beneficiarios WHERE COALESCE(fundacion_id, 1) = ?",
        (fundacion_id,),
    )
    claves_existentes = {
        (str(row['documento']).strip(), normalize_unidad(row['unidad']))
        for row in cursor.fetchall()
        if row['documento'] is not None and normalize_unidad(row['unidad'])
    }
    ahora = datetime.now().isoformat()

    filas_por_clave = {}
    duplicados_archivo = {}
    trazas_sync = []
    for _, fila in df.iterrows():
        documento = str(fila.get('documento', '')).strip()
        unidad = normalize_unidad(fila.get('unidad', ''))
        if not documento or not unidad:
            continue
        clave = (documento, unidad)
        if clave in filas_por_clave:
            duplicados_archivo[clave] = duplicados_archivo.get(clave, 1) + 1
        filas_por_clave[clave] = fila

    for (documento, unidad), repeticiones in duplicados_archivo.items():
        trazas_sync.append({
            'documento': documento, 'unidad': unidad, 'archivo': archivo_origen,
            'operacion': 'CONSOLIDADO_ULTIMO_REGISTRO', 'repeticiones': repeticiones,
            'motivo': 'CLAVE_REPETIDA_DENTRO_DEL_ARCHIVO',
        })

    for (documento, unidad), fila in filas_por_clave.items():
        nombre = str(fila.get('nombre', '')).strip()
        nombres_fallback, apellidos_fallback = dividir_nombre(nombre)
        primer_nombre = limpiar_valor(fila.get('primer_nombre', ''))
        segundo_nombre = limpiar_valor(fila.get('segundo_nombre', ''))
        primer_apellido = limpiar_valor(fila.get('primer_apellido', ''))
        segundo_apellido = limpiar_valor(fila.get('segundo_apellido', ''))
        nombres = unir_partes(primer_nombre, segundo_nombre) or nombres_fallback
        apellidos = unir_partes(primer_apellido, segundo_apellido) or apellidos_fallback
        estado = str(fila['estado']).strip().upper() or EstadoUsuario.ACTIVO
        if estado in ['FALLECIDO', 'FALLECIDA']:
            estado = EstadoUsuario.FALLECIDO
        elif estado in ['RETIRADO', 'RETIRADA']:
            estado = EstadoUsuario.RETIRADO
        elif estado not in EstadoUsuario.ESTADOS_VALIDOS:
            estado = EstadoUsuario.ACTIVO

        clave_existia = (documento, unidad) in claves_existentes

        datos_comunes = {
            'documento': documento,
            'nombres': nombres,
            'apellidos': apellidos,
            'fecha_nacimiento': str(fila.get('fecha_nacimiento', '')).strip(),
            'sexo': str(fila.get('sexo', '')).strip(),
            'unidad': unidad,
            'estado': estado,
            'tipo_beneficiario': str(fila.get('tipo_beneficiario', 'NINO')).strip().upper(),
            'fecha_carga': ahora,
            'fecha_ingreso': str(fila.get('fecha_ingreso', '')).strip() or ahora,
            'nui': str(fila.get('nui', documento)).strip(),
            'tipo_documento': str(fila.get('tipo_documento', '')).strip(),
            'primer_nombre': primer_nombre,
            'segundo_nombre': segundo_nombre,
            'primer_apellido': primer_apellido,
            'segundo_apellido': segundo_apellido,
            'nombre_acudiente': str(fila.get('nombre_acudiente', '')).strip(),
            'documento_acudiente': str(fila.get('documento_acudiente', '')).strip(),
            'tipo_documento_acudiente': str(fila.get('tipo_documento_acudiente', '')).strip(),
            'parentesco': str(fila.get('parentesco', '')).strip(),
            'primer_nombre_acudiente': str(fila.get('primer_nombre_acudiente', '')).strip(),
            'segundo_nombre_acudiente': str(fila.get('segundo_nombre_acudiente', '')).strip(),
            'primer_apellido_acudiente': str(fila.get('primer_apellido_acudiente', '')).strip(),
            'segundo_apellido_acudiente': str(fila.get('segundo_apellido_acudiente', '')).strip(),
            'fecha_modificacion_cuentame': str(fila.get('fecha_modificacion_cuentame', '')).strip(),
            'edad_meses': int(fila.get('edad_meses', 0) or 0),
            'grupo_edad': str(fila.get('grupo_edad', '')).strip(),
            'telefono': str(fila.get('telefono', '')).strip(),
            'regional': str(fila.get('regional', '')).strip(),
            'centro_zonal': str(fila.get('centro_zonal', '')).strip(),
            'municipio': str(fila.get('municipio', '')).strip(),
            'modalidad': str(fila.get('modalidad', '')).strip(),
            'numero_contrato': str(fila.get('numero_contrato', '')).strip(),
            'vigencia': str(fila.get('vigencia', '')).strip(),
            'nombre_eas': str(fila.get('nombre_eas', '')).strip(),
            'nit_eas': str(fila.get('nit_eas', '')).strip(),
            'servicio_atencion': str(fila.get('servicio_atencion', '')).strip(),
            'direccion_unidad': str(fila.get('direccion_unidad', '')).strip(),
            'codigo_unidad_servicio': str(fila.get('codigo_unidad_servicio', '')).strip(),
            'fundacion_id': fundacion_id,
            'usuario_creador_id': usuario_actual_id(),
            'fecha_creacion': ahora,
            'fecha_actualizacion': ahora
        }

        cursor.execute("""
                INSERT INTO beneficiarios
                (documento, nombres, apellidos, fecha_nacimiento, sexo, unidad, estado,
                 tipo_beneficiario, fecha_ingreso, fecha_carga, nui, tipo_documento,
                 primer_nombre, segundo_nombre, primer_apellido, segundo_apellido,
                 nombre_acudiente, documento_acudiente, tipo_documento_acudiente, parentesco,
                 primer_nombre_acudiente, segundo_nombre_acudiente,
                 primer_apellido_acudiente, segundo_apellido_acudiente,
                 fecha_modificacion_cuentame, edad_meses, grupo_edad, telefono,
                 regional, centro_zonal, municipio, modalidad, numero_contrato,
                 vigencia, nombre_eas, nit_eas, servicio_atencion, direccion_unidad, codigo_unidad_servicio,
                 fundacion_id, usuario_creador_id, fecha_creacion, fecha_actualizacion)
                VALUES
                (:documento, :nombres, :apellidos, :fecha_nacimiento, :sexo, :unidad, :estado,
                 :tipo_beneficiario, :fecha_ingreso, :fecha_carga, :nui, :tipo_documento,
                 :primer_nombre, :segundo_nombre, :primer_apellido, :segundo_apellido,
                 :nombre_acudiente, :documento_acudiente, :tipo_documento_acudiente, :parentesco,
                 :primer_nombre_acudiente, :segundo_nombre_acudiente,
                 :primer_apellido_acudiente, :segundo_apellido_acudiente,
                 :fecha_modificacion_cuentame, :edad_meses, :grupo_edad, :telefono,
                 :regional, :centro_zonal, :municipio, :modalidad, :numero_contrato,
                 :vigencia, :nombre_eas, :nit_eas, :servicio_atencion, :direccion_unidad, :codigo_unidad_servicio,
                 :fundacion_id, :usuario_creador_id, :fecha_creacion, :fecha_actualizacion)
                ON CONFLICT DO NOTHING
            """, datos_comunes)
        insertado = bool(getattr(cursor, 'rowcount', 0) == 1)
        cursor.execute("""
            UPDATE beneficiarios
            SET nombres = :nombres, apellidos = :apellidos,
                fecha_nacimiento = :fecha_nacimiento, sexo = :sexo,
                estado = :estado, tipo_beneficiario = :tipo_beneficiario, fecha_carga = :fecha_carga,
                fecha_ingreso = COALESCE(NULLIF(:fecha_ingreso, ''), fecha_ingreso),
                nui = :nui, tipo_documento = :tipo_documento,
                primer_nombre = :primer_nombre, segundo_nombre = :segundo_nombre,
                primer_apellido = :primer_apellido, segundo_apellido = :segundo_apellido,
                nombre_acudiente = :nombre_acudiente, documento_acudiente = :documento_acudiente,
                tipo_documento_acudiente = :tipo_documento_acudiente, parentesco = :parentesco,
                primer_nombre_acudiente = :primer_nombre_acudiente,
                segundo_nombre_acudiente = :segundo_nombre_acudiente,
                primer_apellido_acudiente = :primer_apellido_acudiente,
                segundo_apellido_acudiente = :segundo_apellido_acudiente,
                fecha_modificacion_cuentame = :fecha_modificacion_cuentame,
                edad_meses = :edad_meses, grupo_edad = :grupo_edad, telefono = :telefono,
                regional = :regional, centro_zonal = :centro_zonal, municipio = :municipio,
                modalidad = :modalidad, numero_contrato = :numero_contrato,
                vigencia = :vigencia, nombre_eas = :nombre_eas, nit_eas = :nit_eas,
                servicio_atencion = :servicio_atencion, direccion_unidad = :direccion_unidad,
                codigo_unidad_servicio = :codigo_unidad_servicio,
                usuario_creador_id = COALESCE(usuario_creador_id, :usuario_creador_id),
                fecha_actualizacion = :fecha_actualizacion
            WHERE documento = :documento AND unidad = :unidad
              AND COALESCE(fundacion_id, 1) = :fundacion_id
        """, datos_comunes)
        actualizado = bool(getattr(cursor, 'rowcount', 0) >= 1)
        operacion = 'INSERTADO' if insertado else ('ACTUALIZADO' if actualizado or clave_existia else 'IGNORADO_CONFLICTO_OTRO_TENANT')
        if actualizado:
            claves_existentes.add((documento, unidad))
        trazas_sync.append({
            'documento': documento, 'unidad': unidad, 'archivo': archivo_origen,
            'operacion': operacion, 'motivo': 'UPSERT_IDEMPOTENTE_DOCUMENTO_UNIDAD',
        })

    log_beneficiarios_sincronizacion_batch(trazas_sync)
    log_procesamiento_base_maestra(
        'Sincronización idempotente de beneficiarios finalizada',
        archivo=archivo_origen, claves=len(filas_por_clave),
        duplicados_consolidados=len(duplicados_archivo),
    )

    unidades_detectadas = {normalize_unidad(u) for u in df['unidad'].dropna().unique()}
    unidades_catalogo = sorted({u for u in set(ConfiguracionSistema.UNIDADES) | unidades_detectadas if u})
    for unidad in unidades_catalogo:
        cursor.execute("""
            INSERT INTO unidades (nombre, total_usuarios, total_gestantes, fecha_actualizacion, fundacion_id)
            VALUES (?, (
                SELECT COUNT(*) FROM beneficiarios
                WHERE unidad = ? AND estado = ? AND COALESCE(fundacion_id, 1) = ?
            ), (
                SELECT COUNT(*) FROM beneficiarios
                WHERE unidad = ? AND estado = ? AND UPPER(tipo_beneficiario) LIKE '%GESTANTE%'
                  AND COALESCE(fundacion_id, 1) = ?
            ), ?, ?)
            ON CONFLICT(fundacion_id, nombre) DO UPDATE SET
                total_usuarios = excluded.total_usuarios,
                total_gestantes = excluded.total_gestantes,
                fecha_actualizacion = excluded.fecha_actualizacion
        """, (
            unidad, unidad, EstadoUsuario.ACTIVO, fundacion_id,
            unidad, EstadoUsuario.ACTIVO, fundacion_id, ahora, fundacion_id
        ))

    cursor.execute('UPDATE unidades SET fundacion_id = ?, fecha_actualizacion = ? WHERE fundacion_id IS NULL', (fundacion_id, ahora))
    conn.commit()
    conn.close()

def extraer_texto_documento(ruta_archivo, extension):
    """Indexa texto cuando el formato lo permite sin dependencias externas."""
    try:
        if extension in {'.txt', '.csv'}:
            with open(ruta_archivo, 'r', encoding='utf-8', errors='ignore') as fh:
                return fh.read()[:200000]
        if extension in {'.xlsx', '.xls'}:
            hojas = pd.read_excel(ruta_archivo, sheet_name=None, dtype=str)
            textos = []
            for nombre_hoja, df in hojas.items():
                textos.append(nombre_hoja)
                textos.append(' '.join(df.fillna('').astype(str).values.flatten().tolist()))
            return '\n'.join(textos)[:200000]
    except Exception:
        return ''
    return ''


def periodo_actual():
    return datetime.now().strftime('%Y-%m')


def contar_edad_retiro(cursor):
    cursor.execute("""
        SELECT fecha_nacimiento
        FROM master_ninos
        WHERE activo = 1 AND COALESCE(fundacion_id, 1) = ?
    """, (fundacion_actual_id(),))
    return sum(1 for row in cursor.fetchall() if calcular_edad_meses(row['fecha_nacimiento']) >= 71)


def contar_peso_talla_vencido(cursor):
    limite = (datetime.now() - timedelta(days=AlertaConfiguracion.DIAS_CONTROL_NUTRICION)).date().isoformat()
    fid = fundacion_actual_id()
    cursor.execute("""
        SELECT COUNT(*) as total
        FROM master_ninos b
        WHERE b.activo = 1
          AND COALESCE(b.fundacion_id, 1) = ?
          AND NOT EXISTS (
              SELECT 1 FROM master_salud_nutricion s
              WHERE s.version_id = b.version_id AND s.documento = b.documento
                AND s.activo = 1 AND COALESCE(s.fundacion_id, 1) = ?
                AND date(s.fecha_toma) >= date(?)
          )
    """, (fid, fid, limite))
    return cursor.fetchone()['total']


def contar_entregables(cursor, periodo, tipos):
    placeholders = ','.join(['?'] * len(tipos))
    cursor.execute(f"""
        SELECT COUNT(*) as total
        FROM entregables_operacion
        WHERE periodo = ?
        AND estado = 'cargado'
        AND lower(tipo) IN ({placeholders})
    """, [periodo] + [tipo.lower() for tipo in tipos])
    return cursor.fetchone()['total']


def tipos_entregables_mensuales():
    return [
        {'tipo': 'Planeación mensual', 'categoria': 'Proceso pedagógico', 'dia_limite': 5},
        {'tipo': 'Encuentro en el hogar', 'categoria': 'Familia y comunidad', 'dia_limite': 10},
        {'tipo': 'Encuentro comunitario', 'categoria': 'Familia y comunidad', 'dia_limite': 15},
        {'tipo': 'Acta grupal', 'categoria': 'Familia y comunidad', 'dia_limite': 15},
        {'tipo': 'Lista de chequeo mensual', 'categoria': 'Administrativo', 'dia_limite': 20},
        {'tipo': 'Evidencias fotográficas', 'categoria': 'Evidencias', 'dia_limite': 20},
        {'tipo': 'Informe mensual', 'categoria': 'Informes', 'dia_limite': 25},
        {'tipo': 'Soporte nutricional', 'categoria': 'Salud y nutrición', 'dia_limite': 25},
        {'tipo': 'RAM / Asistencia mensual', 'categoria': 'Formatos', 'dia_limite': 28},
        {'tipo': 'RPP / RAN', 'categoria': 'Formatos', 'dia_limite': 28}
    ]


def fecha_limite_entregable(periodo, dia):
    try:
        anio, mes = [int(x) for x in str(periodo).split('-')[:2]]
        ultimo = calendar.monthrange(anio, mes)[1]
        dia = max(1, min(int(dia), ultimo))
        return f'{anio:04d}-{mes:02d}-{dia:02d}'
    except Exception:
        return ''


def listar_entregables_periodo(cursor, periodo):
    filas = [dict(row) for row in cursor.execute("""
        SELECT id, tipo, periodo, unidad, ruta_archivo, estado, observaciones, fecha_carga,
               fecha_limite, responsable, categoria, documento_analizado
        FROM entregables_operacion
        WHERE periodo = ?
        ORDER BY COALESCE(fecha_limite, fecha_carga) ASC
    """, (periodo,)).fetchall()]
    return filas


def tablero_entregables_periodo(cursor, periodo):
    cargados = listar_entregables_periodo(cursor, periodo)
    cargados_por_tipo = {}
    for item in cargados:
        clave = normalizar_texto_clave(item.get('tipo'))
        cargados_por_tipo.setdefault(clave, []).append(item)

    hoy = datetime.now().date()
    tablero = []
    for esperado in tipos_entregables_mensuales():
        clave = normalizar_texto_clave(esperado['tipo'])
        registros = cargados_por_tipo.get(clave, [])
        fecha_limite = registros[0].get('fecha_limite') if registros and registros[0].get('fecha_limite') else fecha_limite_entregable(periodo, esperado['dia_limite'])
        estado = 'pendiente'
        ruta_archivo = ''
        fecha_carga = ''
        responsable = registros[0].get('responsable', '') if registros else ''
        observaciones = registros[0].get('observaciones', '') if registros else ''
        if registros:
            estado = registros[0].get('estado') or 'cargado'
            ruta_archivo = registros[0].get('ruta_archivo') or ''
            fecha_carga = registros[0].get('fecha_carga') or ''
        try:
            limite_date = datetime.fromisoformat(fecha_limite[:10]).date()
            if estado != 'cargado':
                if limite_date < hoy:
                    estado = 'vencido'
                elif 0 <= (limite_date - hoy).days <= 3:
                    estado = 'proximo'
        except Exception:
            pass
        tablero.append({
            'tipo': esperado['tipo'],
            'categoria': esperado['categoria'],
            'periodo': periodo,
            'fecha_limite': fecha_limite,
            'estado': estado,
            'ruta_archivo': ruta_archivo,
            'fecha_carga': fecha_carga,
            'responsable': responsable,
            'observaciones': observaciones
        })

    resumen = {
        'total': len(tablero),
        'cargados': sum(1 for i in tablero if i['estado'] == 'cargado'),
        'pendientes': sum(1 for i in tablero if i['estado'] == 'pendiente'),
        'proximos': sum(1 for i in tablero if i['estado'] == 'proximo'),
        'vencidos': sum(1 for i in tablero if i['estado'] == 'vencido')
    }
    return tablero, resumen



def evaluar_operacion(periodo=None, usuario='sistema', guardar=True):
    periodo = periodo or periodo_actual()
    conn = get_db_connection()
    cursor = conn.cursor()

    total_beneficiarios = cursor.execute(
        "SELECT COUNT(*) as total FROM master_ninos WHERE activo = 1 AND COALESCE(fundacion_id,1) = ?",
        (fundacion_actual_id(),)
    ).fetchone()['total']
    edad_retiro = contar_edad_retiro(cursor)
    peso_talla_vencido = contar_peso_talla_vencido(cursor)
    informes = cursor.execute("SELECT COUNT(*) as total FROM informes_pedagogicos").fetchone()['total']
    evidencias = cursor.execute("SELECT COUNT(*) as total FROM evidencias").fetchone()['total']
    talento = cursor.execute("SELECT COUNT(DISTINCT documento) as total FROM master_talento_humano WHERE activo=1 AND COALESCE(fundacion_id,1)=?", (fundacion_actual_id(),)).fetchone()['total']
    formatos = len([n for n in os.listdir(OUTPUT_FOLDER) if n.lower().endswith(('.xlsx', '.xls', '.pdf'))])
    documentos = cursor.execute("SELECT COUNT(*) as total FROM documentos_institucionales WHERE estado = 'vigente'").fetchone()['total']
    entregables = listar_entregables_periodo(cursor, periodo)
    planeaciones_cargadas = informes + contar_entregables(cursor, periodo, ['planeación', 'planeacion', 'informe pedagógico', 'informe pedagogico'])
    evidencias_cargadas = evidencias + contar_entregables(cursor, periodo, ['evidencia', 'acta', 'listado de asistencia', 'encuentro comunitario'])
    formatos_cargados = formatos + contar_entregables(cursor, periodo, ['rpp', 'ran', 'asistencia'])

    def score(ok, total=1):
        if total <= 0:
            return 0
        return round(max(0, min(100, (ok / total) * 100)), 2)

    componentes = {
        'Proceso Pedagógico': score(1 if planeaciones_cargadas > 0 else 0),
        'Familia y Comunidad': score(1 if evidencias_cargadas > 0 else 0),
        'Salud y Nutrición': score(max(0, total_beneficiarios - peso_talla_vencido), total_beneficiarios or 1),
        'Talento Humano': score(1 if talento > 0 else 0),
        'Ambientes Educativos': score(1 if documentos > 0 else 0),
        'Administrativo y Gestión': score(1 if formatos_cargados > 0 else 0)
    }

    incumplimientos = []
    if edad_retiro:
        incumplimientos.append({'tipo': 'ALERTA DE RETIRO', 'detalle': f'{edad_retiro} beneficiario(s) con 71 meses o más.', 'nivel': 'CRITICA'})
    if peso_talla_vencido:
        incumplimientos.append({'tipo': 'Peso y talla vencido', 'detalle': f'{peso_talla_vencido} beneficiario(s) sin control vigente.', 'nivel': 'ROJO'})
    if planeaciones_cargadas == 0:
        incumplimientos.append({'tipo': 'No se cargó la planeación/informe pedagógico.', 'detalle': 'Sin registros pedagógicos para validar el componente.', 'nivel': 'AMARILLO'})
    if evidencias_cargadas == 0:
        incumplimientos.append({'tipo': 'No se cargó la evidencia.', 'detalle': 'No hay evidencias o actas registradas.', 'nivel': 'AMARILLO'})
    if formatos_cargados == 0:
        incumplimientos.append({'tipo': 'No se entregó RPP/RAN/asistencia.', 'detalle': 'No hay formatos generados o cargados en archivos actualizados.', 'nivel': 'ROJO'})

    reglas = cursor.execute("""
        SELECT codigo, componente, descripcion, frecuencia, criterio, nivel_alerta
        FROM reglas_cumplimiento
        WHERE activa = 1
        ORDER BY codigo
    """).fetchall()
    reglas_resultado = []
    for regla in reglas:
        criterio = (regla['criterio'] or '').lower()
        cumple = True
        evidencia = 'Validado'
        if 'edad_meses' in criterio:
            cumple = edad_retiro == 0
            evidencia = 'Sin alertas de retiro' if cumple else f'{edad_retiro} alerta(s) de retiro'
        elif 'peso_talla' in criterio:
            cumple = peso_talla_vencido == 0
            evidencia = 'Peso y talla vigente' if cumple else f'{peso_talla_vencido} vencido(s)'
        elif 'planeacion' in criterio or 'planeación' in criterio:
            cumple = planeaciones_cargadas > 0
            evidencia = 'Planeación cargada' if cumple else 'Pendiente'
        elif 'evidencia' in criterio or 'acta' in criterio or 'encuentro' in criterio:
            cumple = evidencias_cargadas > 0
            evidencia = 'Evidencia cargada' if cumple else 'Pendiente'
        elif 'rpp' in criterio or 'ran' in criterio or 'asistencia' in criterio:
            cumple = formatos_cargados > 0
            evidencia = 'Formato cargado/generado' if cumple else 'Pendiente'
        elif 'coordinadores' in criterio or 'docentes' in criterio:
            cumple = talento > 0
            evidencia = 'Talento humano registrado' if cumple else 'Pendiente'
        reglas_resultado.append({
            'codigo': regla['codigo'],
            'componente': regla['componente'],
            'descripcion': regla['descripcion'],
            'frecuencia': regla['frecuencia'],
            'cumple': cumple,
            'evidencia': evidencia,
            'nivel': regla['nivel_alerta']
        })
        if not cumple and not any(item['tipo'] == regla['descripcion'] for item in incumplimientos):
            incumplimientos.append({'tipo': regla['descripcion'], 'detalle': evidencia, 'nivel': regla['nivel_alerta']})

    estandares = []
    for row in cursor.execute("SELECT * FROM estandares_icbf WHERE activo = 1 ORDER BY codigo").fetchall():
        valor = componentes.get(row['componente'], 0)
        estandares.append({
            'estandar': row['codigo'],
            'componente': row['componente'],
            'descripcion': row['descripcion'],
            'cumple': valor >= 80,
            'evidencia': row['evidencia_requerida'] if valor >= 80 else 'Pendiente'
        })

    cumplimiento_general = round(sum(componentes.values()) / len(componentes), 2)
    resultado = {
        'periodo': periodo,
        'cumplimiento_general': cumplimiento_general,
        'componentes': componentes,
        'matriz_estandares': estandares,
        'incumplimientos': incumplimientos,
        'indicadores': {
            'beneficiarios_activos': total_beneficiarios,
            'edad_retiro': edad_retiro,
            'peso_talla_vencido': peso_talla_vencido,
            'formatos_generados': formatos,
            'documentos_institucionales': documentos,
            'entregables_cargados': len(entregables)
        },
        'reglas': reglas_resultado,
        'entregables': entregables
    }

    if guardar:
        cursor.execute("""
            INSERT INTO evaluaciones_cumplimiento
            (periodo, cumplimiento_general, resultado_json, usuario, fecha_evaluacion)
            VALUES (?, ?, ?, ?, ?)
        """, (periodo, cumplimiento_general, json.dumps(resultado, ensure_ascii=False), usuario, datetime.now().isoformat()))
        conn.commit()

    conn.close()
    return resultado


def generar_informe_supervision_excel(resultado):
    nombre = f"INFORME_ICBF_SUPERVISION_{resultado['periodo'].replace('-', '')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    ruta = os.path.join(OUTPUT_FOLDER, nombre)
    planes_mejora = [
        {
            'Hallazgo': item.get('tipo'),
            'Nivel': item.get('nivel'),
            'Accion sugerida': 'Cargar soporte faltante y registrar responsable/fecha de cierre.',
            'Estado': 'Abierto'
        }
        for item in resultado.get('incumplimientos', [])
    ] or [{'Hallazgo': 'Sin hallazgos abiertos', 'Nivel': 'VERDE', 'Accion sugerida': 'Mantener seguimiento mensual.', 'Estado': 'En seguimiento'}]

    with pd.ExcelWriter(ruta, engine='openpyxl') as writer:
        pd.DataFrame([{
            'Periodo': resultado['periodo'],
            'Cumplimiento General': resultado['cumplimiento_general'],
            **resultado['indicadores']
        }]).to_excel(writer, sheet_name='Resumen', index=False)
        pd.DataFrame([{
            'Beneficiarios activos': resultado['indicadores'].get('beneficiarios_activos', 0),
            'Alertas de retiro': resultado['indicadores'].get('edad_retiro', 0),
            'Formatos generados': resultado['indicadores'].get('formatos_generados', 0)
        }]).to_excel(writer, sheet_name='Cobertura', index=False)
        pd.DataFrame([{
            'Peso y talla vencido': resultado['indicadores'].get('peso_talla_vencido', 0),
            'Cumplimiento Salud y Nutricion': resultado['componentes'].get('Salud y Nutrición', 0)
        }]).to_excel(writer, sheet_name='Nutricion', index=False)
        pd.DataFrame([{
            'Cumplimiento': resultado['componentes'].get('Proceso Pedagógico', 0),
            'Estado': 'Cumple' if resultado['componentes'].get('Proceso Pedagógico', 0) >= 80 else 'Pendiente'
        }]).to_excel(writer, sheet_name='Planeaciones', index=False)
        pd.DataFrame([{
            'Cumplimiento': resultado['componentes'].get('Familia y Comunidad', 0),
            'Estado': 'Cumple' if resultado['componentes'].get('Familia y Comunidad', 0) >= 80 else 'Pendiente'
        }]).to_excel(writer, sheet_name='Evidencias', index=False)
        pd.DataFrame([{
            'Cumplimiento': resultado['componentes'].get('Talento Humano', 0),
            'Estado': 'Cumple' if resultado['componentes'].get('Talento Humano', 0) >= 80 else 'Pendiente'
        }]).to_excel(writer, sheet_name='Talento Humano', index=False)
        pd.DataFrame(resultado.get('entregables', [])).to_excel(writer, sheet_name='Entregables', index=False)
        pd.DataFrame([
            {'Componente': k, 'Cumplimiento': v}
            for k, v in resultado['componentes'].items()
        ]).to_excel(writer, sheet_name='Componentes', index=False)
        pd.DataFrame(resultado.get('reglas', [])).to_excel(writer, sheet_name='Reglas', index=False)
        pd.DataFrame(resultado['matriz_estandares']).to_excel(writer, sheet_name='Estandares', index=False)
        pd.DataFrame(resultado['incumplimientos']).to_excel(writer, sheet_name='Hallazgos', index=False)
        pd.DataFrame(planes_mejora).to_excel(writer, sheet_name='Planes de mejora', index=False)
    return nombre


def consultar_usuarios_anteriores():
    # Este flujo se ejecuta dentro de un job y necesita conservar el resultado
    # mientras inspecciona columnas y consulta la base anterior en PostgreSQL.
    conn = database_connection()
    cursor = conn.cursor()
    columnas = table_columns(cursor, 'usuarios')
    columnas_base = [
        'documento', 'nombre', 'unidad', 'estado', 'fecha_nacimiento', 'fecha_ingreso',
        'primer_nombre', 'segundo_nombre', 'primer_apellido', 'segundo_apellido',
        'nui', 'tipo_documento', 'nombre_acudiente', 'documento_acudiente',
        'tipo_documento_acudiente', 'parentesco', 'primer_nombre_acudiente',
        'segundo_nombre_acudiente', 'primer_apellido_acudiente',
        'segundo_apellido_acudiente', 'fecha_modificacion_cuentame',
        'edad_meses', 'grupo_edad', 'regional', 'centro_zonal', 'municipio',
        'modalidad', 'numero_contrato', 'vigencia', 'nombre_eas', 'nit_eas', 'servicio_atencion',
        'direccion_unidad', 'codigo_unidad_servicio'
    ]
    select_cols = [col for col in columnas_base if col in columnas]
    if not select_cols:
        conn.close()
        return {}
    if rol_actual() == 'SUPERADMIN':
        usuarios = conn.execute(f"SELECT {', '.join(select_cols)} FROM usuarios").fetchall()
    else:
        usuarios = conn.execute(
            f"SELECT {', '.join(select_cols)} FROM usuarios WHERE COALESCE(fundacion_id, 1) = ?",
            (fundacion_actual_id(),)
        ).fetchall()
    conn.close()
    return {str(u['documento']).strip(): dict(u) for u in usuarios if u['documento']}


def guardar_usuarios_actuales(df):
    conn = database_connection()
    cursor = conn.cursor()
    columnas_db = table_columns(cursor, 'usuarios')
    cursor.execute('DELETE FROM usuarios WHERE COALESCE(fundacion_id, 1) = ?', (fundacion_actual_id(),))

    columnas_insertar = [
        'documento', 'nombre', 'unidad', 'fecha_nacimiento', 'estado', 'peso_talla_al_dia',
        'docente', 'tipo_beneficiario', 'fecha_carga', 'fecha_ingreso', 'nui', 'tipo_documento',
        'primer_nombre', 'segundo_nombre', 'primer_apellido', 'segundo_apellido',
        'sexo', 'nombre_acudiente', 'documento_acudiente', 'tipo_documento_acudiente',
        'parentesco', 'primer_nombre_acudiente', 'segundo_nombre_acudiente',
        'primer_apellido_acudiente', 'segundo_apellido_acudiente',
        'fecha_modificacion_cuentame', 'edad_meses', 'grupo_edad', 'telefono',
        'regional', 'centro_zonal', 'municipio', 'modalidad', 'numero_contrato',
        'vigencia', 'nombre_eas', 'nit_eas', 'servicio_atencion', 'direccion_unidad', 'codigo_unidad_servicio',
        'fundacion_id', 'usuario_creador_id', 'fecha_creacion', 'fecha_actualizacion'
    ]
    columnas_insertar = [col for col in columnas_insertar if col in columnas_db]
    placeholders = ', '.join(['?'] * len(columnas_insertar))
    sql = f"INSERT INTO usuarios ({', '.join(columnas_insertar)}) VALUES ({placeholders})"

    ahora = datetime.now().isoformat()
    filas_insertar = []
    for _, fila in df.iterrows():
        valores = []
        for col in columnas_insertar:
            if col == 'fecha_carga':
                valores.append(ahora)
            elif col == 'fundacion_id':
                valores.append(fundacion_actual_id())
            elif col == 'usuario_creador_id':
                valores.append(usuario_actual_id())
            elif col in {'fecha_creacion', 'fecha_actualizacion'}:
                valores.append(ahora)
            else:
                valores.append(str(fila.get(col, '')).strip())
        filas_insertar.append(tuple(valores))

    try:
        safe_executemany(
            conn,
            sql,
            filas_insertar,
            batch_size=500,
            logger=lambda total, total_general: log_procesamiento_base_maestra(
                'Lote de usuarios insertado',
                procesados=total,
                total=total_general
            )
        )
        conn.commit()
    except Exception as exc:
        log_procesamiento_base_maestra('Error guardando usuarios actuales', str(exc), total=len(filas_insertar))
        raise
    finally:
        try:
            conn.close()
        except Exception:
            pass

def guardar_auditoria(usuario, archivo, total_registros, cambios_detectados):
    conn = database_connection()
    conn.execute('''
        INSERT INTO auditoria
        (fecha, usuario, accion, archivo, total_registros, cambios_detectados, archivo_cargado, fecha_accion)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        datetime.now().isoformat(),
        usuario,
        'IMPORTAR_CUENTAME',
        archivo,
        total_registros,
        cambios_detectados,
        archivo,
        datetime.now().isoformat()
    ))
    conn.commit()
    conn.close()


def _movimiento_payload(tipo, documento, nombre, unidad_origen, unidad_destino, detalle):
    """Normaliza un movimiento operativo para insertar uno a uno o por lote."""
    return {
        'tipo': limpiar_valor(tipo),
        'documento': limpiar_documento_talento(documento),
        'nombre': limpiar_valor(nombre),
        'unidad_origen': limpiar_valor(unidad_origen),
        'unidad_destino': limpiar_valor(unidad_destino),
        'detalle': limpiar_valor(detalle),
    }


def registrar_movimientos_lote(movimientos):
    """Inserta movimientos en una sola transacción.

    ALPHA34: antes `comparar_con_ultima_base` abría/confirmaba una conexión por
    cada movimiento detectado. En equipos lentos o con SQLite bloqueado esto
    dejaba el job detenido en 30% ("Comparando con la última base maestra").
    Esta función conserva la auditoría, pero la hace en lote y evita el bloqueo.
    """
    movimientos = [m for m in (movimientos or []) if m and limpiar_documento_talento(m.get('documento'))]
    if not movimientos:
        return 0

    ahora = datetime.now().isoformat()
    fundacion_id = fundacion_actual_id()
    conn = database_connection()
    cursor = conn.cursor()
    try:
        ensure_runtime_schema(cursor)
    except Exception:
        pass

    documentos = sorted({limpiar_documento_talento(m.get('documento')) for m in movimientos if limpiar_documento_talento(m.get('documento'))})
    beneficiario_por_clave = {}
    beneficiarios_por_doc = {}
    for i in range(0, len(documentos), 400):
        bloque = documentos[i:i + 400]
        placeholders = ','.join(['?'] * len(bloque))
        try:
            filas = cursor.execute(
                f"""
                SELECT documento, unidad, id
                  FROM beneficiarios
                 WHERE COALESCE(fundacion_id, 1) = ?
                   AND documento IN ({placeholders})
                """,
                [fundacion_id, *bloque],
            ).fetchall()
        except Exception:
            conn.close()
            raise
        for fila in filas:
            documento = limpiar_documento_talento(fila['documento'])
            clave = (documento, normalize_unidad(fila['unidad']))
            beneficiario_por_clave[clave] = int(fila['id'])
            beneficiarios_por_doc.setdefault(documento, []).append(int(fila['id']))

    valores = []
    movimientos_sin_beneficiario = []
    for mov in movimientos:
        documento = limpiar_documento_talento(mov.get('documento'))
        unidad_movimiento = normalize_unidad(
            mov.get('unidad_destino') or mov.get('unidad_origen') or ''
        )
        beneficiario_id = beneficiario_por_clave.get((documento, unidad_movimiento))
        if not beneficiario_id:
            candidatos = beneficiarios_por_doc.get(documento) or []
            if len(candidatos) == 1:
                beneficiario_id = candidatos[0]
        if not beneficiario_id:
            movimientos_sin_beneficiario.append({
                'documento': documento,
                'unidad': unidad_movimiento,
                'tipo': limpiar_valor(mov.get('tipo')),
            })
            continue
        valores.append((
            int(beneficiario_id),
            limpiar_valor(mov.get('tipo')),
            documento,
            limpiar_valor(mov.get('nombre')),
            limpiar_valor(mov.get('unidad_origen')),
            limpiar_valor(mov.get('unidad_destino')),
            ahora,
            limpiar_valor(mov.get('detalle')),
            ahora,
            limpiar_valor(mov.get('detalle')),
            'sistema',
            ahora
        ))

    if movimientos_sin_beneficiario:
        log_procesamiento_base_maestra(
            'Movimientos omitidos sin beneficiario válido',
            total=len(movimientos_sin_beneficiario),
            muestras=json.dumps(movimientos_sin_beneficiario[:10], ensure_ascii=False),
            fundacion_id=fundacion_id,
        )
    if not valores:
        conn.close()
        return 0

    sql_insert_movimientos = '''
        INSERT INTO movimientos
        (beneficiario_id, tipo, documento, nombre, unidad_origen, unidad_destino,
         fecha, detalle, fecha_movimiento, razon, usuario_registra, fecha_registro)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    '''
    try:
        log_procesamiento_base_maestra('Inicio de inserción por lotes', total_movimientos=len(valores))
        insertados = safe_executemany(
            conn,
            sql_insert_movimientos,
            valores,
            batch_size=500,
            logger=lambda total, total_general: log_procesamiento_base_maestra(
                'Lote de movimientos insertado',
                procesados=total,
                total=total_general
            )
        )
        conn.commit()
        log_procesamiento_base_maestra('Fin de inserción por lotes', insertados=insertados)
        return insertados
    except Exception as exc:
        log_procesamiento_base_maestra('Error insertando movimientos en lote', str(exc), total_movimientos=len(valores))
        raise
    finally:
        try:
            conn.close()
        except Exception:
            pass


def registrar_movimiento(tipo, documento, nombre, unidad_origen, unidad_destino, detalle):
    """Compatibilidad para llamadas antiguas: registra un solo movimiento."""
    return registrar_movimientos_lote([_movimiento_payload(tipo, documento, nombre, unidad_origen, unidad_destino, detalle)])


def generar_reporte_unidades(df, cambios_por_documento=None):
    cambios_por_documento = cambios_por_documento or {}
    reporte_unidades = {}
    alertas_cobertura_count = 0
    falta_peso_talla_count = 0
    proximos_retiro_lista = []
    unidades_sin_cobertura = []
    ultimas_valoraciones_salud = obtener_ultimas_valoraciones_salud()

    if 'unidad' in df.columns:
        df = df.copy()
        df['unidad'] = df['unidad'].apply(normalize_unidad)
        df.loc[df['unidad'] == '', 'unidad'] = 'SIN UNIDAD'
        df = df[df['unidad'] != 'SIN UNIDAD']

    grupos_base = {
        '0 A 6 MESES Y GESTANTES': 0,
        '6 A 11 MESES 29 DÍAS': 0,
        '1 A 2 AÑOS 11 MESES': 0,
        '3 A 5 AÑOS 11 MESES': 0,
        '5 AÑOS EN ADELANTE': 0,
        'FUERA DE RANGO': 0
    }
    grupos_totales = grupos_base.copy()

    for unidad_nombre, grupo in df.groupby('unidad'):
        unidad_nombre = normalize_unidad(unidad_nombre) or 'SIN UNIDAD'
        usuarios_lista = []
        usuarios_criticos = []
        nutricion_pendiente = 0
        count_unidad = int(len(grupo))
        grupos_edad = grupos_base.copy()

        if 0 < count_unidad < 20:
            alertas_cobertura_count += 1
            unidades_sin_cobertura.append({
                'unidad': unidad_nombre,
                'total': count_unidad,
                'meta': 20,
                'faltan': 20 - count_unidad
            })

        for _, fila in grupo.iterrows():
            nombre = str(fila.get('nombre', '')).strip()
            doc = str(fila['documento']).strip()
            estado = str(fila.get('estado', '')).strip().lower()
            pt_al_dia = str(fila.get('peso_talla_al_dia', '')).strip().lower()
            edad_m = int(fila.get('edad_meses') or calcular_edad_meses(fila.get('fecha_nacimiento')))
            tipo_beneficiario = str(fila.get('tipo_beneficiario', '')).strip()
            es_gestante = 'gestante' in normalizar_texto_clave(tipo_beneficiario)

            grupo_edad = clasificar_grupo_edad(tipo_beneficiario, fila.get('fecha_nacimiento', ''), edad_m)
            if grupo_edad not in grupos_edad:
                grupos_edad[grupo_edad] = 0
            if grupo_edad not in grupos_totales:
                grupos_totales[grupo_edad] = 0
            grupos_edad[grupo_edad] += 1
            grupos_totales[grupo_edad] += 1

            alerta_edad = False
            motivo_critico = None

            if estado == 'fallecido':
                motivo_critico = 'CRÍTICO: Reportado como Fallecido. Retirar inmediatamente.'
            elif not es_gestante and edad_m >= 71:
                motivo_critico = f'RETIRO OBLIGATORIO: Tiene {edad_m} meses.'
                alerta_edad = True
            elif not es_gestante and edad_m >= 60:
                motivo_critico = f'SEGUIMIENTO POR EDAD: Tiene {edad_m} meses, 5 años o más.'
                alerta_edad = True

            if not es_gestante and edad_m >= 60:
                proximos_retiro_lista.append({
                    'unidad': unidad_nombre,
                    'documento': doc,
                    'nombre': nombre,
                    'edad_meses': edad_m,
                    'edad_completa': formatear_edad_completa(edad_m, tipo_beneficiario),
                    'edad_anios': round(edad_m / 12, 1),
                    'grupo_edad': grupo_edad
                })

            if pt_al_dia in ['no', 'pendiente', 'falta', '0', '']:
                nutricion_pendiente += 1
                falta_peso_talla_count += 1

            if motivo_critico:
                usuarios_criticos.append({'nombre': nombre, 'motivo': motivo_critico})

            cambios_usuario = cambios_por_documento.get(doc, {})
            valoracion_salud = ultimas_valoraciones_salud.get(doc, {})
            usuarios_lista.append({
                'Documento': doc,
                'NUI': str(fila.get('nui', doc)).strip(),
                'TipoDocumento': str(fila.get('tipo_documento', '')).strip(),
                'Nombre': nombre,
                'PrimerNombre': str(fila.get('primer_nombre', '')).strip(),
                'SegundoNombre': str(fila.get('segundo_nombre', '')).strip(),
                'PrimerApellido': str(fila.get('primer_apellido', '')).strip(),
                'SegundoApellido': str(fila.get('segundo_apellido', '')).strip(),
                'Acudiente': str(fila.get('nombre_acudiente', '')).strip(),
                'DocumentoAcudiente': str(fila.get('documento_acudiente', '')).strip(),
                'TipoDocumentoAcudiente': str(fila.get('tipo_documento_acudiente', '')).strip(),
                'Parentesco': str(fila.get('parentesco', '')).strip(),
                'Telefono': str(fila.get('telefono', '')).strip(),
                'FechaNacimiento': str(fila.get('fecha_nacimiento', '')).strip(),
                'FechaIngreso': str(fila.get('fecha_ingreso', '')).strip(),
                'PesoKg': str(valoracion_salud.get('peso_kg', '') or '').strip(),
                'TallaCm': str(valoracion_salud.get('talla_cm', '') or '').strip(),
                'IMC': str(valoracion_salud.get('imc', '') or '').strip(),
                'PerimetroBraquial': str(valoracion_salud.get('perimetro_braquial_cm', '') or '').strip(),
                'PerimetroCefalico': str(valoracion_salud.get('perimetro_cefalico_cm', '') or '').strip(),
                'DiagnosticoNutricional': str(valoracion_salud.get('diagnostico_global', '') or '').strip(),
                'NivelAlertaNutricional': str(valoracion_salud.get('nivel_alerta', '') or '').strip(),
                'EstadoControlNutricional': str(valoracion_salud.get('estado_control', '') or '').strip(),
                'FechaValoracionNutricional': str(valoracion_salud.get('fecha_valoracion', '') or '').strip(),
                'ProximoControlNutricional': str(valoracion_salud.get('proximo_control', '') or '').strip(),
                'Regional': str(fila.get('regional', '')).strip(),
                'CentroZonal': str(fila.get('centro_zonal', '')).strip(),
                'Municipio': str(fila.get('municipio', '')).strip(),
                'Modalidad': str(fila.get('modalidad', '')).strip(),
                'NumeroContrato': str(fila.get('numero_contrato', '')).strip(),
                'Vigencia': str(fila.get('vigencia', '')).strip(),
                'NombreEAS': str(fila.get('nombre_eas', '')).strip(),
                'NIT': str(fila.get('nit_eas', '')).strip(),
                'ServicioAtencion': str(fila.get('servicio_atencion', '')).strip(),
                'DireccionUnidad': str(fila.get('direccion_unidad', '')).strip(),
                'CodigoUnidadServicio': str(fila.get('codigo_unidad_servicio', '')).strip(),
                'EdadMeses': edad_m,
                'EdadCompleta': formatear_edad_completa(edad_m, tipo_beneficiario),
                'GrupoEdad': grupo_edad,
                'TipoBeneficiario': tipo_beneficiario,
                'alerta_edad': alerta_edad,
                'Estado': estado,
                'Cambios': cambios_usuario,
                'TieneCambios': bool(cambios_usuario)
            })

        docente_asignado = obtener_docente_relacion(unidad_nombre) or 'Sin docente asignado'
        reporte_unidades[unidad_nombre or 'Sin unidad'] = {
            'total_usuarios': count_unidad,
            'alerta_cobertura': 0 < count_unidad < 20,
            'usuarios_criticos': usuarios_criticos,
            'nutricion_pendiente': nutricion_pendiente,
            'grupos_edad': grupos_edad,
            'docente_asignado': docente_asignado,
            'datos_completos': usuarios_lista
        }

    proximos_retiro_lista.sort(key=lambda item: (-int(item.get('edad_meses') or 0), item.get('unidad', ''), item.get('nombre', '')))
    unidades_sin_cobertura.sort(key=lambda item: (item['total'], item['unidad']))

    return {
        'unidades': dict(sorted(reporte_unidades.items(), key=lambda item: item[0])),
        'alertas_cobertura': alertas_cobertura_count,
        'unidades_sin_cobertura': unidades_sin_cobertura,
        'proximos_retiros': len(proximos_retiro_lista),
        'proximos_retiros_lista': proximos_retiro_lista,
        'falta_nutricion': falta_peso_talla_count,
        'grupos_edad_totales': grupos_totales
    }

def valor_comparacion(valor):
    return normalizar_texto_clave(limpiar_valor(valor))


def comparar_con_ultima_base(df, registrar=True, update_job=None, alcance='base completa', unidades_alcance=None):
    """Compara la base cargada contra la última base guardada sin bloquear el job.

    ALPHA34:
    - Evita registrar movimientos uno por uno.
    - Permite comparar solo las unidades seleccionadas cuando el usuario no pidió
      procesar todo.
    - Reporta subetapas para que la barra no se quede congelada en 30%.
    """
    def update(progreso=None, etapa=None, log=None):
        if update_job:
            try:
                update_job(progreso=progreso, etapa=etapa, log=log)
            except TypeError:
                # Compatibilidad con callbacks antiguos que solo aceptan kwargs.
                payload = {}
                if progreso is not None:
                    payload['progreso'] = progreso
                if etapa is not None:
                    payload['etapa'] = etapa
                if log is not None:
                    payload['log'] = log
                try:
                    update_job(**payload)
                except Exception:
                    pass
            except Exception:
                pass

    update(30, 'Consultando última base maestra', f'Comparación sobre {len(df) if df is not None else 0} registro(s): {alcance}.')
    log_procesamiento_base_maestra('Inicio de comparación con Base Maestra', alcance=alcance, registros_nuevos=len(df) if df is not None else 0)
    antiguo = consultar_usuarios_anteriores()
    update(40, 'Base anterior cargada para comparación', f'{len(antiguo)} registro(s) anteriores consultados.')
    log_procesamiento_base_maestra('Total de registros anteriores', total=len(antiguo))
    unidades_scope = {normalize_unidad(u) for u in _parse_lista_unidades(unidades_alcance) if normalize_unidad(u)}
    if unidades_scope:
        # Si el usuario seleccionó unidades, la comparación debe acotarse a esas
        # mismas unidades. De lo contrario, todos los niños de unidades no
        # seleccionadas aparecerían falsamente como retirados.
        antiguo = {
            doc: datos for doc, datos in antiguo.items()
            if normalize_unidad(datos.get('unidad')) in unidades_scope
        }

    actual = {}
    duplicados_actuales = []
    if df is not None and not df.empty:
        for _, fila in df.iterrows():
            documento = limpiar_documento_talento(fila.get('documento') if hasattr(fila, 'get') else '')
            if not documento:
                continue
            if documento in actual:
                duplicados_actuales.append(documento)
            actual[documento] = fila

    update(45, 'Nueva base indexada por documento', f'{len(actual)} documento(s) únicos. Duplicados detectados: {len(set(duplicados_actuales))}.')
    log_procesamiento_base_maestra('Total de registros nuevos indexados', total=len(actual), duplicados=len(set(duplicados_actuales)))

    update(50, 'Comparando documentos nuevos y retirados')
    ingresos = [doc for doc in actual if doc not in antiguo]
    retiros = [doc for doc in antiguo if doc not in actual]
    trasladados = []
    fallecidos = []
    transicion = []
    reemplazos = []
    por_documento = {}

    campos_auditar = [
        'nombre', 'unidad', 'estado', 'fecha_nacimiento', 'primer_nombre',
        'segundo_nombre', 'primer_apellido', 'segundo_apellido', 'nui',
        'tipo_documento', 'nombre_acudiente', 'documento_acudiente',
        'tipo_documento_acudiente', 'parentesco', 'primer_nombre_acudiente',
        'segundo_nombre_acudiente', 'primer_apellido_acudiente',
        'segundo_apellido_acudiente', 'edad_meses', 'grupo_edad'
    ]

    update(52, 'Detectando traslados, fallecidos y cambios')
    for documento, fila in actual.items():
        estado = str(fila.get('estado', '')).strip().lower()
        edad_m = int(fila.get('edad_meses') or calcular_edad_meses(fila.get('fecha_nacimiento')) or 0)
        if estado == 'fallecido':
            fallecidos.append(documento)
        if edad_m >= 60 and estado not in ['fallecido', 'retirado']:
            transicion.append(documento)
        if documento in antiguo:
            unidad_anterior = limpiar_valor(antiguo[documento].get('unidad'))
            unidad_nueva = limpiar_valor(fila.get('unidad'))
            if normalize_unidad(unidad_anterior) != normalize_unidad(unidad_nueva):
                trasladados.append(documento)

            cambios_campos = {}
            for campo in campos_auditar:
                anterior = antiguo[documento].get(campo, '')
                nuevo = fila.get(campo, '')
                # Si la base anterior no guardaba ese campo, no forzamos falso positivo.
                if limpiar_valor(anterior) == '':
                    continue
                if valor_comparacion(anterior) != valor_comparacion(nuevo):
                    cambios_campos[campo] = {
                        'anterior': limpiar_valor(anterior),
                        'nuevo': limpiar_valor(nuevo)
                    }
            if cambios_campos:
                por_documento[documento] = cambios_campos

    update(55, 'Preparando movimientos operativos')
    movimientos_pendientes = []

    for doc in ingresos:
        unidad = limpiar_valor(actual[doc].get('unidad'))
        similar_retiro = next((d for d in retiros if normalize_unidad(antiguo[d].get('unidad')) == normalize_unidad(unidad)), None)
        if similar_retiro:
            reemplazos.append(doc)
        por_documento[doc] = {'_tipo': 'INGRESO'}
        movimientos_pendientes.append(_movimiento_payload('Ingreso', doc, actual[doc].get('nombre', ''), None, unidad, 'Nuevo ingreso detectado'))

    for doc in retiros:
        registro = antiguo[doc]
        movimientos_pendientes.append(_movimiento_payload('Retiro', doc, registro.get('nombre', ''), registro.get('unidad', ''), None, 'Usuario desaparece de la base anterior'))

    for doc in fallecidos:
        registro = actual[doc]
        movimientos_pendientes.append(_movimiento_payload('Fallecido', doc, registro.get('nombre', ''), registro.get('unidad', ''), None, 'Usuario reportado como fallecido'))

    for doc in trasladados:
        registro_anterior = antiguo[doc]
        registro_actual = actual[doc]
        movimientos_pendientes.append(_movimiento_payload('Traslado', doc, registro_actual.get('nombre', ''), registro_anterior.get('unidad', ''), registro_actual.get('unidad', ''), 'Cambio de unidad detectado'))

    for doc in transicion:
        registro = actual[doc]
        movimientos_pendientes.append(_movimiento_payload('Transición escolar', doc, registro.get('nombre', ''), registro.get('unidad', ''), registro.get('unidad', ''), 'Cuidadoso seguimiento por edad'))

    movimientos_registrados = 0
    if registrar and movimientos_pendientes:
        update(58, 'Registrando movimientos en lote', f'{len(movimientos_pendientes)} movimiento(s) detectado(s).')
        log_procesamiento_base_maestra('Total de movimientos calculados', total=len(movimientos_pendientes))
        movimientos_registrados = registrar_movimientos_lote(movimientos_pendientes)

    resumen = {
        'ingresos': len(ingresos),
        'retiros': len(retiros),
        'fallecidos': len(fallecidos),
        'traslados': len(trasladados),
        'reemplazos': len(reemplazos),
        'transicion_escolar': len(transicion),
        'duplicados_actuales': len(set(duplicados_actuales)),
        'movimientos_registrados': movimientos_registrados,
        'alcance_comparacion': alcance,
        'por_documento': por_documento
    }
    if not registrar:
        resumen['_movimientos_pendientes'] = movimientos_pendientes
    update(60, 'Comparación finalizada', f'{movimientos_registrados} movimiento(s) registrado(s) en lote.')
    log_procesamiento_base_maestra('Comparación finalizada', movimientos_registrados=movimientos_registrados)
    return resumen


# ==================== RUTAS: AUTENTICACIÓN ====================
# La autenticación principal se registra en backend/modules/seguridad.
# Se conserva este bloque como separador para no duplicar rutas /api/auth/*.


# ==================== RUTAS: ALERTAS ====================
@app.route('/api/alertas/generar', methods=['POST'])
def generar_alertas_sistema():
    """Genera alertas automáticas del sistema"""
    try:
        resultados = motor_alertas.generar_todas_alertas()
        log_auditoria(request.args.get('usuario', 'admin'), 'GENERAR_ALERTAS_SISTEMA',
                     datos_nuevos=resultados)
        
        return jsonify({
            'mensaje': 'Alertas generadas',
            'resumen': resultados
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/alertas', methods=['GET'])
def obtener_alertas():
    """Obtiene alertas del sistema"""
    nivel = request.args.get('nivel')
    resuelta = request.args.get('resuelta', '0')
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if rol_actual() == 'SUPERADMIN':
        query = "SELECT * FROM alertas WHERE resuelta = ?"
        params = [int(resuelta)]
    else:
        query = "SELECT * FROM alertas WHERE resuelta = ? AND COALESCE(fundacion_id, 1) = ?"
        params = [int(resuelta), fundacion_actual_id()]
    
    if nivel:
        query += " AND nivel = ?"
        params.append(nivel)
    
    query += " ORDER BY fecha_generacion DESC LIMIT 100"
    
    cursor.execute(query, params)
    alertas = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    return jsonify(alertas), 200


# ==================== RUTAS: BENEFICIARIOS ====================
@app.route('/api/beneficiarios', methods=['GET'])
def obtener_beneficiarios():
    """Lista la población de la Base Maestra publicada."""
    unidad = request.args.get('unidad')
    estado = request.args.get('estado', EstadoUsuario.ACTIVO)
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if rol_actual() == 'SUPERADMIN':
        query = "SELECT *, unidad_servicio AS unidad, documento AS nui FROM master_ninos WHERE activo = 1"
        params = []
    else:
        query = "SELECT *, unidad_servicio AS unidad, documento AS nui FROM master_ninos WHERE activo = 1 AND COALESCE(fundacion_id, 1) = ?"
        params = [fundacion_actual_id()]
    
    if unidad:
        query += " AND unidad_servicio = ?"
        params.append(unidad)
    
    cursor.execute(query, params)
    beneficiarios = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    return jsonify(beneficiarios), 200


# ==================== RUTAS: NUTRICIÓN ====================
@app.route('/api/nutricion/registrar', methods=['POST'])
def registrar_nutricion():
    """Registra peso y talla de un beneficiario"""
    data = request.get_json()
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        peso = float(data['peso'])
        talla = float(data['talla'])
        benef_id = int(data['beneficiario_id'])
        
        cursor.execute("SELECT fecha_nacimiento FROM beneficiarios WHERE id = ? AND COALESCE(fundacion_id, 1) = ?", (benef_id, fundacion_actual_id()))
        benef = cursor.fetchone()
        
        if not benef:
            return jsonify({'error': 'Beneficiario no encontrado'}), 404
        
        edad_meses = calcular_edad_meses(benef['fecha_nacimiento'])
        estado = clasificar_nutricional(peso, talla, edad_meses)
        fecha_proximo = (datetime.now() + timedelta(days=AlertaConfiguracion.DIAS_CONTROL_NUTRICION)).isoformat()
        
        cursor.execute("""
            INSERT INTO peso_talla
            (beneficiario_id, peso, talla, fecha_medicion, responsable, estado_nutricional,
             fecha_proximo_control, fecha_carga, fundacion_id, usuario_creador_id, fecha_creacion, fecha_actualizacion)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (benef_id, peso, talla, datetime.now().isoformat(),
              request.args.get('usuario', 'sistema'), estado, fecha_proximo,
              datetime.now().isoformat(), fundacion_actual_id(), usuario_actual_id(),
              datetime.now().isoformat(), datetime.now().isoformat()))
        
        conn.commit()
        conn.close()
        
        return jsonify({'mensaje': 'Registro guardado', 'estado': estado}), 201
    
    except Exception as e:
        conn.close()
        return jsonify({'error': str(e)}), 400


# ==================== RUTAS: FORMATOS ====================
@app.route('/api/formatos/generar-mes', methods=['POST'])
def generar_mes_completo():
    """Genera todos los formatos para un mes"""
    data = request.get_json()
    mes = int(data.get('mes'))
    año = int(data.get('año'))
    unidad = data.get('unidad')
    
    crear_backup_operativo('ANTES_GENERAR_FORMATOS', f'Backup antes de generar formatos del mes {mes}/{año}. Unidad: {unidad or "TODAS"}.')

    try:
        archivos = generador.generar_mes_completo(mes, año, unidad)
        
        log_auditoria(request.args.get('usuario', 'sistema'), 'GENERAR_FORMATOS_MES',
                     datos_nuevos={'mes': mes, 'año': año, 'unidad': unidad})
        
        return jsonify({
            'mensaje': 'Formatos generados',
            'archivos': archivos
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== RUTAS: SINCRONIZACIÓN ====================
@app.route('/api/sincronizar', methods=['POST'])
def sincronizar():
    """Sincroniza todas las operaciones del sistema"""
    try:
        resultados = {
            'alertas': motor_alertas.generar_todas_alertas(),
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify(resultados), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== RUTAS: SALUD ====================
@app.route('/health', methods=['GET'])
@app.route('/api/ready', methods=['GET'])
@app.route('/api/health', methods=['GET'])
def health():
    """Estado mínimo y huella de instancia, sin datos personales.

    ``project_instance_id`` permite que los scripts distingan dos copias de la
    plataforma que intentan usar el mismo puerto 5000. Así el túnel no termina
    publicando accidentalmente un backend viejo de otra carpeta.
    """
    instance_id = project_instance_id(app.config)
    log_status = logging_health(app.config)
    try:
        db_status = database.healthcheck()
        if not db_status.get('ok'):
            raise RuntimeError(db_status.get('error') or 'database unavailable')
        return jsonify({
            'status': 'ok',
            'database': 'ok',
            'database_backend': db_status.get('dialect'),
            'database_latency_ms': db_status.get('latency_ms'),
            'version': app.config.get('APP_VERSION', 'unknown'),
            'git_sha': str(os.getenv('RAILWAY_GIT_COMMIT_SHA') or os.getenv('GIT_COMMIT_SHA') or os.getenv('BUILD_COMMIT') or 'unknown'),
            'build_time': str(os.getenv('RAILWAY_DEPLOYMENT_START_TIME') or os.getenv('BUILD_TIME') or 'unknown'),
            'schema_migration_mode': str(os.getenv('APP_SCHEMA_MIGRATION_MODE', '0')).strip().lower() in {'1', 'true', 'yes', 'si', 'sí', 'on'},
            'environment': app.config.get('APP_ENV', 'unknown'),
            'server_mode': app.config.get('SERVER_MODE', 'LOCAL'),
            'public_tunnel_mode': bool(app.config.get('PUBLIC_TUNNEL_MODE', False)),
            'project_instance_id': instance_id,
            'logging': log_status,
            'timestamp': datetime.now().isoformat(timespec='seconds'),
        }), 200
    except Exception as exc:
        app.logger.error('Healthcheck de base falló instance_id=%s: %s', instance_id, exc)
        return jsonify({
            'status': 'error',
            'database': 'unavailable',
            'project_instance_id': instance_id,
            'logging': log_status,
        }), 503


@app.route('/api/system/version', methods=['GET'])
def system_version():
    """Huella exacta del artefacto desplegado, sin secretos."""
    from modules.centro_documental.schema import DOCUMENTS_SCHEMA_VERSION
    return jsonify({
        'app_version': app.config.get('APP_VERSION', 'unknown'),
        'git_sha': str(os.getenv('RAILWAY_GIT_COMMIT_SHA') or os.getenv('GIT_COMMIT_SHA') or os.getenv('BUILD_COMMIT') or 'unknown'),
        'build_time': str(os.getenv('RAILWAY_DEPLOYMENT_START_TIME') or os.getenv('BUILD_TIME') or 'unknown'),
        'environment': app.config.get('APP_ENV', 'unknown'),
        'database_backend': database.dialect_name,
        'project_instance_id': project_instance_id(app.config),
        'schema_runtime_ddl_disabled': str(os.getenv('SKIP_RUNTIME_SCHEMA_DDL', '0')).strip().lower() in {'1', 'true', 'yes', 'si', 'sí', 'on'},
        'documents_schema_version': DOCUMENTS_SCHEMA_VERSION,
        'templates_catalog_version': '1',
    }), 200


def _valor_booleano(valor):
    """Interpreta banderas enviadas desde formularios sin romper compatibilidad."""
    return str(valor or '').strip().lower() in {'1', 'true', 'si', 'sí', 'yes', 'on'}


def _parse_lista_unidades(valor):
    """Normaliza lista de unidades recibida como JSON, CSV o arreglo de formulario."""
    if valor is None:
        return []
    if isinstance(valor, (list, tuple, set)):
        candidatos = list(valor)
    else:
        raw = str(valor or '').strip()
        if not raw:
            return []
        try:
            parsed = json.loads(raw)
            if isinstance(parsed, (list, tuple, set)):
                candidatos = list(parsed)
            else:
                candidatos = [parsed]
        except Exception:
            candidatos = re.split(r'[,;|\n]+', raw)

    unidades = []
    vistas = set()
    for item in candidatos:
        unidad = normalize_unidad(item)
        if unidad and unidad not in vistas and unidad not in UNIDADES_INVALIDAS:
            vistas.add(unidad)
            unidades.append(unidad)
    return unidades


def normalizar_unidades_seleccionadas(req):
    """Lee unidades desde formularios/queries y devuelve una lista canónica.

    Acepta JSON, CSV, campos repetidos y nombres usados por versiones previas
    para que el selector no pierda la selección del usuario.
    """
    claves = (
        'unidad_seleccionada',
        'unidades_seleccionadas[]',
        'unidades_seleccionadas',
        'unidades_seleccionadas_csv',
        'unidades',
        'unidad',
    )
    valores_crudos = []
    for source in (getattr(req, 'form', None), getattr(req, 'args', None)):
        if source is None:
            continue
        for clave in claves:
            try:
                valores_crudos.extend(source.getlist(clave))
            except Exception:
                try:
                    valor = source.get(clave)
                except Exception:
                    valor = None
                if valor:
                    valores_crudos.append(valor)

    unidades = []
    vistas = set()
    for valor in valores_crudos:
        for unidad in _parse_lista_unidades(valor):
            if unidad and unidad not in vistas:
                vistas.add(unidad)
                unidades.append(unidad)
    return unidades


def _resumen_unidades_dataframe(df):
    """Devuelve conteo liviano de unidades detectadas antes de generar formatos."""
    if df is None or df.empty or 'unidad' not in df.columns:
        return []

    tmp = df.copy()
    tmp['unidad'] = tmp['unidad'].apply(normalize_unidad)
    tmp = tmp[(tmp['unidad'] != '') & (~tmp['unidad'].isin(UNIDADES_INVALIDAS))]
    if tmp.empty:
        return []

    tipo_col = 'tipo_beneficiario' if 'tipo_beneficiario' in tmp.columns else None
    estado_col = 'estado' if 'estado' in tmp.columns else None
    resumen = []
    for unidad, grupo in tmp.groupby('unidad'):
        item = {
            'nombre': unidad,
            'total': int(len(grupo)),
            'activos': int(len(grupo)),
            'gestantes': 0,
        }
        if estado_col:
            estados = grupo[estado_col].astype(str).str.upper().str.strip()
            item['activos'] = int(estados.isin(['ACTIVO', 'ACTIVA', 'EN ATENCION', 'EN ATENCIÓN']).sum() or len(grupo))
        if tipo_col:
            tipos = grupo[tipo_col].astype(str).str.upper()
            item['gestantes'] = int(tipos.str.contains('GESTANTE', na=False).sum())
        resumen.append(item)

    resumen.sort(key=lambda item: (item['nombre']))
    return resumen


def _filtrar_dataframe_por_unidades(df, unidades_seleccionadas):
    """Filtra una base normalizada por unidades, preservando columnas originales."""
    unidades = _parse_lista_unidades(unidades_seleccionadas)
    if not unidades:
        return df.copy(), []

    buscadas = set(unidades)
    tmp = df.copy()
    serie_unidad = tmp['unidad'].apply(normalize_unidad) if 'unidad' in tmp.columns else pd.Series([''] * len(tmp), index=tmp.index)
    presentes = set(serie_unidad.dropna().unique())
    filtrado = tmp[serie_unidad.isin(buscadas)].copy()
    no_encontradas = sorted(buscadas - presentes)
    return filtrado, no_encontradas


def _resolver_archivo_cuentame_guardado(archivo_token):
    """Resuelve un archivo ya subido a uploads sin permitir rutas externas."""
    token = secure_filename(os.path.basename(str(archivo_token or '').strip()))
    if not token:
        return None, None
    ruta = os.path.abspath(os.path.join(UPLOAD_FOLDER, token))
    uploads_abs = os.path.abspath(UPLOAD_FOLDER)
    if not (ruta == uploads_abs or ruta.startswith(uploads_abs + os.sep)):
        return None, None
    if not os.path.exists(ruta) or not os.path.isfile(ruta):
        return None, None
    if not es_extension_valida(token, ALLOWED_BASE_EXTENSIONS):
        return None, None
    return ruta, token


def _guardar_archivo_cuentame_desde_request(file_storage):
    """Valida y guarda el archivo Cuéntame en uploads con nombre trazable."""
    if not file_storage or not file_storage.filename:
        raise ValueError('Archivo no seleccionado.')
    if not es_extension_valida(file_storage.filename, ALLOWED_BASE_EXTENSIONS):
        raise ValueError('Extensión no permitida para base de datos. Usa Excel, CSV, TXT, TSV, JSON, HTML, DOCX o PDF tabular.')

    filename = secure_filename(file_storage.filename)
    if not filename:
        raise ValueError('El nombre del archivo no es válido.')

    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    nombre_guardado = f'CUENTAME_{timestamp}_{filename}'
    ruta_cuentame = os.path.join(UPLOAD_FOLDER, nombre_guardado)
    file_storage.save(ruta_cuentame)
    return ruta_cuentame, nombre_guardado


def _detectar_unidades_cuentame(ruta_cuentame, filename):
    """Lee la base una sola vez para mostrar unidades seleccionables sin generar formatos."""
    df = leer_base_datos_flexible(ruta_cuentame, filename)
    df = limpiar_y_normalizar_dataframe(df)
    unidades = _resumen_unidades_dataframe(df)
    if not unidades:
        raise ValueError('No se detectaron unidades de atención válidas en la base cargada.')
    return {
        'modo': 'seleccion_unidades',
        'archivo': filename,
        'archivo_token': os.path.basename(ruta_cuentame),
        'total_usuarios': int(len(df)),
        'total_unidades': int(len(unidades)),
        'unidades': unidades,
        'mensaje': 'Unidades detectadas correctamente. Selecciona una o varias unidades para generar formatos.'
    }


def _opcion_entero(options, key, default, min_value=None, max_value=None):
    try:
        value = int((options or {}).get(key) or default)
    except Exception:
        value = default
    if min_value is not None:
        value = max(min_value, value)
    if max_value is not None:
        value = min(max_value, value)
    return value


def _procesar_base_cuentame_core(ruta_cuentame, filename, options=None, update_job=None):
    """Procesa base Cuéntame/beneficiarios y genera formatos sin depender del request.

    Alpha24: se ejecuta en segundo plano cuando la plataforma está publicada por
    túnel para evitar errores 524 de Cloudflare/ngrok. Devuelve exactamente la
    estructura que el dashboard ya sabe renderizar.
    """
    options = dict(options or {})

    def update(progreso=None, etapa=None, log=None):
        if update_job:
            payload = {}
            if progreso is not None:
                payload['progreso'] = int(max(0, min(100, progreso)))
            if etapa:
                payload['etapa'] = etapa
            if log:
                payload['log'] = log
            update_job(**payload)

    update(5, 'Leyendo archivo cargado', f'Archivo: {filename}')
    log_procesamiento_base_maestra('Archivo recibido', archivo=filename, ruta=ruta_cuentame)
    df = leer_base_datos_flexible(ruta_cuentame, filename)
    log_procesamiento_base_maestra('Total de registros leídos', total=len(df))

    update(12, 'Normalizando columnas y datos')
    df = limpiar_y_normalizar_dataframe(df)
    total_usuarios_base = len(df)

    unidades_seleccionadas = _parse_lista_unidades(options.get('unidades_seleccionadas'))
    procesar_todo = _valor_booleano(options.get('procesar_todo')) or not unidades_seleccionadas
    df_formatos = df
    unidades_no_encontradas = []
    if unidades_seleccionadas and not procesar_todo:
        df_formatos, unidades_no_encontradas = _filtrar_dataframe_por_unidades(df, unidades_seleccionadas)
        if df_formatos.empty:
            seleccion = ', '.join(unidades_seleccionadas)
            raise ValueError(f'Las unidades seleccionadas no tienen registros en la base cargada: {seleccion}.')
        update(18, 'Base filtrada por unidades seleccionadas', f'{len(df_formatos)} registro(s) en {len(unidades_seleccionadas)} unidad(es).')
        log_procesamiento_base_maestra('Unidades seleccionadas normalizadas', unidades=' | '.join(unidades_seleccionadas), registros=len(df_formatos))

    total_usuarios_formatos = len(df_formatos)

    update(20, 'Sincronizando catálogo de unidades')
    log_procesamiento_base_maestra('Total de unidades detectadas', total=len(_resumen_unidades_dataframe(df)))
    # ALPHA34: cuando el usuario selecciona unidades, el diagnóstico y formatos
    # trabajan solo con esa selección. El catálogo puede sincronizarse completo,
    # pero la comparación costosa se limita a lo que se va a procesar para evitar
    # que el job quede congelado en 30%.
    sincronizar_unidades_desde_dataframe(df)

    alcance_comparacion = 'unidades seleccionadas' if unidades_seleccionadas and not procesar_todo else 'base completa'
    df_comparacion = df_formatos if unidades_seleccionadas and not procesar_todo else df
    update(28, 'Preparando comparación con Base Maestra', f'{len(df_comparacion)} registro(s) para {alcance_comparacion}.')
    cambios = comparar_con_ultima_base(
        df_comparacion,
        registrar=False,
        update_job=update,
        alcance=alcance_comparacion,
        unidades_alcance=unidades_seleccionadas if unidades_seleccionadas and not procesar_todo else None
    )

    etapa_diagnostico = 'Construyendo diagnóstico de unidades seleccionadas' if unidades_seleccionadas and not procesar_todo else 'Construyendo diagnóstico operativo'
    update(62, etapa_diagnostico, f'{len(df_formatos)} registro(s) serán usados para tablero y formatos.')
    reporte = generar_reporte_unidades(df_formatos, cambios.get('por_documento', {}))

    update(68, 'Guardando Base Maestra de Cuéntame')
    log_procesamiento_base_maestra('Inicio de guardado de usuarios actuales', total=len(df))
    guardar_usuarios_actuales(df)
    log_procesamiento_base_maestra('Fin de guardado de usuarios actuales', total=len(df))
    update(72, 'Actualizando beneficiarios actuales')
    guardar_beneficiarios_actuales(df, archivo_origen=filename)
    movimientos_pendientes = cambios.pop('_movimientos_pendientes', [])
    if movimientos_pendientes:
        update(74, 'Registrando movimientos operativos', f'{len(movimientos_pendientes)} movimiento(s) después de sincronizar beneficiarios.')
        cambios['movimientos_registrados'] = registrar_movimientos_lote(movimientos_pendientes)
    guardar_auditoria('Operador', filename, total_usuarios_base, str(cambios))

    update(76, 'Actualizando alertas operativas')
    try:
        motor_alertas.generar_todas_alertas()
    except Exception as exc:
        update(log=f'Alertas no bloquearon el proceso: {exc}')

    unidades = list((reporte.get('unidades') or {}).items())
    total_unidades = max(1, len(unidades))
    errores_formatos = []
    for index, (unidad_nombre, unidad_info) in enumerate(unidades, start=1):
        progreso = 80 + int((index / total_unidades) * 15)
        update(progreso, f'Generando formatos de {unidad_nombre}', f'Unidad {index}/{total_unidades}: {unidad_nombre}')
        try:
            log_procesamiento_base_maestra('Inicio de generación de formatos', unidad=unidad_nombre, usuarios=len(unidad_info.get('datos_completos') or []))
            inyectar_datos_en_plantillas(unidad_nombre, unidad_info.get('datos_completos') or [], options=options)
            log_procesamiento_base_maestra('Formatos generados por unidad', unidad=unidad_nombre)
        except Exception as exc:
            errores_formatos.append({'unidad': unidad_nombre, 'error': str(exc)})
            update(log=f'No se generaron todos los formatos de {unidad_nombre}: {exc}')

    update(98, 'Preparando respuesta final')
    resultado = {
        'stats': {
            'total_usuarios': total_usuarios_formatos,
            'total_usuarios_base_maestra': total_usuarios_base,
            'total_usuarios_formatos': total_usuarios_formatos,
            'unidades_procesadas': len(reporte.get('unidades', {}) or {}),
            'alertas_cobertura': reporte.get('alertas_cobertura', 0),
            'unidades_sin_cobertura': reporte.get('unidades_sin_cobertura', []),
            'proximos_retiros': reporte.get('proximos_retiros', 0),
            'proximos_retiros_lista': reporte.get('proximos_retiros_lista', []),
            'falta_nutricion': reporte.get('falta_nutricion', 0),
            'grupos_edad_totales': reporte.get('grupos_edad_totales', {})
        },
        'unidades': reporte.get('unidades', {}),
        'movimientos': cambios,
        'errores_formatos': errores_formatos,
        'archivo': filename,
        'procesamiento': {
            'modo': 'procesar_todo' if procesar_todo else 'unidades_seleccionadas',
            'unidades_solicitadas': unidades_seleccionadas,
            'unidades_no_encontradas': unidades_no_encontradas,
            'total_usuarios_base_maestra': total_usuarios_base,
            'total_usuarios_formatos': total_usuarios_formatos,
            'total_unidades_formatos': len(reporte.get('unidades', {}) or {})
        },
        'procesado_en_segundo_plano': bool(update_job)
    }
    update(100, 'Base procesada y formatos generados')
    log_procesamiento_base_maestra('Proceso finalizado', archivo=filename, total_unidades=len(reporte.get('unidades', {}) or {}))
    return resultado


@app.route('/api/jobs', methods=['GET'])
def api_jobs_list():
    try:
        limit = int(request.args.get('limit') or 50)
    except Exception:
        limit = 50
    return jsonify({'jobs': list_jobs(limit=limit)}), 200


@app.route('/api/jobs/<job_id>', methods=['GET'])
def api_jobs_detail(job_id):
    job = get_job(job_id)
    if not job:
        return jsonify({'error': 'Trabajo operativo no encontrado o expirado.'}), 404
    return jsonify({'job': job}), 200


def _procesamiento_async_explicito(req):
    """True únicamente cuando el cliente solicita expresamente segundo plano."""
    sync_raw = str(req.args.get('sync') or req.form.get('sync') or '').strip().lower()
    modo_raw = str(req.args.get('modo_ejecucion') or req.form.get('modo_ejecucion') or '').strip().lower()
    async_raw = str(
        req.args.get('async') or req.form.get('async')
        or req.args.get('procesamiento_masivo') or req.form.get('procesamiento_masivo')
        or ''
    ).strip().lower()
    valores_verdaderos = {'1', 'true', 'si', 'sí'}
    modos_sincronos = {'sincrono', 'síncrono', 'synchronous'}
    modos_asincronos = {'asincrono', 'asíncrono', 'segundo_plano', 'masivo', 'async', 'asynchronous'}
    sincronico_explicito = sync_raw in valores_verdaderos or modo_raw in modos_sincronos
    return not sincronico_explicito and (
        async_raw in valores_verdaderos or modo_raw in modos_asincronos
    )


@app.route('/api/procesar', methods=['POST'])
def procesar_sistema():
    """Carga Cuéntame, permite detectar unidades y procesa solo las seleccionadas.

    Compatibilidad:
    - Sin banderas especiales y con archivo nuevo conserva el comportamiento histórico: procesa todo.
    - Con solo_detectar_unidades=1 guarda el archivo, detecta unidades y no toca BD.
    - Con archivo_token + unidades_seleccionadas procesa el archivo ya cargado.
    """
    file = request.files.get('file')
    archivo_token = request.form.get('archivo_token') or request.args.get('archivo_token') or ''
    modo_detectar_unidades = _valor_booleano(
        request.form.get('solo_detectar_unidades')
        or request.args.get('solo_detectar_unidades')
        or request.form.get('detectar_unidades')
        or request.args.get('detectar_unidades')
    )

    try:
        if file and file.filename:
            ruta_cuentame, filename = _guardar_archivo_cuentame_desde_request(file)
        elif archivo_token:
            ruta_cuentame, filename = _resolver_archivo_cuentame_guardado(archivo_token)
            if not ruta_cuentame:
                return jsonify({'error': 'No se encontró el archivo Cuéntame cargado. Vuelve a seleccionar y detectar la base.'}), 400
        else:
            return jsonify({'error': 'No se encontró el archivo del Cuéntame.'}), 400
    except ValueError as exc:
        return jsonify({'error': str(exc)}), 400

    if modo_detectar_unidades:
        try:
            return jsonify(_detectar_unidades_cuentame(ruta_cuentame, filename)), 200
        except Exception as exc:
            return jsonify({'error': f'No se pudieron detectar unidades en el archivo: {str(exc)}'}), 500

    crear_backup_operativo('ANTES_IMPORTAR_CUENTAME', f'Backup antes de importar base Cuéntame: {filename}.')

    # Compatibilidad ALPHA34: antes se leía con request.form.getlist('unidad_seleccionada');
    # ahora se centraliza para aceptar JSON, CSV y campos repetidos sin perder selección.
    unidades_normalizadas = normalizar_unidades_seleccionadas(request)
    unidades_payload = json.dumps(unidades_normalizadas, ensure_ascii=False) if unidades_normalizadas else ''
    log_procesamiento_base_maestra(
        'Unidades seleccionadas recibidas',
        unidades=' | '.join(unidades_normalizadas) if unidades_normalizadas else 'SIN SELECCIÓN',
        procesar_todo=request.form.get('procesar_todo') or request.args.get('procesar_todo') or ''
    )

    options = {
        'mes': request.form.get('mes') or request.args.get('mes'),
        'año': request.form.get('año') or request.args.get('año') or request.form.get('anio') or request.args.get('anio'),
        'anio': request.form.get('anio') or request.args.get('anio') or request.form.get('año') or request.args.get('año'),
        'max_usuarios_formato': request.form.get('max_usuarios_formato') or request.args.get('max_usuarios_formato') or '20',
        'bienestarina_por_hoja': request.form.get('bienestarina_por_hoja') or request.args.get('bienestarina_por_hoja') or '14',
        'fecha_entrega_bienestarina': request.form.get('fecha_entrega_bienestarina') or request.args.get('fecha_entrega_bienestarina') or '',
        'dia_entrega_bienestarina': request.form.get('dia_entrega_bienestarina') or request.args.get('dia_entrega_bienestarina') or '',
        'mes_entrega_bienestarina': request.form.get('mes_entrega_bienestarina') or request.args.get('mes_entrega_bienestarina') or '',
        'anio_entrega_bienestarina': request.form.get('anio_entrega_bienestarina') or request.args.get('anio_entrega_bienestarina') or '',
        'lote_bienestarina': request.form.get('lote_bienestarina') or request.args.get('lote_bienestarina') or '',
        'cantidad_bienestarina': request.form.get('cantidad_bienestarina') or request.args.get('cantidad_bienestarina') or '',
        'formatos_seleccionados': request.form.get('formatos_seleccionados') or request.args.get('formatos_seleccionados') or '',
        'unidades_seleccionadas': unidades_payload,
        'procesar_todo': request.form.get('procesar_todo') or request.args.get('procesar_todo') or '',
        'archivo_token': os.path.basename(ruta_cuentame),
    }

    unidades_solicitadas = unidades_normalizadas
    if not unidades_solicitadas and not _valor_booleano(options.get('procesar_todo')):
        # Compatibilidad: una carga directa con archivo nuevo y sin selector conserva
        # el comportamiento histórico. Cuando ya existe archivo_token, sí se exige
        # selección explícita desde la interfaz nueva.
        if archivo_token:
            return jsonify({
                'error': 'Selecciona al menos una unidad de atención o usa la opción Procesar todo.',
                'code': 'UNIDADES_REQUERIDAS'
            }), 400
        options['procesar_todo'] = '1'

    async_explicito = _procesamiento_async_explicito(request)

    # ALPHA77: Railway procesa una operación normal en la misma solicitud.
    # El motor de jobs se conserva para cargas masivas solicitadas expresamente.
    if not async_explicito:
        trace_id = uuid.uuid4().hex[:16]
        try:
            resultado = _procesar_base_cuentame_core(
                ruta_cuentame,
                filename,
                options=options,
                update_job=None,
            )
            resultado['modo'] = 'sincrono'
            resultado['modo_ejecucion'] = 'sincrono'
            return jsonify(resultado), 200
        except Exception as exc:
            log_procesamiento_base_maestra(
                'Error exacto del procesamiento síncrono', str(exc),
                archivo=filename, trace_id=trace_id,
            )
            return jsonify({
                'error': 'No fue posible completar el procesamiento.',
                'detalle': str(exc),
                'trace_id': trace_id,
                'modo': 'sincrono',
            }), 500

    descripcion = 'Procesamiento de base Cuéntame para unidades seleccionadas'
    if _valor_booleano(options.get('procesar_todo')):
        descripcion = 'Procesamiento completo de base Cuéntame y formatos oficiales'

    def ejecutar_procesamiento_cuentame(update):
        # Flask no comparte automáticamente el contexto de aplicación con hilos.
        # El job en segundo plano necesita app_context para que funciones heredadas
        # que usan g/current_app/config/SQLAlchemy puedan ejecutarse sin romperse.
        with app.app_context():
            try:
                return _procesar_base_cuentame_core(
                    ruta_cuentame,
                    filename,
                    options=options,
                    update_job=update
                )
            except Exception as exc:
                log_procesamiento_base_maestra('Error exacto del procesamiento', str(exc), archivo=filename)
                raise

    job = start_job(
        'procesar_base_cuentame',
        ejecutar_procesamiento_cuentame,
        metadata={'archivo': filename, 'ruta': ruta_cuentame, 'opciones': options},
        descripcion=descripcion
    )
    return jsonify({
        'message': 'La operación masiva fue recibida y continuará en segundo plano por solicitud explícita.',
        'job_id': job['id'],
        'job': job,
        'status_url': f"/api/jobs/{job['id']}",
        'modo': 'segundo_plano_explicito'
    }), 202


@app.route('/api/plantillas', methods=['GET', 'POST'])
def manejar_plantillas():
    if request.method == 'GET':
        conn = database_connection()
        cursor = conn.cursor()
        ensure_runtime_schema(cursor)
        conn.commit()
        plantillas = cursor.execute("SELECT * FROM plantillas ORDER BY COALESCE(fecha_carga, fecha_ultima_actualizacion, '') DESC").fetchall()
        conn.close()
        return jsonify({'plantillas': [dict(row) for row in plantillas]})

    if 'file' not in request.files:
        return jsonify({'error': 'Falta el archivo de plantilla.'}), 400
    if 'tipo' not in request.form:
        return jsonify({'error': 'Falta el tipo de plantilla.'}), 400

    file = request.files['file']
    tipo = request.form['tipo'].strip()[:50]
    version = request.form.get('version', '').strip()[:50]

    if file.filename == '':
        return jsonify({'error': 'Archivo no seleccionado.'}), 400
    if not es_extension_valida(file.filename, ALLOWED_TEMPLATE_EXTENSIONS):
        return jsonify({'error': 'Formato no permitido. Puedes cargar plantillas oficiales Excel y documentos comunes: xlsx, xls, docx, pdf, imágenes, zip o rar.'}), 400

    nombre_original = secure_filename(file.filename)
    nombre_guardado = f"{tipo}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{nombre_original}"
    ruta_destino = os.path.join(TEMPLATES_FOLDER, nombre_guardado)
    file.save(ruta_destino)

    conn = database_connection()
    conn.execute('''
        INSERT INTO plantillas
        (nombre, nombre_original, nombre_guardado, tipo, ruta_archivo, fecha_carga, version, estado, activa)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        nombre_original,
        nombre_original,
        nombre_guardado,
        tipo,
        ruta_destino,
        datetime.now().isoformat(),
        version or '1.0',
        'activo',
        1
    ))
    conn.commit()
    conn.close()

    return jsonify({'message': 'Plantilla cargada correctamente.'})


@app.route('/api/plantillas/<int:plantilla_id>', methods=['PUT', 'PATCH', 'DELETE'])
def editar_o_eliminar_plantilla(plantilla_id):
    """Edita, desactiva o borra permanentemente una plantilla registrada.

    - PUT/PATCH: actualiza metadatos (tipo, versión y estado) sin tocar el archivo.
    - DELETE normal: desactiva la plantilla, conservando historial y archivo.
    - DELETE ?hard=1: borra el registro y, si existe, elimina el archivo físico.
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    plantilla = cursor.execute('SELECT * FROM plantillas WHERE id = ?', (plantilla_id,)).fetchone()

    if not plantilla:
        conn.close()
        return jsonify({'error': 'Plantilla no encontrada.'}), 404

    if request.method in ['PUT', 'PATCH']:
        data = request.get_json(silent=True) or {}
        tipo = limpiar_valor(data.get('tipo', plantilla['tipo'] if 'tipo' in plantilla.keys() else 'Otros'))[:50] or 'Otros'
        version = limpiar_valor(data.get('version', plantilla['version'] if 'version' in plantilla.keys() else '1.0'))[:50] or '1.0'
        estado = limpiar_valor(data.get('estado', plantilla['estado'] if 'estado' in plantilla.keys() else 'activo'))[:20].lower() or 'activo'
        activa = 0 if estado in {'inactivo', 'eliminado', 'desactivado'} else 1
        ahora = datetime.now().isoformat()

        cursor.execute("""
            UPDATE plantillas
            SET tipo = ?, version = ?, estado = ?, activa = ?, fecha_ultima_actualizacion = ?
            WHERE id = ?
        """, (tipo, version, estado, activa, ahora, plantilla_id))
        conn.commit()
        actualizada = cursor.execute('SELECT * FROM plantillas WHERE id = ?', (plantilla_id,)).fetchone()
        conn.close()
        return jsonify({
            'message': 'Plantilla actualizada correctamente.',
            'plantilla': dict(actualizada)
        })

    hard_delete = str(request.args.get('hard', '')).lower() in {'1', 'true', 'si', 'sí', 'permanente'}
    ahora = datetime.now().isoformat()

    if hard_delete:
        ruta_archivo = ''
        try:
            ruta_archivo = plantilla['ruta_archivo'] or ''
        except Exception:
            ruta_archivo = ''

        cursor.execute('DELETE FROM plantillas WHERE id = ?', (plantilla_id,))
        conn.commit()
        conn.close()

        eliminado_archivo = False
        if ruta_archivo:
            try:
                ruta_normalizada = os.path.abspath(ruta_archivo)
                templates_abs = os.path.abspath(TEMPLATES_FOLDER)
                if ruta_normalizada.startswith(templates_abs) and os.path.exists(ruta_normalizada):
                    os.remove(ruta_normalizada)
                    eliminado_archivo = True
            except Exception:
                eliminado_archivo = False

        return jsonify({
            'message': 'Plantilla borrada permanentemente.' + (' Archivo eliminado.' if eliminado_archivo else ''),
            'archivo_eliminado': eliminado_archivo
        })

    cursor.execute("""
        UPDATE plantillas
        SET estado = ?, activa = ?, fecha_ultima_actualizacion = ?
        WHERE id = ?
    """, ('inactivo', 0, ahora, plantilla_id))
    conn.commit()
    conn.close()

    return jsonify({'message': 'Plantilla desactivada correctamente. El historial y archivo se conservaron.'})


@app.route('/api/nutricion', methods=['POST'])
def procesar_nutricion():
    if 'file' not in request.files:
        return jsonify({'error': 'Falta el archivo de nutrición.'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Archivo no seleccionado.'}), 400
    if not es_extension_valida(file.filename, ALLOWED_NUTRICION_EXTENSIONS):
        return jsonify({'error': 'Formato no permitido para nutrición. Usa una base tabular .xlsx, .xls, .csv o .txt.'}), 400

    filename = secure_filename(file.filename)
    path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(path)

    try:
        ext = os.path.splitext(filename.lower())[1]
        if ext in {'.xlsx', '.xls', '.xlsm'}:
            df = pd.read_excel(path)
        else:
            df = pd.read_csv(path, sep=None, engine='python')

        df.columns = [str(c).strip() for c in df.columns]
        col_documento = buscar_columna(df, [
            'documento', 'documento del beneficiario', 'numero documento', 'número documento',
            'nui', 'nuip', 'identificacion', 'identificación', 'no documento', 'n doc ident'
        ])
        col_nombre = buscar_columna(df, [
            'nombre', 'nombre completo', 'nombres y apellidos', 'beneficiario', 'nombre beneficiario'
        ])
        col_unidad = buscar_columna(df, [
            'unidad', 'unidad de servicio', 'nombre de la unidad de servicio', 'uds', 'uca'
        ])
        col_peso = buscar_columna(df, ['peso', 'peso kg', 'peso en kg', 'kg'])
        col_talla = buscar_columna(df, ['talla', 'longitud', 'estatura', 'talla cm', 'cm'])
        col_fecha = buscar_columna(df, [
            'fecha_toma', 'fecha toma', 'fecha de toma', 'fecha medicion', 'fecha de medición', 'fecha', 'fecha control'
        ])
        col_responsable = buscar_columna(df, ['responsable', 'docente', 'auxiliar', 'usuario registra'])

        conn = get_db_connection()
        cursor = conn.cursor()
        ensure_runtime_schema(cursor)
        status = {'al_dia': 0, 'proximo_vencer': 0, 'vencido': 0, 'pendiente': 0}
        procesados = 0
        hoy = datetime.now().date()

        for idx, fila in df.iterrows():
            documento = limpiar_valor(fila.get(col_documento)) if col_documento else ''
            nombre = limpiar_valor(fila.get(col_nombre)) if col_nombre else ''
            unidad = normalize_unidad(fila.get(col_unidad)) if col_unidad else ''
            if not documento and not nombre:
                continue
            if not documento:
                documento = f'SIN_DOC_{idx + 1}'
            if not nombre:
                nombre = documento

            parsed_fecha = pd.to_datetime(fila.get(col_fecha), errors='coerce', dayfirst=True) if col_fecha else pd.NaT
            fecha_toma = None if pd.isna(parsed_fecha) else parsed_fecha.date()
            if not fecha_toma:
                estado = 'pendiente'
            else:
                dias = (hoy - fecha_toma).days
                if dias <= 75:
                    estado = 'al_dia'
                elif dias <= AlertaConfiguracion.DIAS_CONTROL_NUTRICION:
                    estado = 'proximo_vencer'
                else:
                    estado = 'vencido'
            status[estado] = status.get(estado, 0) + 1

            benef = cursor.execute(
                "SELECT id, fecha_nacimiento, unidad, nombres, apellidos FROM beneficiarios WHERE documento = ? OR nui = ?",
                (documento, documento)
            ).fetchone()
            if benef:
                beneficiario_id = benef['id']
                edad_meses = calcular_edad_meses(benef['fecha_nacimiento'])
                if not unidad:
                    unidad = normalize_unidad(benef['unidad'])
            else:
                nombres, apellidos = dividir_nombre(nombre)
                cursor.execute("""
                    INSERT INTO beneficiarios
                    (documento, nombres, apellidos, fecha_nacimiento, unidad, estado,
                     tipo_beneficiario, fecha_ingreso, fecha_carga)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    documento, nombres, apellidos, datetime.now().date().isoformat(), unidad,
                    EstadoUsuario.ACTIVO, 'NINO', datetime.now().isoformat(), datetime.now().isoformat()
                ))
                beneficiario_id = cursor.lastrowid
                edad_meses = 0

            def to_float_safe(value):
                txt = limpiar_valor(value)
                if not txt:
                    return None
                txt = txt.replace(',', '.')
                try:
                    return float(txt)
                except Exception:
                    return None

            peso = to_float_safe(fila.get(col_peso)) if col_peso else None
            talla = to_float_safe(fila.get(col_talla)) if col_talla else None
            estado_nutricional = clasificar_nutricional(peso, talla, edad_meses)
            fecha_medicion = fecha_toma.isoformat() if fecha_toma else datetime.now().date().isoformat()
            fecha_proximo = (datetime.now() + timedelta(days=AlertaConfiguracion.DIAS_CONTROL_NUTRICION)).date().isoformat()
            responsable = limpiar_valor(fila.get(col_responsable)) if col_responsable else 'sistema'

            cursor.execute("""
                INSERT INTO peso_talla
                (beneficiario_id, documento, nombre, unidad, peso, talla, fecha_toma, estado,
                 fecha_medicion, responsable, estado_nutricional, fecha_proximo_control, fecha_carga)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                beneficiario_id, documento, nombre, unidad, peso, talla,
                fecha_toma.isoformat() if fecha_toma else None,
                estado, fecha_medicion, responsable, estado_nutricional, fecha_proximo, datetime.now().isoformat()
            ))
            procesados += 1

        conn.commit()
        conn.close()

        return jsonify({
            'message': f'Nutrición registrada correctamente. Registros procesados: {procesados}.',
            'status': status,
            'boa': generar_boa_nutricion(cursor=None),
            'columnas_detectadas': {
                'documento': col_documento, 'nombre': col_nombre, 'unidad': col_unidad,
                'peso': col_peso, 'talla': col_talla, 'fecha': col_fecha
            }
        })
    except Exception as e:
        return jsonify({'error': f'Error al procesar nutrición: {e}'}), 500




def trimestre_fecha(fecha_texto):
    try:
        fecha = pd.to_datetime(fecha_texto, errors='coerce', dayfirst=True)
        if pd.isna(fecha):
            return ''
        mes = int(fecha.month)
        return f'T{((mes - 1) // 3) + 1}-{int(fecha.year)}'
    except Exception:
        return ''


def generar_boa_nutricion(cursor=None):
    """Construye detalle BOA de nutrición y talla con riesgo, vencimiento y trimestre."""
    conn = get_db_connection()
    fid = fundacion_actual_id()
    filas = conn.execute("""
        SELECT p.*, b.nombres, b.apellidos, b.fecha_nacimiento
        FROM peso_talla p
        LEFT JOIN beneficiarios b
          ON b.id = p.beneficiario_id
         AND COALESCE(b.fundacion_id, 1) = ?
        WHERE COALESCE(p.fundacion_id, 1) = ?
        ORDER BY COALESCE(p.fecha_medicion, p.fecha_toma, p.fecha_carga) DESC
    """, (fid, fid)).fetchall()
    detalles = []
    resumen = {'ADECUADO': 0, 'RIESGO': 0, 'DESNUTRICION': 0, 'SOBREPESO': 0, 'PENDIENTE': 0}
    controles = {'al_dia': 0, 'proximo_vencer': 0, 'vencido': 0, 'pendiente': 0}
    unidades = {}
    seen = set()
    for row in filas:
        doc = limpiar_valor(row['documento'])
        # Mostrar el registro más reciente por documento para que la BOA sea legible.
        if doc in seen:
            continue
        seen.add(doc)
        estado_nut = limpiar_valor(row['estado_nutricional']) or EstadoNutricion.PENDIENTE
        estado_control = limpiar_valor(row['estado']) or 'pendiente'
        resumen[estado_nut] = resumen.get(estado_nut, 0) + 1
        controles[estado_control] = controles.get(estado_control, 0) + 1
        unidad = normalize_unidad(row['unidad']) or 'SIN UNIDAD'
        if unidad not in unidades:
            unidades[unidad] = {'unidad': unidad, 'total': 0, 'riesgo': 0, 'desnutricion': 0, 'vencidos': 0, 'pendientes': 0}
        unidades[unidad]['total'] += 1
        if estado_nut == EstadoNutricion.RIESGO:
            unidades[unidad]['riesgo'] += 1
        if estado_nut == EstadoNutricion.DESNUTRICION:
            unidades[unidad]['desnutricion'] += 1
        if estado_control == 'vencido':
            unidades[unidad]['vencidos'] += 1
        if estado_control == 'pendiente':
            unidades[unidad]['pendientes'] += 1
        nombre = limpiar_valor(row['nombre']) or unir_partes(row['nombres'], row['apellidos'])
        detalles.append({
            'documento': doc,
            'nombre': nombre,
            'unidad': unidad,
            'peso': row['peso'],
            'talla': row['talla'],
            'fecha_medicion': limpiar_valor(row['fecha_medicion'] or row['fecha_toma']),
            'fecha_proximo_control': limpiar_valor(row['fecha_proximo_control']),
            'trimestre': trimestre_fecha(row['fecha_medicion'] or row['fecha_toma']),
            'estado_nutricional': estado_nut,
            'estado_control': estado_control,
            'responsable': limpiar_valor(row['responsable']),
            'foto': ''
        })
    conn.close()
    return {
        'resumen': resumen,
        'controles': controles,
        'unidades': sorted(unidades.values(), key=lambda x: x['unidad']),
        'detalles': detalles
    }


@app.route('/api/nutricion/boa', methods=['GET'])
def nutricion_boa():
    return jsonify({'boa': generar_boa_nutricion()})

@app.route('/api/talento', methods=['GET', 'POST'])
def procesar_talento():
    """Talento Humano migrado a servicio SQLAlchemy Core.

    Mantiene compatibilidad con el frontend existente y conserva el parser
    histórico de archivos. La persistencia y sincronización global delegan en
    modules.talento_humano.
    """
    from modules.talento_humano.services import TalentoHumanoService, normalizar_registro

    service = TalentoHumanoService()

    if request.method == 'GET':
        return jsonify({'talento': service.list_talento(), 'integracion': service.resumen_integracion()})

    # Registro manual desde JSON.
    if request.is_json:
        crear_backup_operativo('ANTES_ACTUALIZAR_TALENTO', 'Backup antes de registrar talento humano manual.')
        data = request.get_json(silent=True) or {}
        registro = normalizar_registro(data, archivo='manual')
        if not registro['documento'] or not registro['nombre']:
            return jsonify({'error': 'Nombre y documento son obligatorios.'}), 400
        resultado = service.guardar_registros([registro], origen='registro_manual')
        return jsonify({
            'message': f"Talento humano guardado correctamente. Registros: {resultado.get('total', 0)}. Plataforma sincronizada.",
            'resultado': resultado,
            'integracion': service.resumen_integracion()
        }), 200

    if 'file' not in request.files:
        return jsonify({'error': 'Falta el archivo de talento humano.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Archivo no seleccionado.'}), 400
    if not es_extension_valida(file.filename, ALLOWED_TALENTO_EXTENSIONS):
        return jsonify({'error': 'Formato no permitido para talento humano. Usa .xlsx, .xls, .csv, .txt, .docx o .zip.'}), 400

    filename = secure_filename(file.filename)
    path = os.path.join(UPLOAD_FOLDER, filename)
    crear_backup_operativo('ANTES_ACTUALIZAR_TALENTO', f'Backup antes de importar talento humano: {filename}.')
    file.save(path)

    try:
        extension = os.path.splitext(filename.lower())[1]
        registros = []
        if extension == '.zip':
            import zipfile
            with zipfile.ZipFile(path) as zf:
                for item in zf.namelist():
                    lower = item.lower()
                    base_item = secure_filename(os.path.basename(item))
                    if not base_item:
                        continue
                    destino = os.path.join(UPLOAD_FOLDER, f"TALENTO_{datetime.now().strftime('%Y%m%d%H%M%S')}_{base_item}")
                    if lower.endswith(('.xlsx', '.xls', '.xlsm', '.csv', '.txt')):
                        with zf.open(item) as src, open(destino, 'wb') as out:
                            out.write(src.read())
                        if lower.endswith(('.xlsx', '.xls', '.xlsm')):
                            xls_zip = pd.ExcelFile(destino)
                            for sheet_zip in xls_zip.sheet_names:
                                df_zip = pd.read_excel(destino, sheet_name=sheet_zip, header=None)
                                coord_respaldo = re.sub(r'(?i)^nomina\s+', '', str(sheet_zip)).strip()
                                registros.extend(parsear_talento_dataframe(df_zip, archivo=f'{filename}:{sheet_zip}', coordinador_respaldo=coord_respaldo))
                        else:
                            df_zip = pd.read_csv(destino, header=None, sep=None, engine='python')
                            registros.extend(parsear_talento_dataframe(df_zip, archivo=filename))
                    elif lower.endswith('.docx'):
                        # Los DOCX dentro de ZIP suelen ser plantillas de cuenta de cobro. Se registran allí.
                        with zf.open(item) as src:
                            destino_docx = os.path.join(CUENTAS_COBRO_FOLDER, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{base_item}")
                            with open(destino_docx, 'wb') as out:
                                out.write(src.read())
                            registrar_plantilla_cuenta(destino_docx, base_item)
        elif extension == '.docx':
            registrar_plantilla_cuenta(path, filename)
        elif filename.endswith('.xlsx') or filename.endswith('.xls') or filename.endswith('.xlsm'):
            xls = pd.ExcelFile(path)
            for sheet in xls.sheet_names:
                df = pd.read_excel(path, sheet_name=sheet, header=None)
                coord_respaldo = re.sub(r'(?i)^nomina\s+', '', str(sheet)).strip()
                registros.extend(parsear_talento_dataframe(df, archivo=f'{filename}:{sheet}', coordinador_respaldo=coord_respaldo))
        else:
            df = pd.read_csv(path, header=None, sep=None, engine='python')
            registros = parsear_talento_dataframe(df, archivo=filename)

        if not registros:
            if extension in {'.zip', '.docx'}:
                return jsonify({'message': 'El archivo fue recibido. Los DOCX se registraron como plantillas de cuentas de cobro; no se detectó una tabla de talento humano.'}), 200
            # Segundo intento para archivos CSV/Excel con encabezados en la primera fila.
            if filename.endswith('.xlsx') or filename.endswith('.xls') or filename.endswith('.xlsm'):
                xls2 = pd.ExcelFile(path)
                for sheet2 in xls2.sheet_names:
                    df2 = pd.read_excel(path, sheet_name=sheet2)
                    coord_respaldo2 = re.sub(r'(?i)^nomina\s+', '', str(sheet2)).strip()
                    registros.extend(parsear_talento_dataframe(df2, archivo=f'{filename}:{sheet2}', coordinador_respaldo=coord_respaldo2))
            else:
                df2 = pd.read_csv(path, sep=None, engine='python')
                registros = parsear_talento_dataframe(df2, archivo=filename)

        if not registros:
            return jsonify({'error': 'No se detectaron registros. Usa columnas como NOMBRES Y APELLIDOS, CEDULA, CARGO, COMUNIDAD, DIRECCION y TELEFONO.'}), 400

        resultado = service.guardar_registros(registros, origen=f'importar_archivo:{filename}')
        return jsonify({
            'message': f"Talento humano procesado correctamente. Registros: {resultado.get('total', 0)}. Plataforma sincronizada.",
            'resultado': resultado,
            'integracion': service.resumen_integracion()
        })
    except Exception as e:
        return jsonify({'error': f'Error al procesar talento humano: {e}'}), 500

@app.route('/api/talento/<int:talento_id>', methods=['PUT', 'PATCH', 'DELETE'])
def editar_o_eliminar_talento(talento_id):
    """Editar, desactivar o borrar Talento Humano usando el servicio SQLAlchemy Core."""
    from modules.talento_humano.services import TalentoHumanoService

    service = TalentoHumanoService()

    if request.method in ['PUT', 'PATCH']:
        data = request.get_json(silent=True) or {}
        try:
            actualizado = service.update_talento(talento_id, data)
        except ValueError as exc:
            return jsonify({'error': str(exc)}), 400
        if not actualizado:
            return jsonify({'error': 'Registro de talento humano no encontrado.'}), 404
        return jsonify({
            'message': 'Talento humano actualizado correctamente. Plataforma sincronizada.',
            'talento': actualizado,
            'integracion': service.resumen_integracion()
        })

    hard_delete = str(request.args.get('hard', '')).lower() in {'1', 'true', 'si', 'sí', 'permanente'}
    ok = service.delete_talento(talento_id, hard=hard_delete)
    if not ok:
        return jsonify({'error': 'Registro de talento humano no encontrado.'}), 404
    return jsonify({
        'message': 'Talento humano borrado permanentemente.' if hard_delete else 'Talento humano desactivado correctamente.',
        'integracion': service.resumen_integracion()
    })


@app.route('/api/talento/integracion', methods=['GET'])
def talento_integracion_estado():
    from modules.talento_humano.services import TalentoHumanoService
    return jsonify({'integracion': TalentoHumanoService().resumen_integracion()})


@app.route('/api/talento/sincronizar-global', methods=['POST'])
def talento_sincronizar_global_endpoint():
    from modules.talento_humano.services import TalentoHumanoService
    service = TalentoHumanoService()
    resultado = service.sincronizar_global(origen='endpoint')
    return jsonify({
        'message': 'Talento Humano sincronizado con Gestión Pedagógica, Gestión por Coordinador, Planeación, Relación del Mes y formatos.',
        'resultado': resultado,
        'integracion': service.resumen_integracion()
    })


@app.route('/api/talento/fuente-maestra', methods=['GET'])
def talento_fuente_maestra_estado():
    from modules.talento_humano.services import TalentoHumanoService
    return jsonify(TalentoHumanoService().fuente_maestra())


@app.route('/api/talento/fuente-maestra/sincronizar', methods=['POST'])
def talento_fuente_maestra_sincronizar():
    from modules.talento_humano.services import TalentoHumanoService
    service = TalentoHumanoService()
    resultado = service.sincronizar_global(origen='fuente_maestra_endpoint')
    return jsonify({
        'message': 'Fuente maestra de Talento Humano sincronizada correctamente.',
        'resultado': resultado,
        'resumen': service.resumen_integracion(),
        'integracion': service.resumen_integracion()
    })



# ==================== RUTAS: CUENTAS DE COBRO Y RELACIÓN DEL MES ====================

def mes_nombre_es(mes: int, minuscula: bool = False) -> str:
    nombre = MESES_ES.get(int(mes), str(mes))
    return nombre.lower() if minuscula else nombre


def periodo_cuenta_cobro(mes: int, anio: int) -> str:
    ultimo = calendar.monthrange(int(anio), int(mes))[1]
    return f"del 1 al {ultimo} de {mes_nombre_es(mes, True)} de {anio}"


def fecha_ciudad_cuenta(ciudad: str, mes: int, anio: int) -> str:
    ciudad = limpiar_valor(ciudad) or 'Ciudad de prueba'
    return f"{ciudad}, 1 de {mes_nombre_es(mes, True)} de {anio}"


def reemplazar_texto_docx_seguro(document, reemplazador):
    """Reemplaza texto en párrafos y tablas conservando formato básico de runs."""
    def patch_paragraph(paragraph):
        texto_original = ''.join(run.text for run in paragraph.runs)
        if not texto_original:
            return
        texto_nuevo = reemplazador(texto_original)
        if texto_nuevo == texto_original:
            return
        if paragraph.runs:
            paragraph.runs[0].text = texto_nuevo
            for run in paragraph.runs[1:]:
                run.text = ''

    for paragraph in document.paragraphs:
        patch_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    patch_paragraph(paragraph)


def extraer_numero_cuenta_docx(ruta_archivo: str) -> int:
    try:
        from docx import Document
        doc = Document(ruta_archivo)
        textos = []
        for p in doc.paragraphs:
            textos.append(p.text or '')
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    textos.append(cell.text or '')
        texto = '\n'.join(textos)
        match = re.search(r'cuenta\s+de\s+cobro\s*(?:n[°ºo\.]*|numero|número)?\s*(\d+)', texto, re.I)
        if match:
            return int(match.group(1))
    except Exception:
        pass
    return 0


def generar_cuenta_cobro_desde_docx(ruta_plantilla: str, salida: str, numero: int, mes: int, anio: int, ciudad: str) -> None:
    from docx import Document
    doc = Document(ruta_plantilla)
    periodo = periodo_cuenta_cobro(mes, anio)
    ciudad_fecha = fecha_ciudad_cuenta(ciudad, mes, anio)
    mes_txt = mes_nombre_es(mes, True)

    def repl(texto: str) -> str:
        nuevo = texto
        nuevo = re.sub(
            r'CUENTA\s+DE\s+COBRO\s*(?:N[°ºO\.]*|NUMERO|NÚMERO)?\s*\d+',
            f'CUENTA DE COBRO N° {numero}',
            nuevo,
            flags=re.I,
        )
        nuevo = re.sub(
            r'(?:del\s+)?1\s+al\s+\d{1,2}\s+de\s+[a-záéíóúñ]+\s+(?:de|del)\s+\d{4}',
            periodo,
            nuevo,
            flags=re.I,
        )
        nuevo = re.sub(
            r'(Ciudad\s+y\s+fecha\s*:\s*)[^\n\r]+',
            rf'\1{ciudad_fecha}',
            nuevo,
            flags=re.I,
        )
        nuevo = re.sub(
            r'(mes\s+de\s+)[a-záéíóúñ]+\s+\d{4}',
            rf'\1{mes_txt} {anio}',
            nuevo,
            flags=re.I,
        )
        return nuevo

    reemplazar_texto_docx_seguro(doc, repl)
    doc.save(salida)


def plantilla_cuenta_from_file(ruta_archivo: str, nombre_original: str) -> dict:
    nombre_base = os.path.splitext(os.path.basename(nombre_original))[0]
    nombre_limpio = re.sub(r'[-_]+', ' ', nombre_base).strip()
    # Intento simple: separar nombre y unidad desde el nombre del archivo.
    unidad = ''
    for alias, destino in ALIAS_UNIDADES_CUENTAME.items():
        if normalizar_texto_clave(alias) in normalizar_texto_clave(nombre_limpio):
            unidad = normalize_unidad(destino)
            break
    return {'docente_nombre': nombre_limpio.upper(), 'documento': '', 'unidad': unidad}


def registrar_plantilla_cuenta(ruta_archivo: str, nombre_original: str) -> int:
    datos = plantilla_cuenta_from_file(ruta_archivo, nombre_original)
    conn = get_db_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    nombre_guardado = os.path.basename(ruta_archivo)
    existente = cursor.execute("SELECT id FROM cuentas_cobro_plantillas WHERE nombre_guardado = ?", (nombre_guardado,)).fetchone()
    if existente:
        conn.close()
        return int(existente['id'])
    cursor.execute("""
        INSERT INTO cuentas_cobro_plantillas
        (nombre_original, nombre_guardado, ruta_archivo, docente_nombre, documento, unidad, estado, fecha_carga)
        VALUES (?, ?, ?, ?, ?, ?, 'activo', ?)
    """, (
        nombre_original,
        nombre_guardado,
        ruta_archivo,
        datos.get('docente_nombre', ''),
        datos.get('documento', ''),
        datos.get('unidad', ''),
        datetime.now().isoformat(),
    ))
    conn.commit()
    new_id = cursor.lastrowid
    conn.close()
    return int(new_id)


def sembrar_plantillas_cuenta_desde_carpeta():
    conn = get_db_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    conn.close()
    for nombre in os.listdir(CUENTAS_COBRO_FOLDER):
        if nombre.lower().endswith('.docx'):
            registrar_plantilla_cuenta(os.path.join(CUENTAS_COBRO_FOLDER, nombre), nombre)


def listar_cuentas_generadas(periodo: str | None = None) -> list:
    conn = get_db_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    params = []
    where = '1=1'
    if periodo:
        where += ' AND periodo = ?'
        params.append(periodo)
    filas = cursor.execute(f"""
        SELECT * FROM cuentas_cobro_generadas
        WHERE {where}
        ORDER BY periodo DESC, docente_nombre
    """, params).fetchall()
    conn.close()
    return [dict(row) for row in filas]


@app.route('/api/cuentas-cobro/plantillas', methods=['GET', 'POST'])
def cuentas_cobro_plantillas():
    sembrar_plantillas_cuenta_desde_carpeta()
    if request.method == 'GET':
        conn = get_db_connection()
        cursor = conn.cursor()
        ensure_runtime_schema(cursor)
        filas = cursor.execute("SELECT * FROM cuentas_cobro_plantillas ORDER BY docente_nombre, nombre_original").fetchall()
        conn.close()
        return jsonify({'plantillas': [dict(row) for row in filas]})

    if 'file' not in request.files:
        return jsonify({'error': 'Falta el archivo de cuentas de cobro.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Archivo no seleccionado.'}), 400

    nombre_original = secure_filename(file.filename)
    extension = os.path.splitext(nombre_original.lower())[1]
    guardadas = []

    if extension == '.zip':
        import zipfile
        zip_path = os.path.join(CUENTAS_COBRO_FOLDER, f"ZIP_{datetime.now().strftime('%Y%m%d%H%M%S')}_{nombre_original}")
        file.save(zip_path)
        with zipfile.ZipFile(zip_path) as zf:
            for item in zf.namelist():
                if not item.lower().endswith('.docx'):
                    continue
                base_name = secure_filename(os.path.basename(item))
                if not base_name:
                    continue
                destino = os.path.join(CUENTAS_COBRO_FOLDER, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{base_name}")
                with zf.open(item) as src, open(destino, 'wb') as out:
                    out.write(src.read())
                registrar_plantilla_cuenta(destino, base_name)
                guardadas.append(base_name)
    elif extension == '.docx':
        destino = os.path.join(CUENTAS_COBRO_FOLDER, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{nombre_original}")
        file.save(destino)
        registrar_plantilla_cuenta(destino, nombre_original)
        guardadas.append(nombre_original)
    else:
        return jsonify({'error': 'Para cuentas de cobro se aceptan .docx o .zip con documentos .docx.'}), 400

    return jsonify({'message': f'Plantillas de cuenta cargadas: {len(guardadas)}.', 'archivos': guardadas})


@app.route('/api/cuentas-cobro', methods=['GET'])
def cuentas_cobro_listar():
    periodo = request.args.get('periodo')
    sembrar_plantillas_cuenta_desde_carpeta()
    return jsonify({'generadas': listar_cuentas_generadas(periodo)})


@app.route('/api/cuentas-cobro/generar', methods=['POST'])
def cuentas_cobro_generar():
    data = request.get_json(silent=True) or request.form.to_dict() or {}
    mes = int(data.get('mes') or datetime.now().month)
    anio = int(data.get('anio') or data.get('año') or datetime.now().year)
    ciudad = data.get('ciudad') or 'Ciudad de prueba'
    periodo = f'{anio}-{mes:02d}'
    numero_inicial = data.get('numero_inicial')
    numero_inicial = int(numero_inicial) if limpiar_valor(numero_inicial) else None

    sembrar_plantillas_cuenta_desde_carpeta()
    conn = get_db_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    plantillas = cursor.execute("SELECT * FROM cuentas_cobro_plantillas WHERE estado = 'activo' ORDER BY docente_nombre, nombre_original").fetchall()
    if not plantillas:
        conn.close()
        return jsonify({'error': 'No hay plantillas de cuenta de cobro cargadas.'}), 400

    generadas = []
    for idx, plantilla in enumerate(plantillas):
        ruta = plantilla['ruta_archivo']
        if not ruta or not os.path.exists(ruta):
            continue
        if numero_inicial is not None:
            numero = numero_inicial + idx
        else:
            ultimo = cursor.execute("""
                SELECT MAX(numero_cuenta) AS n FROM cuentas_cobro_generadas WHERE plantilla_id = ?
            """, (plantilla['id'],)).fetchone()['n']
            numero = int(ultimo or extraer_numero_cuenta_docx(ruta) or 0) + 1
        nombre_archivo = f"CUENTA_COBRO_{periodo}_{secure_filename(plantilla['docente_nombre'] or plantilla['nombre_original'])}_{numero}.docx"
        salida = os.path.join(OUTPUT_FOLDER, nombre_archivo)
        generar_cuenta_cobro_desde_docx(ruta, salida, numero, mes, anio, ciudad)
        cursor.execute("""
            INSERT INTO cuentas_cobro_generadas
            (plantilla_id, docente_nombre, documento, unidad, periodo, numero_cuenta, ciudad,
             nombre_archivo, ruta_archivo, fecha_generacion)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            plantilla['id'], plantilla['docente_nombre'], plantilla['documento'], plantilla['unidad'],
            periodo, numero, ciudad, nombre_archivo, salida, datetime.now().isoformat()
        ))
        generadas.append({'docente': plantilla['docente_nombre'], 'archivo': nombre_archivo, 'numero': numero})
    conn.commit()
    conn.close()
    return jsonify({'message': f'Cuentas de cobro generadas: {len(generadas)}.', 'generadas': generadas})


def obtener_docente_relacion(unidad: str) -> str:
    unidad_norm = normalize_unidad(unidad)
    conn = database_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    fila = None
    try:
        fila = cursor.execute("""
            SELECT nombre_completo AS nombre FROM master_talento_humano
            WHERE activo = 1 AND COALESCE(fundacion_id,1) = ?
              AND UPPER(TRIM(COALESCE(unidad_servicio,''))) = UPPER(TRIM(?))
              AND (UPPER(COALESCE(rol_normalizado,'')) LIKE '%DOCENTE%'
                   OR UPPER(COALESCE(rol_normalizado,'')) LIKE '%AGENTE%'
                   OR UPPER(COALESCE(cargo,'')) LIKE '%AGENTE%'
                   OR UPPER(COALESCE(cargo,'')) LIKE '%DOCENTE%')
            ORDER BY nombre_completo LIMIT 1
        """, (fundacion_actual_id(), unidad_norm)).fetchone()
    except Exception:
        fila = None
    if not fila:
        fila = cursor.execute("""
            SELECT nombre FROM coordinadores
            WHERE unidad = ? AND activo = 1
              AND (UPPER(COALESCE(tipo_equipo, '')) LIKE '%DOCENTE%' OR UPPER(COALESCE(cargo, '')) LIKE '%AGENTE%' OR UPPER(COALESCE(cargo, '')) LIKE '%DOCENTE%')
            ORDER BY nombre LIMIT 1
        """, (unidad_norm,)).fetchone()
    conn.close()
    return fila['nombre'] if fila else ''


@app.route('/api/relacion-mes/generar', methods=['GET', 'POST'])
def relacion_mes_generar():
    data = request.get_json(silent=True) or request.form.to_dict() or request.args.to_dict() or {}
    mes = int(data.get('mes') or datetime.now().month)
    anio = int(data.get('anio') or data.get('año') or datetime.now().year)
    periodo = f'{anio}-{mes:02d}'

    conn = get_db_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    filas = cursor.execute("""
        SELECT unidad_servicio AS unidad, grupo_etario, edad_meses, fecha_nacimiento,
               estado, docente, datos_json
        FROM master_ninos
        WHERE activo = 1 AND COALESCE(fundacion_id,1) = ?
    """, (fundacion_actual_id(),)).fetchall()
    conn.close()

    from services.relacion_mes_service import consolidar_por_unidad, docente_mas_frecuente
    resumen = consolidar_por_unidad((dict(fila) for fila in filas), anio, mes)

    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    wb = Workbook()
    ws = wb.active
    ws.title = f'RELACION {mes_nombre_es(mes)}'
    headers = [
        'UNIDAD DE ATENCIÓN', 'DOCENTE', 'GESTANTES', 'MENORES 6 MESES', '6 A 11 MESES',
        '1 A 2 AÑOS 11 MESES', '3 A 5 AÑOS 11 MESES', 'SIN CLASIFICAR / REVISAR',
        'TOTAL USUARIOS', 'HUEVOS PARA GRUPOS DE 30', 'HUEVOS PARA 6 A 11 (15)',
        'TOTAL HUEVOS (UNIDADES)', 'CUBETAS DE 30', 'PAQUETES COMPLETOS (7 CUBETAS)',
        'CUBETAS SUELTAS', 'VERDURAS', 'OLLA COMUNITARIA', 'BIENESTARINA'
    ]
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    ws.cell(1, 1).value = f'RELACIÓN DEL MES DE {mes_nombre_es(mes)} {anio}'
    ws.cell(1, 1).font = Font(bold=True, size=14)
    ws.cell(1, 1).alignment = Alignment(horizontal='center')
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    ws.cell(2, 1).value = 'Regla: 30 huevos por usuario; 6 a 11 meses recibe 15. Una cubeta contiene 30 huevos y un paquete contiene 7 cubetas.'
    ws.cell(2, 1).alignment = Alignment(horizontal='left', wrap_text=True)
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(3, col)
        cell.value = h
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.fill = PatternFill('solid', fgColor='D9EAD3')
    thin = Side(style='thin', color='000000')
    fill_gestantes = PatternFill('solid', fgColor='FFF2CC')  # amarillo
    fill_menores = PatternFill('solid', fgColor='FFF2CC')    # amarillo
    fill_6_11 = PatternFill('solid', fgColor='CFE2F3')       # azul
    fill_1_2 = PatternFill('solid', fgColor='D9EAD3')        # verde
    fill_3_5 = PatternFill('solid', fgColor='F4CCCC')        # rojo
    row = 4
    for unidad in sorted(resumen):
        d = resumen[unidad]
        docente = docente_mas_frecuente(d) or obtener_docente_relacion(unidad) or 'SIN DOCENTE ASIGNADO'
        valores = [
            unidad, docente, d['gestantes'], d['menores_6'], d['seis_11'], d['uno_2'], d['tres_5'],
            d['sin_clasificar'], f'=SUM(C{row}:H{row})', f'=(C{row}+D{row}+F{row}+G{row}+H{row})*30',
            f'=E{row}*15', f'=SUM(J{row}:K{row})', f'=ROUNDUP(L{row}/30,0)', f'=QUOTIENT(M{row},7)',
            f'=MOD(M{row},7)', f'=I{row}', f'=IF(I{row}>0,1,0)', f'=I{row}'
        ]
        for col, v in enumerate(valores, start=1):
            c = ws.cell(row, col)
            c.value = v
            c.alignment = Alignment(horizontal='center' if col >= 3 else 'left', vertical='center')
            c.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if col == 3:
                c.fill = fill_gestantes
            elif col == 4:
                c.fill = fill_menores
            elif col == 5:
                c.fill = fill_6_11
            elif col == 6:
                c.fill = fill_1_2
            elif col == 7:
                c.fill = fill_3_5
        row += 1
    total_row = row
    ws.cell(total_row, 1).value = 'TOTAL GENERAL'
    ws.cell(total_row, 1).font = Font(bold=True)
    for col in range(3, len(headers) + 1):
        cell = ws.cell(total_row, col)
        cell.value = f'=SUM({get_column_letter(col)}4:{get_column_letter(col)}{total_row - 1})'
        cell.font = Font(bold=True)
        cell.fill = PatternFill('solid', fgColor='B6D7A8')
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws.auto_filter.ref = f'A3:{get_column_letter(len(headers))}{total_row - 1}'
    ws.freeze_panes = 'C4'
    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 20 if col > 2 else 30
    try:
        wb.calculation.fullCalcOnLoad = True
        wb.calculation.forceFullCalc = True
        wb.calculation.calcMode = 'auto'
    except Exception:
        pass
    nombre = f'RELACION_MES_{periodo}.xlsx'
    salida = os.path.join(OUTPUT_FOLDER, nombre)
    wb.save(salida)
    return jsonify({'message': 'Relación del mes generada.', 'archivo': nombre, 'url': f'/api/descargar-archivo/{nombre}'})

# ==================== RUTAS: CUMPLIMIENTO V6 ====================
@app.route('/api/documentos-institucionales', methods=['GET', 'POST'])
def documentos_institucionales():
    if request.method == 'GET':
        conn = get_db_connection()
        documentos = conn.execute("""
            SELECT id, tipo, titulo, nombre_original, version, estado, fecha_carga
            FROM documentos_institucionales
            ORDER BY fecha_carga DESC
        """).fetchall()
        conn.close()
        return jsonify({'documentos': [dict(row) for row in documentos]})

    if 'file' not in request.files:
        return jsonify({'error': 'Falta el documento institucional.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Archivo no seleccionado.'}), 400
    if not es_extension_valida(file.filename, ALLOWED_DOCUMENT_EXTENSIONS):
        return jsonify({'error': 'Formato no permitido para centro documental.'}), 400

    tipo = request.form.get('tipo', 'Manual Operativo').strip()[:80]
    titulo = request.form.get('titulo', file.filename).strip()[:160]
    version = request.form.get('version', '1.0').strip()[:40]
    nombre_original = secure_filename(file.filename)
    extension = os.path.splitext(nombre_original.lower())[1]
    nombre_guardado = f"{tipo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{nombre_original}"
    ruta = os.path.join(DOCUMENTOS_FOLDER, nombre_guardado)
    file.save(ruta)
    texto_indexado = extraer_texto_documento(ruta, extension)

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO documentos_institucionales
        (tipo, titulo, nombre_original, nombre_guardado, ruta_archivo, version,
         texto_indexado, estado, fecha_carga)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        tipo, titulo, nombre_original, nombre_guardado, ruta, version,
        texto_indexado, 'vigente', datetime.now().isoformat()
    ))
    documento_id = cursor.lastrowid
    reglas_inferidas = inferir_reglas_desde_texto(texto_indexado, documento_id)
    guardar_reglas_documentales(cursor, reglas_inferidas)
    conn.commit()
    conn.close()

    return jsonify({
        'message': 'Documento institucional cargado e indexado.',
        'indexado': bool(texto_indexado),
        'nombre': nombre_original,
        'reglas_inferidas': len(reglas_inferidas)
    }), 201


@app.route('/api/entregables-operacion', methods=['GET', 'POST'])
def entregables_operacion():
    if request.method == 'GET':
        periodo = request.args.get('periodo') or periodo_actual()
        conn = get_db_connection()
        cursor = conn.cursor()
        ensure_runtime_schema(cursor)
        tablero, resumen = tablero_entregables_periodo(cursor, periodo)
        entregables = listar_entregables_periodo(cursor, periodo)
        conn.close()
        return jsonify({'periodo': periodo, 'entregables': entregables, 'tablero': tablero, 'resumen': resumen}), 200

    if 'file' not in request.files:
        return jsonify({'error': 'Falta el archivo del entregable.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Archivo no seleccionado.'}), 400
    if not es_extension_valida(file.filename, ALLOWED_DOCUMENT_EXTENSIONS):
        return jsonify({'error': 'Formato no permitido para entregables.'}), 400

    tipo = request.form.get('tipo', 'Evidencia').strip()[:80]
    periodo = request.form.get('periodo', periodo_actual()).strip()[:20]
    unidad = request.form.get('unidad', '').strip()[:120] or None
    responsable = request.form.get('responsable', '').strip()[:120]
    categoria = request.form.get('categoria', '').strip()[:80]
    fecha_limite = request.form.get('fecha_limite', '').strip()[:20]
    observaciones = request.form.get('observaciones', '').strip()[:300]
    nombre_original = secure_filename(file.filename)
    nombre_guardado = f"ENTREGABLE_{periodo}_{tipo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{nombre_original}"
    ruta = os.path.join(UPLOAD_FOLDER, nombre_guardado)
    file.save(ruta)

    conn = get_db_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)
    cursor.execute("""
        INSERT INTO entregables_operacion
        (tipo, periodo, unidad, ruta_archivo, estado, observaciones, fecha_carga,
         fecha_limite, responsable, categoria, documento_analizado)
        VALUES (?, ?, ?, ?, 'cargado', ?, ?, ?, ?, ?, ?)
    """, (
        tipo, periodo, unidad, ruta, observaciones, datetime.now().isoformat(),
        fecha_limite, responsable, categoria, 'pendiente_analisis'
    ))
    conn.commit()
    conn.close()

    log_auditoria(
        request.args.get('usuario', 'sistema'),
        'CARGAR_ENTREGABLE_OPERACION',
        tabla='entregables_operacion',
        datos_nuevos={'tipo': tipo, 'periodo': periodo, 'unidad': unidad, 'archivo': nombre_original, 'fecha_limite': fecha_limite}
    )
    return jsonify({'message': 'Entregable cargado y vinculado a cumplimiento.', 'tipo': tipo, 'periodo': periodo}), 201


@app.route('/api/cumplimiento/evaluar', methods=['POST'])
def evaluar_cumplimiento():
    data = request.get_json(silent=True) or {}
    resultado = evaluar_operacion(
        periodo=data.get('periodo') or periodo_actual(),
        usuario=request.args.get('usuario', 'sistema'),
        guardar=True
    )
    log_auditoria(request.args.get('usuario', 'sistema'), 'EVALUAR_OPERACION_ICBF', datos_nuevos=resultado)
    return jsonify(resultado), 200


@app.route('/api/cumplimiento/matriz', methods=['GET'])
def matriz_estandares():
    resultado = evaluar_operacion(periodo=request.args.get('periodo') or periodo_actual(), guardar=False)
    return jsonify({'matriz': resultado['matriz_estandares']}), 200


@app.route('/api/cumplimiento/incumplimientos', methods=['GET'])
def detectar_incumplimientos():
    resultado = evaluar_operacion(periodo=request.args.get('periodo') or periodo_actual(), guardar=False)
    return jsonify({'incumplimientos': resultado['incumplimientos']}), 200


@app.route('/api/asistente-icbf', methods=['POST'])
def asistente_icbf():
    data = request.get_json(silent=True) or {}
    pregunta = str(data.get('pregunta', '')).strip()
    if not pregunta:
        return jsonify({'error': 'Pregunta requerida.'}), 400

    terminos = [t.lower() for t in pregunta.replace('¿', '').replace('?', '').split() if len(t) > 3]
    conn = get_db_connection()
    documentos = conn.execute("""
        SELECT titulo, tipo, version, texto_indexado
        FROM documentos_institucionales
        WHERE estado = 'vigente' AND texto_indexado IS NOT NULL AND texto_indexado != ''
        ORDER BY fecha_carga DESC
    """).fetchall()
    conn.close()

    mejor = None
    mejor_score = 0
    for doc in documentos:
        texto = doc['texto_indexado'] or ''
        texto_lower = texto.lower()
        score = sum(texto_lower.count(t) for t in terminos)
        if score > mejor_score:
            mejor = doc
            mejor_score = score

    if not mejor:
        return jsonify({
            'respuesta': 'No encontré soporte en los documentos institucionales indexados. Carga el Manual Operativo o lineamientos en texto, CSV o Excel para responder con fuente documental.',
            'fuentes': []
        }), 200

    texto = mejor['texto_indexado']
    posicion = min([texto.lower().find(t) for t in terminos if texto.lower().find(t) >= 0] or [0])
    inicio = max(0, posicion - 350)
    fin = min(len(texto), posicion + 700)
    fragmento = ' '.join(texto[inicio:fin].split())
    return jsonify({
        'respuesta': fragmento,
        'fuentes': [{
            'titulo': mejor['titulo'],
            'tipo': mejor['tipo'],
            'version': mejor['version']
        }]
    }), 200


@app.route('/api/informes/supervision', methods=['POST'])
def generar_informe_supervision():
    data = request.get_json(silent=True) or {}
    resultado = evaluar_operacion(
        periodo=data.get('periodo') or periodo_actual(),
        usuario=request.args.get('usuario', 'sistema'),
        guardar=True
    )
    nombre = generar_informe_supervision_excel(resultado)
    log_auditoria(
        request.args.get('usuario', 'sistema'),
        'GENERAR_INFORME_ICBF',
        datos_nuevos={'archivo': nombre, 'periodo': resultado['periodo']}
    )
    return jsonify({
        'message': 'Informe ICBF generado.',
        'archivo': nombre,
        'url': f"/api/descargar-archivo/{nombre}",
        'cumplimiento_general': resultado['cumplimiento_general']
    }), 200



@app.route('/api/unidades', methods=['GET', 'POST'])
def manejar_unidades():
    """Consulta o crea manualmente unidades para el tablero."""
    # Estas rutas operativas usan el adaptador DB-API PostgreSQL, que conserva
    # correctamente el resultado mientras se aplican los controles multi-tenant.
    conn = database_connection()
    cursor = conn.cursor()
    ensure_runtime_schema(cursor)

    if request.method == 'GET':
        filas = cursor.execute("""
            SELECT nombre, NULL AS direccion, NULL AS telefono,
                   total_ninos AS total_usuarios, 0 AS total_gestantes,
                   fecha_consolidacion AS fecha_actualizacion
            FROM master_unidades
            WHERE activo=1 AND COALESCE(fundacion_id,1)=?
            ORDER BY nombre
        """, (fundacion_actual_id(),)).fetchall()
        conn.close()
        unidades = []
        nombres_incluidos = set()
        for fila in filas:
            nombre = normalize_unidad(fila['nombre'])
            if not nombre:
                continue
            nombres_incluidos.add(nombre)
            item = dict(fila)
            item['nombre'] = nombre
            unidades.append(item)

        # No anexar unidades conocidas con 0 usuarios: el tablero debe mostrar
        # únicamente cobertura real detectada o guardada con actividad.
        unidades = [item for item in unidades if int(item.get('total_usuarios') or 0) > 0]
        unidades.sort(key=lambda item: item['nombre'])
        return jsonify({'unidades': unidades})

    data = request.get_json(silent=True) or {}
    nombre = normalize_unidad(data.get('nombre'))
    if not nombre:
        conn.close()
        return jsonify({'error': 'Nombre de unidad inválido. Escribe un nombre real, no ACTIVO/INACTIVO.'}), 400

    direccion = limpiar_valor(data.get('direccion', ''))
    telefono = limpiar_valor(data.get('telefono', ''))
    ahora = datetime.now().isoformat()

    cursor.execute("""
        INSERT INTO unidades (nombre, direccion, telefono, total_usuarios, total_gestantes, fecha_actualizacion, fundacion_id)
        VALUES (?, ?, ?, 0, 0, ?, ?)
        ON CONFLICT(fundacion_id, nombre) DO UPDATE SET
            direccion = excluded.direccion,
            telefono = excluded.telefono,
            fecha_actualizacion = excluded.fecha_actualizacion
    """, (nombre, direccion, telefono, ahora, fundacion_actual_id()))
    conn.commit()
    conn.close()

    return jsonify({
        'message': f'Unidad {nombre} guardada correctamente.',
        'unidad': {
            'nombre': nombre,
            'direccion': direccion,
            'telefono': telefono,
            'total_usuarios': 0,
            'total_gestantes': 0,
            'fecha_actualizacion': ahora
        }
    }), 200


@app.route('/api/historial', methods=['GET'])
def historial():
    conn = get_db_connection()
    auditorias = conn.execute('SELECT * FROM auditoria ORDER BY COALESCE(fecha, fecha_accion) DESC').fetchall()
    conn.close()
    return jsonify({'auditoria': [dict(row) for row in auditorias]})


@app.route('/api/estadisticas', methods=['GET'])
def estadisticas():
    conn = database_connection()
    if rol_actual() == 'SUPERADMIN':
        filtro = '1=1'
        params = []
    else:
        filtro = 'COALESCE(fundacion_id, 1) = ?'
        params = [fundacion_actual_id()]

    total = conn.execute(f"SELECT COUNT(*) as total FROM beneficiarios WHERE estado != ? AND {filtro}", [EstadoUsuario.FALLECIDO] + params).fetchone()['total']
    gestantes = conn.execute(f"SELECT COUNT(*) as total FROM gestantes WHERE estado = ? AND {filtro}", [EstadoUsuario.ACTIVO] + params).fetchone()['total']
    alertas = conn.execute(f"SELECT COUNT(*) as total FROM alertas WHERE resuelta = 0 AND {filtro}", params).fetchone()['total']
    riesgo = conn.execute(f"""
        SELECT COUNT(*) as total
        FROM peso_talla
        WHERE estado_nutricional IN (?, ?) AND {filtro}
    """, [EstadoNutricion.RIESGO, EstadoNutricion.DESNUTRICION] + params).fetchone()['total']
    vencido = conn.execute(f'SELECT COUNT(*) as total FROM peso_talla WHERE estado = ? AND {filtro}', ['vencido'] + params).fetchone()['total']
    proximo = conn.execute(f'SELECT COUNT(*) as total FROM peso_talla WHERE estado = ? AND {filtro}', ['proximo_vencer'] + params).fetchone()['total']
    al_dia = conn.execute(f'SELECT COUNT(*) as total FROM peso_talla WHERE estado = ? AND {filtro}', ['al_dia'] + params).fetchone()['total']
    conn.close()
    return jsonify({
        'total_activos': total,
        'gestantes': gestantes,
        'alertas_pendientes': alertas,
        'riesgo_nutricional': riesgo,
        'peso_talla_al_dia': al_dia,
        'peso_talla_proximo': proximo,
        'peso_talla_vencido': vencido
    })

def inyectar_datos_en_plantillas(unidad_nombre, lista_usuarios, options=None):
    """
    Actualiza plantillas oficiales sin modificar su estructura.

    No inserta filas, no elimina filas, no renombra hojas, no mueve hojas.
    Solo escribe valores en las celdas existentes detectadas por encabezado.
    """
    options = dict(options or {})
    formatos_solicitados_alpha68 = _alpha68_parse_formatos_seleccionados(options)

    def opt_value(*keys, default=''):
        for key in keys:
            value = limpiar_valor(options.get(key))
            if value:
                return value
        if has_request_context():
            for key in keys:
                value = limpiar_valor(request.form.get(key) or request.args.get(key))
                if value:
                    return value
        return default

    try:
        mes = int(opt_value('mes', default=datetime.now().month))
    except Exception:
        mes = datetime.now().month
    try:
        año = int(opt_value('año', 'anio', default=datetime.now().year))
    except Exception:
        año = datetime.now().year

    mes = max(1, min(12, mes))
    mes_nombre = MESES_ES.get(mes, datetime.now().strftime('%B').upper())
    festivos_alpha68 = _alpha68_festivos_configurados(año, mes)

    try:
        max_usuarios_formato = int(opt_value('max_usuarios_formato', default=20))
    except Exception:
        max_usuarios_formato = 20
    max_usuarios_formato = max(1, min(200, max_usuarios_formato))

    try:
        bienestarina_por_hoja = int(opt_value('bienestarina_por_hoja', default=14))
    except Exception:
        bienestarina_por_hoja = 14
    bienestarina_por_hoja = max(1, min(50, bienestarina_por_hoja))

    fecha_entrega = fecha_entrega_bienestarina_desde_request(mes, año, options=options)
    fecha_dia, fecha_mes, fecha_anio = partes_fecha_ddmmaaaa(fecha_entrega)
    lote_bienestarina = limpiar_valor(opt_value('lote_bienestarina'))
    cantidad_bienestarina = limpiar_valor(opt_value('cantidad_bienestarina'))

    unidad_nombre = normalize_unidad(unidad_nombre) or str(unidad_nombre or '').strip().upper()
    # Se conserva la lista completa y el límite se aplica por formato después
    # de ordenar/filtrar. Así RAM/RAN no pierde gestantes o grupos etarios
    # porque hayan quedado después de los primeros registros de la base.
    usuarios_base = list(lista_usuarios or [])

    # ALPHA53 — Minutas RPP versionadas, RAM automático y encabezados.
    # Importación local para no afectar arranque si el módulo no está disponible en versiones antiguas.
    try:
        from services.rpp_minutas_service import (
            obtener_minuta_vigente, productos_para_usuario, obtener_equivalencias,
            obtener_dias_asistencia_usuario, calcular_verificacion_cobertura_ram,
            normalizar_texto as normalizar_minuta_texto,
        )
    except Exception:
        obtener_minuta_vigente = None
        productos_para_usuario = None
        obtener_equivalencias = None
        obtener_dias_asistencia_usuario = None
        calcular_verificacion_cobertura_ram = None
        normalizar_minuta_texto = normalizar_texto_clave

    # ALPHA54 — Motor RPP completo: auditoría, validación de contexto y no-vacío.
    try:
        from services.rpp_motor_completo_service import (
            log_rpp_event, validate_rpp_context, product_items_from_entry,
        )
    except Exception:
        def log_rpp_event(evento, **detalle):
            try:
                print(f'RPP_ALPHA54 {evento}: {detalle}')
            except Exception:
                pass
        def validate_rpp_context(entry, usuarios, minuta, categoria=None):
            errores = []
            if not usuarios:
                errores.append('No hay usuarios para la UDS/grupo seleccionado.')
            if minuta is None:
                errores.append('No existe minuta RPP vigente para el mes seleccionado.')
            return (not errores), errores, {'usuarios': len(usuarios or []), 'categoria_rpp': categoria}
        def product_items_from_entry(entry):
            return []

    try:
        minuta_rpp_vigente = obtener_minuta_vigente(DATABASE_PATH, mes=mes, anio=año) if obtener_minuta_vigente else None
    except Exception as exc:
        print(f'No se pudo consultar minuta RPP vigente: {exc}')
        minuta_rpp_vigente = None

    try:
        equivalencias_minuta = obtener_equivalencias(DATABASE_PATH) if obtener_equivalencias else {}
    except Exception:
        equivalencias_minuta = {}

    def orden_grupo_etario_operativo(user):
        """Orden oficial para RAM/RAN: gestantes, 0-6, 6-11, 1-2 y 3-5.

        El usuario pidió que el formato quede agrupado por edades. Se mantiene
        a niños/as de 0 a 6 meses junto al primer bloque operativo para no
        excluirlos cuando existan en la base.
        """
        tipo = normalizar_texto_clave(user.get('TipoBeneficiario') or user.get('tipo_beneficiario') or '')
        grupo = normalizar_texto_clave(user.get('GrupoEdad') or user.get('grupo_edad') or '')
        try:
            edad = int(float(user.get('EdadMeses') or user.get('edad_meses') or 0))
        except Exception:
            edad = 0

        if 'gestante' in tipo or 'gestante' in grupo:
            grupo_orden = 0
        elif '0 a 6' in grupo or '0 a 5' in grupo or 'menor de seis' in tipo or edad <= 5:
            grupo_orden = 1
        elif '6 a 11' in grupo or 6 <= edad <= 11:
            grupo_orden = 2
        elif '1 a 2' in grupo or 12 <= edad <= 35:
            grupo_orden = 3
        elif '3 a 5' in grupo or 36 <= edad <= 71:
            grupo_orden = 4
        else:
            grupo_orden = 9

        nombre = normalizar_texto_clave(
            unir_partes(
                user.get('PrimerApellido'), user.get('SegundoApellido'),
                user.get('PrimerNombre'), user.get('SegundoNombre')
            ) or user.get('Nombre') or user.get('nombre') or ''
        )
        documento = normalizar_texto_clave(user.get('Documento') or user.get('documento') or user.get('NUI') or '')
        return (grupo_orden, nombre, documento)

    def ordenar_usuarios_para_formato(usuarios, formato_norm, categoria_rpp=None):
        usuarios = list(usuarios or [])
        if any(k in formato_norm for k in ['asistencia', 'ram', 'rram', 'ran', 'run']):
            return sorted(usuarios, key=orden_grupo_etario_operativo)
        return usuarios

    fill_verde = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
    contexto_plantilla_actual = {'oficial': False, 'preservar_estilos': False, 'preservar_impresion': False}
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center', wrap_text=True)

    def celda_es_combinada(ws, row, col):
        coord = ws.cell(row=row, column=col).coordinate
        for merged_range in ws.merged_cells.ranges:
            if coord in merged_range:
                return ws.cell(merged_range.min_row, merged_range.min_col)
        return ws.cell(row=row, column=col)

    def set_cell(ws, row, col, value, center=True, fill=None):
        cell = celda_es_combinada(ws, row, col)
        try:
            # En plantillas oficiales entregadas por el usuario, regla crítica:
            # modificar únicamente el valor. No tocar bordes, colores, rellenos,
            # combinaciones, anchos, altos ni fórmulas. Alpha55 solo ajusta
            # texto largo con wrap/shrink-to-fit cuando el valor escrito es texto.
            cell.value = value

            def _ajustar_texto_largo_seguro(celda, texto):
                txt = str(texto or '')
                if not txt or len(txt) < 22 or str(txt).startswith('='):
                    return
                try:
                    celda.alignment = copy.copy(celda.alignment)
                    celda.alignment = Alignment(
                        horizontal=celda.alignment.horizontal or ('center' if center else 'left'),
                        vertical=celda.alignment.vertical or 'center',
                        text_rotation=celda.alignment.text_rotation,
                        wrap_text=True,
                        shrink_to_fit=True,
                        indent=celda.alignment.indent,
                    )
                except Exception:
                    pass
                try:
                    # Conserva familia/color/negrilla originales; solo reduce tamaño si el texto
                    # se sale de la línea. No baja de 8 puntos.
                    fuente = copy.copy(celda.font)
                    size = fuente.sz or 11
                    if len(txt) > 60:
                        size = max(8, min(size, 8))
                    elif len(txt) > 42:
                        size = max(8, min(size, 9))
                    elif len(txt) > 28:
                        size = max(8, min(size, 10))
                    fuente.sz = size
                    celda.font = fuente
                except Exception:
                    pass

            if contexto_plantilla_actual.get('preservar_estilos'):
                _ajustar_texto_largo_seguro(cell, value)
                return cell
            cell.alignment = align_center if center else align_left
            _ajustar_texto_largo_seguro(cell, value)
            if fill is not None:
                cell.fill = fill
            return cell
        except Exception:
            return cell

    def texto_celda(cell):
        return normalizar_texto_clave(cell.value)

    def sheet_score(ws):
        titulo_raw = str(ws.title or '').strip()
        titulo_norm = normalizar_texto_clave(titulo_raw)
        titulo_unidad = normalize_unidad(titulo_raw)
        unidad_equivalentes = equivalentes_unidad(unidad_nombre)
        titulo_equivalentes = equivalentes_unidad(titulo_unidad or titulo_raw)

        score = 0
        if titulo_equivalentes & unidad_equivalentes:
            score += 100
        if titulo_norm in unidad_equivalentes:
            score += 100
        for eq in unidad_equivalentes:
            if eq and (eq in titulo_norm or titulo_norm in eq):
                score += 35
        unidad_clave = normalizar_texto_clave(unidad_nombre)
        if unidad_clave and titulo_norm:
            score += int(SequenceMatcher(None, unidad_clave, titulo_norm).ratio() * 30)
        return score

    def seleccionar_hojas(wb):
        scored = [(sheet_score(ws), idx, ws) for idx, ws in enumerate(wb.worksheets)]
        scored.sort(key=lambda item: item[0], reverse=True)
        if not scored:
            return []
        mejor = scored[0][0]
        if mejor <= 0:
            return [wb.active]
        return [ws for score, _, ws in scored if score >= max(30, mejor - 10)] or [scored[0][2]]

    def buscar_celda_destino_derecha(ws, row, col, max_offset=10):
        for c in range(col + 1, min(ws.max_column, col + max_offset) + 1):
            txt = normalizar_texto_clave(ws.cell(row=row, column=c).value)
            if txt and any(k in txt for k in ['nombre', 'documento', 'primer', 'segundo', 'apellido', 'fecha', 'lote', 'cantidad', 'tipo']):
                continue
            return c
        return min(col + 1, ws.max_column)

    def valor_metadata(campo):
        for user in usuarios_base:
            valor = limpiar_valor(user.get(campo))
            if valor:
                return valor
        return ''

    def actualizar_encabezados(ws):
        """Diligencia encabezados oficiales desde la fuente maestra de datos.

        Regla Alpha16: los encabezados se actualizan según la UDS real y el
        agente educativo/docente asociado. No se copian datos obsoletos de la
        plantilla y no se cambian estilos, colores, bordes, combinaciones ni
        configuración de impresión.
        """

        def valor_metadata_any(*campos):
            for campo in campos:
                valor = valor_metadata(campo)
                if valor:
                    return valor
            return ''

        def obtener_unidad_db():
            try:
                conn = get_db_connection()
                cursor = conn.cursor()
                ensure_runtime_schema(cursor)
                filas = cursor.execute("SELECT * FROM unidades").fetchall()
                conn.close()
                for fila in filas:
                    datos = dict(fila)
                    if normalize_unidad(datos.get('nombre') or '') == unidad_nombre:
                        return datos
                    if normalizar_texto_clave(datos.get('nombre') or '') in equivalentes_unidad(unidad_nombre):
                        return datos
            except Exception:
                pass
            return {}

        def obtener_suplente():
            for item in obtener_talentos_por_unidad(unidad_nombre):
                cargo = normalizar_texto_clave(item.get('cargo') or item.get('tipo_equipo') or '')
                if any(k in cargo for k in ['suplente', 'apoyo', 'auxiliar']):
                    return item
            return {}

        unidad_db = obtener_unidad_db()
        docente = obtener_talento_por_unidad(unidad_nombre) or {}
        suplente = obtener_suplente()

        valores = {
            'regional': valor_metadata_any('Regional', 'regional') or 'CHOCÓ',
            'centro_zonal': valor_metadata_any('CentroZonal', 'Centro Zonal', 'centro_zonal') or 'CZ Ciudad de prueba',
            'municipio': valor_metadata_any('Municipio', 'municipio') or 'Ciudad de prueba',
            'modalidad': valor_metadata_any('Modalidad', 'modalidad'),
            'servicio_atencion': valor_metadata_any('ServicioAtencion', 'Servicio de Atención', 'Modalidad', 'modalidad'),
            'contrato': valor_metadata_any('NumeroContrato', 'Número de Contrato', 'numero_contrato') or limpiar_valor(docente.get('contrato')) or limpiar_valor(unidad_db.get('contrato')),
            'vigencia': valor_metadata_any('Vigencia', 'vigencia') or str(año),
            'eas': valor_metadata_any('NombreEAS', 'Nombre EAS', 'EntidadAdministradora', 'nombre_eas') or 'FUNDACIÓN PACÍFICO VIVE',
            'unidad': valor_metadata_any('Unidad', 'unidad', 'unidad_servicio', 'nombre_unidad') or limpiar_valor(unidad_db.get('nombre')) or unidad_nombre,
            'unidad_origen': valor_metadata_any('NombreUnidadOrigen', 'Nombre Punto Entrega Origen') or unidad_nombre,
            'codigo_uds': valor_metadata_any('CodigoUnidadServicio', 'Código UDS', 'Codigo UDS', 'codigo_unidad_servicio', 'codigo_unidad') or limpiar_valor(unidad_db.get('codigo_unidad_servicio')),
            'codigo_origen': valor_metadata_any('CodigoUnidadOrigen', 'Código Punto Entrega Origen') or valor_metadata_any('CodigoUnidadServicio', 'codigo_unidad_servicio'),
            'mes': mes_nombre,
            'año': str(año),
            'anio': str(año),
            'fecha_entrega': fecha_entrega,
            'lote': lote_bienestarina,
            'cantidad': cantidad_bienestarina,
            'direccion_unidad': limpiar_valor(unidad_db.get('direccion')) or valor_metadata_any('DireccionUnidad', 'Dirección Unidad', 'direccion_unidad') or limpiar_valor(docente.get('direccion')),
            'telefono_unidad': limpiar_documento_talento(unidad_db.get('telefono')) or valor_metadata_any('TelefonoUnidad', 'Teléfono UDS', 'Telefono UDS') or limpiar_documento_talento(docente.get('telefono')),
            'barrio': valor_metadata_any('Barrio', 'barrio'),
            'docente': limpiar_valor(docente.get('nombre') or unir_partes(docente.get('nombres'), docente.get('apellidos'))).upper(),
            'cedula_docente': limpiar_documento_talento(docente.get('documento')),
            'telefono_docente': limpiar_documento_talento(docente.get('telefono')),
            'suplente': limpiar_valor(suplente.get('nombre') or unir_partes(suplente.get('nombres'), suplente.get('apellidos'))).upper(),
            'telefono_suplente': limpiar_documento_talento(suplente.get('telefono')),
            'coordinador': limpiar_valor(docente.get('coordinador') or unidad_db.get('coordinador_nombre')),
        }

        # Respaldo: si no hay teléfono institucional, usar el del agente educativo.
        if not valores.get('telefono_unidad'):
            valores['telefono_unidad'] = valores.get('telefono_docente')

        etiquetas_rotulo = [
            'regional', 'centro zonal', 'municipio', 'modalidad', 'mes', 'ano', 'año',
            'numero de contrato', 'número de contrato', 'contrato', 'agente educativo',
            'docente', 'responsable', 'suplente', 'unidad de servicio', 'unidad de atencion',
            'unidad de atención', 'punto de entrega', 'direccion', 'dirección', 'telefono',
            'teléfono', 'codigo', 'código', 'entidad administradora', 'servicio de atencion',
            'servicio de atención', 'lugar', 'barrio', 'eas'
        ]

        def es_rotulo(texto):
            t = normalizar_texto_clave(texto)
            return any(e in t for e in etiquetas_rotulo)

        def buscar_destino_encabezado(row, col, max_offset=14):
            # Busca hacia la derecha una celda de valor. Puede estar llena con un
            # dato viejo de la plantilla; si no parece otro rótulo, se reemplaza.
            for c in range(col + 1, min(ws.max_column, col + max_offset) + 1):
                target = celda_es_combinada(ws, row, c)
                raw = str(target.value or '').strip()
                txt = normalizar_texto_clave(raw)
                if raw and es_rotulo(txt):
                    continue
                return target
            # Algunas planillas ubican el valor debajo del rótulo.
            for r in range(row + 1, min(ws.max_row, row + 3) + 1):
                target = celda_es_combinada(ws, r, col)
                raw = str(target.value or '').strip()
                txt = normalizar_texto_clave(raw)
                if raw and es_rotulo(txt):
                    continue
                return target
            return None

        def alias_encabezado_coincide(txt, alias):
            if not txt or not alias:
                return False
            if alias == 'servicio':
                return txt == 'servicio'
            if len(alias) <= 3:
                # Evita falsos positivos con rótulos cortos como CC dentro de
                # palabras como ACCIÓN. Debe aparecer como token independiente.
                return alias in txt.split()
            return alias in txt

        def escribir_por_etiquetas(aliases, valor, max_row=18, max_offset=14):
            if valor in (None, ''):
                return
            aliases_norm = [normalizar_texto_clave(a) for a in aliases if a]
            for row in range(1, min(ws.max_row, max_row) + 1):
                for col in range(1, min(ws.max_column, 60) + 1):
                    cell = celda_es_combinada(ws, row, col)
                    raw = str(cell.value or '').strip()
                    txt = normalizar_texto_clave(raw)
                    if not txt or not any(alias_encabezado_coincide(txt, alias) for alias in aliases_norm):
                        continue
                    # Evitar tablas de usuarios: teléfono/celular/documentos de participantes no son encabezado.
                    fila_txt = ' '.join(normalizar_texto_clave(ws.cell(row=row, column=c).value) for c in range(1, min(ws.max_column, 45) + 1))
                    if any(k in fila_txt for k in ['primer nombre', 'segundo nombre', 'primer apellido', 'control diario', 'documento de identidad']) and row > 8:
                        continue
                    if ':' in raw and len(raw.split(':', 1)[0].strip()) <= 70:
                        prefijo = raw.split(':', 1)[0].strip()
                        cell.value = f'{prefijo}: {valor}'
                    else:
                        destino = buscar_destino_encabezado(row, col, max_offset=max_offset)
                        if destino is not None:
                            destino.value = valor

        reglas_encabezado = [
            (['entidad administradora del servicio', 'entidad administradora', 'nombre de la eas', 'eas'], valores['eas']),
            (['regional'], valores['regional']),
            (['centro zonal'], valores['centro_zonal']),
            (['municipio'], valores['municipio']),
            (['modalidad de atencion', 'modalidad de atención', 'modalidad'], valores['modalidad']),
            (['servicio de atencion', 'servicio de atención', 'servicio'], valores['servicio_atencion']),
            (['numero de contrato', 'número de contrato', 'contrato'], valores['contrato']),
            (['mes de la entrega', 'mes de entrega', 'mes de consumo', 'mes'], valores['mes']),
            (['ano', 'año', 'anio', 'vigencia'], valores['año']),
            (['codigo del punto de entrega o uds', 'código del punto de entrega o uds', 'codigo cuentame de la uds', 'código cuéntame de la uds', 'codigo cuentame uds', 'código cuéntame uds', 'codigo uds', 'código uds'], valores['codigo_uds']),
            (['codigo punto de entrega de origen', 'código punto de entrega de origen'], valores['codigo_origen'] or valores['codigo_uds']),
            (['nombre punto de entrega de origen'], valores['unidad_origen']),
            (['nombre punto de entrega o uds', 'nombre de la unidad de servicio', 'nombre unidad de servicio', 'nombre unidad de atencion', 'nombre unidad de atención', 'unidad de atencion', 'unidad de atención'], valores['unidad']),
            (['responsable punto de entrega o uds', 'responsable punto de entrega', 'nombre agente educativo', 'agente educativo', 'docente'], valores['docente']),
            (['suplente punto de entrega', 'suplente'], valores['suplente']),
            (['direccion punto de entrega o uds', 'dirección punto de entrega o uds', 'direccion uds', 'dirección uds', 'direccion unidad', 'dirección unidad'], valores['direccion_unidad']),
            (['barrio'], valores['barrio']),
            (['telefono uds', 'teléfono uds', 'telefono punto de entrega', 'teléfono punto de entrega', 'telefono', 'teléfono'], valores['telefono_unidad']),
            (['cc', 'cedula docente', 'cédula docente', 'documento docente'], valores['cedula_docente']),
            (['telefono docente', 'teléfono docente', 'telefono agente', 'teléfono agente'], valores['telefono_docente']),
            (['fecha de entrega'], valores['fecha_entrega']),
            (['lote'], valores['lote']),
            (['cantidad'], valores['cantidad']),
        ]
        for aliases, valor in reglas_encabezado:
            escribir_por_etiquetas(aliases, valor)

        # Distribución permanente del encabezado de Bienestarina. Se aplica al
        # final para que ningún escritor genérico vuelva a encerrar los valores
        # largos en una sola celda angosta.
        if 'bienestarina' in normalizar_texto_clave(ws.title) or any(
            'bienestarina' in normalizar_texto_clave(ws.cell(r, c).value)
            for r in range(1, min(ws.max_row, 10) + 1)
            for c in range(1, min(ws.max_column, 20) + 1)
        ):
            _alpha75_aplicar_encabezado_bienestarina(ws, valores)

        # Reemplazo seguro de valores obsoletos en encabezado: corrige casos como
        # formato reutilizado que aún traía datos de otra UDS.
        unidades_actuales = {normalizar_texto_clave(unidad_nombre), *equivalentes_unidad(unidad_nombre)}
        unidades_conocidas = {normalizar_texto_clave(u) for u in list(KNOWN_UNITS) + list(ALIAS_UNIDADES_CUENTAME.keys()) + list(ALIAS_UNIDADES_CUENTAME.values())}
        unidades_conocidas = {u for u in unidades_conocidas if u}
        docente_actual_norm = normalizar_texto_clave(valores.get('docente'))
        nombres_talento = set()
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            ensure_runtime_schema(cursor)
            for tabla in ['coordinadores', 'th_personas']:
                try:
                    for fila in cursor.execute(f"SELECT nombre, nombres, apellidos FROM {tabla}").fetchall():
                        d = dict(fila)
                        nombres_talento.add(normalizar_texto_clave(d.get('nombre') or unir_partes(d.get('nombres'), d.get('apellidos'))))
                except Exception:
                    continue
            conn.close()
        except Exception:
            pass

        for row in range(1, min(ws.max_row, 18) + 1):
            fila_txt = ' '.join(normalizar_texto_clave(ws.cell(row=row, column=c).value) for c in range(1, min(ws.max_column, 60) + 1))
            for col in range(1, min(ws.max_column, 60) + 1):
                cell = celda_es_combinada(ws, row, col)
                raw = str(cell.value or '').strip()
                txt = normalizar_texto_clave(raw)
                if not txt:
                    continue
                if txt in unidades_conocidas and txt not in unidades_actuales and valores.get('unidad'):
                    cell.value = valores['unidad']
                    continue
                if docente_actual_norm and txt in nombres_talento and txt != docente_actual_norm:
                    if any(k in fila_txt for k in ['responsable', 'agente educativo', 'docente', 'suplente']):
                        if 'suplente' in fila_txt and valores.get('suplente'):
                            cell.value = valores['suplente']
                        else:
                            cell.value = valores['docente']

        # En algunos formatos el mes aparece escrito dentro de un título.
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 8)):
            for cell in row:
                if isinstance(cell.value, str) and 'CRONOGRAMA DE ENTREGA MES' in cell.value.upper():
                    cell.value = f'CRONOGRAMA DE ENTREGA MES DE {mes_nombre}'

        limpiar_duplicidad_encabezados(ws)

    def limpiar_duplicidad_encabezados(ws):
        """Evita que el encabezado quede con el mismo dato dos veces.

        Algunas plantillas traen rótulo y valor en celdas separadas. Si el motor
        actualiza el rótulo como "REGIONAL: CHOCÓ" y deja el valor "CHOCÓ"
        al lado, Excel muestra duplicidad. Esta limpieza borra solo duplicados
        cercanos; no toca estructura, colores ni rótulos oficiales.
        """
        etiquetas = [
            'regional', 'centro zonal', 'municipio', 'modalidad', 'mes', 'ano', 'año',
            'numero de contrato', 'número de contrato', 'nombre agente educativo',
            'nombre unidad de servicio', 'unidad de atencion', 'unidad de atención',
            'direccion uds', 'dirección uds', 'codigo cuentame uds', 'código cuentame uds',
            'entidad administradora', 'servicio de atencion', 'servicio de atención'
        ]

        def limpiar_parte_valor(raw):
            valor = str(raw or '').split(':', 1)[-1]
            valor = re.sub(r'[_]+', ' ', valor)
            return normalizar_texto_clave(valor)

        max_col = min(ws.max_column, 45)
        for row in range(1, min(ws.max_row, 25) + 1):
            for col in range(1, max_col + 1):
                cell = celda_es_combinada(ws, row, col)
                raw = str(cell.value or '').strip()
                txt = normalizar_texto_clave(raw)
                if ':' not in raw or not any(e in txt for e in etiquetas):
                    continue
                valor_norm = limpiar_parte_valor(raw)
                if not valor_norm or len(valor_norm) < 2:
                    continue
                for col2 in range(col + 1, min(max_col, col + 8) + 1):
                    cell2 = celda_es_combinada(ws, row, col2)
                    raw2 = str(cell2.value or '').strip()
                    txt2 = normalizar_texto_clave(re.sub(r'[_]+', ' ', raw2))
                    if not txt2 or any(e in txt2 for e in etiquetas):
                        continue
                    if txt2 == valor_norm or (len(txt2) > 3 and (txt2 in valor_norm or valor_norm in txt2)):
                        set_cell(ws, row, col2, '', center=False)

    def mapear_columnas(ws, header_row, nombre_formato):
        """Mapea columnas usando únicamente el bloque de encabezados de la tabla.

        Corrección clave:
        - No toma rótulos superiores como TELÉFONO de la UDS para campos del participante.
        - En RPP, la columna NÚMERO DE CELULAR se detecta donde corresponde y no en kilos/libras.
        - Detecta NUI / Nº DOC. IDENT. / documento del acudiente en RPP, Bienestarina y otros formatos.
        """
        mapa = {}
        formato_norm = normalizar_texto_clave(nombre_formato)
        is_bienestarina = 'bienestarina' in formato_norm
        is_asistencia = any(k in formato_norm for k in ['asistencia', 'ram', 'rram', 'ran', 'run'])
        is_rpp = 'rpp' in formato_norm

        # Solo el bloque de la tabla. Filas superiores se usan para metadatos, no para columnas de usuario.
        # Los rótulos institucionales situados antes del encabezado no son
        # columnas del participante. En la segunda página de Bienestarina, por
        # ejemplo, "Cédula del Responsable" está cuatro filas antes y hacía que
        # la columna A se detectara simultáneamente como documento y consecutivo.
        rows_scan = range(max(1, header_row), min(ws.max_row, header_row + 8) + 1)
        control_cols = set()
        total_mensual_col = None

        for row in rows_scan:
            for col in range(1, ws.max_column + 1):
                txt = normalizar_texto_clave(ws.cell(row=row, column=col).value)
                if not txt:
                    continue

                # Campos del participante
                if txt in {'no', 'n', 'nro', 'nº'} or txt.startswith('no de orden') or txt == 'no de orden':
                    mapa.setdefault('orden', col)
                    continue

                if 'parentesco' in txt:
                    mapa.setdefault('parentesco', col)
                    continue

                # Documento del acudiente/responsable debe evaluarse antes que documento del niño.
                if (
                    ('documento' in txt or 'cedula' in txt or 'cédula' in txt or 'doc ident' in txt or 'doc identidad' in txt)
                    and ('acudiente' in txt or 'responsable' in txt)
                ):
                    mapa.setdefault('documento_acudiente', col)
                    continue

                if 'firma' in txt or 'huella' in txt:
                    # Las columnas/celdas de firma o huella son para diligenciar a mano.
                    # Nunca deben recibir nombres, cédulas ni datos del acudiente.
                    continue

                if 'nombre completo y cedula' in txt or 'nombre completo y cédula' in txt or 'quien recibe' in txt:
                    mapa.setdefault('acudiente_nombre_cedula', col)
                    continue

                if 'acudiente' in txt and 'documento' not in txt and 'cedula' not in txt and 'cédula' not in txt:
                    mapa.setdefault('acudiente_completo', col)
                    continue

                # Teléfono/celular del acudiente. Se fuerza a la columna donde dice NÚMERO DE CELULAR.
                if row >= header_row and ('numero de celular' in txt or 'número de celular' in txt or txt == 'celular' or txt.endswith(' celular') or 'telefono celular' in txt):
                    mapa['telefono'] = col
                    continue
                # Evitar que un rótulo genérico "TELÉFONO" de encabezado institucional se confunda con kilos/libras.
                if row >= header_row and ('telefono' in txt or 'teléfono' in txt) and not is_rpp:
                    mapa.setdefault('telefono', col)
                    continue

                if 'tipo documento' in txt or 'tipo doc' in txt or txt in {'td', 'tipo de documento'}:
                    mapa.setdefault('tipo_documento', col)
                    continue

                # Documento/NUI del niño o participante.
                if (
                    'nui' in txt or 'nuip' in txt or
                    'doc ident' in txt or 'doc identidad' in txt or
                    'documento de identidad' in txt or 'numero documento' in txt or
                    'n documento' in txt or 'no documento' in txt or
                    txt in {'documento', 'cedula', 'cédula', 'no doc ident', 'n doc ident'}
                ) and 'acudiente' not in txt and 'responsable' not in txt and 'docente' not in txt:
                    mapa.setdefault('documento', col)
                    continue

                if 'primer nombre' in txt:
                    mapa.setdefault('primer_nombre', col)
                    continue
                if 'segundo nombre' in txt:
                    mapa.setdefault('segundo_nombre', col)
                    continue
                if 'primer apellido' in txt:
                    mapa.setdefault('primer_apellido', col)
                    continue
                if 'segundo apellido' in txt:
                    mapa.setdefault('segundo_apellido', col)
                    continue
                if 'nombres y apellidos' in txt or 'nombre del participante' in txt or 'nombres apellidos del participante' in txt:
                    mapa.setdefault('nombre', col)
                    continue

                # Edad: no tocar los encabezados "AÑOS" y "MESES"; solo se llenan datos en filas.
                if txt in {'ano', 'anos', 'año', 'años'}:
                    mapa.setdefault('edad_anos', col)
                    continue
                if txt == 'meses' or txt == 'mes':
                    if is_asistencia:
                        mapa.setdefault('edad_meses_resto', col)
                    elif is_bienestarina:
                        mapa.setdefault('fecha_mes', col)
                    continue
                if 'edad' in txt and 'mes' not in txt and 'ano' not in txt and 'año' not in txt:
                    mapa.setdefault('edad_completa', col)
                    continue

                # Datos de Salud y Nutrición Inteligente para plantillas oficiales que los soliciten.
                if 'perimetro braquial' in txt or 'perímetro braquial' in txt or txt == 'pb':
                    mapa.setdefault('perimetro_braquial', col)
                    continue
                if 'perimetro cefalico' in txt or 'perímetro cefálico' in txt:
                    mapa.setdefault('perimetro_cefalico', col)
                    continue
                if ('peso' in txt and 'kilo' not in txt and 'kilogramo' not in txt and
                        'gramo' not in txt and 'libra' not in txt and 'racion' not in txt and 'ración' not in txt):
                    mapa.setdefault('peso_kg', col)
                    continue
                if 'talla' in txt or 'estatura' in txt or 'longitud' in txt:
                    mapa.setdefault('talla_cm', col)
                    continue
                if txt == 'imc' or 'indice de masa corporal' in txt or 'índice de masa corporal' in txt:
                    mapa.setdefault('imc', col)
                    continue
                if 'diagnostico nutricional' in txt or 'diagnóstico nutricional' in txt or 'estado nutricional' in txt:
                    mapa.setdefault('diagnostico_nutricional', col)
                    continue
                if 'fecha valoracion' in txt or 'fecha valoración' in txt or 'fecha de control' in txt:
                    mapa.setdefault('fecha_valoracion_nutricional', col)
                    continue
                if 'proximo control' in txt or 'próximo control' in txt:
                    mapa.setdefault('proximo_control_nutricional', col)
                    continue
                if 'estado control' in txt or 'control nutricional' in txt:
                    mapa.setdefault('estado_control_nutricional', col)
                    continue

                # Bienestarina: fecha/lote/cantidad en columnas ya existentes.
                if is_bienestarina:
                    if txt == 'dia' or txt == 'día':
                        mapa.setdefault('fecha_dia', col)
                        continue
                    if txt in {'anio', 'año', 'ano'}:
                        mapa.setdefault('fecha_anio', col)
                        continue
                    if 'fecha entrega' in txt or 'fecha de entrega' in txt:
                        mapa.setdefault('fecha_entrega', col)
                        continue
                    if txt == 'lote' or ' lote ' in f' {txt} ' or 'n lote' in txt or 'no lote' in txt:
                        mapa.setdefault('lote_bienestarina', col)
                        continue
                    if txt in {'und', 'unidad', 'unidades'} or 'cantidad' in txt:
                        mapa.setdefault('cantidad_bienestarina', col)
                        continue

                # RAM/RAN/asistencia.
                if 'total mensual' in txt:
                    total_mensual_col = col
                    continue
                if 'asistencias' in txt and is_asistencia:
                    mapa.setdefault('total_asistencias', col)
                    continue
                if 'inasistencias' in txt and is_asistencia:
                    mapa.setdefault('inasistencias', col)
                    continue
                if 'casos de retiro' in txt or 'casos retiro' in txt:
                    mapa.setdefault('casos_retiro', col)
                    continue
                if 'control diario' in txt:
                    control_cols.add(col)
                    continue
                if is_asistencia and (
                    txt in {'lunes', 'martes', 'miercoles', 'miércoles', 'jueves', 'viernes', 'l', 'm', 'x', 'j', 'v'}
                    or re.fullmatch(r'\d{1,2}', txt)
                ):
                    if col not in {mapa.get('orden'), mapa.get('documento')}:
                        control_cols.add(col)
                    continue

                # RPP grupos poblacionales.
                if ('0 6 meses' in txt or '0 a 6 meses' in txt or 'gestante' in txt) and 'grupo poblacional' not in txt:
                    mapa.setdefault('grupo_0_6_gestante', col)
                    continue
                if '6 a 11' in txt:
                    mapa.setdefault('grupo_6_11', col)
                    continue
                if '1 a 2' in txt:
                    mapa.setdefault('grupo_1_2', col)
                    continue
                if '3 a 5' in txt:
                    mapa.setdefault('grupo_3_5', col)
                    continue

        # Refuerzo: algunos formatos tienen encabezados verticales o divididos en varias filas.
        # Se acumula el texto por columna para ubicar documento/NUI, documento acudiente,
        # celular y datos de bienestarina sin confundirlos con columnas de alimentos.
        textos_por_columna = {}
        for col in range(1, ws.max_column + 1):
            textos_por_columna[col] = ' '.join(
                normalizar_texto_clave(ws.cell(row=r, column=col).value)
                for r in rows_scan
                if normalizar_texto_clave(ws.cell(row=r, column=col).value)
            )

        def columna_es_alimento(txt_col):
            return any(k in txt_col for k in [
                'kilo', 'kilogramo', 'libra', 'gramo', 'racion', 'ración',
                'harina', 'arroz', 'frijol', 'lenteja', 'aceite', 'azucar', 'azúcar',
                'bienestarina liquida', 'bienestarina en polvo', 'alimento'
            ])

        for col, txt_col in textos_por_columna.items():
            if not txt_col:
                continue

            if 'firma' in txt_col or 'huella' in txt_col:
                # Firma/huella queda siempre en blanco para diligenciamiento físico.
                continue

            if 'acudiente' in txt_col or 'responsable' in txt_col:
                if not mapa.get('documento_acudiente') and any(k in txt_col for k in ['documento', 'cedula', 'cedula de identidad', 'doc identidad', 'doc ident']):
                    mapa['documento_acudiente'] = col
                    continue
                if not mapa.get('acudiente_nombre_cedula') and any(k in txt_col for k in ['nombre completo y cedula', 'nombre y cedula', 'quien recibe', 'recibe']):
                    mapa['acudiente_nombre_cedula'] = col
                    continue

            if not mapa.get('documento') and not ('acudiente' in txt_col or 'responsable' in txt_col) and any(k in txt_col for k in ['nui', 'nuip', 'documento de identidad', 'numero documento', 'n documento', 'no documento', 'doc ident']):
                mapa['documento'] = col
                continue

            if not mapa.get('tipo_documento') and any(k in txt_col for k in ['tipo documento', 'tipo doc']):
                mapa['tipo_documento'] = col
                continue

            if any(k in txt_col for k in ['numero de celular', 'número de celular', 'celular', 'telefono celular', 'teléfono celular']):
                if not columna_es_alimento(txt_col):
                    mapa['telefono'] = col
                continue

            if is_bienestarina:
                if not mapa.get('fecha_entrega') and 'fecha' in txt_col and 'entrega' in txt_col:
                    mapa['fecha_entrega'] = col
                if not mapa.get('lote_bienestarina') and ('lote' in txt_col or 'n lote' in txt_col or 'no lote' in txt_col):
                    mapa['lote_bienestarina'] = col
                if not mapa.get('cantidad_bienestarina') and ('cantidad' in txt_col or txt_col.strip() in {'und', 'unidad', 'unidades'}):
                    mapa['cantidad_bienestarina'] = col

        if is_bienestarina:
            # Regla oficial: la numeración de Bienestarina siempre es consecutiva
            # y nunca debe confundirse con el documento/NUI del beneficiario.
            orden_col = None
            limite_orden = int(mapa.get('documento') or ws.max_column)
            for col in range(1, max(2, limite_orden)):
                for r in rows_scan:
                    t = normalizar_texto_clave(ws.cell(row=r, column=col).value)
                    if t in {'no', 'n', 'nro', 'numero'}:
                        orden_col = col
                        break
                if orden_col:
                    break
            if orden_col:
                mapa['orden'] = orden_col

            if not mapa.get('documento') or mapa.get('documento') == mapa.get('orden'):
                for col, txt_col in textos_por_columna.items():
                    if col == mapa.get('orden'):
                        continue
                    if any(k in txt_col for k in ['no doc ident', 'n doc ident', 'doc ident', 'doc identidad', 'numero documento', 'documento identidad', 'documento de identidad', 'nui', 'nuip']):
                        mapa['documento'] = col
                        break

        if is_rpp:
            # Columnas de minuta patrón/alimentos. Se usan para limpiar valores
            # que queden por fuera del número real de usuarios, sin tocar colores.
            limite_grupos = max([c for c in [
                mapa.get('grupo_0_6_gestante'), mapa.get('grupo_6_11'),
                mapa.get('grupo_1_2'), mapa.get('grupo_3_5')
            ] if c] or [0])
            minuta_cols = []
            for col, txt_col in textos_por_columna.items():
                if col <= limite_grupos:
                    continue
                if any(k in txt_col for k in [
                    'recibe aavn', 'acudiente', 'parentesco', 'documento',
                    'celular', 'telefono', 'firma', 'huella'
                ]):
                    continue
                if columna_es_alimento(txt_col) or any(k in txt_col for k in [
                    'kilos libras', 'kilos libras', 'libra kilos', 'gramos',
                    'no unidades', 'unidades', 'ml litro', 'mililitro', 'litro'
                ]):
                    minuta_cols.append(col)
            if minuta_cols:
                mapa['minuta_cols'] = sorted(set(minuta_cols))

        if is_rpp and mapa.get('telefono'):
            # En RPP el celular debe estar en la zona final de acudiente/contacto.
            # Si la columna detectada queda dentro del bloque de alimentos, se elimina para no escribir teléfonos en kilos/gramos.
            col_tel = mapa.get('telefono')
            txt_tel = textos_por_columna.get(col_tel, '')
            col_contacto_min = min([c for c in [
                mapa.get('acudiente_nombre_cedula'), mapa.get('documento_acudiente'), mapa.get('parentesco')
            ] if c] or [0])
            if columna_es_alimento(txt_tel) or (col_contacto_min and col_tel < col_contacto_min):
                mapa.pop('telefono', None)

        if mapa.get('edad_anos') and mapa.get('edad_meses_resto'):
            mapa.pop('edad_completa', None)

        # Bienestarina/RPP: la columna de numeración debe ser independiente del documento.
        # Nunca usar NUI/documento como consecutivo. Si el detector se confundió, corregirlo.
        if mapa.get('orden') and mapa.get('documento') and mapa['orden'] >= mapa['documento']:
            candidatos_orden = []
            for col, txt_col in textos_por_columna.items():
                if col < mapa['documento'] and txt_col.strip() in {'no', 'n', 'nro', 'n'}:
                    candidatos_orden.append(col)
            if candidatos_orden:
                mapa['orden'] = min(candidatos_orden)
            elif is_bienestarina and mapa.get('documento', 1) > 1:
                mapa['orden'] = 1
            else:
                mapa.pop('orden', None)
        if is_bienestarina and not mapa.get('orden') and mapa.get('documento', 1) > 1:
            mapa['orden'] = 1

        if is_asistencia:
            min_data_col = max([c for c in [
                mapa.get('edad_meses_resto'),
                mapa.get('edad_completa'),
                mapa.get('segundo_apellido'),
                mapa.get('nombre')
            ] if c] or [0])
            cols = sorted(c for c in control_cols if c > min_data_col)
            if total_mensual_col:
                cols = [c for c in cols if c < total_mensual_col]

            # ALPHA56 — RAM completo: si el detector solo encuentra una parte de la
            # matriz por encabezados verticales/mergeados, inferir todas las columnas
            # entre EDAD y TOTAL MENSUAL/ASISTENCIAS. Esto evita llenar solo media matriz.
            if total_mensual_col and min_data_col and len(cols) < 20:
                rango_completo = list(range(min_data_col + 1, total_mensual_col))
                cols = sorted(set(cols).union(rango_completo))
            mapa['dias_asistencia'] = cols

        return mapa

    def fila_es_encabezado(ws, row):
        textos = ' '.join(normalizar_texto_clave(ws.cell(row=row, column=col).value) for col in range(1, ws.max_column + 1))
        claves = [
            'primer nombre', 'segundo nombre', 'primer apellido', 'segundo apellido',
            'anos', 'años', 'meses', 'semana', 'lunes', 'martes', 'miercoles', 'jueves',
            'viernes', 'control diario', 'total mensual', 'asistencias', 'inasistencias',
            'casos retiro', 'fecha entrega', 'lote', 'cantidad'
        ]
        return any(clave in textos for clave in claves)

    def detectar_secciones(ws, nombre_formato):
        secciones = []
        formato_norm = normalizar_texto_clave(nombre_formato)
        is_rpp = 'rpp' in formato_norm
        is_bienestarina = 'bienestarina' in formato_norm

        for row in range(1, min(ws.max_row, 80) + 1):
            textos = [normalizar_texto_clave(ws.cell(row=row, column=col).value) for col in range(1, ws.max_column + 1)]
            score = 0
            if any('primer nombre' in t for t in textos): score += 1
            if any('primer apellido' in t for t in textos): score += 1
            if any('documento' in t or 'nui' in t or 'nuip' in t for t in textos): score += 1
            if any('nombres y apellidos' in t for t in textos): score += 1
            if any('control diario' in t for t in textos): score += 1
            if score >= 2:
                mapa = mapear_columnas(ws, row, nombre_formato)
                last_header = row
                for r in range(row, min(ws.max_row, row + 8) + 1):
                    if fila_es_encabezado(ws, r):
                        last_header = max(last_header, r)
                if is_rpp:
                    fila_inicio = max(last_header + 1, row + 3)
                else:
                    fila_inicio = last_header + 1
                secciones.append({'header_row': row, 'fila_inicio': fila_inicio, 'mapa': mapa})

        # Eliminar duplicados muy cercanos.
        filtradas = []
        for sec in secciones:
            if not filtradas or sec['fila_inicio'] - filtradas[-1]['fila_inicio'] > 3:
                filtradas.append(sec)

        # Determinar capacidad de cada sección.
        for idx, sec in enumerate(filtradas):
            siguiente = filtradas[idx + 1]['header_row'] if idx + 1 < len(filtradas) else ws.max_row + 1
            fin = siguiente - 1
            for row in range(sec['fila_inicio'], min(fin, ws.max_row) + 1):
                textos = ' '.join(normalizar_texto_clave(ws.cell(row=row, column=col).value) for col in range(1, min(ws.max_column, 20) + 1))
                if any(clave in textos for clave in ['total', 'verificacion de cobertura', 'coordinador', 'responsable', 'observacion']):
                    fin = row - 1
                    break
            if is_bienestarina:
                sec['capacidad'] = max(0, min(bienestarina_por_hoja, fin - sec['fila_inicio'] + 1))
            else:
                sec['capacidad'] = max(0, min(max_usuarios_formato, fin - sec['fila_inicio'] + 1))

        return [sec for sec in filtradas if sec.get('capacidad', 0) > 0 and sec.get('mapa')]

    campo_dias_actuales = []

    def valor_usuario(user, campo):
        edad_total = int(user.get('EdadMeses') or user.get('edad_meses') or 0)
        tipo_benef = user.get('TipoBeneficiario') or user.get('tipo_beneficiario') or ''
        grupo = normalizar_texto_clave(user.get('GrupoEdad') or user.get('grupo_edad') or '')
        if campo == 'orden':
            return None
        if campo == 'documento':
            return limpiar_valor(user.get('NUI') or user.get('Documento') or user.get('documento'))
        if campo == 'tipo_documento':
            return abreviar_tipo_documento(user.get('TipoDocumento') or user.get('tipo_documento'))
        if campo == 'nombre':
            return limpiar_valor(user.get('Nombre')) or unir_partes(user.get('PrimerNombre'), user.get('SegundoNombre'), user.get('PrimerApellido'), user.get('SegundoApellido'))
        if campo == 'primer_nombre':
            return limpiar_valor(user.get('PrimerNombre'))
        if campo == 'segundo_nombre':
            return limpiar_valor(user.get('SegundoNombre'))
        if campo == 'primer_apellido':
            return limpiar_valor(user.get('PrimerApellido'))
        if campo == 'segundo_apellido':
            return limpiar_valor(user.get('SegundoApellido'))
        if campo == 'edad_anos':
            return edad_total // 12
        if campo == 'edad_meses_resto':
            return edad_total % 12
        if campo == 'edad_completa':
            return formatear_edad_completa(edad_total, tipo_benef)
        if campo == 'acudiente_completo':
            return limpiar_valor(user.get('Acudiente'))
        if campo == 'acudiente_nombre_cedula':
            return nombre_documento_acudiente(user)
        if campo == 'documento_acudiente':
            return limpiar_valor(user.get('DocumentoAcudiente'))
        if campo == 'telefono':
            return limpiar_valor(user.get('Telefono'))
        if campo == 'parentesco':
            return limpiar_valor(user.get('Parentesco'))
        if campo == 'fecha_entrega':
            return fecha_entrega
        if campo == 'fecha_dia':
            return fecha_dia
        if campo == 'fecha_mes':
            return fecha_mes
        if campo == 'fecha_anio':
            return fecha_anio
        if campo == 'lote_bienestarina':
            return lote_bienestarina
        if campo == 'cantidad_bienestarina':
            return cantidad_bienestarina
        if campo == 'peso_kg':
            return limpiar_valor(user.get('PesoKg'))
        if campo == 'talla_cm':
            return limpiar_valor(user.get('TallaCm'))
        if campo == 'imc':
            return limpiar_valor(user.get('IMC'))
        if campo == 'perimetro_braquial':
            return limpiar_valor(user.get('PerimetroBraquial'))
        if campo == 'perimetro_cefalico':
            return limpiar_valor(user.get('PerimetroCefalico'))
        if campo == 'diagnostico_nutricional':
            return limpiar_valor(user.get('DiagnosticoNutricional'))
        if campo == 'estado_control_nutricional':
            return limpiar_valor(user.get('EstadoControlNutricional'))
        if campo == 'fecha_valoracion_nutricional':
            return limpiar_valor(user.get('FechaValoracionNutricional'))
        if campo == 'proximo_control_nutricional':
            return limpiar_valor(user.get('ProximoControlNutricional'))
        if campo == 'total_asistencias':
            return len(campo_dias_actuales)
        if campo == 'inasistencias':
            return ''
        if campo == 'casos_retiro':
            return ''
        # Los usuarios ya llegan filtrados al RPP correspondiente. Aun así se
        # calcula el grupo por edad para marcar exactamente una X en la columna
        # poblacional correcta, incluso cuando Cuéntame trae el grupo amplio
        # "6 meses a 5 años".
        grupo_calculado = {
            '0 A 6 MESES Y GESTANTES': 'grupo_0_6_gestante',
            '6 A 11 MESES 29 DÍAS': 'grupo_6_11',
            '1 A 2 AÑOS 11 MESES': 'grupo_1_2',
            '3 A 5 AÑOS 11 MESES': 'grupo_3_5',
        }.get(categoria_rpp)
        if grupo_calculado:
            pass
        elif 'gestante' in normalizar_texto_clave(tipo_benef):
            grupo_calculado = 'grupo_0_6_gestante'
        elif 0 <= edad_total <= 5:
            grupo_calculado = 'grupo_0_6_gestante'
        elif 6 <= edad_total <= 11:
            grupo_calculado = 'grupo_6_11'
        elif 12 <= edad_total <= 35:
            grupo_calculado = 'grupo_1_2'
        elif 36 <= edad_total <= 71:
            grupo_calculado = 'grupo_3_5'
        if campo == 'grupo_0_6_gestante':
            return 'X' if grupo_calculado == campo else ''
        if campo == 'grupo_6_11':
            return 'X' if grupo_calculado == campo else ''
        if campo == 'grupo_1_2':
            return 'X' if grupo_calculado == campo else ''
        if campo == 'grupo_3_5':
            return 'X' if grupo_calculado == campo else ''
        return ''

    def cambio_aplica(user, campo):
        cambios = user.get('Cambios') or {}
        if cambios.get('_tipo') == 'INGRESO':
            return True
        equivalencias = {
            'documento': ['documento', 'nui'],
            'tipo_documento': ['tipo_documento'],
            'nombre': ['nombre', 'primer_nombre', 'segundo_nombre', 'primer_apellido', 'segundo_apellido'],
            'primer_nombre': ['primer_nombre', 'nombre'],
            'segundo_nombre': ['segundo_nombre', 'nombre'],
            'primer_apellido': ['primer_apellido', 'nombre'],
            'segundo_apellido': ['segundo_apellido', 'nombre'],
            'edad_anos': ['fecha_nacimiento', 'edad_meses'],
            'edad_meses_resto': ['fecha_nacimiento', 'edad_meses'],
            'edad_completa': ['fecha_nacimiento', 'edad_meses'],
            'acudiente_completo': ['nombre_acudiente'],
            'acudiente_nombre_cedula': ['nombre_acudiente', 'documento_acudiente'],
            'documento_acudiente': ['documento_acudiente'],
            'telefono': ['telefono'],
            'parentesco': ['parentesco']
        }
        return any(c in cambios for c in equivalencias.get(campo, [campo]))

    # ALPHA30 — RPP/DRPP por grupo etario.
    # El botón del tablero envía códigos internos (rpp_0_6_gestantes,
    # rpp_6_11, rpp_1_2, rpp_3_5). La plantilla oficial puede ser una sola,
    # por eso el generador debe clonar la misma plantilla y filtrar ANTES de
    # escribir usuarios, evitando que todos queden en una sola hoja/descarga.
    RPP_CATEGORIA_CONFIG = {
        'rpp_0_6_gestantes': {
            'nombre': '0 A 6 MESES Y GESTANTES',
            'slug': '0_A_6_MESES_Y_GESTANTES',
            'alias': ['0 a 6 meses y gestantes', '0 6 meses y gestantes', 'menor de seis meses', 'gestante']
        },
        'rpp_6_11': {
            'nombre': '6 A 11 MESES 29 DÍAS',
            'slug': '6_A_11_MESES_29_DIAS',
            'alias': ['6 a 11 meses 29 dias', '6 11 meses', '6 a 11']
        },
        'rpp_1_2': {
            'nombre': '1 A 2 AÑOS 11 MESES',
            'slug': '1_A_2_ANOS_11_MESES',
            'alias': ['1 a 2 anos 11 meses', '1 2 anos', '1 a 2']
        },
        'rpp_3_5': {
            'nombre': '3 A 5 AÑOS 11 MESES',
            'slug': '3_A_5_ANOS_11_MESES',
            'alias': ['3 a 5 anos 11 meses', '3 5 anos', '3 a 5']
        }
    }
    CATEGORIAS_RPP = {cfg['nombre']: cfg for cfg in RPP_CATEGORIA_CONFIG.values()}
    CATEGORIAS_RPP_NOMBRES = [cfg['nombre'] for cfg in RPP_CATEGORIA_CONFIG.values()]

    def categoria_rpp_desde_nombre(nombre_formato):
        raw = str(nombre_formato or '').strip().lower().replace('-', '_')
        if raw in RPP_CATEGORIA_CONFIG:
            return RPP_CATEGORIA_CONFIG[raw]['nombre']

        nombre = normalizar_texto_clave(nombre_formato)
        if not nombre:
            return None

        # Los códigos de UI llegan como rpp_0_6_gestantes, pero el normalizador
        # los convierte en "rpp 0 6 gestantes".
        if 'rpp 0 6' in nombre or '0 a 6' in nombre or ('rpp' in nombre and 'gestante' in nombre):
            return '0 A 6 MESES Y GESTANTES'
        if 'rpp 6 11' in nombre or '6 a 11' in nombre or '6 11 meses' in nombre:
            return '6 A 11 MESES 29 DÍAS'
        if 'rpp 1 2' in nombre or '1 a 2' in nombre or '1 2 anos' in nombre or '1 2 anios' in nombre:
            return '1 A 2 AÑOS 11 MESES'
        if 'rpp 3 5' in nombre or '3 a 5' in nombre or '3 5 anos' in nombre or '3 5 anios' in nombre:
            return '3 A 5 AÑOS 11 MESES'
        return None

    def slug_categoria_rpp(categoria):
        if not categoria:
            return ''
        return CATEGORIAS_RPP.get(categoria, {}).get(
            'slug',
            re.sub(r'[^A-Z0-9]+', '_', normalizar_texto_clave(categoria).upper()).strip('_')
        )

    def documento_usuario_rpp(user):
        documento = limpiar_valor(user.get('NUI') or user.get('Documento') or user.get('documento') or user.get('numero_documento'))
        documento = re.sub(r'\.0$', '', str(documento).strip())
        return documento

    def usuario_activo_para_rpp(user):
        estado = normalizar_texto_clave(user.get('Estado') or user.get('estado') or user.get('estado_beneficiario') or '')
        if not estado:
            return True
        estados_excluidos = {
            'retirado', 'retirada', 'fallecido', 'fallecida', 'inactivo', 'inactiva',
            'trasladado', 'trasladada', 'egresado', 'egresada'
        }
        return estado not in estados_excluidos

    def edad_meses_usuario_rpp(user):
        for key in ('EdadMeses', 'edad_meses', 'edadMeses'):
            value = user.get(key)
            if value not in (None, ''):
                try:
                    return int(float(value))
                except Exception:
                    pass
        fecha = user.get('FechaNacimiento') or user.get('fecha_nacimiento') or user.get('fechaNacimiento')
        if fecha:
            try:
                return int(calcular_edad_meses(fecha))
            except Exception:
                return None
        return None

    def clasificar_grupo_etario_rpp(user):
        tipo = normalizar_texto_clave(user.get('TipoBeneficiario') or user.get('tipo_beneficiario') or '')
        grupo = normalizar_texto_clave(user.get('GrupoEdad') or user.get('grupo_edad') or '')

        if 'gestante' in tipo or 'gestante' in grupo:
            return '0 A 6 MESES Y GESTANTES'
        if '0 a 6' in grupo or '0 a 5' in grupo or 'menor de seis' in tipo or 'menor de seis' in grupo:
            return '0 A 6 MESES Y GESTANTES'
        if '6 a 11' in grupo:
            return '6 A 11 MESES 29 DÍAS'
        if '1 a 2' in grupo:
            return '1 A 2 AÑOS 11 MESES'
        if '3 a 5' in grupo:
            return '3 A 5 AÑOS 11 MESES'

        edad = edad_meses_usuario_rpp(user)
        if edad is None:
            return 'INCONSISTENTE'
        # Se mantiene la misma lógica operativa del tablero: menores de 6 meses
        # van con gestantes; desde 6 meses pasan al bloque 6 a 11.
        if 0 <= edad <= 5:
            return '0 A 6 MESES Y GESTANTES'
        if 6 <= edad <= 11:
            return '6 A 11 MESES 29 DÍAS'
        if 12 <= edad <= 35:
            return '1 A 2 AÑOS 11 MESES'
        if 36 <= edad <= 71:
            return '3 A 5 AÑOS 11 MESES'
        return 'FUERA_DE_RANGO'

    def usuario_pertenece_a_categoria(user, categoria):
        if not categoria:
            return True
        if not usuario_activo_para_rpp(user):
            return False
        return clasificar_grupo_etario_rpp(user) == categoria

    def filtrar_usuarios_rpp_por_categoria(usuarios, categoria):
        filtrados = []
        documentos_vistos = set()
        duplicados = 0
        inconsistentes = 0
        for user in usuarios or []:
            categoria_usuario = clasificar_grupo_etario_rpp(user)
            if categoria_usuario in {'INCONSISTENTE', 'FUERA_DE_RANGO'}:
                inconsistentes += 1
                continue
            if categoria_usuario != categoria or not usuario_activo_para_rpp(user):
                continue
            documento = documento_usuario_rpp(user)
            if documento:
                clave = normalizar_texto_clave(documento)
                if clave in documentos_vistos:
                    duplicados += 1
                    continue
                documentos_vistos.add(clave)
            filtrados.append(user)
        if duplicados or inconsistentes:
            print(
                f'RPP {unidad_nombre} / {categoria}: {len(filtrados)} usuario(s), '
                f'{duplicados} duplicado(s) omitido(s), {inconsistentes} inconsistente(s) excluido(s).'
            )
        return filtrados

    def nombre_archivo_salida_formato(formato, categoria_rpp=None):
        unidad_sanitizada = re.sub(r'[^A-Za-z0-9_]+', '_', unidad_nombre.replace(' ', '_')).strip('_')
        formato_norm = normalizar_texto_clave(formato)
        if categoria_rpp:
            return f"{unidad_sanitizada}_RPP_{slug_categoria_rpp(categoria_rpp)}.xlsx"
        if 'bienestarina' in formato_norm:
            return f"{unidad_sanitizada}_Bienestarina_{año}_{mes:02d}.xlsx"
        if 'rpp' in formato_norm:
            return f"{unidad_sanitizada}_RPP_OFICIAL_{año}_{mes:02d}.xlsx"
        if 'ram' in formato_norm or 'rram' in formato_norm or 'asistencia' in formato_norm:
            return f"{unidad_sanitizada}_RAM_ASISTENCIA_MENSUAL_{año}_{mes:02d}.xlsx"
        formato_sanitizado = secure_filename(formato) or formato
        return f"{unidad_sanitizada}_{formato_sanitizado}"

    def limpiar_seccion(ws, sec, minimo_filas=None):
        mapa = sec['mapa']
        campos_datos = [c for k, c in mapa.items() if isinstance(c, int) and k not in {'dias_asistencia'}]
        campos_datos += list(mapa.get('dias_asistencia') or [])
        campos_datos = sorted({c for c in campos_datos if c})
        filas_a_limpiar = sec['capacidad']
        if minimo_filas:
            filas_a_limpiar = max(filas_a_limpiar, int(minimo_filas))
        for row in range(sec['fila_inicio'], min(ws.max_row, sec['fila_inicio'] + filas_a_limpiar - 1) + 1):
            # No borrar pie de página ni firmas oficiales aunque el cálculo de capacidad falle.
            textos = ' '.join(normalizar_texto_clave(ws.cell(row=row, column=col).value) for col in range(1, min(ws.max_column, 20) + 1))
            if row >= sec['fila_inicio'] + sec['capacidad'] and any(k in textos for k in ['firma del responsable', 'cedula del responsable', 'cédula del responsable', 'antes de imprimir', 'observacion']):
                break
            for col in campos_datos:
                set_cell(ws, row, col, '', center=True)

    def _valor_no_vacio(valor):
        return valor is not None and str(valor).strip() != ''

    def asegurar_minuta_patron_en_filas(ws, sec, filas_usadas):
        """Copia la minuta patrón solo dentro del perímetro de usuarios.

        El formato oficial trae cantidades base en las columnas de alimentos.
        La plataforma no debe escribir cantidades en filas sin usuario.
        """
        mapa = sec.get('mapa') or {}
        minuta_cols = [c for c in mapa.get('minuta_cols', []) if isinstance(c, int)]
        if not minuta_cols or filas_usadas <= 0:
            return

        fila_inicio = int(sec.get('fila_inicio') or 1)
        capacidad = int(sec.get('capacidad') or 0)
        fila_fin = min(ws.max_row, fila_inicio + max(capacidad, filas_usadas) - 1)
        patron = {}
        for col in minuta_cols:
            for row in range(fila_inicio, fila_fin + 1):
                valor = celda_es_combinada(ws, row, col).value
                if _valor_no_vacio(valor):
                    patron[col] = valor
                    break

        for row in range(fila_inicio, min(ws.max_row, fila_inicio + filas_usadas - 1) + 1):
            for col, valor in patron.items():
                cell = celda_es_combinada(ws, row, col)
                if not _valor_no_vacio(cell.value):
                    set_cell(ws, row, col, valor, center=True)

    def limpiar_minuta_fuera_perimetro(ws, sec, filas_usadas):
        """Borra únicamente valores de minuta en filas sin usuario.

        No cambia colores, bordes, combinaciones ni estructura oficial.
        """
        mapa = sec.get('mapa') or {}
        minuta_cols = [c for c in mapa.get('minuta_cols', []) if isinstance(c, int)]
        if not minuta_cols:
            return
        fila_inicio = int(sec.get('fila_inicio') or 1)
        capacidad = int(sec.get('capacidad') or 0)
        inicio_limpieza = fila_inicio + max(0, int(filas_usadas or 0))
        fin_limpieza = min(ws.max_row, fila_inicio + max(0, capacidad) - 1)
        if inicio_limpieza > fin_limpieza:
            return
        for row in range(inicio_limpieza, fin_limpieza + 1):
            for col in minuta_cols:
                set_cell(ws, row, col, '', center=True)

    def _texto_columna_contexto(ws, col, fila_inicio):
        filas = range(max(1, int(fila_inicio or 1) - 8), min(ws.max_row, int(fila_inicio or 1) + 2) + 1)
        return ' '.join(normalizar_texto_clave(ws.cell(row=r, column=col).value) for r in filas if normalizar_texto_clave(ws.cell(row=r, column=col).value))

    def _aliases_producto_minuta(nombre_producto):
        clave = normalizar_minuta_texto(nombre_producto)
        aliases = []
        if clave in equivalencias_minuta:
            aliases.extend(equivalencias_minuta.get(clave) or [])
        aliases.append(clave)
        for token in clave.split():
            if len(token) >= 4:
                aliases.append(token)
        return {normalizar_texto_clave(a) for a in aliases if a}

    def _buscar_columna_producto_minuta(ws, sec, producto):
        mapa = sec.get('mapa') or {}
        minuta_cols = [c for c in mapa.get('minuta_cols', []) if isinstance(c, int)]
        if not minuta_cols:
            return None
        aliases = _aliases_producto_minuta(producto.get('nombre_producto'))
        if not aliases:
            return None
        fila_inicio = int(sec.get('fila_inicio') or 1)
        best = (0, None)
        for col in minuta_cols:
            contexto = _texto_columna_contexto(ws, col, fila_inicio)
            if not contexto:
                continue
            score = 0
            for alias in aliases:
                if alias and alias in contexto:
                    score += 10 + len(alias)
            comp = normalizar_texto_clave(producto.get('componente'))
            if comp and comp in contexto:
                score += 3
            if score > best[0]:
                best = (score, col)
        return best[1]

    def aplicar_minuta_rpp_versionada_en_filas(ws, sec, bloque):
        """Escribe cantidades de la minuta vigente por grupo etario sin alterar estilos."""
        if not minuta_rpp_vigente or not productos_para_usuario:
            return False
        fila_inicio = int(sec.get('fila_inicio') or 1)
        escritos = 0
        pendientes = []
        for offset, user in enumerate(bloque or []):
            row = fila_inicio + offset
            productos = productos_para_usuario(minuta_rpp_vigente, user) or []
            for producto in productos:
                col = _buscar_columna_producto_minuta(ws, sec, producto)
                if not col:
                    pendientes.append(producto.get('nombre_producto'))
                    continue
                cantidad = str(producto.get('cantidad') or '').strip()
                unidad = str(producto.get('unidad_medida') or '').strip()
                valor = f'{cantidad}{unidad}' if cantidad and unidad and not cantidad.lower().endswith(unidad.lower()) else (cantidad or '')
                if valor:
                    set_cell(ws, row, col, valor, center=True)
                    escritos += 1
        if pendientes:
            print(f'RPP minuta vigente: productos sin columna mapeada para {unidad_nombre}: {sorted(set(pendientes))[:8]}')
        return escritos > 0

    def _weekday_from_col(ws, col, data_row):
        """Detecta día de semana de una columna RAM/RAN a partir de encabezados cercanos."""
        for r in range(max(1, int(data_row or 1) - 14), int(data_row or 1)):
            txt = normalizar_texto_clave(ws.cell(row=r, column=col).value)
            if not txt:
                continue
            for dia in ['lunes', 'martes', 'miercoles', 'jueves', 'viernes', 'sabado', 'domingo']:
                if dia in txt:
                    return dia
            if txt in {'lun', 'lu', 'l'}: return 'lunes'
            if txt in {'mar', 'ma', 'm'}: return 'martes'
            if txt in {'mie', 'mier', 'mi', 'x'}: return 'miercoles'
            if txt in {'jue', 'ju', 'j'}: return 'jueves'
            if txt in {'vie', 'vi', 'v'}: return 'viernes'
        for r in range(max(1, int(data_row or 1) - 14), int(data_row or 1)):
            raw = str(ws.cell(row=r, column=col).value or '').strip()
            if raw.isdigit():
                try:
                    dia_num = int(raw)
                    if 1 <= dia_num <= 31:
                        idx = datetime(int(año), int(mes), dia_num).weekday()
                        return {0:'lunes',1:'martes',2:'miercoles',3:'jueves',4:'viernes',5:'sabado',6:'domingo'}[idx]
                except Exception:
                    pass
        # ALPHA56 fallback: si la columna no tiene encabezado legible por celdas
        # combinadas/verticales, usar su posición dentro de la matriz diaria.
        try:
            if campo_dias_actuales and col in campo_dias_actuales:
                pos = campo_dias_actuales.index(col)
                dia_num = pos + 1
                if 1 <= dia_num <= 31:
                    idx = datetime(int(año), int(mes), dia_num).weekday()
                    return {0:'lunes',1:'martes',2:'miercoles',3:'jueves',4:'viernes',5:'sabado',6:'domingo'}[idx]
        except Exception:
            pass
        return None

    def _es_formato_ran_actual():
        """Identifica RAN/RRAN sin tocar RAM ni otros formatos de asistencia."""
        try:
            texto = normalizar_texto_clave(formato_norm)
        except Exception:
            texto = ''
        return bool(re.search(r'\b(rran|ran)\b', texto)) and not bool(re.search(r'\b(rram|ram)\b', texto))

    def _semana_control_diario_desde_columna(ws, col, data_row):
        """Devuelve 1..5 según el bloque semanal de la columna de control diario.

        Las plantillas RAN/RRAN manejan columnas por semanas (PRIMERA, SEGUNDA,
        TERCERA, CUARTA, QUINTA). No equivalen a días consecutivos 1..31; por
        eso se calcula semana+weekday antes de decidir si el lunes pertenece al mes.
        """
        aliases = {
            'primera': 1, 'primer': 1, '1 semana': 1, 'semana 1': 1,
            'segunda': 2, '2 semana': 2, 'semana 2': 2,
            'tercera': 3, '3 semana': 3, 'semana 3': 3,
            'cuarta': 4, '4 semana': 4, 'semana 4': 4,
            'quinta': 5, '5 semana': 5, 'semana 5': 5,
        }
        for r in range(max(1, int(data_row or 1) - 14), int(data_row or 1)):
            try:
                txt = normalizar_texto_clave(celda_es_combinada(ws, r, col).value)
            except Exception:
                txt = normalizar_texto_clave(ws.cell(row=r, column=col).value)
            if not txt:
                continue
            for alias, numero in aliases.items():
                if alias in txt:
                    return numero

        try:
            cols = list(campo_dias_actuales or [])
            if col not in cols:
                return None
            semana = 1
            previo = None
            weekday_idx = {'lunes': 0, 'martes': 1, 'miercoles': 2, 'jueves': 3, 'viernes': 4, 'sabado': 5, 'domingo': 6}
            for c in cols[:cols.index(col) + 1]:
                dia = _weekday_from_col(ws, c, data_row)
                actual = weekday_idx.get(dia)
                if previo is not None and actual is not None and actual <= previo:
                    semana += 1
                if actual is not None:
                    previo = actual
            return semana
        except Exception:
            return None

    def _fecha_real_ran_desde_columna(ws, col, data_row):
        """Calcula la fecha real de una celda semanal RAN.

        Si la celda cae fuera del mes seleccionado retorna None, evitando marcar
        lunes inexistentes al inicio/final de meses incompletos.
        """
        weekday_idx = {'lunes': 0, 'martes': 1, 'miercoles': 2, 'jueves': 3, 'viernes': 4, 'sabado': 5, 'domingo': 6}
        dia = _weekday_from_col(ws, col, data_row)
        if dia not in weekday_idx:
            return None
        semana = _semana_control_diario_desde_columna(ws, col, data_row)
        if not semana:
            return None
        try:
            primer_dia_mes = datetime(int(año), int(mes), 1)
            inicio_semana_uno = primer_dia_mes - timedelta(days=primer_dia_mes.weekday())
            fecha = inicio_semana_uno + timedelta(days=(int(semana) - 1) * 7 + weekday_idx[dia])
            if fecha.month != int(mes):
                return None
            return fecha
        except Exception:
            return None

    def escribir_usuario(ws, row, mapa, user, indice):
        nonlocal campo_dias_actuales
        campo_dias_actuales = list(mapa.get('dias_asistencia') or [])
        if 'orden' in mapa:
            set_cell(ws, row, mapa['orden'], indice + 1)
        campos = [
            'tipo_documento', 'documento', 'nombre', 'primer_nombre', 'segundo_nombre',
            'primer_apellido', 'segundo_apellido', 'edad_anos', 'edad_meses_resto',
            'edad_completa', 'acudiente_completo', 'acudiente_nombre_cedula',
            'documento_acudiente', 'telefono', 'parentesco', 'fecha_entrega',
            'fecha_dia', 'fecha_mes', 'fecha_anio', 'lote_bienestarina',
            'cantidad_bienestarina', 'total_asistencias', 'inasistencias',
            'casos_retiro', 'grupo_0_6_gestante', 'grupo_6_11', 'grupo_1_2', 'grupo_3_5'
        ]
        for campo in campos:
            col = mapa.get(campo)
            if not col:
                continue
            val = valor_usuario(user, campo)
            cell = set_cell(ws, row, col, val, center=True)
            # No cambiar colores oficiales del formato.
            # Antes se pintaban cambios en verde, pero eso alteraba columnas y celdas
            # oficiales de ICBF. La trazabilidad queda en datos/auditoría, no en rellenos.
            # if cambio_aplica(user, campo):
            #     cell.fill = fill_verde

        # ALPHA53 — RAM/RAN: asistencia automática con A mayúscula según grupo etario
        # y calendario real del mes. Ya no se marca X en todos los días.
        # RELEASE RAN LUNES: para RAN/RRAN, por solicitud puntual del release,
        # se marca únicamente lunes real del mes seleccionado, sin tocar RAM.
        total_asistencias_marcadas = 0
        es_ran_lunes = _es_formato_ran_actual()
        dias_permitidos = {'lunes'} if es_ran_lunes else set()
        if not es_ran_lunes:
            try:
                dias_permitidos = obtener_dias_asistencia_usuario(user) if obtener_dias_asistencia_usuario else set()
            except Exception:
                dias_permitidos = set()
        for col in campo_dias_actuales:
            dia_col = _weekday_from_col(ws, col, row)
            fecha_col = None
            es_festivo = False
            if es_ran_lunes:
                fecha_real = _fecha_real_ran_desde_columna(ws, col, row)
                fecha_col = fecha_real.strftime('%Y-%m-%d') if fecha_real else None
                marca = 'A' if dia_col == 'lunes' and fecha_real is not None else ''
            else:
                try:
                    if campo_dias_actuales and col in campo_dias_actuales:
                        dia_num = campo_dias_actuales.index(col) + 1
                        fecha_col = f'{int(año):04d}-{int(mes):02d}-{int(dia_num):02d}'
                except Exception:
                    fecha_col = None
                es_festivo = bool(fecha_col and fecha_col in festivos_alpha68)
                marca = 'A' if dia_col and dia_col in dias_permitidos and not es_festivo else ''
            if marca:
                total_asistencias_marcadas += 1
            set_cell(ws, row, col, marca, center=True)
        if mapa.get('total_asistencias'):
            set_cell(ws, row, mapa['total_asistencias'], total_asistencias_marcadas, center=True)

    def actualizar_verificacion_cobertura(ws, usuarios):
        """Actualiza valores de verificación de cobertura sin borrar textos oficiales.

        Prioridad:
        1. Escribir en la celda amarilla disponible de la misma fila.
        2. Escribir en una celda vacía cercana a la derecha.
        3. No escribir si no se encuentra un destino seguro.
        """
        # ALPHA53 — Verificación de cobertura RAM desde usuarios únicos y reglas de asistencia.
        if calcular_verificacion_cobertura_ram:
            cobertura = calcular_verificacion_cobertura_ram(list(usuarios or []), {})
            menores = cobertura.get('menores_6_meses_asistentes', 0)
            mayores = cobertura.get('mayores_6_meses_asistentes', 0)
            gestantes = cobertura.get('gestantes_inscritas', 0)
            gestantes_asistentes = cobertura.get('gestantes_asistentes', 0)
            total = menores + mayores + gestantes_asistentes
        else:
            menores = 0
            mayores = 0
            gestantes = 0
            for user in usuarios:
                tipo = normalizar_texto_clave(user.get('TipoBeneficiario', ''))
                edad = int(user.get('EdadMeses') or 0)
                if 'gestante' in tipo:
                    gestantes += 1
                elif edad < 6:
                    menores += 1
                else:
                    mayores += 1
            total = len(usuarios)

        def es_celda_amarilla(cell):
            try:
                fill = cell.fill
                if not fill or fill.fill_type is None:
                    return False
                rgb = ''
                if fill.fgColor and fill.fgColor.type == 'rgb':
                    rgb = (fill.fgColor.rgb or '').upper()
                elif fill.start_color and fill.start_color.type == 'rgb':
                    rgb = (fill.start_color.rgb or '').upper()
                # Amarillos frecuentes: FFFF00, FFC000, FFE699, FFF2CC.
                return any(token in rgb for token in ['FFFF00', 'FFC000', 'FFE699', 'FFF2CC', 'FFFF99'])
            except Exception:
                return False

        def destino_seguro(row, label_col):
            # Buscar primero cuadros amarillos en la misma fila, sin texto oficial.
            for col in range(label_col + 1, min(ws.max_column, label_col + 12) + 1):
                cell = celda_es_combinada(ws, row, col)
                texto = normalizar_texto_clave(cell.value)
                if es_celda_amarilla(cell) and not texto:
                    return col
            # Luego cualquier celda vacía cercana, evitando columnas con texto oficial.
            for col in range(label_col + 1, min(ws.max_column, label_col + 12) + 1):
                cell = celda_es_combinada(ws, row, col)
                texto = normalizar_texto_clave(cell.value)
                if not texto:
                    return col
            return None

        filas_bloque = []
        for row in range(1, ws.max_row + 1):
            textos = ' '.join(
                normalizar_texto_clave(ws.cell(row=row, column=col).value)
                for col in range(1, min(ws.max_column, 35) + 1)
            )
            if 'verificacion de cobertura' in textos or 'verificación de cobertura' in textos:
                filas_bloque = list(range(row, min(ws.max_row, row + 10) + 1))
                break

        if not filas_bloque:
            return

        reglas = [
            (['menores de seis', 'menor de seis', '0 a 6', 'menores de 6'], menores),
            (['mayores de seis', 'mayor de seis', 'mayores de 6'], mayores),
            (['gestantes', 'madres gestantes', 'mujeres gestantes'], gestantes),
            (['total participantes', 'total asistentes', 'total usuarios', 'total general'], total)
        ]

        for row in filas_bloque:
            fila_textos = [
                normalizar_texto_clave(ws.cell(row=row, column=col).value)
                for col in range(1, min(ws.max_column, 35) + 1)
            ]
            texto_fila = ' '.join(fila_textos)
            for aliases, valor in reglas:
                if not any(alias in texto_fila for alias in aliases):
                    continue
                for col, txt in enumerate(fila_textos, start=1):
                    if any(alias in txt for alias in aliases):
                        destino = destino_seguro(row, col)
                        if destino:
                            set_cell(ws, row, destino, valor, center=True)
                        break

    def metadata_ram_v3():
        """Construye encabezado RAM desde las mismas fuentes maestras existentes."""
        def meta_any(*keys):
            for key in keys:
                value = valor_metadata(key)
                if value:
                    return value
            return ''

        def pertenece_tenant(item):
            fid = item.get('fundacion_id')
            if fid in (None, ''):
                return True
            try:
                return int(fid) == int(fundacion_actual_id())
            except (TypeError, ValueError):
                return False

        def coincide_unidad(item):
            valor = item.get('nombre') or item.get('unidad_servicio') or item.get('unidad') or ''
            return normalize_unidad(valor) == unidad_nombre

        def datos_json(item):
            raw = item.get('datos_json')
            if isinstance(raw, dict):
                return raw
            if isinstance(raw, str) and raw.strip():
                try:
                    value = json.loads(raw)
                    return value if isinstance(value, dict) else {}
                except Exception:
                    return {}
            return {}

        def first_value(item, *keys):
            extra = datos_json(item)
            for source in (item, extra):
                for key in keys:
                    value = limpiar_valor(source.get(key))
                    if value:
                        return value
            return ''

        unidad_db = {}
        try:
            conn = get_db_connection()
            # Fuente canónica primero; la tabla histórica queda solo como respaldo.
            for tabla in ('master_unidades', 'unidades'):
                try:
                    filas = conn.execute(f"SELECT * FROM {tabla}").fetchall()
                except Exception:
                    filas = []
                for row in filas:
                    item = dict(row)
                    activo = item.get('activo')
                    if tabla == 'master_unidades' and activo not in (None, '', 1, True, '1'):
                        continue
                    if pertenece_tenant(item) and coincide_unidad(item):
                        unidad_db = item
                        break
                if unidad_db:
                    break
            conn.close()
        except Exception:
            unidad_db = {}

        def agente_ram():
            candidatos = []
            try:
                conn = get_db_connection()
                consultas = [
                    "SELECT * FROM master_talento_humano",
                    """SELECT p.*, a.unidad AS unidad_asignada,
                              COALESCE(a.rol, a.cargo, p.rol_normalizado, p.cargo) AS cargo_asignado
                       FROM th_asignaciones a
                       JOIN th_personas p ON p.id=a.persona_id
                       WHERE UPPER(COALESCE(a.estado,'ACTIVO'))='ACTIVO'
                         AND COALESCE(p.activo,1)=1""",
                ]
                for index, sql in enumerate(consultas):
                    try:
                        filas = conn.execute(sql).fetchall()
                    except Exception:
                        filas = []
                    for row in filas:
                        item = dict(row)
                        if not pertenece_tenant(item):
                            continue
                        if index == 0 and item.get('activo') not in (None, '', 1, True, '1'):
                            continue
                        if item.get('unidad_asignada') and not item.get('unidad'):
                            item['unidad'] = item.get('unidad_asignada')
                        if item.get('cargo_asignado'):
                            item['cargo'] = item.get('cargo_asignado')
                        if not _talento_coincide_unidad(item, unidad_nombre):
                            continue
                        cargo = normalizar_texto_clave(item.get('cargo') or item.get('rol_normalizado') or '')
                        if not any(token in cargo for token in ('agente', 'docente', 'educativo')):
                            continue
                        if not item.get('nombre'):
                            item['nombre'] = item.get('nombre_completo') or unir_partes(item.get('nombres'), item.get('apellidos'))
                        candidatos.append((index, normalizar_texto_clave(item.get('nombre')), item))
                conn.close()
            except Exception:
                candidatos = []
            if candidatos:
                candidatos.sort(key=lambda value: (value[0], value[1]))
                return candidatos[0][2]
            return obtener_talento_por_unidad(unidad_nombre) or {}

        agente = agente_ram()
        eas = meta_any('NombreEAS', 'nombre_eas', 'EntidadAdministradora')
        nit = opt_value('nit', 'NIT', 'nit_eas', default='') or meta_any('NIT', 'Nit', 'nit_eas', 'nit')
        modalidad = meta_any('Modalidad', 'modalidad')
        nui_uds = (
            meta_any('NUIUDS', 'NuiUds', 'nui_uds', 'NUI_UCA', 'nui_uca')
            or first_value(unidad_db, 'nui_uds', 'nui_uca', 'nui', 'NUI')
        )
        codigo_cuentame = (
            meta_any('CodigoCuentame', 'CodigoCUENTAME', 'codigo_cuentame', 'Código CUENTAME UDS')
            or first_value(unidad_db, 'codigo_cuentame', 'codigo_uds', 'codigo_unidad_servicio', 'codigo_unidad')
        )
        codigo_uds = codigo_cuentame or nui_uds or (
            meta_any('CodigoUnidadServicio', 'codigo_unidad_servicio', 'codigo_unidad', 'CodigoUnidad')
            or first_value(unidad_db, 'codigo_unidad', 'codigo_unidad_servicio')
        )
        agente_nombre = (
            limpiar_valor(agente.get('nombre') or unir_partes(agente.get('nombres'), agente.get('apellidos')))
            or first_value(unidad_db, 'docente_asignado', 'agente_educativo', 'docente')
            or meta_any('Docente', 'docente', 'AgenteEducativo', 'agente_educativo')
        )
        documento_agente = (
            limpiar_documento_talento(agente.get('documento'))
            or first_value(unidad_db, 'docente_documento', 'documento_docente', 'documento_agente')
            or meta_any('DocumentoDocente', 'documento_docente', 'CedulaDocente', 'cedula_docente', 'documento_agente')
        )
        telefono_agente = (
            limpiar_documento_talento(agente.get('telefono'))
            or first_value(unidad_db, 'telefono_docente', 'telefono_agente')
            or meta_any('TelefonoDocente', 'telefono_docente', 'TelefonoAgente', 'telefono_agente')
        )
        return {
            'eas_pds': eas,
            'eas': eas,
            'nit': re.sub(r'[.\-\s]', '', str(nit or '')),
            'contrato': meta_any('NumeroContrato', 'numero_contrato') or limpiar_valor(agente.get('contrato')) or limpiar_valor(unidad_db.get('contrato')),
            'regional': meta_any('Regional', 'regional'),
            'centro_zonal': meta_any('CentroZonal', 'centro_zonal'),
            'municipio': meta_any('Municipio', 'municipio'),
            'mes_nombre': mes_nombre,
            'mes_numero': mes,
            'anio': año,
            'agente_educativo': agente_nombre.upper(),
            'documento_agente': documento_agente,
            'modalidad': modalidad,
            'nui_uds': nui_uds or codigo_uds,
            'codigo_cuentame': codigo_cuentame or codigo_uds,
            'codigo_uds': codigo_uds,
            'unidad': unidad_nombre,
            'servicio_atencion': meta_any('ServicioAtencion', 'servicio_atencion'),
            'direccion_uds': opt_value('direccion_uds', 'direccion_unidad', default='') or limpiar_valor(unidad_db.get('direccion')),
            # Regla funcional RAM: este rótulo recibe el teléfono del agente
            # responsable; solo si falta se usa el teléfono institucional.
            'telefono_uds': telefono_agente or opt_value('telefono_uds', 'telefono_unidad', default='') or limpiar_documento_talento(unidad_db.get('telefono')),
        }

    def dias_no_atencion_ram_v3():
        fechas = set(festivos_alpha68 or set())
        raw = opt_value('dias_no_atencion', 'fechas_no_atencion', default='')
        if raw:
            try:
                parsed = json.loads(raw) if isinstance(raw, str) else raw
                if isinstance(parsed, list):
                    fechas.update(str(v).strip() for v in parsed if str(v).strip())
                elif isinstance(parsed, str):
                    fechas.update(x.strip() for x in parsed.split(',') if x.strip())
            except Exception:
                fechas.update(x.strip() for x in str(raw).split(',') if x.strip())
        return fechas

    def plantilla_es_de_otra_unidad(nombre_formato):
        """Evita usar una plantilla específica de otra UCA/UDS para la unidad actual.

        Cuando se cargan varias plantillas RAM/RAN una por unidad, la plataforma debe
        tomar solo la que corresponde. Las plantillas generales sin nombre de unidad sí
        se siguen usando.
        """
        nombre_norm = normalizar_texto_clave(nombre_formato)
        unidades_en_nombre = set()

        for alias, destino in ALIAS_UNIDADES_CUENTAME.items():
            alias_norm = normalizar_texto_clave(alias)
            if alias_norm and alias_norm in nombre_norm:
                unidad_detectada = normalize_unidad(destino)
                if unidad_detectada:
                    unidades_en_nombre.add(unidad_detectada)

        for unidad_conocida in KNOWN_UNITS:
            unidad_norm = normalizar_texto_clave(unidad_conocida)
            if unidad_norm and unidad_norm in nombre_norm:
                unidades_en_nombre.add(normalize_unidad(unidad_conocida))

        if not unidades_en_nombre:
            return False
        return unidad_nombre not in unidades_en_nombre

    formatos_entries = []
    tipos_oficiales_activos = set()
    try:
        for entry in iter_plantillas_oficiales_para_generacion(TEMPLATES_FOLDER, mes=mes, anio=año):
            formatos_entries.append(entry)
            if entry.get('tipo'):
                tipos_oficiales_activos.add(entry.get('tipo'))
    except Exception as exc:
        print(f'No se pudieron leer plantillas oficiales: {exc}')

    for f in os.listdir(TEMPLATES_FOLDER):
        if not f.lower().endswith(('.xlsx', '.xls', '.xlsm')):
            continue
        tipo_detectado = infer_print_format(f)
        # Si existe plantilla oficial aplicable, se evita generar duplicados
        # desde plantillas antiguas de la raíz. RAM V3 solo aplica desde su vigencia.
        nombre_raiz_norm = normalizar_texto_clave(f)
        if tipo_detectado in tipos_oficiales_activos and tipo_detectado in {'rpp', 'bienestarina', 'ram'}:
            continue
        if 'ram' in tipos_oficiales_activos and (
            re.search(r'\b(rram|ram)\b', nombre_raiz_norm)
            or 'registro mensual' in nombre_raiz_norm
            or 'listado de asistencia' in nombre_raiz_norm
        ):
            # Evita duplicar el RAM oficial V3 con una plantilla RAM/listado antigua,
            # pero conserva RAN y RRAN como formatos independientes.
            continue
        formatos_entries.append({
            'nombre': f,
            'ruta': os.path.join(TEMPLATES_FOLDER, f),
            'tipo': tipo_detectado,
            'oficial': False,
            'preservar_estilos': False,
            'preservar_impresion': False,
        })

    def expandir_rpp_por_categoria(entries):
        entries_expandidos = []
        for entry in entries:
            formato_entry = entry.get('nombre') or os.path.basename(entry.get('ruta') or '')
            formato_norm_entry = normalizar_texto_clave(formato_entry)
            tipo_entry = entry.get('tipo') or infer_print_format(formato_entry)
            es_rpp = tipo_entry == 'rpp' or 'rpp' in formato_norm_entry
            if not es_rpp:
                entries_expandidos.append(entry)
                continue

            categoria_detectada = entry.get('categoria_rpp') or categoria_rpp_desde_nombre(formato_entry)
            if categoria_detectada:
                copia = dict(entry)
                copia['categoria_rpp'] = categoria_detectada
                entries_expandidos.append(copia)
                continue

            # Plantilla RPP general/oficial: generar cuatro archivos separados,
            # uno por cada grupo etario. Esto corrige la descarga donde todos
            # los usuarios terminaban en una sola hoja.
            for codigo, cfg in RPP_CATEGORIA_CONFIG.items():
                copia = dict(entry)
                copia['categoria_rpp'] = cfg['nombre']
                copia['codigo_categoria_rpp'] = codigo
                entries_expandidos.append(copia)
        return entries_expandidos

    formatos_entries = expandir_rpp_por_categoria(formatos_entries)

    for entry in formatos_entries:
        formato = entry.get('nombre') or os.path.basename(entry.get('ruta') or '')
        if plantilla_es_de_otra_unidad(formato):
            continue
        ruta_plantilla = entry.get('ruta') or os.path.join(TEMPLATES_FOLDER, formato)
        formato_norm = normalizar_texto_clave(formato)
        categoria_rpp = entry.get('categoria_rpp') or categoria_rpp_desde_nombre(formato)
        es_rpp_formato = (entry.get('tipo') == 'rpp') or ('rpp' in formato_norm)
        if not _alpha68_formato_entry_permitido(entry, formato_norm, categoria_rpp, formatos_solicitados_alpha68):
            continue
        if categoria_rpp and es_rpp_formato:
            usuarios_formato = filtrar_usuarios_rpp_por_categoria(usuarios_base, categoria_rpp)
        elif categoria_rpp:
            usuarios_formato = [u for u in usuarios_base if usuario_pertenece_a_categoria(u, categoria_rpp)]
        else:
            usuarios_formato = list(usuarios_base)
        usuarios_ordenados = ordenar_usuarios_para_formato(usuarios_formato, formato_norm, categoria_rpp)
        es_ram_v3_formato = bool(
            entry.get('tipo') == 'ram'
            and (str(entry.get('version') or '').strip() == '3' or 'plantilla ram oficial v3' in formato_norm or 'registro mensual v3' in formato_norm)
        )
        es_ram_oficial = entry.get('tipo') == 'ram'
        # Todo RAM oficial pagina en bloques de 20; no se trunca la población.
        usuarios_formato = usuarios_ordenados if es_ram_oficial else usuarios_ordenados[:max_usuarios_formato]

        if es_rpp_formato:
            ok_ctx, errores_ctx, resumen_ctx = validate_rpp_context(entry, usuarios_formato, minuta_rpp_vigente, categoria_rpp)
            log_rpp_event(
                'PREVALIDACION_RPP',
                ruta_funcion='inyectar_datos_en_plantillas',
                unidad=unidad_nombre,
                mes=mes,
                anio=año,
                categoria_rpp=categoria_rpp,
                plantilla=os.path.basename(str(ruta_plantilla)),
                ruta_plantilla=ruta_plantilla,
                source=entry.get('source'),
                version_id=entry.get('version_id') or entry.get('plantilla_oficial_version_id'),
                resumen=resumen_ctx,
                errores=errores_ctx,
            )
            if not ok_ctx:
                print(f'No se generó RPP para {unidad_nombre}/{categoria_rpp}: ' + '; '.join(errores_ctx))
                continue

        try:
            if es_ram_v3_formato:
                from services.ram_v3_service import generate_ram_v3
                nombre_salida = secure_filename(f"RAM_V3_{unidad_nombre}_{int(año):04d}_{int(mes):02d}.xlsx")
                salida_path = os.path.join(OUTPUT_FOLDER, nombre_salida)
                resultado_ram = generate_ram_v3(
                    ruta_plantilla,
                    salida_path,
                    usuarios_formato,
                    año,
                    mes,
                    metadata=metadata_ram_v3(),
                    attendance_provider=obtener_dias_asistencia_usuario,
                    non_service_dates=dias_no_atencion_ram_v3(),
                    expected_sha256=entry.get('hash_sha256'),
                )
                registrar_archivo_generado_alpha57(
                    'ram', unidad_nombre, nombre_salida, salida_path,
                    mes=mes, anio=año, grupo_etario=None,
                    extra={
                        'plantilla': os.path.basename(str(ruta_plantilla)),
                        'version': '3',
                        'usuarios': resultado_ram.get('total_participantes', 0),
                        'paginas_ram': resultado_ram.get('paginas_ram', 1),
                        'advertencias': resultado_ram.get('warnings', [])[:20],
                    }
                )
                print(
                    f"RAM V3 generado para {unidad_nombre}: {nombre_salida}; "
                    f"participantes={resultado_ram.get('total_participantes')}; "
                    f"paginas={resultado_ram.get('paginas_ram')}; "
                    f"advertencias={len(resultado_ram.get('warnings') or [])}"
                )
                continue

            if es_ram_oficial:
                from services.ram_historical_service import generate_ram_historical
                nombre_salida = secure_filename(f"{unidad_nombre}_RAM_ASISTENCIA_MENSUAL_{int(año):04d}_{int(mes):02d}.xlsx")
                salida_path = os.path.join(OUTPUT_FOLDER, nombre_salida)
                resultado_ram = generate_ram_historical(
                    ruta_plantilla,
                    salida_path,
                    usuarios_formato,
                    año,
                    mes,
                    metadata=metadata_ram_v3(),
                    attendance_provider=obtener_dias_asistencia_usuario,
                    non_service_dates=dias_no_atencion_ram_v3(),
                    expected_sha256=entry.get('hash_sha256'),
                )
                registrar_archivo_generado_alpha57(
                    'ram', unidad_nombre, nombre_salida, salida_path,
                    mes=mes, anio=año, grupo_etario=None,
                    extra={'plantilla': os.path.basename(str(ruta_plantilla)), 'version': '2', 'usuarios': resultado_ram.get('total_participantes', 0), 'paginas_ram': resultado_ram.get('paginas_ram', 1)},
                )
                continue

            contexto_plantilla_actual.update({
                'oficial': bool(entry.get('oficial')),
                'preservar_estilos': bool(entry.get('preservar_estilos')),
                'preservar_impresion': bool(entry.get('preservar_impresion')),
            })
            wb = load_workbook(ruta_plantilla)
            hojas = seleccionar_hojas(wb)
            if not hojas:
                continue

            is_bienestarina = 'bienestarina' in formato_norm or entry.get('tipo') == 'bienestarina'
            pendientes = list(usuarios_formato)
            paginas_usadas = 0
            rpp_usuarios_escritos = 0
            rpp_productos_escritos = 0

            for ws in hojas:
                actualizar_encabezados(ws)
                secciones = detectar_secciones(ws, formato)
                if not secciones:
                    continue

                for sec in secciones:
                    if not pendientes:
                        limpiar_seccion(ws, sec, bienestarina_por_hoja if is_bienestarina else None)
                        if 'rpp' in formato_norm:
                            limpiar_minuta_fuera_perimetro(ws, sec, 0)
                        continue

                    capacidad = sec['capacidad']
                    if is_bienestarina:
                        capacidad = min(capacidad, bienestarina_por_hoja)

                    bloque = pendientes[:capacidad]
                    pendientes = pendientes[capacidad:]
                    limpiar_seccion(ws, sec, bienestarina_por_hoja if is_bienestarina else None)

                    for i, user in enumerate(bloque):
                        escribir_usuario(ws, sec['fila_inicio'] + i, sec['mapa'], user, paginas_usadas * bienestarina_por_hoja + i)

                    if 'rpp' in formato_norm:
                        # ALPHA54: el RPP no debe usar cantidades quemadas ni fallback silencioso
                        # cuando existe Motor de Minutas. Si no se pudo aplicar la minuta vigente,
                        # cancelar esta salida para evitar archivos nuevos vacíos o con formato viejo.
                        minuta_aplicada = aplicar_minuta_rpp_versionada_en_filas(ws, sec, bloque)
                        if not minuta_aplicada:
                            log_rpp_event(
                                'ERROR_RPP_MINUTA_NO_APLICADA',
                                unidad=unidad_nombre, mes=mes, anio=año, categoria_rpp=categoria_rpp,
                                plantilla=os.path.basename(str(ruta_plantilla)),
                                usuarios_bloque=len(bloque),
                                mensaje='La minuta vigente no pudo mapear productos/cantidades a columnas de la plantilla.'
                            )
                            raise ValueError('No se puede generar RPP porque la minuta vigente no se pudo aplicar al mapeo de la plantilla.')
                        rpp_usuarios_escritos += len(bloque)
                        productos_entry = product_items_from_entry(entry)
                        rpp_productos_escritos += max(1, len(productos_entry)) * len(bloque)
                        limpiar_minuta_fuera_perimetro(ws, sec, len(bloque))

                    actualizar_verificacion_cobertura(ws, usuarios_formato)
                    paginas_usadas += 1

                    if is_bienestarina and not pendientes:
                        # Limpiar secciones/páginas sobrantes de la unidad.
                        continue

            def validar_numeracion_bienestarina_libro():
                if not is_bienestarina:
                    return []
                errores = []
                esperado = 1
                total = len(usuarios_formato or [])
                for ws_val in hojas:
                    secs_val = detectar_secciones(ws_val, formato)
                    for sec_val in secs_val:
                        mapa_val = sec_val.get('mapa') or {}
                        col_orden = mapa_val.get('orden')
                        col_doc = mapa_val.get('documento')
                        if not col_orden:
                            continue
                        fila_inicio = int(sec_val.get('fila_inicio') or 1)
                        capacidad_val = int(sec_val.get('capacidad') or bienestarina_por_hoja or 14)
                        for rr in range(fila_inicio, fila_inicio + capacidad_val):
                            if esperado > total:
                                break
                            val = celda_es_combinada(ws_val, rr, col_orden).value
                            doc_val = celda_es_combinada(ws_val, rr, col_doc).value if col_doc else None
                            val_txt = str(val or '').strip()
                            doc_txt = str(doc_val or '').strip()
                            if val_txt == '':
                                errores.append(f'Consecutivo vacío en {ws_val.title}!{rr},{col_orden}; esperado {esperado}')
                            elif val_txt == doc_txt and len(val_txt) >= 5:
                                errores.append(f'Consecutivo parece documento en {ws_val.title}!{rr},{col_orden}: {val_txt}')
                            else:
                                try:
                                    val_int = int(float(val_txt))
                                    if val_int != esperado:
                                        errores.append(f'Consecutivo incorrecto en {ws_val.title}!{rr},{col_orden}: {val_int}; esperado {esperado}')
                                except Exception:
                                    errores.append(f'Consecutivo no numérico en {ws_val.title}!{rr},{col_orden}: {val_txt}')
                            esperado += 1
                        if esperado > total:
                            break
                    if esperado > total:
                        break
                if esperado <= total:
                    errores.append(f'Faltan consecutivos: solo se validaron {esperado-1} de {total}')
                return errores

            if is_bienestarina:
                errores_num = validar_numeracion_bienestarina_libro()
                if errores_num:
                    log_rpp_event('BIENESTARINA_NUMERACION_INVALIDA', unidad=unidad_nombre, mes=mes, anio=año, errores=errores_num[:10])
                    raise ValueError('No se generó Bienestarina porque la numeración de consecutivos no es válida: ' + '; '.join(errores_num[:3]))

            nombre_salida = nombre_archivo_salida_formato(formato, categoria_rpp)
            if es_rpp_formato:
                if rpp_usuarios_escritos <= 0:
                    raise ValueError('No se generó el RPP porque quedaría vacío: no se escribieron usuarios.')
                if rpp_productos_escritos <= 0:
                    raise ValueError('No se generó el RPP porque quedaría vacío: no se escribieron productos/cantidades.')
                log_rpp_event(
                    'RPP_LISTO_PARA_GUARDAR',
                    unidad=unidad_nombre, mes=mes, anio=año, categoria_rpp=categoria_rpp,
                    plantilla=os.path.basename(str(ruta_plantilla)),
                    ruta_plantilla=ruta_plantilla,
                    version_id=entry.get('version_id') or entry.get('plantilla_oficial_version_id'),
                    minuta_id=(minuta_rpp_vigente or {}).get('id') if isinstance(minuta_rpp_vigente, dict) else None,
                    total_usuarios_encontrados=len(usuarios_formato),
                    total_usuarios_escritos=rpp_usuarios_escritos,
                    total_productos_escritos=rpp_productos_escritos,
                    archivo=nombre_salida,
                )
            tipo_impresion = entry.get('tipo') or infer_print_format(formato)
            if tipo_impresion and not contexto_plantilla_actual.get('preservar_impresion'):
                aplicar_configuracion_impresion_libro(wb, tipo_impresion, source_name=formato)
            elif contexto_plantilla_actual.get('preservar_impresion'):
                # Las plantillas oficiales deben conservar la configuración de impresión.
                # Si no traen área de impresión, se fija únicamente el área segura conocida.
                area_default = 'A1:AA42' if tipo_impresion == 'rpp' else ('A1:T50' if tipo_impresion == 'bienestarina' else None)
                if area_default:
                    for ws in wb.worksheets:
                        try:
                            if not ws.print_area:
                                ws.print_area = area_default
                        except Exception:
                            pass
            salida_path = os.path.join(OUTPUT_FOLDER, nombre_salida)
            wb.save(salida_path)
            if not os.path.exists(salida_path):
                raise ValueError(f'El formato {formato} no se guardó correctamente en disco: {salida_path}')
            if os.path.getsize(salida_path) <= 0:
                raise ValueError(f'El formato {formato} quedó vacío en disco: {salida_path}')
            registrar_archivo_generado_alpha57(
                'rpp' if es_rpp_formato else ('bienestarina' if is_bienestarina else ('ram' if any(k in formato_norm for k in ['asistencia', 'ram', 'rram', 'ran', 'run']) else formato)),
                unidad_nombre, nombre_salida, salida_path, mes=mes, anio=año, grupo_etario=categoria_rpp,
                extra={'plantilla': os.path.basename(str(ruta_plantilla)), 'usuarios': len(usuarios_formato or [])}
            )
            if es_rpp_formato:
                log_rpp_event(
                    'RPP_GENERADO_OK',
                    unidad=unidad_nombre, mes=mes, anio=año, categoria_rpp=categoria_rpp,
                    plantilla=os.path.basename(str(ruta_plantilla)), archivo=nombre_salida,
                    salida_path=salida_path, total_usuarios_escritos=rpp_usuarios_escritos,
                    total_productos_escritos=rpp_productos_escritos
                )
        except Exception as e:
            if 'rpp' in formato_norm or (entry.get('tipo') == 'rpp'):
                log_rpp_event(
                    'RPP_GENERACION_CANCELADA',
                    unidad=unidad_nombre, mes=mes, anio=año, categoria_rpp=categoria_rpp,
                    plantilla=os.path.basename(str(ruta_plantilla or '')),
                    ruta_plantilla=ruta_plantilla, error=str(e)
                )
            print(f'No se pudo rellenar el formato {formato} para {unidad_nombre}: {e}')
        finally:
            contexto_plantilla_actual.update({'oficial': False, 'preservar_estilos': False, 'preservar_impresion': False})

    _alpha68_generar_complementarios_formato(
        unidad_nombre,
        usuarios_base,
        options,
        formatos_solicitados_alpha68,
        mes,
        año,
    )




# -----------------------------------------------------------------------------
# ALPHA68 — Selector de formatos, complementarios y festivos sin tocar plantillas.
# -----------------------------------------------------------------------------
def _alpha68_log(evento, **datos):
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha68_mejoras_formatos.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha68_parse_formatos_seleccionados(options=None):
    raw = ''
    try:
        options = options or {}
        raw = options.get('formatos_seleccionados') or ''
        if not raw and has_request_context():
            raw = request.form.get('formatos_seleccionados') or request.args.get('formatos_seleccionados') or ''
    except Exception:
        raw = ''
    if isinstance(raw, (list, tuple, set)):
        items = list(raw)
    else:
        items = re.split(r'[,;|\s]+', str(raw or ''))
    normalizados = set()
    aliases = {
        'rpp': 'rpp', 'bienestarina': 'bienestarina', 'ram': 'ram',
        # RAN/RRAN fueron retirados. Se conserva un marcador para que clientes
        # antiguos no conviertan una selección obsoleta en "generar todo".
        'ran': 'formato_no_disponible', 'rran': 'formato_no_disponible',
        'relacion': 'relacion_mensual',
        'relacion_mensual': 'relacion_mensual', 'listado': 'listado_usuarios',
        'listado_usuarios': 'listado_usuarios', 'usuarios': 'listado_usuarios',
        'listado_asistencia_usuarios': 'listado_asistencia_usuarios',
        'asistencia_usuarios': 'listado_asistencia_usuarios',
        'distribucion': 'distribucion_alimentos', 'distribucion_alimentos': 'distribucion_alimentos',
        'alimentos': 'distribucion_alimentos', 'paquete': 'paquete_completo',
        'paquete_completo': 'paquete_completo', 'todo': 'paquete_completo', 'todos': 'paquete_completo',
    }
    for item in items:
        key = normalizar_texto_clave(item).replace(' ', '_')
        if not key:
            continue
        normalizados.add(aliases.get(key, key))
    if 'paquete_completo' in normalizados:
        return set()
    return normalizados


def _alpha68_formato_entry_permitido(entry, formato_norm, categoria_rpp, seleccionados):
    if not seleccionados:
        return True
    tipo = normalizar_texto_clave(entry.get('tipo') or infer_print_format(entry.get('nombre') or '') or '')
    nombre = normalizar_texto_clave(entry.get('nombre') or os.path.basename(entry.get('ruta') or '') or '')
    claves = set()
    if categoria_rpp or tipo == 'rpp' or 'rpp' in formato_norm or 'rpp' in nombre:
        claves.add('rpp')
    if tipo == 'bienestarina' or 'bienestarina' in formato_norm or 'bienestarina' in nombre:
        claves.add('bienestarina')
    es_rran = bool(re.search(r'\brran\b', formato_norm + ' ' + nombre)) or tipo == 'rran'
    es_ran = (bool(re.search(r'\bran\b', formato_norm + ' ' + nombre)) or tipo == 'ran') and not es_rran
    es_ram = (
        tipo == 'ram'
        or bool(re.search(r'\b(rram|ram)\b', formato_norm + ' ' + nombre))
        or 'registro mensual' in formato_norm
        or 'registro mensual' in nombre
    ) and not es_ran and not es_rran
    if es_ram:
        claves.add('ram')
    if es_ran:
        claves.add('ran')
    if es_rran:
        claves.add('rran')
    # Si no se pudo clasificar una plantilla vieja, solo se genera cuando no hay selección.
    permitido = bool(claves & set(seleccionados))
    if not permitido:
        _alpha68_log('FORMATO_OMITIDO_POR_SELECTOR', plantilla=entry.get('nombre'), tipo=tipo, categoria_rpp=categoria_rpp, seleccion=list(seleccionados), claves=list(claves))
    return permitido


def _alpha68_should_generar(seleccionados, clave):
    return bool(seleccionados and clave in set(seleccionados))


def _alpha68_valor_usuario(user, *keys):
    for k in keys:
        try:
            val = user.get(k)
        except Exception:
            val = ''
        if val not in (None, ''):
            return val
    return ''


def _alpha68_usuario_nombre(user):
    nombre = _alpha68_valor_usuario(user, 'Nombre', 'nombre', 'nombre_completo')
    if nombre:
        return nombre
    return unir_partes(
        _alpha68_valor_usuario(user, 'PrimerNombre', 'primer_nombre'),
        _alpha68_valor_usuario(user, 'SegundoNombre', 'segundo_nombre'),
        _alpha68_valor_usuario(user, 'PrimerApellido', 'primer_apellido'),
        _alpha68_valor_usuario(user, 'SegundoApellido', 'segundo_apellido'),
    )


def _alpha68_grupo_usuario(user):
    grupo = _alpha68_valor_usuario(user, 'GrupoEdad', 'grupo_edad', 'grupo_etario')
    if grupo:
        return str(grupo)
    try:
        edad = int(float(_alpha68_valor_usuario(user, 'EdadMeses', 'edad_meses') or 0))
    except Exception:
        edad = 0
    tipo = normalizar_texto_clave(_alpha68_valor_usuario(user, 'TipoBeneficiario', 'tipo_beneficiario'))
    if 'gestante' in tipo:
        return 'GESTANTE'
    if edad <= 6:
        return '0 A 6 MESES'
    if edad <= 11:
        return '6 A 11 MESES'
    if edad <= 35:
        return '1 A 2 ANOS 11 MESES'
    if edad <= 71:
        return '3 A 5 ANOS 11 MESES'
    return 'SIN GRUPO'


def _alpha68_docente_por_usuarios(unidad, usuarios):
    for user in usuarios or []:
        for key in ['Docente', 'docente', 'AgenteEducativo', 'agente_educativo', 'agente', 'Responsable', 'Coordinador', 'coordinador']:
            val = _alpha68_valor_usuario(user, key)
            if val:
                return str(val)
    return 'SIN DOCENTE ASIGNADO'


def _alpha68_guardar_workbook_registrado(wb, nombre_archivo, formato, unidad, mes, anio, grupo_etario=None, extra=None):
    nombre_seguro = secure_filename(nombre_archivo)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    ruta = os.path.join(OUTPUT_FOLDER, nombre_seguro)
    wb.save(ruta)
    if os.path.exists(ruta) and os.path.getsize(ruta) > 0:
        try:
            registrar_archivo_generado_alpha57(formato, unidad, nombre_seguro, ruta, mes=mes, anio=anio, grupo_etario=grupo_etario, extra=extra or {})
        except Exception as exc:
            _alpha68_log('REGISTRO_ARCHIVO_ERROR', formato=formato, unidad=unidad, archivo=nombre_seguro, error=str(exc))
    return ruta


def _alpha68_generar_listado_usuarios(unidad, usuarios, mes, anio, options=None):
    # Si la corporación cargó su formato oficial Word, se diligencia esa copia.
    # La salida Excel histórica permanece como fallback compatible.
    try:
        from services.listado_usuarios_docx_service import generate_list, template_info
        data_dir = app.config.get('DATA_DIR')
        info = template_info(data_dir)
        if info.get('existe'):
            nombre_docx = secure_filename(f"{_alpha67_unidad_slug(unidad)}_LISTADO_USUARIOS_{int(anio)}_{int(mes):02d}.docx")
            ruta_docx = os.path.join(OUTPUT_FOLDER, nombre_docx)
            opts = dict(options or {})
            metadata = {
                'unidad': unidad,
                'uca': unidad,
                'mes': mes,
                'anio': anio,
                'tema': opts.get('tema') or '',
                'fecha': opts.get('fecha') or '',
                'hora_inicio': opts.get('hora_inicio') or '',
                'hora_final': opts.get('hora_final') or '',
                'profesional': opts.get('profesional') or _alpha68_docente_por_usuarios(unidad, usuarios),
                'docente': _alpha68_docente_por_usuarios(unidad, usuarios),
                'cargo': opts.get('cargo') or '',
                'modalidad': opts.get('modalidad') or '',
                'servicio': opts.get('servicio') or '',
            }
            os.makedirs(OUTPUT_FOLDER, exist_ok=True)
            generate_list(data_dir, ruta_docx, usuarios or [], metadata=metadata)
            registrar_archivo_generado_alpha57(
                'listado_usuarios', unidad, nombre_docx, ruta_docx,
                mes=mes, anio=anio, extra={'usuarios': len(usuarios or []), 'plantilla_oficial': True},
            )
            return ruta_docx
    except (FileNotFoundError, ModuleNotFoundError):
        pass
    except Exception as exc:
        _alpha68_log('LISTADO_OFICIAL_DOCX_ERROR', unidad=unidad, error=str(exc), traceback=traceback.format_exc())
        raise RuntimeError(f'No se pudo diligenciar el listado oficial Word de {unidad}: {exc}') from exc
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()
    ws = wb.active
    ws.title = 'Listado usuarios'
    headers = ['Documento', 'Tipo documento', 'Nombres', 'Edad', 'Edad meses', 'Grupo etario', 'UDS', 'Docente', 'Coordinador', 'Estado', 'Observaciones']
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='1F2937')
        cell.alignment = Alignment(horizontal='center')
    colores = {
        '0 a 6': 'E0F2FE', '6 a 11': 'DCFCE7', '1 a 2': 'FEF9C3',
        '3 a 5': 'FCE7F3', 'gestante': 'EDE9FE'
    }
    for user in usuarios or []:
        grupo = _alpha68_grupo_usuario(user)
        row = [
            _alpha68_valor_usuario(user, 'Documento', 'documento', 'NUI', 'nui'),
            _alpha68_valor_usuario(user, 'TipoDocumento', 'tipo_documento'),
            _alpha68_usuario_nombre(user),
            _alpha68_valor_usuario(user, 'EdadCompleta', 'edad_completa'),
            _alpha68_valor_usuario(user, 'EdadMeses', 'edad_meses'),
            grupo,
            unidad,
            _alpha68_valor_usuario(user, 'Docente', 'docente', 'AgenteEducativo', 'agente_educativo'),
            _alpha68_valor_usuario(user, 'Coordinador', 'coordinador'),
            _alpha68_valor_usuario(user, 'Estado', 'estado'),
            _alpha68_valor_usuario(user, 'Observacion', 'observacion', 'observaciones'),
        ]
        ws.append(row)
        color = 'FFFFFF'
        grupo_norm = normalizar_texto_clave(grupo)
        for clave, color_hex in colores.items():
            if clave in grupo_norm:
                color = color_hex
                break
        for cell in ws[ws.max_row]:
            cell.fill = PatternFill('solid', fgColor=color)
    for col in range(1, len(headers) + 1):
        ws.column_dimensions[chr(64 + col)].width = 18
    nombre = f"{_alpha67_unidad_slug(unidad)}_LISTADO_USUARIOS_{int(anio)}_{int(mes):02d}.xlsx"
    return _alpha68_guardar_workbook_registrado(wb, nombre, 'listado_usuarios', unidad, mes, anio, extra={'usuarios': len(usuarios or [])})


def _alpha68_generar_relacion_mensual(unidad, usuarios, mes, anio):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    wb = Workbook()
    ws = wb.active
    ws.title = 'Relación mensual'
    docente = _alpha68_docente_por_usuarios(unidad, usuarios)
    total = len(usuarios or [])
    conteos = {}
    for u in usuarios or []:
        grupo = _alpha68_grupo_usuario(u)
        conteos[grupo] = conteos.get(grupo, 0) + 1
    rows = [
        ['UDS', unidad], ['Docente / Agente educativo', docente], ['Mes', mes], ['Año', anio], ['Total usuarios', total],
        [], ['Grupo etario', 'Total']
    ]
    for grupo, val in sorted(conteos.items()):
        rows.append([grupo, val])
    for row in rows:
        ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for cell in ws[7]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='334155')
    ws.column_dimensions['A'].width = 34
    ws.column_dimensions['B'].width = 28
    nombre = f"{_alpha67_unidad_slug(unidad)}_RELACION_MENSUAL_{int(anio)}_{int(mes):02d}.xlsx"
    return _alpha68_guardar_workbook_registrado(wb, nombre, 'relacion_mensual', unidad, mes, anio, extra={'usuarios': total})


def _alpha68_generar_listado_asistencia_usuarios(unidad, usuarios, mes, anio):
    from services.listado_asistencia_usuarios_service import generate_list, template_info
    data_dir = app.config.get('DATA_DIR')
    info = template_info(data_dir)
    if not info.get('existe'):
        raise FileNotFoundError('Carga primero la planilla oficial de asistencia en Plantillas Oficiales.')
    nombre = secure_filename(f"{_alpha67_unidad_slug(unidad)}_LISTADO_ASISTENCIA_USUARIOS_{int(anio)}_{int(mes):02d}.xlsx")
    ruta = os.path.join(OUTPUT_FOLDER, nombre)
    generate_list(data_dir, ruta, usuarios or [], metadata={'unidad': unidad, 'mes': mes, 'anio': anio})
    registrar_archivo_generado_alpha57(
        'listado_asistencia_usuarios', unidad, nombre, ruta, mes=mes, anio=anio,
        extra={'usuarios': len(usuarios or []), 'plantilla_oficial': True, 'mapeo': info.get('mapeo') or {}},
    )
    return ruta


def _alpha68_cargar_config_distribucion():
    ruta = os.path.join(_project_path('backend'), 'config', 'distribucion_alimentos.json')
    os.makedirs(os.path.dirname(ruta), exist_ok=True)
    if not os.path.exists(ruta):
        data = {
            'nota': 'Configure cantidad_por_usuario y unidad_medida por grupo_etario. Cero significa pendiente de configuración.',
            'olla_comunitaria_por_uds': 1,
            'items': [
                {'grupo_etario': '0 A 6 MESES', 'alimento': 'Bienestarina', 'cantidad_por_usuario': 0, 'unidad_medida': 'unidad', 'frecuencia': 'mensual', 'activo': True},
                {'grupo_etario': '6 A 11 MESES', 'alimento': 'Bienestarina', 'cantidad_por_usuario': 0, 'unidad_medida': 'unidad', 'frecuencia': 'mensual', 'activo': True},
                {'grupo_etario': '1 A 2 ANOS 11 MESES', 'alimento': 'Huevos', 'cantidad_por_usuario': 0, 'unidad_medida': 'unidad', 'frecuencia': 'mensual', 'activo': True},
                {'grupo_etario': '3 A 5 ANOS 11 MESES', 'alimento': 'Panela', 'cantidad_por_usuario': 0, 'unidad_medida': 'unidad', 'frecuencia': 'mensual', 'activo': True},
                {'grupo_etario': 'TODOS', 'alimento': 'Olla comunitaria', 'cantidad_por_usuario': 1, 'unidad_medida': 'olla por UDS', 'frecuencia': 'mensual', 'activo': True},
            ]
        }
        with open(ruta, 'w', encoding='utf-8') as fh:
            json.dump(data, fh, ensure_ascii=False, indent=2)
    with open(ruta, 'r', encoding='utf-8') as fh:
        return json.load(fh)


def _alpha68_generar_distribucion_alimentos(unidad, usuarios, mes, anio):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    cfg = _alpha68_cargar_config_distribucion()
    conteos = {}
    for user in usuarios or []:
        grupo = _alpha68_grupo_usuario(user).upper()
        conteos[grupo] = conteos.get(grupo, 0) + 1
    wb = Workbook()
    ws = wb.active
    ws.title = 'Distribución alimentos'
    ws.append(['UDS', unidad])
    ws.append(['Mes/Año', f'{int(mes):02d}/{int(anio)}'])
    ws.append(['Olla comunitaria por UDS', cfg.get('olla_comunitaria_por_uds', 1)])
    ws.append([])
    headers = ['Grupo etario', 'Usuarios', 'Alimento', 'Cantidad por usuario', 'Unidad medida', 'Frecuencia', 'Total calculado', 'Observación']
    ws.append(headers)
    for cell in ws[5]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='14532D')
    for item in cfg.get('items', []):
        if not item.get('activo', True):
            continue
        grupo = str(item.get('grupo_etario') or 'TODOS').upper()
        usuarios_grupo = sum(conteos.values()) if grupo == 'TODOS' else conteos.get(grupo, 0)
        cantidad = float(item.get('cantidad_por_usuario') or 0)
        total = usuarios_grupo * cantidad
        obs = '' if cantidad else 'Cantidad pendiente de configurar'
        ws.append([grupo, usuarios_grupo, item.get('alimento'), cantidad, item.get('unidad_medida'), item.get('frecuencia'), total, obs])
    ws.column_dimensions['A'].width = 28
    ws.column_dimensions['C'].width = 24
    ws.column_dimensions['H'].width = 34
    nombre = f"{_alpha67_unidad_slug(unidad)}_DISTRIBUCION_ALIMENTOS_{int(anio)}_{int(mes):02d}.xlsx"
    return _alpha68_guardar_workbook_registrado(wb, nombre, 'distribucion_alimentos', unidad, mes, anio, extra={'usuarios': len(usuarios or [])})


def _alpha68_generar_complementarios_formato(unidad, usuarios, options, seleccionados, mes, anio):
    # Una selección vacía significa "todos". Ese es el contrato histórico de
    # la pantalla y también el valor normalizado de ``paquete_completo``.
    # Antes se retornaba aquí y el paquete omitía justamente estos tres XLSX.
    claves = set(seleccionados or {
        'listado_usuarios',
        'listado_asistencia_usuarios',
        'relacion_mensual',
        'distribucion_alimentos',
    })
    generados = []
    try:
        if _alpha68_should_generar(claves, 'listado_usuarios'):
            generados.append(_alpha68_generar_listado_usuarios(unidad, usuarios, mes, anio, options))
        if _alpha68_should_generar(claves, 'listado_asistencia_usuarios'):
            try:
                generados.append(_alpha68_generar_listado_asistencia_usuarios(unidad, usuarios, mes, anio))
            except FileNotFoundError:
                # En el paquete historico la nueva plantilla es opcional hasta
                # que la corporacion la cargue. Si el usuario la selecciono de
                # forma expresa, se conserva el error claro de configuracion.
                if seleccionados:
                    raise
        if _alpha68_should_generar(claves, 'relacion_mensual'):
            generados.append(_alpha68_generar_relacion_mensual(unidad, usuarios, mes, anio))
        if _alpha68_should_generar(claves, 'distribucion_alimentos'):
            generados.append(_alpha68_generar_distribucion_alimentos(unidad, usuarios, mes, anio))
        if generados:
            _alpha68_log('COMPLEMENTARIOS_GENERADOS', unidad=unidad, mes=mes, anio=anio, archivos=generados, seleccion=list(claves))
    except Exception as exc:
        _alpha68_log('COMPLEMENTARIOS_ERROR', unidad=unidad, mes=mes, anio=anio, error=str(exc), traceback=traceback.format_exc())
        raise RuntimeError(f'No se pudieron generar los formatos complementarios de {unidad}: {exc}') from exc
    return generados


def _alpha68_festivos_configurados(anio=None, mes=None):
    rutas = [
        os.path.join(_project_path('backend'), 'config', 'festivos.json'),
        os.path.join(_project_path('backend'), 'config', 'calendario_no_laboral.json'),
    ]
    fechas = set()
    for ruta in rutas:
        try:
            if not os.path.exists(ruta):
                continue
            with open(ruta, 'r', encoding='utf-8') as fh:
                data = json.load(fh)
            items = data.get('festivos') if isinstance(data, dict) else data
            for item in items or []:
                fecha = item.get('fecha') if isinstance(item, dict) else item
                activo = item.get('activo', True) if isinstance(item, dict) else True
                if not activo:
                    continue
                fecha = str(fecha or '').strip()[:10]
                if re.match(r'^\d{4}-\d{2}-\d{2}$', fecha):
                    if anio and mes and not fecha.startswith(f'{int(anio):04d}-{int(mes):02d}-'):
                        continue
                    fechas.add(fecha)
        except Exception as exc:
            _alpha68_log('FESTIVOS_CONFIG_ERROR', ruta=ruta, error=str(exc))
    return fechas

def log_alpha56_formato(evento, **datos):
    """Log liviano para diagnosticar descarga Bienestarina y RAM sin afectar el flujo."""
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha56_descarga_bienestarina_ram.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass



def _alpha57_registro_path():
    return os.path.join(OUTPUT_FOLDER, '_registro_formatos_generados.json')


def _alpha57_normalizar_formato_descarga(formato):
    raw = str(formato or '').strip()
    txt = normalizar_texto_clave(raw)
    if 'bienestarina' in txt or 'bienesterina' in txt:
        return 'bienestarina'
    # RAN/RRAN es un formato distinto de RAM. Debe evaluarse antes de buscar
    # "ram", pues de otro modo RRAN contiene el substring RAM/RAN según el alias.
    tokens = set(txt.split())
    if tokens.intersection({'ran', 'rran'}) or raw.lower() in {'ran', 'rran'}:
        return 'ran'
    if 'asistencia' in txt or txt in {'ram', 'rram'}:
        return 'ram'
    if 'rpp' in txt or raw.lower().startswith('rpp_'):
        return raw.lower().replace('-', '_') if raw.lower().startswith('rpp_') else 'rpp'
    return raw.lower()


def _alpha57_safe_join_output(nombre_archivo):
    nombre = secure_filename(os.path.basename(str(nombre_archivo or '')))
    if not nombre:
        return None, None
    ruta = os.path.abspath(os.path.join(OUTPUT_FOLDER, nombre))
    base = os.path.abspath(OUTPUT_FOLDER)
    if not (ruta == base or ruta.startswith(base + os.sep)):
        return None, None
    return nombre, ruta


def _alpha57_leer_registro():
    ruta = _alpha57_registro_path()
    try:
        if os.path.exists(ruta):
            with open(ruta, 'r', encoding='utf-8') as fh:
                data = json.load(fh)
            return data if isinstance(data, list) else []
    except Exception as exc:
        log_alpha56_formato('ALPHA57_REGISTRO_LECTURA_ERROR', error=str(exc))
    return []


def _alpha57_guardar_registro(registros):
    try:
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        with open(_alpha57_registro_path(), 'w', encoding='utf-8') as fh:
            json.dump(registros[-500:], fh, ensure_ascii=False, indent=2, default=str)
    except Exception as exc:
        log_alpha56_formato('ALPHA57_REGISTRO_ESCRITURA_ERROR', error=str(exc))


def registrar_archivo_generado_alpha57(formato, unidad, archivo, ruta_archivo, mes=None, anio=None, grupo_etario=None, estado='generado', extra=None):
    try:
        nombre = os.path.basename(str(archivo or ruta_archivo or ''))
        ruta_abs = os.path.abspath(str(ruta_archivo or os.path.join(OUTPUT_FOLDER, nombre)))
        existe = os.path.exists(ruta_abs)
        item = {
            'formato': _alpha57_normalizar_formato_descarga(formato),
            'formato_original': str(formato or ''),
            'unidad': str(unidad or ''),
            'unidad_normalizada': normalize_unidad(unidad) or str(unidad or '').upper(),
            'mes': int(mes) if str(mes or '').isdigit() else mes,
            'anio': int(anio) if str(anio or '').isdigit() else anio,
            'grupo_etario': grupo_etario,
            'archivo': nombre,
            'ruta_absoluta': ruta_abs,
            'ruta_relativa': os.path.relpath(ruta_abs, _project_path('backend')) if ruta_abs else '',
            'download_url': f"/api/descargar-archivo/{secure_filename(nombre)}" if nombre else '',
            'fecha_generacion': datetime.now().isoformat(timespec='seconds'),
            'existe': bool(existe),
            'tamano_bytes': os.path.getsize(ruta_abs) if existe else 0,
            'estado': estado,
        }
        if isinstance(extra, dict):
            item.update(extra)
        registros = _alpha57_leer_registro()
        registros.append(item)
        _alpha57_guardar_registro(registros)
        log_alpha56_formato('ALPHA57_ARCHIVO_REGISTRADO', **item)
        return item
    except Exception as exc:
        log_alpha56_formato('ALPHA57_REGISTRO_ERROR', formato=formato, unidad=unidad, archivo=archivo, error=str(exc))
        return None


def _alpha57_formato_match(registro_formato, solicitado):
    reg = _alpha57_normalizar_formato_descarga(registro_formato)
    sol = _alpha57_normalizar_formato_descarga(solicitado)
    if sol in {'bienestarina', 'ram', 'ran'}:
        return reg == sol
    if str(sol).startswith('rpp_'):
        return reg == 'rpp' or str(reg).startswith('rpp')
    if sol == 'rpp':
        return reg == 'rpp' or str(reg).startswith('rpp')
    return reg == sol or sol in str(reg)


def buscar_archivo_registrado_alpha57(unidad, formato):
    registros = _alpha57_leer_registro()
    if not registros:
        return None
    unidad_norm = normalize_unidad(unidad) or str(unidad or '').upper()
    unidad_txt = normalizar_texto_clave(unidad_norm)
    formato_sol = _alpha57_normalizar_formato_descarga(formato)
    candidatos = []
    for item in registros:
        if not _alpha57_formato_match(item.get('formato') or item.get('formato_original'), formato_sol):
            continue
        item_unidad = normalize_unidad(item.get('unidad') or item.get('unidad_normalizada')) or str(item.get('unidad') or '').upper()
        item_unidad_txt = normalizar_texto_clave(item_unidad)
        if unidad_txt and item_unidad_txt and unidad_txt != item_unidad_txt and unidad_txt not in item_unidad_txt and item_unidad_txt not in unidad_txt:
            continue
        archivo = item.get('archivo') or ''
        nombre, ruta = _alpha57_safe_join_output(archivo)
        if nombre and ruta and os.path.exists(ruta):
            candidatos.append((os.path.getmtime(ruta), nombre))
    candidatos.sort(reverse=True)
    return candidatos[0][1] if candidatos else None



def buscar_archivo_generado(unidad, formato):
    """Ubica el archivo generado para una unidad y formato.

    Alpha30: para solicitudes RPP por categoría no se permite devolver el
    RPP general/oficial. Si la categoría no existe, se responde 404 para evitar
    entregar un archivo con todos los usuarios mezclados.
    """
    unidad_norm = normalize_unidad(unidad) or str(unidad or '').strip().upper()
    formato_raw = str(formato or '').strip().lower().replace('-', '_')
    formato_lower = normalizar_texto_clave(formato)
    # ALPHA57: los botones históricos enviaban nombres de plantilla
    # (plantilla_bienestarina.xlsx / plantilla_asistencia.xlsx).
    # Esos nombres NO son archivos generados. Se normalizan al formato real.
    if 'plantilla_bienestarina' in formato_raw or 'bienestarina' in formato_lower or 'bienesterina' in formato_lower:
        formato_raw = 'bienestarina'
        formato_lower = 'bienestarina'
    elif 'plantilla_asistencia' in formato_raw or 'plantilla_ram' in formato_raw or 'asistencia' in formato_lower or 'ram' in formato_lower or 'rram' in formato_lower:
        formato_raw = 'ram'
        formato_lower = 'ram'
    elif 'plantilla_rpp' in formato_raw:
        formato_raw = 'rpp'
        formato_lower = 'rpp'

    # ALPHA63 — búsqueda estricta primero.
    # Evita que RPP descargue otro grupo, Bienestarina otra UDS o RAM/RAN otro archivo.
    try:
        grupo_cod = _alpha61_normalizar_grupo_rpp(formato_raw)
        if grupo_cod:
            exacto = _alpha61_buscar_archivo_rpp_exacto(unidad_norm, grupo_cod)
            if exacto:
                return exacto
            return None
        if formato_raw == 'bienestarina' or formato_lower == 'bienestarina':
            exacto = _alpha61_buscar_archivo_bienestarina_exacto(unidad_norm)
            if exacto:
                return exacto
            return None
        if formato_raw == 'ram' or formato_lower == 'ram':
            exacto = _alpha63_buscar_archivo_ram_exacto(unidad_norm)
            if exacto:
                return exacto
            # No se retorna aún: se permite que el flujo de generación intente crear RAM.
    except Exception as exc:
        try:
            log_alpha56_formato('ALPHA63_BUSQUEDA_ESTRICTA_ERROR', unidad=unidad, formato=formato, error=str(exc))
        except Exception:
            pass

    registrado = buscar_archivo_registrado_alpha57(unidad_norm, formato_raw)
    if registrado:
        # ALPHA63: para Bienestarina/RPP por grupo no se aceptan registros amplios.
        if formato_raw not in {'bienestarina'} and not str(formato_raw).startswith('rpp_'):
            return registrado

    def sanitizar_unidad(valor):
        return re.sub(r'[^A-Za-z0-9_]+', '_', str(valor or '').replace(' ', '_')).strip('_').upper()

    unidad_sanitizada = sanitizar_unidad(unidad_norm)
    unidades_validas = {unidad_norm}
    for eq in equivalentes_unidad(unidad_norm):
        unidades_validas.add(normalize_unidad(eq) or str(eq).upper())
    unidades_sanitizadas = {sanitizar_unidad(u) for u in unidades_validas if u}
    unidades_texto = {normalizar_texto_clave(u) for u in unidades_validas if u}

    filtros_alias = {
        'rpp_0_6_gestantes': ['rpp 0 a 6 meses y gestantes', 'rpp 0 6 meses y gestantes', 'rpp nino o nina de 0 6 meses y gestante', '0 a 6 meses y gestantes'],
        'rpp_6_11': ['rpp 6 a 11 meses 29 dias', '6 a 11 meses 29 dias'],
        'rpp_1_2': ['rpp 1 a 2 anos 11 meses', 'ninas y ninos de 1 a 2 anos 11 meses', '1 a 2 anos 11 meses'],
        'rpp_3_5': ['rpp 3 a 5 anos 11 meses', 'ninas y ninos de 3 a 5 anos 11 meses', '3 a 5 anos 11 meses'],
        'rpp_general': ['rpp oficial', 'rpp ofical', 'rpp 2026', 'plantilla rpp', 'rpp'],
        'asistencia': ['asistencia', 'ram', 'rram', 'registro asistencia mensual', 'registro de asistencia mensual'],
        'ram': ['ram', 'rram', 'registro asistencia mensual', 'registro de asistencia mensual', 'asistencia'],
        'bienestarina': ['bienestarina', 'bienesterina'],
        'ran': ['ran', 'ram', 'asistencia']
    }

    solicitud_categoria_rpp = formato_raw in {'rpp_0_6_gestantes', 'rpp_6_11', 'rpp_1_2', 'rpp_3_5'}

    if formato_raw in filtros_alias:
        filtros = filtros_alias[formato_raw]
    elif '0 6' in formato_lower or '0 a 6' in formato_lower or 'gestante' in formato_lower:
        filtros = filtros_alias['rpp_0_6_gestantes']
        solicitud_categoria_rpp = True
    elif '6 a 11' in formato_lower:
        filtros = filtros_alias['rpp_6_11']
        solicitud_categoria_rpp = True
    elif '1 a 2' in formato_lower:
        filtros = filtros_alias['rpp_1_2']
        solicitud_categoria_rpp = True
    elif '3 a 5' in formato_lower:
        filtros = filtros_alias['rpp_3_5']
        solicitud_categoria_rpp = True
    elif 'ram' in formato_lower or 'rram' in formato_lower or 'registro asistencia' in formato_lower:
        filtros = filtros_alias['ram']
    elif 'asistencia' in formato_lower:
        filtros = filtros_alias['asistencia']
    elif 'bienestarina' in formato_lower or 'bienesterina' in formato_lower:
        filtros = filtros_alias['bienestarina']
    elif 'rpp' in formato_lower:
        filtros = filtros_alias['rpp_general']
    elif 'ran' in formato_lower:
        filtros = filtros_alias['ran']
    else:
        filtros = [formato_lower] if formato_lower else []

    def nombre_corresponde_unidad(nombre):
        nombre_upper = nombre.upper()
        if any(nombre_upper.startswith(f"{u}_") for u in unidades_sanitizadas if u):
            return True
        nombre_txt = normalizar_texto_clave(nombre)
        if any(u and u in nombre_txt for u in unidades_texto):
            return True
        # ALPHA56 — búsqueda tolerante: algunas UDS se consultan como
        # "UNIDAD DEMO 02" pero el archivo se nombra con descripción/dirección.
        tokens_unidad = {tok for txt in unidades_texto for tok in str(txt).split() if len(tok) >= 4}
        if tokens_unidad and any(tok in nombre_txt for tok in tokens_unidad):
            return True
        return False

    def nombre_corresponde_formato(nombre, filtros_busqueda):
        nombre_norm = normalizar_texto_clave(nombre)
        return (not filtros_busqueda) or any(normalizar_texto_clave(f) in nombre_norm for f in filtros_busqueda)

    def buscar(filtros_busqueda):
        candidatos = []
        if not os.path.isdir(OUTPUT_FOLDER):
            return []
        for nombre in os.listdir(OUTPUT_FOLDER):
            if not nombre.lower().endswith(('.xlsx', '.xls', '.xlsm', '.pdf', '.docx')):
                continue
            if not nombre_corresponde_unidad(nombre):
                continue
            if not nombre_corresponde_formato(nombre, filtros_busqueda):
                continue
            candidatos.append(nombre)
        candidatos.sort(key=lambda n: os.path.getmtime(os.path.join(OUTPUT_FOLDER, n)), reverse=True)
        return candidatos

    candidatos = buscar(filtros)

    # ALPHA30: no devolver RPP general cuando se pidió una categoría.
    # Si no existe archivo de la categoría, es preferible avisar 404 y obligar
    # a regenerar formatos antes que entregar una hoja con todos los usuarios.
    if not candidatos and (formato_raw.startswith('rpp') or 'rpp' in formato_lower) and not solicitud_categoria_rpp:
        candidatos = buscar(['rpp oficial', 'rpp'])

    # Fallback para RAM/RAN: algunos botones llaman RAN y el archivo se nombra RAM_ASISTENCIA.
    if not candidatos and ('ran' in formato_lower or formato_raw == 'ran'):
        candidatos = buscar(['ram asistencia', 'asistencia', 'ram', 'ran'])

    if not candidatos:
        # ALPHA63: prohibido devolver Bienestarina de otra UDS o RPP de otro grupo.
        # Si no hay coincidencia exacta, el endpoint debe generar el archivo correcto
        # o devolver error controlado. Nunca se usa el último archivo global.
        return None
    return candidatos[0]



# -----------------------------------------------------------------------------
# ALPHA61: validación estricta de UDS + grupo para descargas RPP/Bienestarina
# -----------------------------------------------------------------------------
GRUPOS_RPP_ALPHA61 = {
    '0_6_GESTANTES': {
        'legacy': 'rpp_0_6_gestantes',
        'nombre': '0 A 6 MESES Y GESTANTES',
        'archivo_tag': 'RPP_0_A_6_MESES_Y_GESTANTES',
        'min_meses': 0,
        'max_meses': 6,
        'incluye_gestantes': True,
        'aliases': {'0_6_GESTANTES', 'rpp_0_6_gestantes', 'rpp_0_6', '0_6', '0 a 6', '0 a 6 meses', 'gestantes'},
    },
    '6_11_MESES': {
        'legacy': 'rpp_6_11',
        'nombre': '6 A 11 MESES 29 DIAS',
        'archivo_tag': 'RPP_6_A_11_MESES_29_DIAS',
        'min_meses': 7,
        'max_meses': 11,
        'incluye_gestantes': False,
        'aliases': {'6_11_MESES', 'rpp_6_11', '6_11', '6 a 11', '6 a 11 meses'},
    },
    '1_2_ANOS': {
        'legacy': 'rpp_1_2',
        'nombre': '1 A 2 ANOS 11 MESES',
        'archivo_tag': 'RPP_1_A_2_ANOS_11_MESES',
        'min_meses': 12,
        'max_meses': 35,
        'incluye_gestantes': False,
        'aliases': {'1_2_ANOS', 'rpp_1_2', '1_2', '1 a 2', '1 a 2 anos', '1 a 2 años'},
    },
    '3_5_ANOS': {
        'legacy': 'rpp_3_5',
        'nombre': '3 A 5 ANOS 11 MESES',
        'archivo_tag': 'RPP_3_A_5_ANOS_11_MESES',
        'min_meses': 36,
        'max_meses': 71,
        'incluye_gestantes': False,
        'aliases': {'3_5_ANOS', 'rpp_3_5', '3_5', '3 a 5', '3 a 5 anos', '3 a 5 años'},
    },
}


def _alpha61_log(evento, **datos):
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha61_fix_mapeo_uds_grupo_descargas.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha61_slug_archivo(valor):
    base = normalizar_texto_clave(valor or '').upper().replace(' ', '_')
    return re.sub(r'[^A-Z0-9_]+', '_', base).strip('_')


def _alpha61_normalizar_grupo_rpp(grupo):
    raw = str(grupo or '').strip()
    raw_norm = raw.lower().replace('-', '_').replace(' ', '_')
    txt = normalizar_texto_clave(raw)
    for codigo, cfg in GRUPOS_RPP_ALPHA61.items():
        aliases = set(cfg.get('aliases') or set()) | {codigo, cfg.get('legacy', ''), cfg.get('nombre', ''), cfg.get('archivo_tag', '')}
        for alias in aliases:
            alias_raw = str(alias or '')
            if not alias_raw:
                continue
            if raw_norm == alias_raw.lower().replace('-', '_').replace(' ', '_'):
                return codigo
            alias_txt = normalizar_texto_clave(alias_raw)
            if alias_txt and (txt == alias_txt or alias_txt in txt or txt in alias_txt):
                # Evitar que "RPP" solo haga match con todo.
                if len(alias_txt) >= 4:
                    return codigo
    if '0 6' in txt or '0 a 6' in txt or 'gestante' in txt:
        return '0_6_GESTANTES'
    if '6 11' in txt or '6 a 11' in txt:
        return '6_11_MESES'
    if '1 2' in txt or '1 a 2' in txt:
        return '1_2_ANOS'
    if '3 5' in txt or '3 a 5' in txt:
        return '3_5_ANOS'
    return ''


def _alpha61_registro_candidatos(unidad, formato, grupo=None):
    registros = _alpha57_leer_registro()
    unidad_norm = normalize_unidad(unidad) or str(unidad or '').upper()
    unidad_slug = _alpha61_slug_archivo(unidad_norm)
    formato_norm = _alpha57_normalizar_formato_descarga(formato)
    grupo_cod = _alpha61_normalizar_grupo_rpp(grupo or formato) if (formato_norm == 'rpp' or str(formato or '').lower().startswith('rpp') or grupo) else ''
    candidatos = []
    for item in registros:
        item_formato = _alpha57_normalizar_formato_descarga(item.get('formato') or item.get('formato_original'))
        if formato_norm == 'bienestarina':
            if item_formato != 'bienestarina':
                continue
        elif grupo_cod:
            if item_formato != 'rpp' and not str(item_formato).startswith('rpp'):
                continue
        else:
            if item_formato != formato_norm:
                continue
        item_unidad_norm = normalize_unidad(item.get('unidad') or item.get('unidad_normalizada')) or str(item.get('unidad') or '').upper()
        if _alpha61_slug_archivo(item_unidad_norm) != unidad_slug:
            continue
        if grupo_cod:
            item_grupo = _alpha61_normalizar_grupo_rpp(item.get('grupo') or item.get('grupo_etario') or item.get('formato_original') or item.get('archivo'))
            if item_grupo and item_grupo != grupo_cod:
                continue
            # Si el registro viejo no tiene grupo, exigir que el nombre de archivo tenga el tag exacto.
            tag = GRUPOS_RPP_ALPHA61[grupo_cod]['archivo_tag']
            if not _alpha61_archivo_tiene_grupo_rpp(item.get('archivo') or '', grupo_cod):
                continue
        archivo = item.get('archivo') or ''
        nombre, ruta = _alpha57_safe_join_output(archivo)
        if nombre and ruta and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
            candidatos.append((os.path.getmtime(ruta), nombre, 'registro'))
    candidatos.sort(reverse=True)
    return candidatos


def _alpha61_archivo_valido_por_uds(nombre, unidad):
    unidad_norm = normalize_unidad(unidad) or str(unidad or '').upper()
    unidad_slug = _alpha61_slug_archivo(unidad_norm)
    nombre_slug = _alpha61_slug_archivo(nombre)
    return bool(unidad_slug and (nombre_slug.startswith(unidad_slug + '_') or ('_' + unidad_slug + '_') in ('_' + nombre_slug + '_')))


def _alpha61_archivo_tiene_grupo_rpp(nombre, grupo_cod):
    if grupo_cod not in GRUPOS_RPP_ALPHA61:
        return False
    nombre_slug = _alpha61_slug_archivo(nombre)
    cfg = GRUPOS_RPP_ALPHA61[grupo_cod]
    tags = {
        _alpha61_slug_archivo(cfg.get('archivo_tag')),
        _alpha61_slug_archivo(cfg.get('legacy')),
        _alpha61_slug_archivo(cfg.get('nombre')),
    }
    tags = {t for t in tags if t}
    return any(t in nombre_slug for t in tags)


def _alpha61_buscar_archivo_rpp_exacto(unidad, grupo):
    grupo_cod = _alpha61_normalizar_grupo_rpp(grupo)
    if not grupo_cod:
        return None
    candidatos = _alpha61_registro_candidatos(unidad, 'rpp', grupo_cod)
    if candidatos:
        return candidatos[0][1]
    tag_slug = _alpha61_slug_archivo(GRUPOS_RPP_ALPHA61[grupo_cod]['archivo_tag'])
    candidatos_fs = []
    if os.path.isdir(OUTPUT_FOLDER):
        for nombre in os.listdir(OUTPUT_FOLDER):
            if not nombre.lower().endswith(('.xlsx', '.xlsm', '.xls')):
                continue
            nombre_slug = _alpha61_slug_archivo(nombre)
            if 'RPP' not in nombre_slug:
                continue
            if not _alpha61_archivo_tiene_grupo_rpp(nombre, grupo_cod):
                continue
            if not _alpha61_archivo_valido_por_uds(nombre, unidad):
                continue
            ruta = os.path.join(OUTPUT_FOLDER, nombre)
            if os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                candidatos_fs.append((os.path.getmtime(ruta), nombre, 'fs'))
    candidatos_fs.sort(reverse=True)
    return candidatos_fs[0][1] if candidatos_fs else None


def _alpha61_buscar_archivo_bienestarina_exacto(unidad):
    candidatos = _alpha61_registro_candidatos(unidad, 'bienestarina')
    if candidatos:
        return candidatos[0][1]
    candidatos_fs = []
    if os.path.isdir(OUTPUT_FOLDER):
        for nombre in os.listdir(OUTPUT_FOLDER):
            if not nombre.lower().endswith(('.xlsx', '.xlsm', '.xls')):
                continue
            nombre_slug = _alpha61_slug_archivo(nombre)
            if 'BIENESTARINA' not in nombre_slug:
                continue
            if 'PLANTILLA' in nombre_slug or 'OFICIAL' in nombre_slug:
                continue
            if not _alpha61_archivo_valido_por_uds(nombre, unidad):
                continue
            ruta = os.path.join(OUTPUT_FOLDER, nombre)
            if os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                candidatos_fs.append((os.path.getmtime(ruta), nombre, 'fs'))
    candidatos_fs.sort(reverse=True)
    return candidatos_fs[0][1] if candidatos_fs else None


def _alpha63_buscar_archivo_ram_exacto(unidad):
    """Busca RAM/RAN/asistencia únicamente para la UDS solicitada.

    No devuelve archivos de otra UDS, no devuelve plantillas y mantiene compatibilidad
    con nombres históricos de archivos por UDS.
    """
    candidatos = _alpha61_registro_candidatos(unidad, 'ram')
    if candidatos:
        return candidatos[0][1]
    candidatos_fs = []
    if os.path.isdir(OUTPUT_FOLDER):
        for nombre in os.listdir(OUTPUT_FOLDER):
            if not nombre.lower().endswith(('.xlsx', '.xlsm', '.xls', '.pdf')):
                continue
            nombre_slug = _alpha61_slug_archivo(nombre)
            if 'PLANTILLA' in nombre_slug or 'OFICIAL' in nombre_slug:
                continue
            if not any(k in nombre_slug for k in ['RAM', 'RAN', 'RRAN', 'ASISTENCIA']):
                continue
            if not _alpha61_archivo_valido_por_uds(nombre, unidad):
                continue
            ruta = os.path.join(OUTPUT_FOLDER, nombre)
            if os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                candidatos_fs.append((os.path.getmtime(ruta), nombre, 'fs'))
    candidatos_fs.sort(reverse=True)
    return candidatos_fs[0][1] if candidatos_fs else None


def _alpha63_validar_archivo_descarga(nombre_archivo, unidad, formato, grupo=None):
    """Validación común antes de descargar formatos oficiales."""
    nombre, ruta = _alpha57_safe_join_output(nombre_archivo)
    if not nombre or not ruta or not os.path.exists(ruta) or os.path.getsize(ruta) <= 0:
        return False, 'archivo_no_existe_o_vacio'
    slug = _alpha61_slug_archivo(nombre)
    formato_norm = _alpha57_normalizar_formato_descarga(formato)
    if 'PLANTILLA' in slug or 'OFICIAL' in slug:
        return False, 'archivo_es_plantilla'
    if not _alpha61_archivo_valido_por_uds(nombre, unidad):
        return False, 'archivo_de_otra_uds'
    if formato_norm == 'bienestarina':
        if 'BIENESTARINA' not in slug:
            return False, 'no_es_bienestarina'
    elif formato_norm == 'ram':
        if not any(k in slug for k in ['RAM', 'RAN', 'RRAN', 'ASISTENCIA']):
            return False, 'no_es_ram_ran'
    elif formato_norm == 'rpp' or _alpha61_normalizar_grupo_rpp(grupo or formato):
        if 'RPP' not in slug:
            return False, 'no_es_rpp'
        grupo_cod = _alpha61_normalizar_grupo_rpp(grupo or formato)
        if grupo_cod and not _alpha61_archivo_tiene_grupo_rpp(nombre, grupo_cod):
            return False, 'rpp_grupo_incorrecto'
    return True, 'ok'


def _alpha61_generar_formatos_unidad(unidad):
    usuarios = _alpha59_obtener_usuarios_unidad(unidad)
    if not usuarios:
        _alpha61_log('GENERAR_FORMATOS_SIN_USUARIOS', unidad=unidad)
        return False
    try:
        inyectar_datos_en_plantillas(
            unidad,
            usuarios,
            options={'mes': datetime.now().month, 'anio': datetime.now().year, 'año': datetime.now().year}
        )
        _alpha61_log('GENERAR_FORMATOS_OK', unidad=unidad, total_usuarios=len(usuarios))
        return True
    except Exception as exc:
        _alpha61_log('GENERAR_FORMATOS_ERROR', unidad=unidad, error=str(exc), traceback=traceback.format_exc())
        return False


def _alpha61_generar_rpp_grupo(unidad, grupo):
    grupo_cod = _alpha61_normalizar_grupo_rpp(grupo)
    if not grupo_cod:
        return None
    nombre = _alpha61_buscar_archivo_rpp_exacto(unidad, grupo_cod)
    if nombre:
        return nombre
    # Generación directa: nunca producir todo el paquete de la UDS para descargar
    # un único grupo RPP. Esto mantiene rápidas las fundaciones nuevas, que aún no
    # tienen artefactos previamente generados en su carpeta tenant.
    legacy = GRUPOS_RPP_ALPHA61[grupo_cod]['legacy']
    try:
        _alpha59_generar_oficial_desde_template('rpp', unidad, grupo=legacy)
    except Exception as exc:
        _alpha61_log('RPP_OFICIAL_RESPALDO_ERROR', unidad=unidad, grupo=grupo_cod, error=str(exc))
    return _alpha61_buscar_archivo_rpp_exacto(unidad, grupo_cod)


def _alpha61_generar_bienestarina(unidad):
    # ALPHA62: Bienestarina se genera de forma directa y liviana.
    # Antes se intentaba generar todo el paquete de formatos de la UDS y eso podía
    # dejar la interfaz lenta o devolver JSON en el navegador. Esta función solo
    # toca Bienestarina y valida UDS estrictamente.
    return _alpha62_generar_bienestarina_solo(unidad)


def _alpha62_generar_bienestarina_solo(unidad):
    """Genera únicamente Bienestarina para la UDS solicitada.

    No ejecuta generación completa de RPP/RAM/RAN. No devuelve archivos de otra
    UDS. No devuelve plantillas. Se usa únicamente desde el endpoint de descarga
    de Bienestarina.
    """
    try:
        nombre = _alpha60_generar_bienestarina_directa(unidad)
        if nombre and _alpha62_es_bienestarina_valida_para_unidad(nombre, unidad):
            _alpha62_log_bienestarina('BIENESTARINA_GENERADA_DIRECTA_OK', unidad=unidad, archivo=nombre)
            return nombre
        _alpha62_log_bienestarina('BIENESTARINA_GENERADA_DIRECTA_INVALIDA', unidad=unidad, archivo=nombre)
    except Exception as exc:
        _alpha62_log_bienestarina('BIENESTARINA_GENERADA_DIRECTA_ERROR', unidad=unidad, error=str(exc), traceback=traceback.format_exc())
    return None


def _alpha62_log_bienestarina(evento, **datos):
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha62_solo_bienestarina_descarga.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha62_es_bienestarina_valida_para_unidad(nombre_archivo, unidad):
    try:
        if not nombre_archivo:
            return False
        nombre = os.path.basename(str(nombre_archivo))
        nombre_slug = _alpha61_slug_archivo(nombre)
        if not nombre.lower().endswith(('.xlsx', '.xlsm', '.xls')):
            return False
        if 'BIENESTARINA' not in nombre_slug:
            return False
        if 'PLANTILLA' in nombre_slug or 'OFICIAL' in nombre_slug:
            return False
        if not _alpha61_archivo_valido_por_uds(nombre, unidad):
            return False
        ruta = os.path.abspath(os.path.join(OUTPUT_FOLDER, secure_filename(nombre)))
        base = os.path.abspath(OUTPUT_FOLDER)
        if not (ruta == base or ruta.startswith(base + os.sep)):
            return False
        return os.path.exists(ruta) and os.path.getsize(ruta) > 0
    except Exception:
        return False


def _alpha62_ultimos_archivos_bienestarina():
    try:
        if not os.path.isdir(OUTPUT_FOLDER):
            return []
        archivos = [n for n in os.listdir(OUTPUT_FOLDER) if n.lower().endswith(('.xlsx','.xls','.xlsm','.pdf'))]
        return sorted(archivos, key=lambda n: os.path.getmtime(os.path.join(OUTPUT_FOLDER, n)), reverse=True)[:10]
    except Exception:
        return []


def _alpha62_respuesta_error_bienestarina(unidad, mensaje, status=404, extra=None):
    payload = {
        'ok': False,
        'formato': 'bienestarina',
        'error': 'No se pudo descargar la Bienestarina solicitada.',
        'mensaje': mensaje,
        'unidad': unidad,
        'unidad_normalizada': normalize_unidad(unidad),
        'output_folder': os.fspath(OUTPUT_FOLDER),
        'ultimos_archivos_generados': _alpha62_ultimos_archivos_bienestarina(),
    }
    if isinstance(extra, dict):
        payload.update(extra)
    _alpha62_log_bienestarina('BIENESTARINA_RESPUESTA_ERROR', **payload)
    return jsonify(payload), status

# =============================================================
# ALPHA59 — Fallback real de generación para descargas de formatos
# Objetivo: si Bienestarina/RPP no existen al descargar, generar el archivo
# final desde la ruta real o desde plantilla oficial; nunca descargar plantilla.
# =============================================================

def _alpha59_slug(valor):
    try:
        base = normalizar_texto_clave(valor).upper().replace(' ', '_')
    except Exception:
        base = str(valor or '').upper().replace(' ', '_')
    base = re.sub(r'[^A-Z0-9_]+', '_', base).strip('_')
    return base or 'SIN_UNIDAD'


def _alpha59_row_to_dict(row):
    try:
        return dict(row)
    except Exception:
        try:
            return {k: row[k] for k in row.keys()}
        except Exception:
            return {}


def _alpha59_split_nombre_completo(nombre):
    partes = [p for p in str(nombre or '').strip().split() if p]
    return {
        'primer_nombre': partes[0] if len(partes) > 0 else '',
        'segundo_nombre': partes[1] if len(partes) > 1 else '',
        'primer_apellido': partes[-2] if len(partes) > 2 else (partes[1] if len(partes) > 1 else ''),
        'segundo_apellido': partes[-1] if len(partes) > 3 else '',
    }


def _alpha59_usuario_normalizado(row, unidad_consulta=''):
    data = _alpha59_row_to_dict(row)

    # La consolidación conserva el registro documental completo dentro de
    # datos_json (en algunas versiones hay un segundo datos_json anidado).
    # Promover esos campos evita perder acudiente, contacto y encabezados al
    # generar formatos desde master_ninos.
    documental = {}
    pendiente = [data.get('datos_json')]
    for _nivel in range(3):
        if not pendiente:
            break
        raw_json = pendiente.pop(0)
        if not raw_json:
            continue
        try:
            parsed = json.loads(raw_json) if isinstance(raw_json, str) else dict(raw_json)
        except Exception:
            parsed = {}
        if not isinstance(parsed, dict):
            continue
        for clave, valor in parsed.items():
            if clave != 'datos_json' and valor not in (None, ''):
                documental.setdefault(clave, valor)
        if parsed.get('datos_json'):
            pendiente.append(parsed.get('datos_json'))

    def dato(*claves, default=''):
        for clave in claves:
            valor = data.get(clave)
            if valor not in (None, ''):
                return valor
            valor = documental.get(clave)
            if valor not in (None, ''):
                return valor
        return default

    nombres = (
        dato('nombres', 'nombre_completo', 'nombre', 'Nombre') or ''
    )
    apellidos = dato('apellidos') or ''
    primer_nombre_doc = dato('primer_nombre', 'primer_nombre_del_beneficiario')
    segundo_nombre_doc = dato('segundo_nombre', 'segundo_nombre_del_beneficiario')
    primer_apellido_doc = dato('primer_apellido', 'primer_apellido_del_beneficiario')
    segundo_apellido_doc = dato('segundo_apellido', 'segundo_apellido_del_beneficiario')
    if primer_nombre_doc or primer_apellido_doc:
        primer_nombre = primer_nombre_doc or ''
        segundo_nombre = segundo_nombre_doc or ''
        primer_apellido = primer_apellido_doc or ''
        segundo_apellido = segundo_apellido_doc or ''
    else:
        split = _alpha59_split_nombre_completo(f'{nombres} {apellidos}'.strip())
        primer_nombre = split['primer_nombre']
        segundo_nombre = split['segundo_nombre']
        primer_apellido = split['primer_apellido']
        segundo_apellido = split['segundo_apellido']
    documento = dato('documento', 'nui', 'NUI', 'documento_del_beneficiario') or ''
    unidad = (
        dato('unidad', 'unidad_servicio', 'nombre_unidad', 'unidad_atencion',
             'uds', 'nombre_de_la_unidad_de_servicio') or unidad_consulta or ''
    )
    edad_meses = data.get('edad_meses')
    if edad_meses in (None, ''):
        fecha_nacimiento = dato('fecha_nacimiento', 'FechaNacimiento', 'fechaNacimiento', 'fecha_de_nacimiento_del_beneficiario')
        if fecha_nacimiento:
            try:
                edad_meses = calcular_edad_meses(fecha_nacimiento)
            except Exception:
                edad_meses = None
    item = dict(data)
    # Alias institucionales utilizados por los encabezados de RPP,
    # Bienestarina y RAN/RAM.
    item.update({
        'regional': dato('regional', 'regional_del_contrato', 'nombre_de_la_regional_de_la_unidad_de_servicio'),
        'centro_zonal': dato('centro_zonal', 'nombre_del_centro_zonal'),
        'municipio': dato('municipio', 'nombre_municipio_de_la_unidad_de_servicio'),
        'contrato': dato('contrato', 'numero_contrato', 'numero_del_contrato'),
        'nombre_eas': dato('nombre_eas', 'nombre_de_la_entidad_contratista'),
        'codigo_unidad': dato('codigo_unidad', 'codigo_unidad_servicio', 'codigo_de_la_unidad_de_servicio'),
        'modalidad': dato('modalidad'),
        'telefono': dato('telefono', 'celular', 'telefono_del_beneficiario'),
        'documento_acudiente': dato('documento_acudiente', 'numero_de_documento_del_acudiente_o_responsable'),
        'tipo_documento_acudiente': dato('tipo_documento_acudiente', 'tipo_de_documento_del_acudiente_o_responsable'),
        'parentesco': dato('parentesco', 'tipo_de_responsable'),
    })
    acudiente = ' '.join(str(x).strip() for x in [
        dato('primer_nombre_acudiente', 'primer_nombre_del_acudiente_o_responsable'),
        dato('segundo_nombre_acudiente', 'segundo_nombre_del_acudiente_o_responsable'),
        dato('primer_apellido_acudiente', 'primer_apellido_del_acudiente_o_responsable'),
        dato('segundo_apellido_acudiente', 'segundo_apellido_del_acudiente_o_responsable'),
    ] if str(x or '').strip())
    item['nombre_acudiente'] = dato('nombre_acudiente') or acudiente
    item.update({
        'PrimerNombre': primer_nombre,
        'SegundoNombre': segundo_nombre,
        'PrimerApellido': primer_apellido,
        'SegundoApellido': segundo_apellido,
        'Nombre': ' '.join([x for x in [primer_nombre, segundo_nombre, primer_apellido, segundo_apellido] if x]).strip(),
        'Documento': documento,
        'NUI': data.get('nui') or documento,
        'TipoDocumento': dato('tipo_documento', 'TipoDocumento', 'tipo_de_documento_del_beneficiario'),
        # No convertir un dato ausente en cero: eso clasificaba a todos los niños
        # de Base Maestra como menores de seis meses e impedía generar los otros RPP.
        'EdadMeses': edad_meses,
        'GrupoEdad': dato('grupo_edad', 'grupo_etario', 'nombre_tipo_de_beneficiario'),
        'TipoBeneficiario': dato('tipo_beneficiario', 'nombre_tipo_de_beneficiario'),
        'FechaNacimiento': dato('fecha_nacimiento', 'FechaNacimiento', 'fecha_de_nacimiento_del_beneficiario'),
        'FechaIngreso': dato('fecha_ingreso', 'FechaIngreso', 'fecha_de_atencion_del_beneficiario_a_la_uds'),
        'FechaRetiro': dato('fecha_retiro', 'FechaRetiro'),
        'MotivoRetiro': dato('motivo_retiro', 'MotivoRetiro'),
        'Unidad': unidad,
        'Telefono': dato('telefono', 'celular', 'telefono_del_beneficiario'),
        'Acudiente': dato('nombre_acudiente') or acudiente,
        'DocumentoAcudiente': dato('documento_acudiente', 'numero_de_documento_del_acudiente_o_responsable'),
        'TipoDocumentoAcudiente': dato('tipo_documento_acudiente', 'tipo_de_documento_del_acudiente_o_responsable'),
        'Parentesco': dato('parentesco', 'tipo_de_responsable'),
    })
    return item


def _alpha59_documento_key(user):
    raw = str(user.get('Documento') or user.get('documento') or user.get('NUI') or '').strip()
    doc = re.sub(r'\D+', '', raw.replace('.0', ''))
    if doc:
        return doc
    nombre = normalizar_texto_clave(user.get('Nombre') or user.get('nombre') or '')
    fnac = str(user.get('fecha_nacimiento') or '').strip()
    return f'{nombre}|{fnac}'


def _alpha59_deduplicar_usuarios(usuarios):
    vistos = set()
    salida = []
    for user in usuarios or []:
        key = _alpha59_documento_key(user)
        if key and key in vistos:
            continue
        if key:
            vistos.add(key)
        salida.append(user)
    return salida


def _alpha59_obtener_usuarios_unidad(unidad):
    """Obtiene la población consolidada de una UDS.

    ``master_ninos`` es la fuente autoritativa. Las tablas históricas se usan
    únicamente cuando la Base Maestra no contiene población para la UDS; así
    evitamos mezclar versiones antiguas o contar dos veces al mismo niño.
    """
    unidad_norm = normalize_unidad(unidad) or str(unidad or '').upper()
    unidad_txt = normalizar_texto_clave(unidad_norm)
    tenant_id = fundacion_actual_id()
    resultado = []
    conn = None

    def pertenece_unidad(data):
        unidad_fila = (
            data.get('unidad_servicio') or data.get('unidad') or data.get('Unidad') or
            data.get('nombre_unidad') or data.get('unidad_atencion') or data.get('uds') or ''
        )
        fila_norm = normalize_unidad(unidad_fila) or str(unidad_fila or '').upper()
        fila_txt = normalizar_texto_clave(fila_norm)
        return bool(
            unidad_txt and fila_txt and
            (unidad_txt == fila_txt or unidad_txt in fila_txt or fila_txt in unidad_txt)
        )

    def fila_activa(data, fuente):
        try:
            fid = data.get('fundacion_id')
            if fid not in (None, '') and int(fid) != int(tenant_id):
                return False
        except (TypeError, ValueError):
            return False
        if fuente == 'master_ninos' and data.get('activo') not in (None, '', 1, True, '1', 'true', 'TRUE'):
            return False
        estado = normalizar_texto_clave(data.get('estado') or '')
        return not any(token in estado for token in ('retir', 'fallec', 'inactiv'))

    try:
        conn = get_db_connection()

        # Primero la fuente consolidada; solo si está vacía se consultan legados.
        for fuentes in (('master_ninos',), ('usuarios', 'beneficiarios')):
            usuarios = []
            for tabla in fuentes:
                # No usar PRAGMA/sqlite_master aquí: la plataforma opera sobre
                # PostgreSQL y esas comprobaciones abortan la transacción antes de
                # poder leer la Base Maestra. La consulta directa es compatible
                # con ambos motores; las instalaciones vigentes siempre crean
                # estas tablas mediante las migraciones de arranque.
                try:
                    rows = conn.execute(f'SELECT * FROM {tabla}').fetchall()
                except Exception:
                    rows = []
                for row in rows:
                    data = _alpha59_row_to_dict(row)
                    if fila_activa(data, tabla) and pertenece_unidad(data):
                        usuarios.append(_alpha59_usuario_normalizado(data, unidad))
            usuarios = _alpha59_deduplicar_usuarios(usuarios)
            if usuarios:
                resultado = usuarios
                break
    except Exception as exc:
        log_alpha56_formato('ALPHA59_USUARIOS_UNIDAD_ERROR', unidad=unidad, error=str(exc))
    finally:
        try:
            if conn is not None:
                conn.close()
        except Exception:
            pass
    return resultado



@app.route('/api/formatos/diagnostico', methods=['GET'])
def formatos_diagnostico_previo():
    """Preflight sin datos personales para RPP, Bienestarina y RAM."""
    unidad_solicitada = str(request.args.get('unidad') or '').strip()
    try:
        mes = max(1, min(12, int(request.args.get('mes') or datetime.now().month)))
        anio = max(2020, min(2100, int(request.args.get('anio') or request.args.get('año') or datetime.now().year)))
    except (TypeError, ValueError):
        return jsonify({'ok': False, 'error': 'Mes o año inválido.'}), 400

    unidad_normalizada = normalize_unidad(unidad_solicitada)
    unidades_conocidas = set(uds_canonical_units())
    unidad_conocida = bool(unidad_normalizada and unidad_normalizada in unidades_conocidas)
    usuarios = _alpha59_obtener_usuarios_unidad(unidad_normalizada) if unidad_normalizada else []

    disponibles = {}
    try:
        entries = iter_plantillas_oficiales_para_generacion(TEMPLATES_FOLDER, mes=mes, anio=anio)
    except Exception as exc:
        entries = []
        plantillas_error = str(exc)
    else:
        plantillas_error = ''
    for tipo in ('rpp', 'bienestarina', 'ram'):
        entry = next((item for item in entries if item.get('tipo') == tipo), None)
        disponibles[tipo] = {
            'disponible': bool(entry and entry.get('ruta') and os.path.exists(entry.get('ruta'))),
            'version': entry.get('version') if entry else None,
            'archivo': os.path.basename(str(entry.get('ruta') or '')) if entry else None,
            'source': entry.get('source') if entry else None,
            'fechaVigencia': entry.get('fecha_vigencia') if entry else None,
            'fechaVigenciaFin': entry.get('fecha_vigencia_fin') if entry else None,
        }

    try:
        from services.rpp_minutas_service import obtener_minuta_vigente
        minuta = obtener_minuta_vigente(DATABASE_PATH, mes=mes, anio=anio)
    except Exception as exc:
        minuta = None
        minuta_error = str(exc)
    else:
        minuta_error = ''

    minuta_mes = int(minuta.get('mes') or 0) if minuta else None
    minuta_anio = int(minuta.get('anio') or 0) if minuta else None
    minuta_aplicable = bool(minuta)
    grupos = len(minuta.get('grupos') or []) if minuta else 0
    productos = sum(len(grupo.get('productos') or []) for grupo in (minuta.get('grupos') or [])) if minuta else 0

    razones = []
    if not unidad_solicitada:
        razones.append('Debe indicar una UDS.')
    elif not unidad_conocida:
        razones.append('La UDS no pertenece al catálogo operativo central.')
    if not usuarios:
        razones.append('No se encontraron participantes asociados a la UDS.')
    for tipo, info in disponibles.items():
        if not info['disponible']:
            razones.append(f'No existe plantilla {tipo.upper()} aplicable al periodo.')
    if not minuta:
        razones.append('No existe una minuta RPP iniciada antes o durante el período solicitado.')

    ready_common = unidad_conocida and bool(usuarios)
    storage_diagnostic = diagnostico_almacenamiento()
    result = {
        'ok': True,
        'periodo': {'mes': mes, 'anio': anio},
        'unidad': {
            'solicitada': unidad_solicitada,
            'normalizada': unidad_normalizada,
            'conocida': unidad_conocida,
        },
        'participantes': {
            'total': len(usuarios),
            'incluyeDatosPersonales': False,
        },
        'plantillas': disponibles,
        'rppMinuta': {
            'disponible': bool(minuta),
            'periodoExacto': bool(minuta and minuta_mes == mes and minuta_anio == anio),
            'aplicableAlPeriodo': minuta_aplicable,
            'mes': minuta_mes,
            'anio': minuta_anio,
            'version': minuta.get('version') if minuta else None,
            'codigo': minuta.get('codigo') if minuta else None,
            'grupos': grupos,
            'productos': productos,
            'errorTecnico': minuta_error or None,
        },
        'preparado': {
            'bienestarina': bool(ready_common and disponibles['bienestarina']['disponible']),
            'ram': bool(ready_common and disponibles['ram']['disponible']),
            'rpp': bool(ready_common and disponibles['rpp']['disponible'] and minuta_aplicable),
        },
        'razones': razones,
        'errorPlantillas': plantillas_error or None,
        'storage': {
            'dataDir': str(app.config.get('DATA_DIR') or ''),
            'databaseInsideDataDir': storage_diagnostic.get('databaseInsideDataDir'),
            'persistentVolumeDeclared': storage_diagnostic.get('persistentVolumeDeclared'),
            'dataDirTargetsExpectedMount': storage_diagnostic.get('dataDirTargetsExpectedMount'),
            'volumeStatus': storage_diagnostic.get('volumeStatus'),
        },
    }
    return jsonify(result)


def _alpha59_edad_meses(user):
    try:
        val = user.get('EdadMeses') if user.get('EdadMeses') not in (None, '') else user.get('edad_meses')
        return int(float(val or 0))
    except Exception:
        return 0


def _alpha59_filtrar_rpp_grupo(usuarios, grupo):
    grupo = str(grupo or '').strip().lower().replace('-', '_')
    filtrados = []
    for user in usuarios or []:
        tipo = normalizar_texto_clave(user.get('TipoBeneficiario') or user.get('tipo_beneficiario') or '')
        gtxt = normalizar_texto_clave(user.get('GrupoEdad') or user.get('grupo_edad') or '')
        edad = _alpha59_edad_meses(user)
        es_gestante = 'gestante' in tipo or 'gestante' in gtxt
        # Rangos disjuntos: menores de 6 meses (0-5), luego 6-11.
        if grupo == 'rpp_0_6_gestantes' and (es_gestante or 0 <= edad <= 5 or '0 a 6' in gtxt):
            filtrados.append(user)
        elif grupo == 'rpp_6_11' and (6 <= edad <= 11 or '6 a 11' in gtxt):
            filtrados.append(user)
        elif grupo == 'rpp_1_2' and (12 <= edad <= 35 or '1 a 2' in gtxt):
            filtrados.append(user)
        elif grupo == 'rpp_3_5' and (36 <= edad <= 71 or '3 a 5' in gtxt):
            filtrados.append(user)
    return filtrados


def _alpha59_metadata_formato(unidad, usuarios, mes=None, anio=None):
    base = (usuarios or [{}])[0] if usuarios else {}
    unidad_db = {}
    try:
        conn = get_db_connection()
        for row in conn.execute('SELECT * FROM unidades').fetchall():
            candidato = _alpha59_row_to_dict(row)
            nombre = candidato.get('nombre') or ''
            if normalize_unidad(nombre) == normalize_unidad(unidad) or normalizar_texto_clave(nombre) in equivalentes_unidad(unidad):
                unidad_db = candidato
                break
        conn.close()
    except Exception:
        unidad_db = {}
    try:
        talento = obtener_talento_por_unidad(unidad) or {}
    except Exception:
        talento = {}
    mes_val = int(mes or request.args.get('mes') or datetime.now().month) if has_request_context() else int(mes or datetime.now().month)
    anio_val = int(anio or request.args.get('anio') or request.args.get('año') or datetime.now().year) if has_request_context() else int(anio or datetime.now().year)
    return {
        'unidad': unidad,
        'Unidad': unidad,
        'mes': MESES_ES.get(mes_val, str(mes_val)).upper() if 'MESES_ES' in globals() else str(mes_val),
        'Mes': MESES_ES.get(mes_val, str(mes_val)).upper() if 'MESES_ES' in globals() else str(mes_val),
        'anio': anio_val,
        'año': anio_val,
        'regional': base.get('regional') or 'CHOCÓ',
        'centro_zonal': base.get('centro_zonal') or '',
        'municipio': base.get('municipio') or '',
        'modalidad': base.get('modalidad') or '',
        'codigo_unidad': base.get('codigo_unidad_servicio') or base.get('codigo_unidad') or unidad_db.get('codigo_unidad_servicio') or '',
        'codigo_uds': base.get('codigo_unidad_servicio') or base.get('codigo_unidad') or unidad_db.get('codigo_unidad_servicio') or '',
        'codigo_origen': base.get('codigo_unidad_servicio') or base.get('codigo_unidad') or unidad_db.get('codigo_unidad_servicio') or '',
        'unidad_origen': base.get('Unidad') or base.get('unidad_servicio') or unidad_db.get('nombre') or unidad,
        'responsable': base.get('docente') or base.get('agente_educativo') or talento.get('nombre') or unidad_db.get('docente_asignado') or '',
        'docente': base.get('docente') or talento.get('nombre') or unidad_db.get('docente_asignado') or '',
        'direccion': base.get('direccion_unidad') or unidad_db.get('direccion') or talento.get('direccion') or '',
        'direccion_unidad': base.get('direccion_unidad') or unidad_db.get('direccion') or talento.get('direccion') or '',
        'telefono': unidad_db.get('telefono') or talento.get('telefono') or base.get('telefono') or base.get('Telefono') or '',
        'telefono_docente': talento.get('telefono') or unidad_db.get('telefono') or '',
        'contrato': base.get('contrato') or unidad_db.get('contrato') or talento.get('contrato') or '',
        'eas': base.get('nombre_eas') or '',
        'fecha_entrega': request.args.get('fecha_entrega') if has_request_context() else '',
        'lote': request.args.get('lote') if has_request_context() else '',
        'cantidad': request.args.get('cantidad') if has_request_context() else 1,
    }


def _alpha59_generar_oficial_desde_template(tipo, unidad, grupo=None):
    try:
        usuarios = _alpha59_obtener_usuarios_unidad(unidad)
        if not usuarios:
            log_alpha56_formato('ALPHA59_GENERAR_SIN_USUARIOS', formato=tipo, unidad=unidad, grupo=grupo)
            return None
        formato_norm = _alpha57_normalizar_formato_descarga(tipo)
        if formato_norm == 'rpp' or str(tipo).startswith('rpp_'):
            usuarios = _alpha59_filtrar_rpp_grupo(usuarios, grupo or tipo)
            if not usuarios:
                log_alpha56_formato('ALPHA59_RPP_GRUPO_SIN_USUARIOS', unidad=unidad, grupo=grupo or tipo)
                return None
            slug = _alpha59_slug(grupo or tipo).replace('RPP_', '')
            nombre = f"{_alpha59_slug(unidad)}_RPP_{slug}.xlsx"
            tipo_generador = 'rpp'
        elif formato_norm == 'bienestarina':
            nombre = f"{_alpha59_slug(unidad)}_BIENESTARINA_{datetime.now().year}_{datetime.now().month:02d}.xlsx"
            tipo_generador = 'bienestarina'
        else:
            return None
        salida = os.path.join(OUTPUT_FOLDER, secure_filename(nombre))
        try:
            from modules.plantillas_oficiales import generar_desde_plantilla_oficial
            generar_desde_plantilla_oficial(
                tipo_generador,
                {'metadata': _alpha59_metadata_formato(unidad, usuarios), 'usuarios': usuarios},
                salida,
                TEMPLATES_FOLDER,
            )
        except Exception as exc:
            log_alpha56_formato('ALPHA59_GENERAR_OFICIAL_ERROR', formato=tipo, unidad=unidad, grupo=grupo, error=str(exc))
            return None
        if os.path.exists(salida) and os.path.getsize(salida) > 0:
            registrar_archivo_generado_alpha57(tipo_generador, unidad, os.path.basename(salida), salida, grupo_etario=grupo)
            log_alpha56_formato('ALPHA59_GENERAR_OFICIAL_OK', formato=tipo_generador, unidad=unidad, grupo=grupo, archivo=os.path.basename(salida), ruta=salida)
            return os.path.basename(salida)
    except Exception as exc:
        log_alpha56_formato('ALPHA59_GENERAR_OFICIAL_FATAL', formato=tipo, unidad=unidad, grupo=grupo, error=str(exc))
    return None


# -----------------------------------------------------------------------------
# ALPHA60: generación robusta de Bienestarina antes de descargar
# -----------------------------------------------------------------------------
def _alpha60_usuarios_bienestarina_unidad(unidad):
    """Obtiene usuarios de una UDS para generar Bienestarina aunque el flujo
    histórico no haya dejado un archivo previamente generado.

    La función prioriza beneficiarios, pero también conserva compatibilidad con
    usuarios. No modifica la base de datos; solo lee y normaliza datos para llenar
    la plantilla oficial.
    """
    unidad_norm = normalize_unidad(unidad) or str(unidad or '').strip().upper()
    unidad_txt = normalizar_texto_clave(unidad_norm)
    usuarios = []
    try:
        conn = get_db_connection()
        try:
            tablas_db = [r[0] for r in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
        except Exception:
            tablas_db = []
        for tabla in ['beneficiarios', 'usuarios']:
            if tabla not in tablas_db:
                continue
            try:
                rows = conn.execute(f'SELECT * FROM {tabla}').fetchall()
            except Exception as exc:
                log_alpha56_formato('ALPHA60_BIENESTARINA_QUERY_ERROR', tabla=tabla, unidad=unidad, error=str(exc))
                rows = []
            for row in rows:
                data = _alpha59_row_to_dict(row)
                unidad_raw = data.get('unidad') or data.get('Unidad') or data.get('unidad_servicio') or data.get('nombre_unidad') or ''
                item_unidad = normalize_unidad(unidad_raw) or str(unidad_raw or '').strip().upper()
                item_txt = normalizar_texto_clave(item_unidad)
                if unidad_txt and item_txt and not (unidad_txt == item_txt or unidad_txt in item_txt or item_txt in unidad_txt):
                    continue
                usuarios.append(_alpha59_usuario_normalizado(data, unidad))
        try:
            conn.close()
        except Exception:
            pass
    except Exception as exc:
        log_alpha56_formato('ALPHA60_BIENESTARINA_USUARIOS_ERROR', unidad=unidad, error=str(exc))
    return _alpha59_deduplicar_usuarios(usuarios)


def _alpha60_generar_bienestarina_directa(unidad, mes=None, anio=None):
    """Genera Bienestarina directamente desde la plantilla oficial.

    Corrige el caso reportado donde Descargar buscaba plantilla_bienestarina.xlsx
    o no encontraba archivo generado, aunque existieran datos de la UDS. Esta
    función siempre valida existencia física antes de devolver el nombre.
    """
    try:
        usuarios = _alpha59_obtener_usuarios_unidad(unidad)
        if not usuarios:
            log_alpha56_formato('ALPHA60_BIENESTARINA_DIRECTA_SIN_USUARIOS', unidad=unidad)
            return None
        mes_val = mes
        anio_val = anio
        if has_request_context():
            mes_val = mes_val or request.args.get('mes') or request.args.get('month')
            anio_val = anio_val or request.args.get('anio') or request.args.get('año') or request.args.get('year')
        try:
            mes_int = int(mes_val or datetime.now().month)
        except Exception:
            mes_int = datetime.now().month
        try:
            anio_int = int(anio_val or datetime.now().year)
        except Exception:
            anio_int = datetime.now().year
        nombre = secure_filename(f"{_alpha59_slug(unidad)}_BIENESTARINA_{anio_int}_{mes_int:02d}.xlsx")
        ruta = os.path.join(OUTPUT_FOLDER, nombre)
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        metadata = _alpha59_metadata_formato(unidad, usuarios, mes=mes_int, anio=anio_int)
        try:
            from modules.plantillas_oficiales import generar_desde_plantilla_oficial
            generar_desde_plantilla_oficial(
                'bienestarina',
                {'metadata': metadata, 'usuarios': usuarios},
                ruta,
                TEMPLATES_FOLDER,
            )
        except Exception as exc:
            log_alpha56_formato('ALPHA60_BIENESTARINA_DIRECTA_GENERADOR_ERROR', unidad=unidad, archivo=nombre, ruta=ruta, error=str(exc), traceback=traceback.format_exc())
            return None
        if os.path.exists(ruta) and os.path.getsize(ruta) > 0:
            registrar_archivo_generado_alpha57('bienestarina', unidad, nombre, ruta, mes=mes_int, anio=anio_int, estado='generado_alpha60')
            log_alpha56_formato('ALPHA60_BIENESTARINA_DIRECTA_OK', unidad=unidad, archivo=nombre, ruta=ruta, total_usuarios=len(usuarios), tamano=os.path.getsize(ruta))
            return nombre
        log_alpha56_formato('ALPHA60_BIENESTARINA_DIRECTA_NO_EXISTE', unidad=unidad, archivo=nombre, ruta=ruta, existe=os.path.exists(ruta))
    except Exception as exc:
        log_alpha56_formato('ALPHA60_BIENESTARINA_DIRECTA_FATAL', unidad=unidad, error=str(exc), traceback=traceback.format_exc())
    return None


def _alpha59_intentar_generar_faltante(unidad, formato):
    formato_norm = _alpha57_normalizar_formato_descarga(formato)
    try:
        usuarios = _alpha59_obtener_usuarios_unidad(unidad)
        if not usuarios:
            log_alpha56_formato('ALPHA76_GENERACION_DIRECTA_SIN_USUARIOS', unidad=unidad, formato=formato)
            return None

        # ALPHA76: los formatos solicitables desde el tablero tienen generadores
        # directos. Atenderlos antes del fallback histórico evita generar RPP,
        # RAM, Bienestarina y complementarios juntos para una sola descarga.
        try:
            mes = int(request.args.get('mes') or request.args.get('month') or datetime.now().month) if has_request_context() else datetime.now().month
        except Exception:
            mes = datetime.now().month
        try:
            anio = int(request.args.get('anio') or request.args.get('año') or request.args.get('year') or datetime.now().year) if has_request_context() else datetime.now().year
        except Exception:
            anio = datetime.now().year
        mes = max(1, min(12, mes))
        anio = max(2020, min(2100, anio))

        directos = {
            'listado_usuarios': lambda: _alpha68_generar_listado_usuarios(unidad, usuarios, mes, anio),
            'listado_asistencia_usuarios': lambda: _alpha68_generar_listado_asistencia_usuarios(unidad, usuarios, mes, anio),
            'relacion_mensual': lambda: _alpha68_generar_relacion_mensual(unidad, usuarios, mes, anio),
            'distribucion_alimentos': lambda: _alpha68_generar_distribucion_alimentos(unidad, usuarios, mes, anio),
        }
        if formato_norm in directos:
            ruta = directos[formato_norm]()
            nombre = os.path.basename(os.fspath(ruta)) if ruta else None
            log_alpha56_formato(
                'ALPHA76_GENERACION_DIRECTA_OK', unidad=unidad, formato=formato_norm,
                archivo=nombre, usuarios=len(usuarios), mes=mes, anio=anio,
            )
            return nombre
        if str(formato).startswith('rpp_') or formato_norm == 'rpp':
            return _alpha59_generar_oficial_desde_template('rpp', unidad, grupo=formato)
        if formato_norm == 'bienestarina':
            return _alpha65_generar_bienestarina_para_uds(unidad, mes=mes, anio=anio)

        if usuarios:
            try:
                inyectar_datos_en_plantillas(unidad, usuarios, options={'mes': datetime.now().month, 'anio': datetime.now().year, 'año': datetime.now().year})
                nombre = buscar_archivo_generado(unidad, formato)
                if nombre:
                    log_alpha56_formato('ALPHA59_GENERAR_HISTORICO_OK', unidad=unidad, formato=formato, archivo=nombre)
                    return nombre
            except Exception as exc:
                log_alpha56_formato('ALPHA59_GENERAR_HISTORICO_ERROR', unidad=unidad, formato=formato, error=str(exc))
    except Exception as exc:
        log_alpha56_formato('ALPHA59_GENERAR_FALTANTE_ERROR', unidad=unidad, formato=formato, error=str(exc))
    return None



# -----------------------------------------------------------------------------
# ALPHA64: recuperación segura de descargas RPP/Bienestarina con traceback
# -----------------------------------------------------------------------------
def _alpha64_log(evento, **datos):
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha64_descargas_rpp_bienestarina.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha64_listar_ultimos_generados(limite=12):
    try:
        if not os.path.isdir(OUTPUT_FOLDER):
            return []
        archivos = [n for n in os.listdir(OUTPUT_FOLDER) if n.lower().endswith(('.xlsx', '.xlsm', '.xls', '.pdf'))]
        archivos.sort(key=lambda n: os.path.getmtime(os.path.join(OUTPUT_FOLDER, n)), reverse=True)
        return archivos[:limite]
    except Exception:
        return []


def _alpha64_send_output(nombre_archivo, unidad, formato, grupo=None):
    """Descarga un archivo validado de OUTPUT_FOLDER sin permitir rutas externas."""
    nombre, ruta = _alpha57_safe_join_output(nombre_archivo)
    valido, motivo = _alpha63_validar_archivo_descarga(nombre or '', unidad, formato, grupo=grupo)
    _alpha64_log('ALPHA64_SEND_VALIDACION', unidad=unidad, formato=formato, grupo=grupo, archivo=nombre, ruta=ruta, valido=valido, motivo=motivo)
    if not valido:
        return None, motivo
    try:
        return send_from_directory(OUTPUT_FOLDER, nombre, as_attachment=True), 'ok'
    except Exception as exc:
        _alpha64_log('ALPHA64_SEND_EXCEPTION', unidad=unidad, formato=formato, grupo=grupo, archivo=nombre, error=str(exc), traceback=traceback.format_exc())
        return None, f'send_exception:{exc}'


# -----------------------------------------------------------------------------
# ALPHA65: generación real de Bienestarina por UDS, sin depender del paquete
# completo y sin devolver 404 si existen usuarios y plantilla oficial.
# -----------------------------------------------------------------------------
def _alpha65_log(evento, **datos):
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha65_bienestarina_404_fix.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha65_nombre_bienestarina(unidad, mes=None, anio=None):
    try:
        mes_int = int(mes or (request.args.get('mes') if has_request_context() else 0) or datetime.now().month)
    except Exception:
        mes_int = datetime.now().month
    try:
        anio_int = int(anio or (request.args.get('anio') or request.args.get('año') if has_request_context() else 0) or datetime.now().year)
    except Exception:
        anio_int = datetime.now().year
    return secure_filename(f"{_alpha59_slug(unidad)}_BIENESTARINA_{anio_int}_{mes_int:02d}.xlsx"), mes_int, anio_int


def _alpha65_plantillas_bienestarina_disponibles():
    candidatos = []
    posibles = [
        os.path.join(TEMPLATES_FOLDER, 'oficiales', 'plantilla_bienestarina_oficial.xlsx'),
        os.path.join(TEMPLATES_FOLDER, 'oficiales', 'plantilla_bienestarina_oficial_v2026.xlsx'),
        os.path.join(TEMPLATES_FOLDER, 'plantilla_bienestarina.xlsx'),
        os.path.join(TEMPLATES_FOLDER, 'plantilla_bienestarina_oficial.xlsx'),
    ]
    try:
        for root, _dirs, files in os.walk(TEMPLATES_FOLDER):
            for fname in files:
                low = fname.lower()
                if low.endswith(('.xlsx', '.xlsm')) and 'bienestarina' in low and 'backup' not in low:
                    posibles.append(os.path.join(root, fname))
    except Exception:
        pass
    vistos = set()
    for ruta in posibles:
        try:
            ruta_abs = os.path.abspath(ruta)
            if ruta_abs not in vistos and os.path.exists(ruta_abs) and os.path.getsize(ruta_abs) > 0:
                vistos.add(ruta_abs)
                candidatos.append(ruta_abs)
        except Exception:
            continue
    return candidatos


def _alpha65_generar_bienestarina_para_uds(unidad, mes=None, anio=None):
    """Genera SOLO Bienestarina para la UDS solicitada.

    Si el generador oficial falla por mapeo/manifest, usa un fallback mínimo pero
    válido: copia la plantilla oficial y diligencia encabezado + usuarios, sin
    modificar la plantilla original ni generar otros formatos.
    """
    nombre, mes_int, anio_int = _alpha65_nombre_bienestarina(unidad, mes=mes, anio=anio)
    ruta = os.path.join(OUTPUT_FOLDER, nombre)
    try:
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        usuarios = _alpha59_obtener_usuarios_unidad(unidad)
        usuarios = _alpha59_deduplicar_usuarios(usuarios)
        _alpha65_log('INICIO_GENERAR_BIENESTARINA', unidad=unidad, archivo_esperado=nombre, usuarios=len(usuarios or []))
        if not usuarios:
            _alpha65_log('SIN_USUARIOS', unidad=unidad, archivo_esperado=nombre)
            return None

        metadata = _alpha59_metadata_formato(unidad, usuarios, mes=mes_int, anio=anio_int)

        # Camino A: generador oficial existente.
        try:
            from modules.plantillas_oficiales import generar_desde_plantilla_oficial
            generar_desde_plantilla_oficial(
                'bienestarina',
                {'metadata': metadata, 'usuarios': usuarios},
                ruta,
                TEMPLATES_FOLDER,
            )
            if os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                registrar_archivo_generado_alpha57('bienestarina', unidad, nombre, ruta, mes=mes_int, anio=anio_int, estado='generado_alpha65_oficial')
                _alpha65_log('GENERADA_OFICIAL_OK', unidad=unidad, archivo=nombre, ruta=ruta, size=os.path.getsize(ruta))
                return nombre
        except Exception as exc:
            _alpha65_log('GENERADOR_OFICIAL_ERROR', unidad=unidad, archivo=nombre, error=str(exc), traceback=traceback.format_exc())

        # Camino B: fallback propio sobre plantilla de Bienestarina.
        plantillas = _alpha65_plantillas_bienestarina_disponibles()
        if not plantillas:
            _alpha65_log('SIN_PLANTILLA', unidad=unidad, archivo=nombre, templates_folder=TEMPLATES_FOLDER)
            return None
        plantilla = plantillas[0]
        wb = load_workbook(plantilla, data_only=False, keep_vba=plantilla.lower().endswith('.xlsm'))
        ws = wb['plantilla de bienestarina '] if 'plantilla de bienestarina ' in wb.sheetnames else wb[wb.sheetnames[0]]

        def put(cell, value):
            try:
                if value not in (None, ''):
                    ws[cell] = value
            except Exception:
                pass

        # Encabezado básico, sin tocar diseño.
        put('J2', unidad)
        put('R1', f'AÑO: {anio_int}')
        mes_txt = MESES_ES.get(mes_int, str(mes_int)).upper() if 'MESES_ES' in globals() else str(mes_int)
        put('N1', f'MES DE CONSUMO: {mes_txt}')
        put('C2', metadata.get('regional'))
        put('C3', metadata.get('centro_zonal'))
        put('C4', metadata.get('municipio'))
        put('C5', metadata.get('modalidad'))
        put('J1', metadata.get('codigo_uds') or metadata.get('codigo_unidad'))
        put('J3', metadata.get('responsable') or metadata.get('docente'))
        put('J4', metadata.get('direccion'))
        put('S4', metadata.get('telefono'))

        filas = list(range(10, 24)) + list(range(31, 47))
        cols = list('ABCDEFGHIJKLMNOPQRS')
        for idx, row in enumerate(filas):
            for col in cols:
                try:
                    ws[f'{col}{row}'] = None
                except Exception:
                    pass
            put(f'A{row}', idx + 1)
            if idx >= len(usuarios):
                continue
            user = usuarios[idx]
            put(f'B{row}', user.get('PrimerNombre') or '')
            put(f'C{row}', user.get('SegundoNombre') or '')
            put(f'D{row}', user.get('PrimerApellido') or '')
            put(f'E{row}', user.get('SegundoApellido') or '')
            put(f'F{row}', user.get('TipoDocumento') or user.get('tipo_documento') or '')
            put(f'G{row}', user.get('NUI') or user.get('Documento') or user.get('documento') or '')
            put(f'H{row}', metadata.get('fecha_entrega') or '')
            put(f'I{row}', metadata.get('lote') or '')
            put(f'J{row}', metadata.get('cantidad') or 1)
            acud = str(user.get('Acudiente') or user.get('nombre_acudiente') or '').strip()
            doc_acud = str(user.get('DocumentoAcudiente') or user.get('documento_acudiente') or '').strip()
            put(f'Q{row}', f'{acud} {doc_acud}'.strip())
            put(f'R{row}', user.get('Parentesco') or user.get('parentesco') or '')

        # No imponer área de impresión: se conserva exactamente la definida en
        # la plantilla oficial (incluido el caso en que esté sin configurar).
        wb.save(ruta)
        wb.close()

        if os.path.exists(ruta) and os.path.getsize(ruta) > 0:
            registrar_archivo_generado_alpha57('bienestarina', unidad, nombre, ruta, mes=mes_int, anio=anio_int, estado='generado_alpha65_fallback')
            _alpha65_log('GENERADA_FALLBACK_OK', unidad=unidad, archivo=nombre, ruta=ruta, plantilla=plantilla, usuarios=len(usuarios), size=os.path.getsize(ruta))
            return nombre
        _alpha65_log('FALLBACK_NO_CREO_ARCHIVO', unidad=unidad, archivo=nombre, ruta=ruta)
    except Exception as exc:
        _alpha65_log('GENERAR_BIENESTARINA_FATAL', unidad=unidad, archivo=nombre, error=str(exc), traceback=traceback.format_exc())
    return None


def _alpha64_generar_bienestarina_resiliente(unidad):
    """ALPHA65: genera SOLO Bienestarina para la UDS solicitada.

    No descarga ni devuelve archivos de otra UDS. No modifica plantillas oficiales.
    """
    try:
        existente = _alpha61_buscar_archivo_bienestarina_exacto(unidad)
        if existente:
            return existente
        return _alpha65_generar_bienestarina_para_uds(unidad)
    except Exception as exc:
        _alpha65_log('BIENESTARINA_RESILIENTE_FATAL', unidad=unidad, error=str(exc), traceback=traceback.format_exc())
        return None


def _alpha64_generar_rpp_resiliente(unidad, grupo):
    """Intenta obtener/generar RPP exacto por UDS + grupo, sin fallback cruzado."""
    grupo_cod = _alpha61_normalizar_grupo_rpp(grupo)
    if not grupo_cod:
        return None
    try:
        existente = _alpha61_buscar_archivo_rpp_exacto(unidad, grupo_cod)
        if existente:
            return existente

        # Ruta oficial directa: genera solo el grupo solicitado. El camino
        # histórico generaba todos los formatos de la UDS y era la diferencia
        # de rendimiento observable en fundaciones sin archivos previos.
        try:
            legacy = GRUPOS_RPP_ALPHA61[grupo_cod]['legacy']
            _alpha59_generar_oficial_desde_template('rpp', unidad, grupo=legacy)
            nombre = _alpha61_buscar_archivo_rpp_exacto(unidad, grupo_cod)
            if nombre:
                _alpha64_log('RPP_GENERADO_CAMINO_OFICIAL', unidad=unidad, grupo=grupo_cod, archivo=nombre)
                return nombre
        except Exception as exc:
            _alpha64_log('RPP_CAMINO_OFICIAL_ERROR', unidad=unidad, grupo=grupo_cod, error=str(exc), traceback=traceback.format_exc())

        return _alpha61_buscar_archivo_rpp_exacto(unidad, grupo_cod)
    except Exception as exc:
        _alpha64_log('RPP_RESILIENTE_FATAL', unidad=unidad, grupo=grupo_cod, error=str(exc), traceback=traceback.format_exc())
        return None



def _alpha69_periodo_descarga_ram():
    """Obtiene un periodo válido para generar/descargar RAM sin adivinarlo."""
    try:
        mes = int(request.args.get('mes') or request.args.get('month') or datetime.now().month)
    except Exception:
        mes = datetime.now().month
    try:
        anio = int(request.args.get('anio') or request.args.get('año') or request.args.get('year') or datetime.now().year)
    except Exception:
        anio = datetime.now().year
    return max(1, min(12, mes)), max(2020, min(2100, anio))


def _alpha69_buscar_ram_periodo(unidad, mes, anio):
    """Busca únicamente un RAM de la UDS y periodo solicitados."""
    unidad_norm = normalize_unidad(unidad) or str(unidad or '').strip().upper()
    candidatos = []

    # Primero, registro estructurado con periodo explícito.
    for item in _alpha57_leer_registro():
        if _alpha57_normalizar_formato_descarga(item.get('formato') or item.get('formato_original')) != 'ram':
            continue
        try:
            if int(item.get('mes')) != int(mes) or int(item.get('anio')) != int(anio):
                continue
        except Exception:
            continue
        item_unidad = normalize_unidad(item.get('unidad') or item.get('unidad_normalizada')) or str(item.get('unidad') or '').strip().upper()
        if normalizar_texto_clave(item_unidad) != normalizar_texto_clave(unidad_norm):
            continue
        nombre, ruta = _alpha57_safe_join_output(item.get('archivo') or '')
        if nombre and ruta and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
            valido, _ = _alpha63_validar_archivo_descarga(nombre, unidad_norm, 'ram')
            if valido:
                candidatos.append((os.path.getmtime(ruta), nombre))

    # Respaldo físico para instalaciones que aún no tienen registro JSON.
    if os.path.isdir(OUTPUT_FOLDER):
        periodo_tags = {f'{int(anio):04d}_{int(mes):02d}', f'{int(anio):04d}{int(mes):02d}'}
        for nombre in os.listdir(OUTPUT_FOLDER):
            if not nombre.lower().endswith(('.xlsx', '.xlsm', '.xls')):
                continue
            slug = _alpha61_slug_archivo(nombre)
            if not any(tag in slug for tag in periodo_tags):
                continue
            if not any(tag in slug for tag in ('RAM', 'ASISTENCIA')):
                continue
            valido, _ = _alpha63_validar_archivo_descarga(nombre, unidad_norm, 'ram')
            ruta = os.path.join(OUTPUT_FOLDER, nombre)
            if valido and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                candidatos.append((os.path.getmtime(ruta), nombre))

    candidatos.sort(reverse=True)
    return candidatos[0][1] if candidatos else None


def _alpha69_generar_ram_directo(unidad, mes, anio):
    """Genera únicamente RAM para la UDS/periodo solicitado y devuelve su archivo."""
    try:
        usuarios = _alpha59_obtener_usuarios_unidad(unidad)
        if not usuarios:
            log_alpha56_formato('ALPHA69_RAM_SIN_USUARIOS', unidad=unidad, mes=mes, anio=anio)
            return None
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        inyectar_datos_en_plantillas(
            unidad,
            usuarios,
            options={
                'mes': int(mes), 'anio': int(anio), 'año': int(anio),
                'formatos_seleccionados': 'ram'
            },
        )
        nombre = _alpha69_buscar_ram_periodo(unidad, mes, anio)
        log_alpha56_formato(
            'ALPHA69_RAM_GENERACION_RESULTADO', unidad=unidad, mes=mes, anio=anio,
            archivo=nombre, ok=bool(nombre), usuarios=len(usuarios)
        )
        return nombre
    except Exception as exc:
        log_alpha56_formato(
            'ALPHA69_RAM_GENERACION_ERROR', unidad=unidad, mes=mes, anio=anio,
            error=str(exc), traceback=traceback.format_exc()
        )
        return None


@app.route('/api/descargar/<unidad>/<formato>', methods=['GET'])
def descargar_formato(unidad, formato):
    """ALPHA63: descarga común segura para RAM/RAN y formatos genéricos.

    Valida UDS + formato antes de enviar el archivo. Si falla, devuelve JSON
    controlado para que el frontend no abandone la plataforma.
    """
    formato_norm = _alpha57_normalizar_formato_descarga(formato)
    if formato_norm == 'ran':
        return jsonify({
            'ok': False,
            'error': 'RAN/RRAN no tiene una plantilla oficial registrada.',
            'mensaje': (
                'RAN es independiente de RAM. Cargue y versione la plantilla '
                'institucional RAN/RRAN antes de solicitar su generación.'
            ),
            'unidad': unidad,
            'formato': formato,
            'requierePlantillaOficial': True,
        }), 422
    if formato_norm == 'ram':
        mes_ram, anio_ram = _alpha69_periodo_descarga_ram()
        nombre_archivo = _alpha69_buscar_ram_periodo(unidad, mes_ram, anio_ram)
        if not nombre_archivo:
            nombre_archivo = _alpha69_generar_ram_directo(unidad, mes_ram, anio_ram)
    else:
        nombre_archivo = buscar_archivo_generado(unidad, formato)
        if not nombre_archivo:
            nombre_archivo = _alpha59_intentar_generar_faltante(unidad, formato)
    if nombre_archivo:
        valido, motivo = _alpha63_validar_archivo_descarga(nombre_archivo, unidad, formato)
        ruta_real = os.path.join(OUTPUT_FOLDER, secure_filename(os.path.basename(nombre_archivo)))
        log_alpha56_formato(
            'DESCARGA_FORMATO_VALIDACION', unidad=unidad, formato=formato,
            archivo=nombre_archivo, ruta_real=ruta_real, valido=valido, motivo=motivo
        )
        if valido:
            return send_from_directory(OUTPUT_FOLDER, secure_filename(os.path.basename(nombre_archivo)), as_attachment=True)
    ultimos = []
    try:
        if os.path.isdir(OUTPUT_FOLDER):
            ultimos = sorted([n for n in os.listdir(OUTPUT_FOLDER) if n.lower().endswith(('.xlsx','.xls','.xlsm','.pdf'))], key=lambda n: os.path.getmtime(os.path.join(OUTPUT_FOLDER, n)), reverse=True)[:10]
    except Exception:
        ultimos = []
    log_alpha56_formato('DESCARGA_FORMATO_404', unidad=unidad, formato=formato, output_folder=OUTPUT_FOLDER, ultimos=ultimos)
    return jsonify({
        'ok': False,
        'error': 'El archivo solicitado no existe o no corresponde a la UDS/formato solicitados.',
        'mensaje': 'Genere el formato para esta UDS y vuelva a intentar. No se descargará un archivo de otra unidad.',
        'unidad': unidad,
        'formato': formato,
        'periodo': ({'mes': mes_ram, 'anio': anio_ram} if formato_norm == 'ram' else None),
        'output_folder': os.fspath(OUTPUT_FOLDER),
        'ultimos_archivos_generados': ultimos
    }), 404



# -----------------------------------------------------------------------------
# ALPHA67: Fix real y aislado de Bienestarina por UDS
# Objetivo: el endpoint /api/bienestarina/descargar debe generar el archivo
# final de la UDS solicitada si no existe, validar físicamente y descargarlo.
# No toca RPP, RAM/RAN, Base Maestra, login ni plantillas oficiales.
# -----------------------------------------------------------------------------
def _alpha67_log_bienestarina(evento, **datos):
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha67_bienestarina_real.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha67_unidad_slug(unidad):
    try:
        return _alpha59_slug(normalize_unidad(unidad) or unidad)
    except Exception:
        txt = unicodedata.normalize('NFKD', str(unidad or '').upper())
        txt = ''.join(ch for ch in txt if not unicodedata.combining(ch))
        return re.sub(r'[^A-Z0-9_]+', '_', txt.replace(' ', '_')).strip('_') or 'SIN_UNIDAD'


def _alpha67_nombre_bienestarina(unidad, mes=None, anio=None):
    try:
        mes_val = mes or (request.args.get('mes') if has_request_context() else None) or datetime.now().month
        mes_int = max(1, min(12, int(mes_val)))
    except Exception:
        mes_int = datetime.now().month
    try:
        anio_val = anio or ((request.args.get('anio') or request.args.get('año')) if has_request_context() else None) or datetime.now().year
        anio_int = int(anio_val)
    except Exception:
        anio_int = datetime.now().year
    nombre = f"{_alpha67_unidad_slug(unidad)}_BIENESTARINA_{anio_int}_{mes_int:02d}.xlsx"
    return secure_filename(nombre), mes_int, anio_int


def _alpha67_unidades_equivalentes_txt(unidad):
    valores = set()
    try:
        valores.add(normalizar_texto_clave(unidad))
        valores.add(normalizar_texto_clave(normalize_unidad(unidad)))
        for eq in equivalentes_unidad(unidad):
            valores.add(normalizar_texto_clave(eq))
    except Exception:
        pass
    return {v for v in valores if v}


def _alpha67_obtener_usuarios_bienestarina(unidad):
    """Usa la misma fuente canónica de participantes que RPP y RAM."""
    return _alpha59_obtener_usuarios_unidad(unidad)


def _alpha67_plantillas_bienestarina():
    candidatos = []
    try:
        # 1. Manifest oficial actual.
        try:
            from modules.plantillas_oficiales import get_plantilla_oficial
            info = get_plantilla_oficial(TEMPLATES_FOLDER, 'bienestarina')
            if info and info.get('ruta'):
                candidatos.append(info.get('ruta'))
        except Exception as exc:
            _alpha67_log_bienestarina('MANIFEST_PLANTILLA_ERROR', error=str(exc))
        # 2. Rutas conocidas y búsqueda profunda en templates_originales.
        for ruta in [
            os.path.join(TEMPLATES_FOLDER, 'oficiales', 'plantilla_bienestarina_oficial.xlsx'),
            os.path.join(TEMPLATES_FOLDER, 'oficiales', 'plantilla_bienestarina_oficial_v2026.xlsx'),
            os.path.join(TEMPLATES_FOLDER, 'plantilla_bienestarina.xlsx'),
            os.path.join(TEMPLATES_FOLDER, 'Bienestarina.xlsx'),
        ]:
            candidatos.append(ruta)
        for root, _dirs, files in os.walk(TEMPLATES_FOLDER):
            for fname in files:
                low = fname.lower()
                if low.endswith(('.xlsx', '.xlsm')) and 'bienestarina' in low and 'backup' not in low:
                    candidatos.append(os.path.join(root, fname))
    except Exception as exc:
        _alpha67_log_bienestarina('BUSCAR_PLANTILLAS_FATAL', error=str(exc), traceback=traceback.format_exc())
    salida, vistos = [], set()
    for ruta in candidatos:
        try:
            ruta_abs = os.path.abspath(str(ruta))
            if ruta_abs in vistos:
                continue
            vistos.add(ruta_abs)
            if os.path.exists(ruta_abs) and os.path.getsize(ruta_abs) > 0:
                salida.append(ruta_abs)
        except Exception:
            pass
    _alpha67_log_bienestarina('PLANTILLAS_DETECTADAS', total=len(salida), plantillas=salida[:5])
    return salida


def _alpha67_es_archivo_bienestarina_valido(nombre_archivo, unidad):
    try:
        nombre = os.path.basename(str(nombre_archivo or ''))
        if not nombre.lower().endswith(('.xlsx', '.xlsm', '.xls')):
            return False, 'extension_invalida'
        slug = _alpha61_slug_archivo(nombre)
        if 'BIENESTARINA' not in slug:
            return False, 'no_contiene_bienestarina'
        if 'PLANTILLA' in slug or 'OFICIAL' in slug:
            return False, 'es_plantilla'
        if not _alpha61_archivo_valido_por_uds(nombre, unidad):
            return False, 'uds_no_coincide'
        nombre_seguro, ruta = _alpha57_safe_join_output(nombre)
        if not nombre_seguro or not ruta or not os.path.exists(ruta):
            return False, 'no_existe'
        if os.path.getsize(ruta) <= 0:
            return False, 'vacio'
        return True, 'ok'
    except Exception as exc:
        return False, f'excepcion:{exc}'


def _alpha67_buscar_bienestarina_exacta(unidad):
    """Busca solo archivos de Bienestarina de la UDS solicitada; nunca otra UDS."""
    candidatos = []
    # Primero registro, luego carpeta física.
    try:
        for item in _alpha57_leer_registro():
            fmt = _alpha57_normalizar_formato_descarga(item.get('formato') or item.get('formato_original'))
            if fmt != 'bienestarina':
                continue
            nombre = item.get('archivo') or ''
            ok, motivo = _alpha67_es_archivo_bienestarina_valido(nombre, unidad)
            if ok:
                _nombre, ruta = _alpha57_safe_join_output(nombre)
                candidatos.append((os.path.getmtime(ruta), _nombre, 'registro'))
            else:
                _alpha67_log_bienestarina('REGISTRO_DESCARTADO', unidad=unidad, archivo=nombre, motivo=motivo)
    except Exception as exc:
        _alpha67_log_bienestarina('BUSCAR_REGISTRO_ERROR', unidad=unidad, error=str(exc))
    try:
        if os.path.isdir(OUTPUT_FOLDER):
            for nombre in os.listdir(OUTPUT_FOLDER):
                ok, motivo = _alpha67_es_archivo_bienestarina_valido(nombre, unidad)
                if ok:
                    ruta = os.path.join(OUTPUT_FOLDER, secure_filename(os.path.basename(nombre)))
                    candidatos.append((os.path.getmtime(ruta), secure_filename(os.path.basename(nombre)), 'carpeta'))
                elif 'bienestarina' in str(nombre).lower():
                    _alpha67_log_bienestarina('ARCHIVO_CARPETA_DESCARTADO', unidad=unidad, archivo=nombre, motivo=motivo)
    except Exception as exc:
        _alpha67_log_bienestarina('BUSCAR_CARPETA_ERROR', unidad=unidad, error=str(exc))
    candidatos.sort(reverse=True)
    elegido = candidatos[0][1] if candidatos else None
    _alpha67_log_bienestarina('BUSQUEDA_EXACTA_RESULTADO', unidad=unidad, encontrado=elegido, total=len(candidatos))
    return elegido


def _alpha67_put(ws, ref, value):
    if value in (None, ''):
        return
    try:
        ws[ref] = value
    except Exception:
        pass


def _alpha67_generar_bienestarina_para_uds(unidad, mes=None, anio=None):
    nombre, mes_int, anio_int = _alpha67_nombre_bienestarina(unidad, mes=mes, anio=anio)
    ruta = os.path.join(OUTPUT_FOLDER, nombre)
    try:
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        usuarios = _alpha59_obtener_usuarios_unidad(unidad)
        if not usuarios:
            _alpha67_log_bienestarina('NO_GENERA_SIN_USUARIOS', unidad=unidad, archivo=nombre)
            return {'ok': False, 'archivo': None, 'ruta': ruta, 'causa': 'No se encontraron usuarios para la UDS solicitada.'}

        metadata = _alpha59_metadata_formato(unidad, usuarios, mes=mes_int, anio=anio_int)

        # Camino 1: generador oficial vigente.
        try:
            from modules.plantillas_oficiales import generar_desde_plantilla_oficial
            generar_desde_plantilla_oficial('bienestarina', {'metadata': metadata, 'usuarios': usuarios}, ruta, TEMPLATES_FOLDER)
            ok, motivo = _alpha67_es_archivo_bienestarina_valido(nombre, unidad)
            if ok:
                registrar_archivo_generado_alpha57('bienestarina', unidad, nombre, ruta, mes=mes_int, anio=anio_int, estado='generado_alpha67_oficial')
                _alpha67_log_bienestarina('GENERADO_OFICIAL_OK', unidad=unidad, archivo=nombre, ruta=ruta, usuarios=len(usuarios), size=os.path.getsize(ruta))
                return {'ok': True, 'archivo': nombre, 'ruta': ruta, 'causa': ''}
            _alpha67_log_bienestarina('GENERADOR_OFICIAL_ARCHIVO_INVALIDO', unidad=unidad, archivo=nombre, motivo=motivo, ruta=ruta)
        except Exception as exc:
            _alpha67_log_bienestarina('GENERADOR_OFICIAL_FALLO', unidad=unidad, archivo=nombre, error=str(exc), traceback=traceback.format_exc())

        # Camino 2: fallback directo sobre plantilla detectada, sin modificar plantilla original.
        plantillas = _alpha67_plantillas_bienestarina()
        if not plantillas:
            return {'ok': False, 'archivo': None, 'ruta': ruta, 'causa': 'No se encontró plantilla oficial de Bienestarina.'}
        plantilla = plantillas[0]
        wb = load_workbook(plantilla, data_only=False, keep_vba=str(plantilla).lower().endswith('.xlsm'))
        ws = wb['plantilla de bienestarina '] if 'plantilla de bienestarina ' in wb.sheetnames else wb[wb.sheetnames[0]]

        mes_txt = MESES_ES.get(mes_int, str(mes_int)).upper() if 'MESES_ES' in globals() else str(mes_int)
        _alpha67_put(ws, 'C2', metadata.get('regional') or 'CHOCÓ')
        _alpha67_put(ws, 'C3', metadata.get('centro_zonal') or '')
        _alpha67_put(ws, 'C4', metadata.get('municipio') or '')
        _alpha67_put(ws, 'C5', metadata.get('modalidad') or '')
        _alpha67_put(ws, 'J1', metadata.get('codigo_uds') or metadata.get('codigo_unidad') or '')
        _alpha67_put(ws, 'J2', unidad)
        _alpha67_put(ws, 'J3', metadata.get('responsable') or metadata.get('docente') or '')
        _alpha67_put(ws, 'J4', metadata.get('direccion') or '')
        _alpha67_put(ws, 'S4', metadata.get('telefono') or '')
        _alpha67_put(ws, 'N1', f'MES DE CONSUMO: {mes_txt}')
        _alpha67_put(ws, 'R1', f'AÑO: {anio_int}')
        _alpha75_aplicar_encabezado_bienestarina(ws, {
            **metadata, 'unidad': unidad, 'mes': mes_txt,
            'anio': anio_int, 'año': anio_int,
        })

        filas = list(range(10, 24)) + list(range(31, 47))
        columnas = list('ABCDEFGHIJKLMNOPQRS')
        for idx, row in enumerate(filas):
            for col in columnas:
                try:
                    ws[f'{col}{row}'] = None
                except Exception:
                    pass
            _alpha67_put(ws, f'A{row}', idx + 1)
            if idx >= len(usuarios):
                continue
            user = usuarios[idx]
            _alpha67_put(ws, f'B{row}', user.get('PrimerNombre') or '')
            _alpha67_put(ws, f'C{row}', user.get('SegundoNombre') or '')
            _alpha67_put(ws, f'D{row}', user.get('PrimerApellido') or '')
            _alpha67_put(ws, f'E{row}', user.get('SegundoApellido') or '')
            _alpha67_put(ws, f'F{row}', user.get('TipoDocumento') or user.get('tipo_documento') or '')
            _alpha67_put(ws, f'G{row}', user.get('NUI') or user.get('Documento') or user.get('documento') or '')
            _alpha67_put(ws, f'H{row}', metadata.get('fecha_entrega') or '')
            _alpha67_put(ws, f'I{row}', metadata.get('lote') or '')
            _alpha67_put(ws, f'J{row}', metadata.get('cantidad') or 1)
            acudiente = str(user.get('Acudiente') or user.get('nombre_acudiente') or '').strip()
            doc_acudiente = str(user.get('DocumentoAcudiente') or user.get('documento_acudiente') or '').strip()
            _alpha67_put(ws, f'Q{row}', f'{acudiente} {doc_acudiente}'.strip())
            _alpha67_put(ws, f'R{row}', user.get('Parentesco') or user.get('parentesco') or '')
        # Conservar exactamente la configuración de impresión oficial.
        wb.save(ruta)
        try:
            wb.close()
        except Exception:
            pass
        ok, motivo = _alpha67_es_archivo_bienestarina_valido(nombre, unidad)
        if ok:
            registrar_archivo_generado_alpha57('bienestarina', unidad, nombre, ruta, mes=mes_int, anio=anio_int, estado='generado_alpha67_fallback')
            _alpha67_log_bienestarina('GENERADO_FALLBACK_OK', unidad=unidad, archivo=nombre, ruta=ruta, plantilla=plantilla, usuarios=len(usuarios), size=os.path.getsize(ruta))
            return {'ok': True, 'archivo': nombre, 'ruta': ruta, 'causa': ''}
        _alpha67_log_bienestarina('FALLBACK_INVALIDO', unidad=unidad, archivo=nombre, ruta=ruta, motivo=motivo)
        return {'ok': False, 'archivo': None, 'ruta': ruta, 'causa': f'El archivo se intentó generar, pero no pasó validación: {motivo}'}
    except Exception as exc:
        _alpha67_log_bienestarina('GENERACION_FATAL', unidad=unidad, archivo=nombre, ruta=ruta, error=str(exc), traceback=traceback.format_exc())
        return {'ok': False, 'archivo': None, 'ruta': ruta, 'causa': str(exc)}


def _alpha67_error_bienestarina(unidad, mensaje, status=409, extra=None):
    payload = {
        'ok': False,
        'formato': 'bienestarina',
        'error': 'No se pudo generar o descargar la Bienestarina para esta UDS.',
        'mensaje': mensaje,
        'unidad': unidad,
        'unidad_normalizada': normalize_unidad(unidad),
        'output_folder': os.fspath(OUTPUT_FOLDER),
        'ultimos_archivos_generados': _alpha64_listar_ultimos_generados(),
        'log': 'backend/logs/alpha67_bienestarina_real.log',
    }
    if isinstance(extra, dict):
        payload.update(extra)
    _alpha67_log_bienestarina('RESPUESTA_ERROR', **payload)
    try:
        _alpha69_log('bienestarina_auditoria.log', evento='RESPUESTA_ERROR', **payload)
    except Exception:
        pass
    # Alpha69: error controlado para que el frontend no navegue fuera de la plataforma.
    # Se conserva la causa en JSON y en logs, pero se evita 404 genérico.
    return jsonify(payload), 200


# -----------------------------------------------------------------------------
# ALPHA74: generador garantizado de Bienestarina por UDS
# -----------------------------------------------------------------------------
def _alpha74_log_bienestarina(evento, **datos):
    """Log aislado para el fix real de descarga/generación de Bienestarina."""
    try:
        ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha74_bienestarina_generacion_garantizada.log')
        os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
        payload.update(datos or {})
        with open(ruta_log, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha74_template_bienestarina_seguro():
    """Devuelve la mejor plantilla de Bienestarina disponible, evitando archivos vacíos."""
    candidatos = []
    try:
        # Reutiliza el detector existente si está disponible.
        if '_alpha67_plantillas_bienestarina' in globals():
            candidatos.extend(_alpha67_plantillas_bienestarina() or [])
    except Exception as exc:
        _alpha74_log_bienestarina('PLANTILLAS_ALPHA67_ERROR', error=str(exc))

    try:
        for ruta in [
            os.path.join(TEMPLATES_FOLDER, 'oficiales', 'plantilla_bienestarina_oficial.xlsx'),
            os.path.join(TEMPLATES_FOLDER, 'oficiales', 'plantilla_bienestarina_oficial_v2026.xlsx'),
            os.path.join(TEMPLATES_FOLDER, 'Bienestarina_20260608212904_BIENESTARINA_DEL_MES_DE_MAYO.xlsx'),
            os.path.join(TEMPLATES_FOLDER, 'plantilla_bienestarina.xlsx'),
        ]:
            candidatos.append(ruta)
        for root, _dirs, files in os.walk(TEMPLATES_FOLDER):
            for fname in files:
                low = str(fname).lower()
                if low.endswith(('.xlsx', '.xlsm')) and 'bienestarina' in low and 'backup' not in low:
                    candidatos.append(os.path.join(root, fname))
    except Exception as exc:
        _alpha74_log_bienestarina('PLANTILLAS_SCAN_ERROR', error=str(exc))

    vistos = set()
    validas = []
    for ruta in candidatos:
        try:
            ruta_abs = os.path.abspath(str(ruta))
            if ruta_abs in vistos:
                continue
            vistos.add(ruta_abs)
            if os.path.exists(ruta_abs) and os.path.getsize(ruta_abs) > 0:
                validas.append(ruta_abs)
        except Exception:
            continue
    _alpha74_log_bienestarina('PLANTILLAS_VALIDAS', total=len(validas), plantillas=validas[:10])
    return validas[0] if validas else None


def _alpha74_usuarios_bienestarina_seguro(unidad):
    """Obtiene exclusivamente la población canónica compartida por los formatos."""
    return _alpha59_obtener_usuarios_unidad(unidad)


def _alpha74_put_cell(ws, ref, value):
    try:
        if value is None:
            value = ''
        ws[ref] = value
    except Exception:
        pass


# Encabezado oficial Bienestarina: exclusivamente filas 1 a 5. Los rangos se
# centralizan para que todas las rutas de generación produzcan el mismo diseño.
BIENESTARINA_HEADER_LAYOUT = {
    'lugar_label': 'A1:B1', 'lugar': 'C1:E1',
    'codigo_label': 'F1:I1', 'codigo_uds': 'J1:M1',
    'mes_consumo': 'N1:Q1', 'anio': 'R1:T1',
    'regional_label': 'A2:B2', 'regional': 'C2:E2',
    'nombre_uds_label': 'F2:I2', 'nombre_uds': 'J2:T2',
    'centro_zonal_label': 'A3:B3', 'centro_zonal': 'C3:E3',
    'responsable_label': 'F3:I3', 'responsable': 'J3:M3',
    'suplente_label': 'N3:P3', 'suplente': 'Q3:T3',
    'municipio_label': 'A4:B4', 'municipio': 'C4:E4',
    'direccion_label': 'F4:I4', 'direccion': 'J4:M4',
    'barrio_label': 'N4:Q4', 'telefono': 'R4:T4',
    'modalidad_label': 'A5:B5', 'modalidad': 'C5:E5',
    'codigo_origen_label': 'F5:I5', 'codigo_origen': 'J5:M5',
    'nombre_origen_label': 'N5:Q5', 'nombre_origen': 'R5:T5',
}


def _alpha75_aplicar_encabezado_bienestarina(ws, metadata):
    """Reorganiza solo A1:T5 y conserva intactas tabla e impresión oficial."""
    from openpyxl.cell.cell import MergedCell
    from openpyxl.styles import Alignment
    from openpyxl.utils.cell import range_boundaries

    valores = {
        'lugar_label': 'LUGAR:',
        'lugar': metadata.get('lugar') or metadata.get('tipo_punto_entrega') or 'PUNTO DE ENTREGA UDS',
        'codigo_label': 'CÓDIGO DEL PUNTO DE ENTREGA O UDS:',
        'codigo_uds': metadata.get('codigo_uds') or metadata.get('codigo_unidad') or '',
        'mes_consumo': f"MES DE CONSUMO: {metadata.get('mes') or ''}",
        'anio': f"AÑO: {metadata.get('año') or metadata.get('anio') or ''}",
        'regional_label': 'REGIONAL:', 'regional': metadata.get('regional') or '',
        'nombre_uds_label': 'NOMBRE PUNTO DE ENTREGA O UDS:',
        'nombre_uds': metadata.get('unidad') or metadata.get('Unidad') or '',
        'centro_zonal_label': 'CENTRO ZONAL:', 'centro_zonal': metadata.get('centro_zonal') or '',
        'responsable_label': 'RESPONSABLE PUNTO DE ENTREGA O UDS:',
        'responsable': metadata.get('docente') or metadata.get('responsable') or '',
        'suplente_label': 'SUPLENTE / PROVEEDOR:', 'suplente': metadata.get('suplente') or metadata.get('proveedor') or '',
        'municipio_label': 'MUNICIPIO:', 'municipio': metadata.get('municipio') or '',
        'direccion_label': 'DIRECCIÓN PUNTO DE ENTREGA O UDS:',
        'direccion': metadata.get('direccion_unidad') or metadata.get('direccion') or '',
        'barrio_label': f"BARRIO: {metadata.get('barrio') or ''}",
        'telefono': f"TELÉFONO: {metadata.get('telefono_unidad') or metadata.get('telefono_docente') or metadata.get('telefono') or ''}",
        'modalidad_label': 'MODALIDAD:', 'modalidad': metadata.get('modalidad') or '',
        'codigo_origen_label': 'CÓDIGO PUNTO DE ENTREGA DE ORIGEN:',
        'codigo_origen': metadata.get('codigo_origen') or metadata.get('codigo_uds') or metadata.get('codigo_unidad') or '',
        'nombre_origen_label': 'NOMBRE PUNTO DE ENTREGA DE ORIGEN:',
        'nombre_origen': metadata.get('unidad_origen') or metadata.get('unidad') or metadata.get('Unidad') or '',
    }

    # Descombinar únicamente rangos que tocan A1:T5. Nunca se interviene desde
    # la fila 8, donde comienza la tabla de beneficiarios.
    for merged in list(ws.merged_cells.ranges):
        if merged.min_row <= 5 and merged.max_row >= 1 and merged.min_col <= 20:
            ws.unmerge_cells(str(merged))

    for campo, rango in BIENESTARINA_HEADER_LAYOUT.items():
        min_col, min_row, max_col, max_row = range_boundaries(rango)
        anchor = ws.cell(min_row, min_col)
        # Se conserva fuente, relleno, borde y formato numérico del ancla oficial;
        # solo se ajusta alineación para lectura completa.
        anchor.value = valores.get(campo, '')
        alineacion = copy.copy(anchor.alignment)
        alineacion.wrap_text = True
        alineacion.vertical = 'center'
        alineacion.horizontal = 'center' if campo.endswith('_label') or campo in {'mes_consumo', 'anio', 'telefono', 'barrio_label'} else 'left'
        alineacion.shrink_to_fit = False
        anchor.alignment = alineacion
        # Extender el estilo oficial del ancla al bloque antes de combinar evita
        # huecos visuales y mantiene fuente, relleno y bordes del encabezado.
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                celda = ws.cell(row, col)
                celda._style = copy.copy(anchor._style)
                celda.alignment = copy.copy(anchor.alignment)
        if min_col != max_col or min_row != max_row:
            ws.merge_cells(rango)

    # Alturas suficientes sin reducir la fuente. La tabla (fila 8 en adelante)
    # mantiene exactamente sus alturas originales.
    for row, height in {1: 34, 2: 34, 3: 42, 4: 42, 5: 42}.items():
        ws.row_dimensions[row].height = max(float(ws.row_dimensions[row].height or 0), height)

    # Comprobación estructural previa al guardado: todo texto largo debe vivir
    # en un rango combinado y todas las celdas no-ancla deben quedar vacías.
    errores = []
    for campo, rango in BIENESTARINA_HEADER_LAYOUT.items():
        min_col, min_row, max_col, max_row = range_boundaries(rango)
        valor = str(ws.cell(min_row, min_col).value or '')
        if len(valor) > 18 and min_col == max_col:
            errores.append(f'{campo}: texto largo sin combinar')
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                if row == min_row and col == min_col:
                    continue
                celda = ws.cell(row, col)
                if not isinstance(celda, MergedCell) and celda.value not in (None, ''):
                    errores.append(f'{campo}: contenido superpuesto en {celda.coordinate}')
    if errores:
        raise ValueError('Encabezado Bienestarina inválido: ' + '; '.join(errores))
    return {'rangos': dict(BIENESTARINA_HEADER_LAYOUT), 'campos': valores}


def _alpha75_actualizar_archivo_bienestarina(ruta, unidad, mes=None, anio=None):
    """Aplica el encabezado permanente a archivos nuevos o ya existentes."""
    if not ruta or not os.path.exists(ruta):
        return False
    wb = load_workbook(ruta, data_only=False, keep_vba=str(ruta).lower().endswith('.xlsm'))
    try:
        usuarios = _alpha59_obtener_usuarios_unidad(unidad)
        metadata = _alpha59_metadata_formato(unidad, usuarios, mes=mes, anio=anio)
        ws = wb['plantilla de bienestarina '] if 'plantilla de bienestarina ' in wb.sheetnames else wb.active
        _alpha75_aplicar_encabezado_bienestarina(ws, metadata)
        wb.save(ruta)
        return True
    finally:
        wb.close()


def _alpha74_generar_bienestarina_garantizada(unidad, mes=None, anio=None):
    """Genera un Excel real de Bienestarina por UDS aunque fallen los generadores previos.

    No modifica plantillas ni base de datos. Copia una plantilla válida si existe; si no,
    crea un libro mínimo descargable con diagnóstico. El objetivo es que la ruta de
    Bienestarina produzca un archivo físico para la UDS solicitada, o reporte causa clara.
    """
    try:
        nombre, mes_int, anio_int = _alpha67_nombre_bienestarina(unidad, mes=mes, anio=anio) if '_alpha67_nombre_bienestarina' in globals() else (secure_filename(f"{_alpha59_slug(unidad)}_BIENESTARINA_{datetime.now().year}_{datetime.now().month:02d}.xlsx"), datetime.now().month, datetime.now().year)
    except Exception:
        mes_int, anio_int = datetime.now().month, datetime.now().year
        nombre = secure_filename(f"{_alpha59_slug(unidad)}_BIENESTARINA_{anio_int}_{mes_int:02d}.xlsx")
    ruta = os.path.join(OUTPUT_FOLDER, nombre)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    try:
        usuarios = _alpha74_usuarios_bienestarina_seguro(unidad)
        if not usuarios:
            # Se genera un archivo diagnóstico mínimo para no devolver 404 genérico,
            # pero se informa claramente que no hubo usuarios.
            _alpha74_log_bienestarina('SIN_USUARIOS_SE_GENERA_DIAGNOSTICO', unidad=unidad, archivo=nombre)

        metadata = {}
        try:
            metadata = _alpha59_metadata_formato(unidad, usuarios, mes=mes_int, anio=anio_int) or {}
        except Exception as exc:
            metadata = {'unidad': unidad, 'mes': mes_int, 'anio': anio_int}
            _alpha74_log_bienestarina('METADATA_ERROR', unidad=unidad, error=str(exc))

        plantilla = _alpha74_template_bienestarina_seguro()
        wb = None
        if plantilla:
            try:
                wb = load_workbook(plantilla, data_only=False, keep_vba=str(plantilla).lower().endswith('.xlsm'))
            except Exception as exc:
                _alpha74_log_bienestarina('LOAD_TEMPLATE_ERROR', unidad=unidad, plantilla=plantilla, error=str(exc), traceback=traceback.format_exc())
                wb = None
        if wb is None:
            from openpyxl import Workbook
            wb = Workbook()
            ws0 = wb.active
            ws0.title = 'Bienestarina'
            ws0.append(['No se encontró plantilla oficial válida; archivo diagnóstico generado.'])

        try:
            ws = wb['plantilla de bienestarina '] if 'plantilla de bienestarina ' in wb.sheetnames else wb[wb.sheetnames[0]]
        except Exception:
            ws = wb.active

        # Encabezado básico, protegido contra celdas combinadas.
        mes_txt = MESES_ES.get(mes_int, str(mes_int)).upper() if 'MESES_ES' in globals() else str(mes_int)
        for ref, value in {
            'C2': metadata.get('regional') or metadata.get('Regional') or 'CHOCÓ',
            'C3': metadata.get('centro_zonal') or metadata.get('CentroZonal') or '',
            'C4': metadata.get('municipio') or metadata.get('Municipio') or '',
            'C5': metadata.get('modalidad') or metadata.get('Modalidad') or '',
            'J1': metadata.get('codigo_uds') or metadata.get('codigo_unidad') or metadata.get('CodigoUnidadServicio') or '',
            'J2': unidad,
            'J3': metadata.get('responsable') or metadata.get('docente') or metadata.get('agente_educativo') or '',
            'J4': metadata.get('direccion') or metadata.get('direccion_unidad') or '',
            'S4': metadata.get('telefono') or metadata.get('telefono_docente') or '',
            'J5': metadata.get('codigo_origen') or metadata.get('codigo_uds') or metadata.get('codigo_unidad') or '',
            'R5': metadata.get('unidad_origen') or unidad,
            'N1': f'MES DE CONSUMO: {mes_txt}',
            'R1': f'AÑO: {anio_int}',
        }.items():
            _alpha74_put_cell(ws, ref, value)

        _alpha75_aplicar_encabezado_bienestarina(ws, {
            **metadata,
            'unidad': unidad,
            'mes': mes_txt,
            'anio': anio_int,
            'año': anio_int,
        })

        # Filas oficiales conocidas: hoja 1 = 1..14, hoja 2 = 15..30 aprox.
        filas = list(range(10, 24)) + list(range(31, 47))
        data_cols = list('ABCDEFGHIJKLMNOPQRS')
        for idx, row in enumerate(filas):
            # limpiar solo valores de captura, conservar estilos; ignorar celdas combinadas.
            for col in data_cols:
                try:
                    ws[f'{col}{row}'] = None
                except Exception:
                    pass
            _alpha74_put_cell(ws, f'A{row}', idx + 1)
            if idx >= len(usuarios):
                continue
            user = usuarios[idx] or {}
            _alpha74_put_cell(ws, f'B{row}', user.get('PrimerNombre') or user.get('primer_nombre') or '')
            _alpha74_put_cell(ws, f'C{row}', user.get('SegundoNombre') or user.get('segundo_nombre') or '')
            _alpha74_put_cell(ws, f'D{row}', user.get('PrimerApellido') or user.get('primer_apellido') or '')
            _alpha74_put_cell(ws, f'E{row}', user.get('SegundoApellido') or user.get('segundo_apellido') or '')
            _alpha74_put_cell(ws, f'F{row}', user.get('TipoDocumento') or user.get('tipo_documento') or '')
            _alpha74_put_cell(ws, f'G{row}', user.get('NUI') or user.get('Documento') or user.get('documento') or '')
            _alpha74_put_cell(ws, f'H{row}', metadata.get('fecha_entrega') or '')
            _alpha74_put_cell(ws, f'I{row}', metadata.get('lote') or metadata.get('lote_bienestarina') or '')
            _alpha74_put_cell(ws, f'J{row}', metadata.get('cantidad') or metadata.get('cantidad_bienestarina') or 1)
            acudiente = str(user.get('Acudiente') or user.get('nombre_acudiente') or '').strip()
            doc_acudiente = str(user.get('DocumentoAcudiente') or user.get('documento_acudiente') or '').strip()
            _alpha74_put_cell(ws, f'Q{row}', f'{acudiente} {doc_acudiente}'.strip())
            _alpha74_put_cell(ws, f'R{row}', user.get('Parentesco') or user.get('parentesco') or '')

        # Conservar sin cambios el área de impresión de la plantilla oficial.
        wb.save(ruta)
        try:
            wb.close()
        except Exception:
            pass

        existe = os.path.exists(ruta) and os.path.getsize(ruta) > 0
        if existe:
            try:
                registrar_archivo_generado_alpha57('bienestarina', unidad, nombre, ruta, mes=mes_int, anio=anio_int, estado='generado_alpha74_garantizado')
            except Exception:
                pass
            _alpha74_log_bienestarina('GENERADO_GARANTIZADO_OK', unidad=unidad, archivo=nombre, ruta=ruta, usuarios=len(usuarios), plantilla=plantilla, size=os.path.getsize(ruta))
            return {'ok': True, 'archivo': nombre, 'ruta': ruta, 'usuarios': len(usuarios), 'plantilla': plantilla}
        _alpha74_log_bienestarina('GENERADO_GARANTIZADO_NO_EXISTE', unidad=unidad, archivo=nombre, ruta=ruta)
        return {'ok': False, 'archivo': None, 'ruta': ruta, 'causa': 'No se pudo guardar el archivo final.'}
    except Exception as exc:
        _alpha74_log_bienestarina('GENERADO_GARANTIZADO_FATAL', unidad=unidad, archivo=nombre, ruta=ruta, error=str(exc), traceback=traceback.format_exc())
        return {'ok': False, 'archivo': None, 'ruta': ruta, 'causa': str(exc)}

@app.route('/api/bienestarina/descargar', methods=['GET'])
def descargar_bienestarina_alpha57():
    """ALPHA71: descarga Bienestarina restaurando el camino funcional legacy.

    Auditoría diferencial: la versión estable descargaba Bienestarina usando
    /api/descargar/<unidad>/plantilla_bienestarina.xlsx. Este endpoint conserva
    la ruta nueva, pero internamente usa la misma búsqueda/generación histórica
    antes de cualquier fallback nuevo, para evitar 404 persistente.

    No toca RPP, RAM/RAN, Base Maestra, login ni plantillas oficiales.
    """
    unidad = request.args.get('unidad') or request.args.get('nombre_unidad') or ''
    archivo = request.args.get('archivo') or request.args.get('filename') or ''

    def log_alpha71(evento, **datos):
        try:
            ruta_log = os.path.join(_project_path('backend'), 'logs', 'alpha71_bienestarina_diff_fix.log')
            os.makedirs(os.path.dirname(ruta_log), exist_ok=True)
            payload = {'fecha': datetime.now().isoformat(timespec='seconds'), 'evento': evento}
            payload.update(datos or {})
            with open(ruta_log, 'a', encoding='utf-8') as fh:
                fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
        except Exception:
            pass

    def error_controlado(mensaje, status=200, **extra):
        payload = {
            'ok': False,
            'formato': 'bienestarina',
            'unidad': unidad,
            'error': mensaje,
            'mensaje': mensaje,
            'output_folder': os.fspath(OUTPUT_FOLDER),
            'ultimos_archivos_generados': _alpha64_listar_ultimos_generados() if '_alpha64_listar_ultimos_generados' in globals() else [],
            'log': 'backend/logs/alpha71_bienestarina_diff_fix.log',
        }
        payload.update(extra or {})
        log_alpha71('RESPUESTA_ERROR', **payload)
        return jsonify(payload), status

    def enviar_con_encabezado_permanente(nombre, ruta):
        """Toda descarga, incluso histórica, sale con el encabezado vigente."""
        mes = request.args.get('mes') or request.args.get('month')
        anio = request.args.get('anio') or request.args.get('año') or request.args.get('year')
        _alpha75_actualizar_archivo_bienestarina(ruta, unidad, mes=mes, anio=anio)
        return send_from_directory(os.path.dirname(ruta), os.path.basename(ruta), as_attachment=True)

    try:
        if not unidad:
            return error_controlado('Debe indicar la UDS para descargar Bienestarina.', status=400)

        log_alpha71('INICIO', unidad=unidad, archivo_parametro=archivo)

        # 1) Si el frontend envía archivo explícito, aceptarlo solo si es real, no plantilla,
        # y corresponde a la UDS solicitada.
        if archivo:
            nombre, ruta = _alpha57_safe_join_output(archivo)
            if nombre and ruta and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                ok = False
                motivo = 'no_validado'
                try:
                    valido, motivo = _alpha63_validar_archivo_descarga(nombre, unidad, 'bienestarina')
                    ok = bool(valido)
                except Exception as exc:
                    motivo = f'validacion_exception:{exc}'
                log_alpha71('ARCHIVO_EXPLICITO', unidad=unidad, archivo=nombre, ruta=ruta, ok=ok, motivo=motivo)
                if ok:
                    return enviar_con_encabezado_permanente(nombre, ruta)

        # 2) Camino legacy estable: buscar como lo hacía la versión que sí descargaba.
        # En esa versión el botón enviaba plantilla_bienestarina.xlsx al descargador genérico.
        formatos_busqueda = ['plantilla_bienestarina.xlsx', 'bienestarina']
        for formato_busqueda in formatos_busqueda:
            try:
                nombre_archivo = buscar_archivo_generado(unidad, formato_busqueda)
                log_alpha71('BUSQUEDA_LEGACY', unidad=unidad, formato_busqueda=formato_busqueda, archivo=nombre_archivo)
                if nombre_archivo:
                    nombre, ruta = _alpha57_safe_join_output(nombre_archivo)
                    if nombre and ruta and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                        valido, motivo = _alpha63_validar_archivo_descarga(nombre, unidad, 'bienestarina')
                        log_alpha71('VALIDACION_LEGACY', unidad=unidad, archivo=nombre, ruta=ruta, valido=valido, motivo=motivo)
                        if valido:
                            return enviar_con_encabezado_permanente(nombre, ruta)
            except Exception as exc:
                log_alpha71('BUSQUEDA_LEGACY_ERROR', unidad=unidad, formato_busqueda=formato_busqueda, error=str(exc), traceback=traceback.format_exc())

        # 3) Generar usando el camino histórico del motor actual antes del fallback manual.
        for formato_generacion in ['plantilla_bienestarina.xlsx', 'bienestarina']:
            try:
                generado = _alpha59_intentar_generar_faltante(unidad, formato_generacion)
                log_alpha71('GENERACION_HISTORICA', unidad=unidad, formato_generacion=formato_generacion, generado=generado)
                if generado:
                    nombre, ruta = _alpha57_safe_join_output(generado)
                    if nombre and ruta and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                        valido, motivo = _alpha63_validar_archivo_descarga(nombre, unidad, 'bienestarina')
                        log_alpha71('VALIDACION_GENERADO_HISTORICO', unidad=unidad, archivo=nombre, ruta=ruta, valido=valido, motivo=motivo)
                        if valido:
                            return enviar_con_encabezado_permanente(nombre, ruta)
            except Exception as exc:
                log_alpha71('GENERACION_HISTORICA_ERROR', unidad=unidad, formato_generacion=formato_generacion, error=str(exc), traceback=traceback.format_exc())

        # 4) Último recurso: fallback aislado de Bienestarina, si existe en esta versión.
        if '_alpha67_generar_bienestarina_para_uds' in globals():
            try:
                resultado = _alpha67_generar_bienestarina_para_uds(unidad)
                log_alpha71('GENERACION_ALPHA67', unidad=unidad, resultado=resultado)
                if isinstance(resultado, dict) and resultado.get('ok') and resultado.get('archivo'):
                    nombre, ruta = _alpha57_safe_join_output(resultado.get('archivo'))
                    if nombre and ruta and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                        valido, motivo = _alpha63_validar_archivo_descarga(nombre, unidad, 'bienestarina')
                        log_alpha71('VALIDACION_ALPHA67', unidad=unidad, archivo=nombre, ruta=ruta, valido=valido, motivo=motivo)
                        if valido:
                            return enviar_con_encabezado_permanente(nombre, ruta)
            except Exception as exc:
                log_alpha71('GENERACION_ALPHA67_ERROR', unidad=unidad, error=str(exc), traceback=traceback.format_exc())

        # 5) ALPHA74: generación garantizada aislada de Bienestarina.
        # Si los caminos históricos no producen archivo, se crea una copia real
        # desde plantilla válida o libro diagnóstico mínimo para evitar 404 genérico.
        try:
            resultado74 = _alpha74_generar_bienestarina_garantizada(unidad)
            log_alpha71('GENERACION_ALPHA74_GARANTIZADA', unidad=unidad, resultado=resultado74)
            if isinstance(resultado74, dict) and resultado74.get('ok') and resultado74.get('archivo'):
                nombre, ruta = _alpha57_safe_join_output(resultado74.get('archivo'))
                if nombre and ruta and os.path.exists(ruta) and os.path.getsize(ruta) > 0:
                    return enviar_con_encabezado_permanente(nombre, ruta)
        except Exception as exc:
            log_alpha71('GENERACION_ALPHA74_ERROR', unidad=unidad, error=str(exc), traceback=traceback.format_exc())

        # 6) Diagnóstico, no 404 genérico. Se devuelve 200 con JSON para que el frontend
        # muestre el mensaje sin salir de la plataforma.
        return error_controlado(
            'No se pudo encontrar ni generar Bienestarina para esta UDS. Revise usuarios cargados, plantilla oficial y log alpha74_bienestarina_generacion_garantizada.log.',
            status=200,
            unidad_normalizada=normalize_unidad(unidad),
            log_alpha74='backend/logs/alpha74_bienestarina_generacion_garantizada.log',
        )
    except Exception as exc:
        log_alpha71('EXCEPCION_FATAL', unidad=unidad, error=str(exc), traceback=traceback.format_exc())
        return error_controlado('Error interno controlado descargando Bienestarina. Revise el log alpha71_bienestarina_diff_fix.log.', status=200, detalle=str(exc))


@app.route('/api/rpp/descargar', methods=['GET'])
def descargar_rpp_por_categoria():
    """ALPHA64: descarga RPP por UDS + grupo exacto, sin error 500 opaco."""
    unidad = request.args.get('unidad') or request.args.get('unidad_id') or request.args.get('nombre_unidad')
    archivo = request.args.get('archivo') or request.args.get('filename') or ''
    grupo_raw = request.args.get('grupo') or request.args.get('grupo_etario_rpp') or request.args.get('formato')
    try:
        grupo = _alpha61_normalizar_grupo_rpp(grupo_raw)
        if not unidad:
            return jsonify({'ok': False, 'formato': 'rpp', 'error': 'Debes indicar la unidad de atención.'}), 400
        if not grupo:
            return jsonify({'ok': False, 'formato': 'rpp', 'error': 'Grupo etario RPP inválido.', 'grupo_recibido': grupo_raw, 'grupos_validos': sorted(GRUPOS_RPP_ALPHA61.keys())}), 400

        tag_esperado = GRUPOS_RPP_ALPHA61[grupo]['archivo_tag']
        _alpha64_log('RPP_ENDPOINT_INICIO', unidad=unidad, grupo_raw=grupo_raw, grupo=grupo, archivo_parametro=archivo, tag_esperado=tag_esperado)

        # Archivo explícito: aceptar solo si coincide UDS + grupo exacto.
        if archivo:
            nombre, ruta = _alpha57_safe_join_output(archivo)
            if nombre and ruta and os.path.exists(ruta) and _alpha61_archivo_valido_por_uds(nombre, unidad) and _alpha61_archivo_tiene_grupo_rpp(nombre, grupo):
                resp, motivo = _alpha64_send_output(nombre, unidad, 'rpp', grupo=grupo)
                if resp:
                    return resp
                _alpha64_log('RPP_ARCHIVO_EXPLICITO_SEND_FALLO', unidad=unidad, grupo=grupo, archivo=nombre, motivo=motivo)
            else:
                _alpha64_log('RPP_ARCHIVO_EXPLICITO_RECHAZADO_ALPHA64', unidad=unidad, grupo=grupo, archivo=archivo, ruta=ruta, existe=bool(ruta and os.path.exists(ruta)))

        nombre_archivo = _alpha61_buscar_archivo_rpp_exacto(unidad, grupo)
        if not nombre_archivo:
            nombre_archivo = _alpha64_generar_rpp_resiliente(unidad, grupo)

        if nombre_archivo:
            resp, motivo = _alpha64_send_output(nombre_archivo, unidad, 'rpp', grupo=grupo)
            if resp:
                _alpha64_log('RPP_DESCARGA_OK_ALPHA64', unidad=unidad, grupo=grupo, archivo=nombre_archivo)
                return resp
            _alpha64_log('RPP_DESCARGA_VALIDACION_FALLO_ALPHA64', unidad=unidad, grupo=grupo, archivo=nombre_archivo, motivo=motivo)

        payload = {
            'ok': False,
            'formato': 'rpp',
            'error': 'No existe un RPP generado para esa unidad y grupo exactos.',
            'mensaje': 'No se descargará un RPP de otro grupo o de otra UDS. Genere nuevamente el grupo solicitado.',
            'unidad': unidad,
            'grupo_recibido': grupo_raw,
            'grupo_normalizado': grupo,
            'tag_esperado': tag_esperado,
            # TenantPath protege el aislamiento por fundación, pero no es un
            # tipo JSON. Convertirlo a texto evita transformar un 404 funcional
            # (RPP aún no generado) en un 409 por serialización.
            'output_folder': os.fspath(OUTPUT_FOLDER),
            'ultimos_archivos_generados': _alpha64_listar_ultimos_generados(),
        }
        _alpha64_log('RPP_DESCARGA_NO_DISPONIBLE_ALPHA64', **payload)
        return jsonify(payload), 404
    except Exception as exc:
        _alpha64_log('RPP_ENDPOINT_EXCEPTION_ALPHA64', unidad=unidad, grupo_recibido=grupo_raw, error=str(exc), traceback=traceback.format_exc())
        return jsonify({
            'ok': False,
            'formato': 'rpp',
            'error': 'Error interno controlado descargando RPP.',
            'mensaje': 'La plataforma no se cerró. Revise backend/logs/alpha64_descargas_rpp_bienestarina.log para ver el detalle.',
            'unidad': unidad,
            'grupo_recibido': grupo_raw,
            'detalle': str(exc),
        }), 409


@app.route('/api/descargar-archivo/<nombre_archivo>', methods=['GET'])
def descargar_archivo_generico(nombre_archivo):
    nombre_seguro = secure_filename(nombre_archivo)
    ruta = os.path.join(OUTPUT_FOLDER, nombre_seguro)
    if os.path.exists(ruta):
        return send_from_directory(OUTPUT_FOLDER, nombre_seguro, as_attachment=True)
    return jsonify({'error': 'Archivo no encontrado.'}), 404


def detectar_ip_local():
    """Detecta la IP local más útil para compartir en la misma red WiFi."""
    try:
        sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        sock.settimeout(0.2)
        sock.connect(('8.8.8.8', 80))
        ip = sock.getsockname()[0]
        sock.close()
        if ip and not ip.startswith('127.'):
            return ip
    except Exception:
        pass
    try:
        hostname = socket.gethostname()
        for ip in socket.gethostbyname_ex(hostname)[2]:
            if ip and not ip.startswith('127.'):
                return ip
    except Exception:
        pass
    return '127.0.0.1'


def modo_acceso_actual(host):
    """Determina el modo de acceso sin confundir Railway con red local."""
    if str(app.config.get('APP_ENV') or '').lower() == 'production':
        return 'RAILWAY_PUBLICO' if resolver_url_publica() else 'PRODUCCION_SIN_DOMINIO'
    modo_env = str(os.environ.get('SERVER_MODE', '')).strip().upper()
    if modo_env:
        return modo_env
    if host in {'0.0.0.0', '::'}:
        return 'RED_LOCAL'
    return 'LOCAL'


def _normalizar_url_publica(valor):
    raw = str(valor or '').strip()
    if not raw:
        return ''
    if not re.match(r'^https?://', raw, flags=re.I):
        raw = 'https://' + raw.lstrip('/')
    return raw.rstrip('/')


def resolver_url_publica():
    """Resuelve el enlace público estable de Railway sin exponer localhost."""
    candidatos = [
        app.config.get('PUBLIC_APP_URL'),
        os.environ.get('PUBLIC_APP_URL'),
        app.config.get('FRONTEND_ORIGIN'),
        os.environ.get('FRONTEND_ORIGIN'),
        app.config.get('RAILWAY_PUBLIC_DOMAIN'),
        os.environ.get('RAILWAY_PUBLIC_DOMAIN'),
    ]
    for valor in candidatos:
        url = _normalizar_url_publica(valor)
        if url and not re.search(r'(^|//)(localhost|127\.0\.0\.1)(:|/|$)', url, flags=re.I):
            return url
    if has_request_context():
        try:
            url = _normalizar_url_publica(request.url_root)
            if url and not re.search(r'(^|//)(localhost|127\.0\.0\.1)(:|/|$)', url, flags=re.I):
                return url
        except Exception:
            pass
    return ''


def leer_enlace_publico_tunel():
    """Compatibilidad histórica: Railway tiene prioridad sobre túneles locales."""
    url_railway = resolver_url_publica()
    if url_railway:
        return url_railway
    posibles = [
        _project_path('ENLACE_PUBLICO_TUNEL.txt') if '_project_path' in globals() else os.path.abspath(os.path.join(BASE_DIR, '..', 'ENLACE_PUBLICO_TUNEL.txt')),
        os.path.abspath(os.path.join(BASE_DIR, '..', 'ENLACE_PUBLICO_TUNEL.txt')),
    ]
    patron = re.compile(r'https://(?:[a-zA-Z0-9-]+\.trycloudflare\.com|[a-zA-Z0-9-]+\.ngrok-free\.app|[a-zA-Z0-9-]+\.ngrok\.io)')
    for ruta in posibles:
        try:
            if ruta and os.path.exists(ruta):
                contenido = open(ruta, 'r', encoding='utf-8', errors='ignore').read()
                match = patron.search(contenido)
                if match:
                    return match.group(0).rstrip('/')
        except Exception:
            continue
    return ''


@app.route('/api/acceso/config', methods=['GET'])
def configuracion_acceso_compartido():
    user = usuario_actual() or {}
    if user.get('rol') not in {'SUPERADMIN', 'GERENTE'}:
        return jsonify({'error': 'No tienes permiso para ver la configuración de acceso.'}), 403

    es_produccion = str(app.config.get('APP_ENV') or '').lower() == 'production'
    backend_host = os.environ.get('FLASK_HOST') or os.environ.get('HOST') or '127.0.0.1'
    backend_port = int(os.environ.get('FLASK_PORT') or os.environ.get('PORT') or 5000)
    frontend_port = int(os.environ.get('FRONTEND_PORT') or 8080)
    ip_local = detectar_ip_local()
    modo = modo_acceso_actual(backend_host)
    allowed_origins = os.environ.get('ALLOWED_ORIGINS', '').strip()
    url_publica = resolver_url_publica()
    url_tunel = '' if es_produccion else leer_enlace_publico_tunel()

    frontend_local = '' if es_produccion else f'http://127.0.0.1:{frontend_port}/frontend/index.html'
    frontend_red = '' if es_produccion else f'http://{ip_local}:{frontend_port}/frontend/index.html'
    backend_local = '' if es_produccion else f'http://127.0.0.1:{backend_port}'
    backend_red = '' if es_produccion else f'http://{ip_local}:{backend_port}'
    enlace_principal = url_publica or url_tunel or frontend_red or frontend_local
    tunnel_active = bool(url_tunel and not es_produccion)

    if es_produccion:
        instrucciones = [
            'Comparte únicamente el enlace HTTPS público de Railway mostrado en este panel.',
            'El frontend y la API funcionan bajo el mismo dominio; no uses localhost ni el puerto 5000 desde otros equipos.',
            'Crea una cuenta individual para cada usuario y asigna el rol mínimo necesario.',
            'Comprueba que Railway tenga un volumen persistente montado en /data antes de cargar información operativa.',
        ]
    elif tunnel_active:
        instrucciones = [
            'Comparte el enlace HTTPS trycloudflare.com mostrado en este panel.',
            'Mantén abiertas la ventana del backend y la ventana del túnel durante toda la prueba.',
            'El enlace es temporal y cambia al reiniciar el túnel.',
            'No compartas el usuario SUPERADMIN; crea usuarios individuales de prueba.',
        ]
    else:
        instrucciones = [
            'En este computador usa el enlace local.',
            'En la misma red WiFi usa el enlace de red local y verifica el firewall.',
            'Para probar desde internet, ejecuta INICIAR_PLATAFORMA_TUNEL_ONLINE.bat.',
            'No compartas el usuario administrador; crea un usuario por compañero.',
        ]

    return jsonify({
        'ok': True,
        'entorno': app.config.get('APP_ENV'),
        'esProduccion': es_produccion,
        'modo': modo,
        'ipLocal': '' if es_produccion else ip_local,
        'backendHost': backend_host,
        'backendPort': backend_port,
        'frontendPort': frontend_port,
        'publicAppUrl': url_publica,
        'urlPrincipalCompartir': enlace_principal,
        'backendUrlPublico': url_publica or url_tunel,
        'backendUrlLocal': backend_local,
        'backendUrlRedLocal': backend_red,
        'frontendUrlLocal': frontend_local,
        'frontendUrlRedLocal': frontend_red,
        'urlCompartirWifi': frontend_red,
        # Campo histórico conservado para no romper el frontend previo.
        'urlTunelPublico': url_publica or url_tunel,
        'notaTunel': (
            'Railway publica la plataforma mediante un dominio HTTPS estable.'
            if es_produccion else
            ('Quick Tunnel Cloudflare activo y verificado.' if tunnel_active else
             'Ejecuta INICIAR_PLATAFORMA_TUNEL_ONLINE.bat para generar un enlace temporal de internet.')
        ),
        'tunnelActive': tunnel_active,
        'cors': {
            'allowedOriginsEnv': allowed_origins,
            'origenesActivos': parse_allowed_origins(),
        },
        'login': {
            'usuario': user.get('username'),
            'rol': user.get('rol'),
            'fundacion': user.get('fundacion_nombre'),
            'estado': 'AUTENTICADO',
        },
        'instruccionesRapidas': instrucciones,
    })


@app.route('/api/acceso/ping', methods=['GET'])
def acceso_ping():
    return jsonify({
        'ok': True,
        'backend': 'online',
        'publicAppUrl': resolver_url_publica(),
        'ipLocal': '' if str(app.config.get('APP_ENV') or '').lower() == 'production' else detectar_ip_local(),
        'fecha': datetime.now().isoformat(timespec='seconds')
    })



def _directorio_escribible(path_value):
    path = Path(str(path_value or '')).expanduser()
    try:
        path.mkdir(parents=True, exist_ok=True)
        probe = path / '.primera_infancia_write_probe'
        probe.write_text('ok', encoding='utf-8')
        probe.unlink(missing_ok=True)
        return True
    except Exception:
        return False


def diagnostico_almacenamiento():
    data_dir = Path(str(app.config.get('DATA_DIR') or '')).resolve()
    database_path = Path(str(app.config.get('DATABASE_PATH') or '')).resolve()
    declared_mount = str(os.environ.get('RAILWAY_VOLUME_MOUNT_PATH') or '').strip()
    try:
        database_inside = database_path.is_relative_to(data_dir)
    except Exception:
        database_inside = str(database_path).startswith(str(data_dir) + os.sep)
    dirs = {
        'data': data_dir,
        'templates': Path(str(app.config.get('TEMPLATES_FOLDER') or '')),
        'uploads': Path(str(app.config.get('UPLOAD_FOLDER') or '')),
        'outputs': Path(str(app.config.get('OUTPUT_FOLDER') or '')),
        'backups': Path(str(app.config.get('BACKUPS_FOLDER') or '')),
        'logs': Path(str(app.config.get('LOG_FOLDER') or '')),
    }
    writable = {key: _directorio_escribible(value) for key, value in dirs.items()}
    expected_mount = str(data_dir) == '/data'
    try:
        volume_declared = bool(declared_mount) and Path(declared_mount).resolve() == data_dir
    except Exception:
        volume_declared = False
    volume_status = 'detected' if volume_declared else ('expected_path_unverified' if expected_mount else 'not_detected')
    return {
        'dataDir': str(data_dir),
        'databasePath': str(database_path),
        'databaseInsideDataDir': database_inside,
        'directoriesWritable': writable,
        'allRequiredDirectoriesWritable': all(writable.values()),
        'railwayVolumeMountPath': declared_mount,
        'persistentVolumeDeclared': volume_declared,
        'dataDirTargetsExpectedMount': expected_mount,
        'volumeStatus': volume_status,
        'initializationMarkerPresent': (data_dir / '.primera_infancia_initialized.json').exists(),
        'managedSeedStatePresent': (data_dir / '.primera_infancia_seed_state.json').exists(),
        'requiresRedeployPersistenceTest': True,
        'persistenceVerified': False,
        'nota': (
            'La ruta y la escritura pueden comprobarse automáticamente. La variable de montaje solo confirma el volumen cuando Railway la expone; '
            'la persistencia real debe verificarse creando un registro ficticio y ejecutando un redeploy.'
        )
    }


@app.route('/api/acceso/storage-health', methods=['GET'])
def acceso_storage_health():
    user = usuario_actual() or {}
    if user.get('rol') not in {'SUPERADMIN', 'GERENTE'}:
        return jsonify({'error': 'No tienes permiso para revisar el almacenamiento.'}), 403
    result = diagnostico_almacenamiento()
    status = 200 if result.get('allRequiredDirectoriesWritable') and result.get('databaseInsideDataDir') else 503
    return jsonify({'ok': status == 200, 'storage': result}), status


# ------------------------------------------------------------
# Alpha20: frontend servido por el backend para túneles online.
# ------------------------------------------------------------
def _project_path(*parts):
    base = app.config.get('PROJECT_DIR') or os.path.abspath(os.path.join(BASE_DIR, '..'))
    return os.path.abspath(os.path.join(base, *parts))


@app.route('/')
@app.route('/frontend')
@app.route('/frontend/')
@app.route('/frontend/index.html')
def servir_frontend_index():
    """Sirve la interfaz desde Flask para que un solo túnel publique toda la app.

    Antes el túnel exponía solo el puerto 8080 y el navegador remoto intentaba
    llamar al backend por :5000, lo que falla fuera del computador principal.
    Con esta ruta, Cloudflare/ngrok puede apuntar al puerto 5000 y tanto HTML
    como API quedan en el mismo origen público.
    """
    frontend_dir = _project_path('frontend')
    index_path = os.path.join(frontend_dir, 'index.html')
    if not os.path.exists(index_path):
        return jsonify({'error': 'No se encontró frontend/index.html en la plataforma.'}), 404
    return send_from_directory(frontend_dir, 'index.html')



# =============================================================
# ALPHA69 — Buscador global de beneficiarios y auditoría Bienestarina
# Cambios no invasivos: endpoints nuevos + auditoría de descarga Bienestarina.
# No modifica Base Maestra, RPP, RAM/RAN ni plantillas oficiales.
# =============================================================

def _alpha69_log(nombre, **datos):
    try:
        logs_dir = LOG_FOLDER
        os.makedirs(logs_dir, exist_ok=True)
        ruta = os.path.join(logs_dir, nombre)
        payload = {'fecha': datetime.now().isoformat(timespec='seconds')}
        payload.update(datos or {})
        with open(ruta, 'a', encoding='utf-8') as fh:
            fh.write(json.dumps(payload, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass


def _alpha69_row_dict(row):
    try:
        return dict(row)
    except Exception:
        try:
            return {k: row[k] for k in row.keys()}
        except Exception:
            return {}


def _alpha69_doc_key(valor):
    texto = str(valor or '').strip().replace('.0', '')
    return re.sub(r'\D+', '', texto)


def _alpha69_valor(data, *keys):
    for k in keys:
        try:
            v = data.get(k)
        except Exception:
            v = None
        if v not in (None, ''):
            return v
    return ''


def _alpha69_nombre(data):
    nombre = _alpha69_valor(data, 'nombre_completo', 'NombreCompleto', 'nombre', 'Nombre', 'nombres', 'Nombres')
    apellidos = _alpha69_valor(data, 'apellidos', 'Apellidos')
    if nombre and apellidos and str(apellidos).strip().lower() not in str(nombre).strip().lower():
        return f"{nombre} {apellidos}".strip()
    if nombre:
        return str(nombre).strip()
    partes = [
        _alpha69_valor(data, 'primer_nombre', 'PrimerNombre'),
        _alpha69_valor(data, 'segundo_nombre', 'SegundoNombre'),
        _alpha69_valor(data, 'primer_apellido', 'PrimerApellido'),
        _alpha69_valor(data, 'segundo_apellido', 'SegundoApellido'),
    ]
    return ' '.join(str(x).strip() for x in partes if str(x or '').strip()).strip()


def _alpha69_unidad(data):
    return _alpha69_valor(data, 'unidad', 'Unidad', 'unidad_servicio', 'nombre_unidad', 'unidad_atencion', 'uds', 'UDS', 'nombre_uds')


def _alpha69_documento(data):
    return _alpha69_valor(data, 'documento', 'Documento', 'numero_documento', 'nui', 'NUI', 'identificacion', 'identificación', 'documento_nino', 'documento_normalizado')


def _alpha69_grupo_etario(data):
    grupo = _alpha69_valor(data, 'grupo_etario', 'grupo_edad', 'GrupoEdad', 'grupo', 'grupo_poblacional')
    if grupo:
        return str(grupo)
    try:
        edad_meses = int(float(_alpha69_valor(data, 'edad_meses', 'EdadMeses') or 0))
    except Exception:
        edad_meses = 0
    tipo = normalizar_texto_clave(_alpha69_valor(data, 'tipo_beneficiario', 'TipoBeneficiario', 'estado_gestacion'))
    if 'gestante' in tipo:
        return 'GESTANTE'
    if edad_meses <= 6:
        return '0 A 6 MESES'
    if edad_meses <= 11:
        return '6 A 11 MESES'
    if edad_meses <= 35:
        return '1 A 2 AÑOS 11 MESES'
    if edad_meses <= 71:
        return '3 A 5 AÑOS 11 MESES'
    return ''


def _alpha69_item_beneficiario(data, fuente=''):
    unidad = _alpha69_unidad(data)
    return {
        'fuente': fuente,
        'documento': str(_alpha69_documento(data) or '').strip(),
        'documento_normalizado': _alpha69_doc_key(_alpha69_documento(data)),
        'nui': str(_alpha69_valor(data, 'nui', 'NUI') or '').strip(),
        'tipo_documento': _alpha69_valor(data, 'tipo_documento', 'TipoDocumento', 'tipo_doc'),
        'nombre_completo': _alpha69_nombre(data),
        'fecha_nacimiento': _alpha69_valor(data, 'fecha_nacimiento', 'FechaNacimiento', 'nacimiento'),
        'edad': _alpha69_valor(data, 'edad', 'Edad'),
        'edad_meses': _alpha69_valor(data, 'edad_meses', 'EdadMeses'),
        'grupo_etario': _alpha69_grupo_etario(data),
        'unidad': unidad,
        'unidad_normalizada': normalize_unidad(unidad),
        'codigo_uds': _alpha69_valor(data, 'codigo_unidad', 'codigo_unidad_servicio', 'codigo_uds'),
        'docente': _alpha69_valor(data, 'docente', 'Docente', 'agente_educativo', 'AgenteEducativo', 'agente'),
        'coordinador': _alpha69_valor(data, 'coordinador', 'Coordinador'),
        'estado': _alpha69_valor(data, 'estado', 'Estado'),
        'telefono': _alpha69_valor(data, 'telefono', 'Telefono'),
        'direccion': _alpha69_valor(data, 'direccion', 'Direccion', 'dirección'),
    }


def _alpha69_tablas_disponibles(conn):
    try:
        rows = conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
        return {r[0] for r in rows}
    except Exception:
        return set()


def _alpha69_fetch_table(conn, table, limit=7000):
    try:
        return [_alpha69_row_dict(r) for r in conn.execute(f'SELECT * FROM {table} LIMIT {int(limit)}').fetchall()]
    except Exception as exc:
        _alpha69_log('buscador_global.log', evento='FETCH_TABLE_ERROR', tabla=table, error=str(exc))
        return []


def _alpha69_score_match(item, q_text, q_doc):
    valores = ' '.join(str(item.get(k, '') or '') for k in ['documento','nui','tipo_documento','nombre_completo','unidad','codigo_uds','docente','coordinador','grupo_etario','estado']).lower()
    valores_norm = normalizar_texto_clave(valores)
    if q_doc and (q_doc in _alpha69_doc_key(item.get('documento')) or q_doc in _alpha69_doc_key(item.get('nui'))):
        return 100
    if q_text and q_text in valores_norm:
        return 80
    # Coincidencia por tokens para nombres parciales.
    tokens = [t for t in q_text.split() if len(t) >= 2]
    if tokens and all(t in valores_norm for t in tokens):
        return 60
    return 0


def _alpha69_buscar_beneficiarios(q='', limit=20):
    q = str(q or '').strip()
    q_norm = normalizar_texto_clave(q)
    q_doc = _alpha69_doc_key(q)
    resultados = []
    vistos = set()
    conn = None
    try:
        conn = get_db_connection()
        tenant_id = int(fundacion_actual_id() or 1)
        try:
            rows = conn.execute(
                "SELECT * FROM master_ninos WHERE activo = 1 AND COALESCE(fundacion_id, 1) = ? ORDER BY nombre_completo, documento",
                (tenant_id,),
            ).fetchall()
        except Exception as exc:
            _alpha69_log('buscador_global.log', evento='MASTER_NINOS_QUERY_ERROR', fundacion_id=tenant_id, error=str(exc))
            rows = []
        for row in rows:
            data = _alpha69_row_dict(row)
            item = _alpha69_item_beneficiario(data, fuente='master_ninos')
            score = _alpha69_score_match(item, q_norm, q_doc)
            if score <= 0:
                continue
            key = item.get('documento_normalizado') or normalizar_texto_clave(item.get('nombre_completo')) + '|' + normalizar_texto_clave(item.get('unidad'))
            if key in vistos:
                continue
            vistos.add(key)
            item['score'] = score
            resultados.append(item)
        resultados.sort(key=lambda x: (-int(x.get('score') or 0), x.get('nombre_completo') or ''))
        return resultados[:int(limit)]
    finally:
        try:
            conn.close()
        except Exception:
            pass


def _alpha69_ficha_beneficiario(documento):
    base = None
    coincidencias = _alpha69_buscar_beneficiarios(documento, limit=10)
    doc_key = _alpha69_doc_key(documento)
    for item in coincidencias:
        if doc_key and item.get('documento_normalizado') == doc_key:
            base = item
            break
    if base is None and coincidencias:
        base = coincidencias[0]
    if not base:
        return None
    salud, asistencia, formatos, alertas, equipo = [], [], [], [], []
    conn = None
    try:
        conn = get_db_connection()
        # Salud/nutrición/peso y talla.
        for table in ['master_salud_nutricion', 'sn_valoraciones', 'peso_talla']:
            for data in _alpha69_fetch_table(conn, table):
                d = _alpha69_doc_key(_alpha69_valor(data, 'documento_nino','documento','Documento','nui','NUI','documento_normalizado'))
                if doc_key and d == doc_key:
                    salud.append({
                        'fuente': table,
                        'peso': _alpha69_valor(data, 'peso', 'Peso'),
                        'talla': _alpha69_valor(data, 'talla', 'Talla'),
                        'diagnostico': _alpha69_valor(data, 'diagnostico_nutricional', 'diagnostico', 'estado_nutricional'),
                        'vacunas': _alpha69_valor(data, 'vacunas', 'esquema_vacunacion'),
                        'crecimiento_desarrollo': _alpha69_valor(data, 'crecimiento_desarrollo', 'cy_d'),
                        'perimetro_braquial': _alpha69_valor(data, 'perimetro_braquial'),
                        'fecha': _alpha69_valor(data, 'fecha_valoracion', 'fecha')
                    })
        # Talento humano/equipo por unidad.
        unidad_key = normalizar_texto_clave(base.get('unidad_normalizada') or base.get('unidad'))
        for table in ['master_talento_humano']:
            for data in _alpha69_fetch_table(conn, table):
                u = normalizar_texto_clave(_alpha69_valor(data, 'unidad', 'unidad_servicio', 'nombre_unidad', 'uds'))
                if unidad_key and (unidad_key == u or unidad_key in u or u in unidad_key):
                    equipo.append({
                        'fuente': table,
                        'nombre': _alpha69_nombre(data),
                        'documento': _alpha69_documento(data),
                        'rol': _alpha69_valor(data, 'rol', 'cargo', 'perfil'),
                        'unidad': _alpha69_unidad(data),
                        'telefono': _alpha69_valor(data, 'telefono', 'Telefono'),
                    })
        # Archivos relacionados.
        try:
            unidad_slug = _alpha59_slug(base.get('unidad') or '')
            if os.path.isdir(OUTPUT_FOLDER):
                for fname in sorted(os.listdir(OUTPUT_FOLDER), reverse=True):
                    slug = _alpha59_slug(fname)
                    if unidad_slug and unidad_slug in slug and fname.lower().endswith(('.xlsx','.xlsm','.pdf')):
                        formatos.append(fname)
                        if len(formatos) >= 12:
                            break
        except Exception:
            pass
    except Exception as exc:
        alertas.append({'tipo': 'ERROR_FICHA', 'mensaje': str(exc)})
        _alpha69_log('buscador_global.log', evento='FICHA_ERROR', documento=documento, error=str(exc), traceback=traceback.format_exc())
    finally:
        try:
            conn.close()
        except Exception:
            pass
    if not salud:
        alertas.append({'tipo': 'SIN_SALUD_NUTRICION', 'mensaje': 'No se encontró información de salud/nutrición asociada.'})
    if not equipo:
        alertas.append({'tipo': 'SIN_EQUIPO_UDS', 'mensaje': 'No se encontró talento humano asociado a la UDS.'})
    return {
        'beneficiario': base,
        'salud_nutricion': salud[:10],
        'equipo_interdisciplinario': equipo[:20],
        'asistencia': asistencia,
        'formatos_relacionados': formatos,
        'alertas': alertas,
        'acciones': {
            'base_maestra': f"/api/base-maestra/diagnostico-nino?documento={base.get('documento_normalizado') or base.get('documento')}",
            'salud_nutricion': f"/api/base-maestra/modulos/salud-nutricion?documento={base.get('documento_normalizado') or base.get('documento')}",
            'bienestarina': f"/api/bienestarina/descargar?unidad={base.get('unidad')}",
            'ram': f"/api/descargar/{base.get('unidad')}/ram",
        }
    }


@app.route('/api/buscador/beneficiarios', methods=['GET'])
def api_alpha69_buscar_beneficiarios():
    q = request.args.get('q') or request.args.get('query') or ''
    try:
        if len(str(q).strip()) < 2:
            return jsonify({'ok': True, 'query': q, 'total': 0, 'resultados': [], 'mensaje': 'Escriba al menos 2 caracteres.'})
        resultados = _alpha69_buscar_beneficiarios(q, limit=int(request.args.get('limit') or 20))
        return jsonify({'ok': True, 'query': q, 'total': len(resultados), 'resultados': resultados})
    except Exception as exc:
        _alpha69_log('buscador_global.log', evento='BUSCAR_ENDPOINT_ERROR', q=q, error=str(exc), traceback=traceback.format_exc())
        return jsonify({'ok': False, 'error': 'No se pudo ejecutar el buscador global.', 'detalle': str(exc)}), 200


@app.route('/api/buscador/beneficiarios/<path:documento>', methods=['GET'])
def api_alpha69_ficha_beneficiario(documento):
    try:
        ficha = _alpha69_ficha_beneficiario(documento)
        if not ficha:
            return jsonify({'ok': False, 'error': 'No se encontró beneficiario con ese documento o criterio.', 'documento': documento}), 200
        return jsonify({'ok': True, 'ficha': ficha})
    except Exception as exc:
        _alpha69_log('buscador_global.log', evento='FICHA_ENDPOINT_ERROR', documento=documento, error=str(exc), traceback=traceback.format_exc())
        return jsonify({'ok': False, 'error': 'No se pudo cargar la ficha del beneficiario.', 'detalle': str(exc)}), 200


@app.route('/api/bienestarina/auditoria', methods=['GET'])
def api_alpha69_auditoria_bienestarina():
    unidad = request.args.get('unidad') or ''
    try:
        plantillas = _alpha67_plantillas_bienestarina() if '_alpha67_plantillas_bienestarina' in globals() else []
        usuarios = _alpha59_obtener_usuarios_unidad(unidad) if unidad else []
        archivo = _alpha67_buscar_bienestarina_exacta(unidad) if unidad and '_alpha67_buscar_bienestarina_exacta' in globals() else None
        ruta = os.path.join(OUTPUT_FOLDER, secure_filename(os.path.basename(archivo))) if archivo else ''
        payload = {
            'ok': True,
            'unidad': unidad,
            'unidad_normalizada': normalize_unidad(unidad),
            'usuarios_encontrados': len(usuarios),
            'plantilla_encontrada': bool(plantillas),
            'plantillas': plantillas[:5],
            'archivo_generado_detectado': archivo,
            'ruta_archivo': ruta,
            'existe_archivo': bool(ruta and os.path.exists(ruta)),
            'size_bytes': os.path.getsize(ruta) if ruta and os.path.exists(ruta) else 0,
            'output_folder': os.fspath(OUTPUT_FOLDER),
        }
        _alpha69_log('bienestarina_auditoria.log', evento='AUDITORIA', **payload)
        return jsonify(payload)
    except Exception as exc:
        _alpha69_log('bienestarina_auditoria.log', evento='AUDITORIA_ERROR', unidad=unidad, error=str(exc), traceback=traceback.format_exc())
        return jsonify({'ok': False, 'error': 'No se pudo auditar Bienestarina.', 'detalle': str(exc), 'unidad': unidad}), 200


@app.route('/frontend/<path:filename>')
def servir_frontend_archivo(filename):
    frontend_dir = _project_path('frontend')
    return send_from_directory(frontend_dir, filename)


@app.route('/css/<path:filename>')
def servir_css(filename):
    return send_from_directory(_project_path('frontend', 'css'), filename)


@app.route('/js/<path:filename>')
def servir_js(filename):
    return send_from_directory(_project_path('frontend', 'js'), filename)


@app.route('/docs/<path:filename>')
def servir_docs(filename):
    return send_from_directory(_project_path('docs'), filename)


@app.route('/assets/<path:filename>')
def servir_assets(filename):
    return send_from_directory(_project_path('frontend', 'assets'), filename)


@app.route('/<path:client_path>', methods=['GET'])
def servir_ruta_frontend(client_path):
    """Fallback de navegación para enlaces directos de la interfaz web.

    Railway y los navegadores pueden abrir rutas como ``/login`` o
    ``/dashboard`` directamente. La aplicación es una SPA y debe entregar su
    index en esos casos. Rutas API o archivos inexistentes conservan un 404
    explícito para no ocultar endpoints mal escritos ni devolver HTML como JS.
    """
    normalized = str(client_path or '').strip('/')
    if normalized == 'api' or normalized.startswith('api/'):
        return jsonify({'error': 'Endpoint no encontrado.', 'path': f'/{normalized}'}), 404
    if normalized.rsplit('/', 1)[-1].find('.') >= 0:
        return jsonify({'error': 'Archivo no encontrado.', 'path': f'/{normalized}'}), 404
    frontend_dir = _project_path('frontend')
    index_path = os.path.join(frontend_dir, 'index.html')
    if not os.path.isfile(index_path):
        return jsonify({'error': 'No se encontró frontend/index.html en la plataforma.'}), 404
    return send_from_directory(frontend_dir, 'index.html')


@app.errorhandler(405)
def metodo_no_permitido(error):
    """Respuesta útil y segura cuando la URL existe pero el método no coincide."""
    path = str(request.path or '/')
    allowed = sorted(set(getattr(error, 'valid_methods', None) or []))
    app.logger.warning(
        'Método HTTP no permitido method=%s path=%s allowed=%s',
        request.method,
        path,
        ','.join(allowed),
    )
    if path == '/' or path in {'/frontend', '/frontend/', '/frontend/index.html', '/login', '/dashboard'}:
        # Un formulario nativo puede enviarse antes de que app.js instale su
        # preventDefault. PRG evita repetir el POST y vuelve a cargar la SPA.
        return redirect('/', code=303)
    if path == '/api' or path.startswith('/api/'):
        response = jsonify({
            'error': 'Método HTTP no permitido para este endpoint.',
            'method': request.method,
            'path': path,
            'allowed_methods': allowed,
        })
        response.status_code = 405
        if allowed:
            response.headers['Allow'] = ', '.join(allowed)
        return response
    return jsonify({
        'error': 'Método HTTP no permitido.',
        'method': request.method,
        'path': path,
        'allowed_methods': allowed,
    }), 405


@app.route('/api/acceso/tunel-info', methods=['GET'])
def acceso_tunel_info():
    """Compatibilidad de diagnóstico: informa primero el dominio público de Railway."""
    es_produccion = str(app.config.get('APP_ENV') or '').lower() == 'production'
    url_publica = resolver_url_publica()
    return jsonify({
        'ok': True,
        'modo': modo_acceso_actual(app.config.get('FLASK_HOST', '127.0.0.1')),
        'publicTunnelMode': False if es_produccion else os.environ.get('PUBLIC_TUNNEL_MODE', '').strip().lower() in {'1', 'true', 'si', 'sí'},
        'frontendServidoPorBackend': True,
        'publicAppUrl': url_publica,
        'urlLocalUnificada': '' if es_produccion else 'http://127.0.0.1:5000/frontend/index.html',
        'enlacePublicoActual': url_publica or leer_enlace_publico_tunel(),
        'mensaje': (
            'Comparta el dominio HTTPS de Railway; frontend y API usan el mismo origen.'
            if es_produccion else
            'En equipos remotos use un enlace HTTPS público, nunca 127.0.0.1.'
        )
    })


@app.errorhandler(Exception)
def manejar_error_global(exc):
    """Respuesta JSON controlada y reporte técnico no vacío.

    El detalle se guarda en la carpeta operativa real ``data/logs`` (o en el
    volumen configurado), no en el marcador vacío ``backend/logs``. Además se
    imprime en la consola como respaldo si Windows impide escribir el archivo.
    """
    if isinstance(exc, HTTPException):
        if has_request_context() and str(request.path or '').startswith('/api/'):
            return jsonify({'error': exc.description or exc.name, 'codigo': exc.code}), exc.code
        return exc

    trace_id = datetime.now().strftime('%Y%m%d%H%M%S%f')
    report = write_exception_report(
        exc,
        trace_id,
        app.config,
        request_obj=request if has_request_context() else None,
        g_obj=g if has_request_context() else None,
    )
    try:
        app.logger.error(
            'Error API trace_id=%s instance_id=%s log=%s',
            trace_id,
            report.get('instance_id'),
            report.get('reference') or 'stderr',
            exc_info=(type(exc), exc, exc.__traceback__),
        )
    except Exception:
        pass

    if has_request_context() and str(request.path or '').startswith('/api/'):
        payload = {
            'error': 'Error técnico del servidor.',
            'code': 'INTERNAL_SERVER_ERROR',
            'trace_id': trace_id,
            'instance_id': report.get('instance_id'),
        }
        # En desarrollo local/túnel se muestra únicamente una referencia
        # relativa; nunca la ruta absoluta del computador o del volumen.
        if str(app.config.get('APP_ENV') or '').lower() != 'production':
            payload['log_file'] = report.get('reference') or 'consola del backend'
            payload['log_written'] = bool(report.get('written'))
        response = jsonify(payload)
        response.status_code = 500
        response.headers['X-Trace-Id'] = trace_id
        return response
    raise exc

# Última barrera de aislamiento: se instala después de registrar y migrar
# todos los módulos, de modo que las consultas SQLite ejecutadas durante una
# petición multi-fundación queden automáticamente acotadas o fallen cerradas.
try:
    from modules.seguridad.tenant_sql_guard import install_sqlite_tenant_guard
    install_sqlite_tenant_guard()
except Exception as exc:
    if not bool(app.config.get('SINGLE_TENANT_MODE', True)):
        raise RuntimeError('No se pudo activar el cortafuegos SQL multi-fundación.') from exc
    print(f'Cortafuegos SQL multi-fundación no pudo activarse: {exc}')

# Alias WSGI compatible. Producción debe importar backend/wsgi.py.
application = app


if __name__ == '__main__':
    local_app = create_app(os.environ.get('APP_ENV') or os.environ.get('FLASK_ENV'))
    init_db()
    flask_host = local_app.config.get('FLASK_HOST', '127.0.0.1')
    flask_port = int(local_app.config.get('FLASK_PORT', 5000))
    debug = bool(local_app.config.get('DEBUG', False))
    print(f'PrimeraInfancia backend escuchando en http://{flask_host}:{flask_port}')
    local_app.run(host=flask_host, port=flask_port, debug=debug, use_reloader=False)
